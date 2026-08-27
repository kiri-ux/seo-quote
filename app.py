#!/usr/bin/env python3
"""
SSG / adtini — SEO Quote Tool (Render-ready Flask app)

Partner fills the product form -> backend runs the keyword pull, rank check,
and pricing formula against DataForSEO -> quote renders on screen with the full
staged breakdown for a human to review before sending.

ENV (set in Render dashboard -> Environment):
    DFS_LOGIN      DataForSEO account email
    DFS_PASSWORD   DataForSEO API password (from dashboard, not portal login)

Local run:
    pip install -r requirements.txt
    DFS_LOGIN=... DFS_PASSWORD=... python app.py
    -> http://localhost:5000
"""
import os, json, base64, statistics, time, re, threading, io, hashlib
import html
from concurrent.futures import ThreadPoolExecutor
from html.parser import HTMLParser
import requests
from flask import Flask, render_template, request, jsonify, send_file
import time as _time
import storage

app = Flask(__name__)

def _json_error_guard(fn):
    """Any unhandled exception in an API route produces Flask's HTML error page,
    which the frontend can only report as the opaque 'Server 500 (timeout or
    non-JSON)'. This wrapper returns the real cause as JSON instead, turning
    "it broke" into a fixable report, and logs the traceback for Render.

    Written for the save routes and originally applied only there — which meant
    a failure anywhere in the quoting pipeline was still a black box
    (2026-07-26). It is now applied to every /api/ route. abort()/HTTP errors
    are re-raised untouched so 404s stay 404s.
    """
    from functools import wraps
    from werkzeug.exceptions import HTTPException

    @wraps(fn)
    def inner(*a, **k):
        try:
            return fn(*a, **k)
        except HTTPException:
            raise
        except Exception as e:
            import traceback
            traceback.print_exc()
            try:
                app.logger.exception("API route failed: %s", fn.__name__)
            except Exception:
                pass
            return jsonify({"error": f"{type(e).__name__}: {e}",
                            "route": fn.__name__}), 500
    return inner



# Build stamp shown in the header — derives from the deploy commit so it
# updates automatically with every GitHub upload (falls back to boot date).
import datetime as _dt
BUILD_ID = (os.environ.get("RENDER_GIT_COMMIT", "")[:7]
            or _dt.datetime.utcnow().strftime("dev-%m%d"))

# SOURCE FINGERPRINT (2026-07-27).
# The commit hash and the deploy time both describe the DEPLOY, not the code —
# so there was no way to tell, before opening the app, whether the files you
# uploaded are the ones running. Re-uploading the same file, or dropping one in
# the wrong folder, still produces a fresh hash and a fresh timestamp. This
# hashes the actual file contents instead: the same files always produce the
# same six characters, and any change to any of them produces different ones.
# Whoever hands over a build can state the expected value in advance.
# reputation.html added 2026-08-04: it was the one shipped template NOT covered,
# so a rep-mgmt-only change left the fingerprint identical and there was no way
# to confirm from the header that the deploy had taken. rep_pricing/rep_scan
# included for the same reason — they carry the rep quote's actual maths.
FINGERPRINT_FILES = ("app.py", "storage.py", "templates/index.html",
                     "templates/reputation.html", "rep_pricing.py", "rep_scan.py")

def _source_fingerprint():
    import hashlib
    here = os.path.dirname(os.path.abspath(__file__))
    h = hashlib.sha256()
    for rel in FINGERPRINT_FILES:
        try:
            with open(os.path.join(here, rel), "rb") as fh:
                # Normalise line endings so a checkout on a different platform
                # doesn't change the answer for identical content.
                h.update(fh.read().replace(b"\r\n", b"\n"))
        except Exception:
            h.update(b"<missing>")
        h.update(b"\x00")
    return h.hexdigest()[:6]

SOURCE_FP = _source_fingerprint()


def source_file_hashes():
    """Per-file hash for every fingerprinted file.

    The combined `src` hash answers "is this the build I was given" and nothing
    else. When it disagrees it cannot say WHICH file failed to land, so a handover
    turns into re-uploading everything and hoping — twice, on 2026-08-10, on a
    one-file change that was correct in the zip and never reached the app.
    Six hashes make the stale file obvious at a glance.
    """
    import hashlib
    here = os.path.dirname(os.path.abspath(__file__))
    out = []
    for rel in FINGERPRINT_FILES:
        path = os.path.join(here, rel)
        try:
            with open(path, "rb") as fh:
                data = fh.read().replace(b"\r\n", b"\n")
            out.append({"file": rel,
                        "sha": hashlib.sha256(data).hexdigest()[:6],
                        "bytes": len(data),
                        "mtime": _dt.datetime.utcfromtimestamp(
                            os.path.getmtime(path)).strftime("%Y-%m-%d %H:%M:%SZ")})
        except Exception as e:
            out.append({"file": rel, "sha": "MISSING", "bytes": 0,
                        "mtime": "", "error": str(e)[:80]})
    return out


def model_is_snapshot(model_id):
    """Is this model ID a FIXED snapshot, or an alias that can move under us?

    Per Anthropic's model-versioning docs: from the 4.6 generation onward the
    dateless ID *is* the canonical snapshot — claude-sonnet-4-6 does not float.
    Only pre-4.6 dateless IDs (claude-sonnet-4-5 and earlier) are convenience
    aliases resolving to the latest dated snapshot. Anything carrying an
    explicit date is pinned by definition.
    """
    import re as _re
    mid = (model_id or "").strip().lower()
    if _re.search(r"-\d{8}$", mid):
        return True                                   # explicit dated snapshot
    m = _re.match(r"^claude-(?:sonnet|opus|haiku|fable|mythos)-(\d+)(?:-(\d+))?$", mid)
    if not m:
        return False                                  # unrecognised: treat as unpinned
    return (int(m.group(1)), int(m.group(2) or 0)) >= (4, 6)
def _build_stamp():
    """Build time in US Eastern (EST/EDT handled by the tzdb). Falls back to a
    fixed -05:00 if the container image ships without tzdata."""
    try:
        from zoneinfo import ZoneInfo
        now = _dt.datetime.now(ZoneInfo("America/New_York"))
        label = now.strftime("%Z") or "ET"
    except Exception:
        now = _dt.datetime.now(_dt.timezone(_dt.timedelta(hours=-5)))
        label = "ET"
    return (f"build {now.strftime('%Y-%m-%d %I:%M %p').lstrip('0')} {label} "
            f"\u00b7 {BUILD_ID} \u00b7 src {SOURCE_FP}")

BUILD_STR = _build_stamp()
BASE = "https://api.dataforseo.com/v3"

# ---------------------------------------------------------------------------
# CONFIG — every tunable constant. Brendan-calibration items live here only.
# Spring ladder ($1,450–$4,250 by geo scope), per decision.
# ---------------------------------------------------------------------------
CFG = {
    # Geo dropdown (5 options) -> 4 price anchors.
    # Non-contiguous shares the $2,950 (multi-region) anchor with statewide.
    # HARD COST anchors = CEIL50(0.65 × former client anchor). All internal
    # calculations start from hard cost; client price = hard × (1 + markup).
    # HARD COST anchors = CEIL50(client anchor / 1.35). Client anchors blended
    # from the spring ladder uplifted toward the June ~$3,950 pricing. No floor —
    # the raised bases carry the new pricing level directly.
    # Calibrated 2026-07-20 against Brendan's three actuals (Keller Builds,
    # Red Shoes, Waytek): anchors trimmed $250 and the tier step flattened, which
    # lands the formula within ~0-5% of all nine quoted tier prices.
    # Media Venue datapoint (2026-07-20, RFP bid): Brendan $2,925/$4,040/$5,150
    # vs formula $3,450/$4,400/$5,350 (+18/+9/+4%). His base sits BELOW his own
    # $2,950 card and his steps run ~$1,110 (vs his usual ~$1,000) — consistent
    # with a sharpened competitive-RFP base. Root cause of the +18%: the top-20
    # rank check scored his page-3-5 footholds as "not ranking" and fired the
    # +14% zero-ranking uplift. Fix: top-N deepened to 100 (see below) — without
    # the uplift the formula lands $3,037/$3,983/$4,928, within ~4% per tier.
    # ================= EVERY DOLLAR BELOW IS PARTNER COST =================
    # Converted 2026-08-13. These used to be a "calibration basis" — an
    # intermediate unit, back-solved from Brendan's quoted CLIENT prices, that
    # was multiplied by 0.8775 to reach partner cost and then divided by
    # (1 - margin) to reach retail. Nobody could reason about it: the doc said
    # "a flat $750 on partner cost" when 750 was basis and $650 was the partner
    # figure. Each constant is now the partner-cost dollar itself, and the
    # client price is that divided by (1 - margin). The conversion was applied
    # as value x 0.8775 rounded to $50, then verified against 1,875 input
    # combinations and the thirteen-proposal bench.
    "geo_anchor": {
        # single_city raised to match contiguous after the Dental Excellence
        # datapoint (2026-07-20): Brendan's single-city Philadelphia quote was
        # his HIGHEST base ($3,350) — he prices the market, not the pin count.
        # A genuinely tiny single-town client may deserve less; no datapoint
        # yet — use the manual hard-base override until one exists.
        "single_city":          1850,
        "contiguous_region":    1850,
        "non_contiguous_region":2050,
        "statewide":            2050,
        # RECALIBRATED 2026-07-25 (2,900 -> 2,050). The 2,900 figure was
        # back-solved from Brendan's national card WITH the extras muted, so
        # it silently contained the volume add and the zero-ranking uplift —
        # it WAS the "national client with nothing ranking" price, not the
        # starting point. Now that extras are live (Brendan: volume,
        # competition and visibility are what separate national clients), the
        # anchor has to be the bare floor those signals build up FROM, or the
        # scope is charged twice. Validated on both national actuals:
        #   Skidmore (adder 50, vol 24k, 90% not ranking)
        #     -> 4,000 / 5,450 / 6,950 vs actual 3,950 / 5,450 / 6,950
        #   MPG      (adder  0, vol 25k+, 100% not ranking)
        #     -> 3,900 / 5,400 / 6,900 vs actual 3,950 / 5,450 / 6,950
        # An established national brand that already ranks now prices BELOW
        # the card, which is the behaviour Brendan described and the old
        # constant made impossible.
        "nationwide":           1800,
    },
    "competitive_adder": {0: 0, 1: 150, 2: 250},   # partner $. FLAT fallback (used when no bid data)
    "bid_score_breaks": [5.0, 15.0],          # <5->0, 5-15->1, >=15->2 (for the fallback)
    # Organic-difficulty breaks, used ONLY when no bid data exists anywhere (a
    # Google Ads restricted vertical). KD scores onto the SAME 0/1/2 ladder as
    # bids, so the suggestion reuses the existing calibration instead of
    # inventing a second one: <30 -> 0, 30-60 -> 1, >60 -> 2.
    "kd_score_breaks": [30, 60],
    # --- CPC-scaled competitive adder ---
    # The competitive adder scales with the median top-of-page bid (CPC), because
    # CPC is the market's own measure of how valuable a click is: high-CPC verticals
    # (e.g. insurance ~$150) mean ranking organically replaces huge ad spend, so the
    # SEO is worth more. adder = median_cpc × cpc_adder_mult, rounded to $50, capped.
    # When there's NO bid data, fall back to the flat score buckets above.
    "cpc_adder_enabled": True,
    "cpc_adder_mult": 2.6,                     # partner $ of adder per $1 of median CPC (up to the knee)
    # CONFIDENCE FLOOR (2026-07-27). The adder scales off the MEDIAN
    # top-of-page bid, and above the knee each extra dollar of CPC adds $14 —
    # so an outlier is amplified enormously. That is fine on a real median and
    # dangerous on a sample of one: Rockingham priced a +$1,000 adder off a
    # single term's bid estimate (1 of 3 head terms returned data). Below this
    # many samples the adder still applies but the quote is flagged, because
    # the number is an extrapolation rather than a measurement. Raise this to
    # make thin samples fall back to the flat score buckets instead.
    # Grounding filter safety valve (2026-07-27). Requiring every word of a
    # model-invented service to appear in the client's own text catches
    # competitor names cleanly when the client HAS a rich vocabulary — Keller
    # says commercial/industrial/agricultural and never says Turner. It fails
    # badly when they don't: MPG's whole description is "energy and electrolyte
    # gummies for athletes", so hydration, caffeine, pre-workout and b12 all
    # read as alien and 100% of the invented services were removed. A filter
    # that deletes everything is measuring the description's length, not the
    # services' legitimacy — so above this drop ratio it stands down entirely
    # and says so rather than gutting the list.
    # A suggested region name must clear this monthly volume WITH a service
    # attached before it can enter the grid. Below it the name may be real but
    # nobody searches it, and an unsearched keyword in a proposal is a promise
    # to rank for something with no demand behind it.
    "region_min_volume": 10,
    "grounding_max_drop_ratio": 0.5,
    # Rank-check batch budget. Was hardcoded at 24s "to stay well under the
    # ~30s platform kill" — which was right when gunicorn's default 30s
    # timeout applied, and is now far too tight: any keyword still waiting
    # when the budget runs out is recorded as a FAILED lookup, which is what
    # produces "7 lookups failed and were excluded" on a healthy account
    # (2026-07-28). They aren't permanent failures; they're keywords that ran
    # out of clock. Derived from REQUEST_BUDGET_S so raising the server
    # timeout actually buys more lookups instead of nothing.
    "rank_batch_budget_s": 0,           # 0 = derive from REQUEST_BUDGET_S
    "cpc_adder_min_samples": 1,          # apply the CPC adder at/above this n
    "cpc_adder_low_confidence_n": 3,     # warn below this n
    "cpc_adder_knee": 62.0,                    # CPC above this earns the premium rate (just above Waytek's $60 — the highest "normal" client observed)
    "cpc_adder_mult_high": 12.3,               # partner $/CPC above the knee (insurance-carrier tier)
    "cpc_adder_cap": 1300,                     # max adder (partner $) so a freak CPC can't explode price
    "cpc_adder_free_below": 5.0,               # CPC at/below this adds nothing (normal-value clicks)
    "zero_ranking_bonus": 400,                # (legacy flat; superseded by tiers below)
    # Now a MARGIN OF GROSS (agency share of retail), matching rep_pricing and
    # the SSG/Vici grid: retail = hard / (1 - margin). Was a markup-on-cost
    # (x1.35 = a 25.9% margin). Retail output at 35% is unchanged.
    "default_markup_pct": 35,
    # top-N deepened 20 -> 100 (2026-07-20, Media Venue): a client with page-3-5
    # footholds (ranks 25/27/33/51 in Brendan's own table) was scoring "80% not
    # ranking" and drawing the +14% uplift, +18% over his base. "Not in top 20"
    # and "starting from scratch" are different claims — the uplift keys off the
    # latter. Depth <=100 is the same DataForSEO billing unit, so no cost change.
    # Tier thresholds unchanged; re-run Serene Health to confirm its fit holds.
    "zero_ranking_top_n": 100,
    "zero_ranking_frac": 0.10,
    # --- Brendan #5: TIERED zero-ranking. % of head terms NOT ranking in top-N
    # maps to a % uplift on the hard base. Each tier: [min_pct_not_ranking, uplift_pct].
    # Evaluated high-to-low; first threshold met wins. Replaces the flat bonus.
    # (2026-07-20) Serene Health RECLASSIFIED out of the auto-fit ledger: its
    # $3,950/$5,450/$6,950 is the same ladder as Skidmore's national card —
    # Brendan's premium/big-org card (multi-site telehealth), not a computed
    # response to keywords. Honest per-city volumes total ~2k/mo (the original
    # "fit" dated from the inflated-volume lookup bug). Handle via the manual
    # hard-base override (~$2,930 -> his card, ratio steps apply). The tiers
    # below remain calibrated on the zero-ranking signal itself.
    # Visibility moves the price BOTH ways. Brendan: "it's more about their
    # current visibility + search volume + competition." Until now visibility
    # could only add — a client with established rankings paid the same as an
    # average one. The negative bottom tier is the discount for a client that
    # already ranks; fit on Susquehanna (60% of head terms ranking, quoted
    # ~7% under the statewide anchor). ONE datapoint — confirm with Brendan
    # before trusting it on a second well-ranked client.
    # RECALIBRATED ON 12 BE PROPOSALS (2026-08-10). Ranking coverage barely moves
    # his price: r = +0.10 ex-insurance, and the two clearest cases contradict the
    # old ladder outright — Nob Hill Dental at 80% not ranking and Visit Central
    # PA at 5% were both quoted $2,950, and Junk Bee Gone at 5% not ranking (it
    # ranks 1-8 on twenty of twenty-one terms) was also $2,950. He never
    # discounts for good visibility.
    #
    # The old tiers spanned -3% to +14%, swinging the client price $350-$550 —
    # up to 88% of his ENTIRE observed spread of $625 — on the one input his book
    # shows no response to. So: the discount is gone, and the uplifts are halved.
    # Direction is kept because a genuinely greenfield build IS more work; the
    # magnitude is no longer inventing variance he does not have.
    #
    # Superseded values, for the record: [[80,14],[65,9],[50,5],[45,0],[0,-3]].
    "zero_ranking_tiers": [
        [80, 7],    # 80%+ not ranking -> +7%
        [65, 4],    # 65-80% -> +4%
        [50, 2],    # 50-65% -> +2%
        [0,  0],    # below half -> par. Never a discount.
    ],
    # --- VOLUME-based pricing (fixed $ per additional search, declining marginal
    # rate, like tax brackets). Base price assumes a "normalized" volume up to
    # vol_free_below. Above that, each bracket adds $/search for volume WITHIN that
    # bracket; brackets stack. Each: [lo, hi, dollars_per_search]. Open-ended top
    # bracket uses hi = null. Added to the hard base. Admin-editable.
    # NOTE: rates are the lever to calibrate. Brendan's example used $0.50/search,
    # but that produces very large adds (a 15k client would gain ~$2,600 on the hard
    # base, roughly doubling the quote). These starting rates (~$0.05-0.08) keep a
    # normal-volume client near its real proposal while still escalating hard for
    # 100k+ clients. Tune live; no high-volume proposals exist to fit against.
    # HEAD-TERM PINNING (2026-07-25). The service list comes from a Claude
    # call, so the same client can produce a different list run to run. That is
    # fine for long-tail colour and NOT fine for pricing: Skidmore's first run
    # carried "brand identity design" (18,100/mo = 76% of its total volume),
    # the second dropped it for "brand identity agency" (320), and the quote
    # fell $700/tier on nothing the client did. Volume is a pricing input, so
    # the highest-demand terms the search API actually returned are FORCED into
    # the list regardless of what the model picks. Set pin_head_terms to 0 to
    # go back to a purely model-chosen list.
    "pin_head_terms": 3,                # how many top-volume candidates to force
    "pin_min_volume": 300,              # ignore anything thinner than this
    "pin_as_ultra": 2,                  # the top N pins are the ultra-competitive tier
    # VOLUME IS OPPORTUNITY, NOT DEMAND (2026-07-25).
    # Susquehanna River Valley VB has the largest raw volume of any calibration
    # client (135k/mo) and the largest value-weighted demand after Skidmore,
    # and Brendan priced it BELOW the statewide anchor. It also already ranks
    # for 60% of its head terms. The reconciliation: a client that already
    # ranks for its demand has nothing to buy — you are not selling them that
    # volume, you are maintaining it. The volume add prices the OPPORTUNITY,
    # so it only applies where the demand is still uncaptured.
    #
    # Checked against every signal available on three clients with published
    # actuals; only this one splits them:
    #   raw volume    135k -> $0,  25k -> $500,  24k -> $500   non-monotonic
    #   volume x CPC  100k -> $0,  24k -> $500, 471k -> $500   non-monotonic
    #   % not ranking  40% -> $0, 100% -> $500,  88% -> $500   clean at 50%
    # (This is why the value-weighted model was NOT adopted: it fails on MPG,
    # which has the lowest click value of the three and needs the full add.)
    # The 50 pivot is the one zero_ranking_tiers already turns on.
    # A hard gate at 50% put a $900/tier cliff between a client at 49% and one
    # at 51%, which no operator could defend in a room. Ramped instead: the
    # volume add scales linearly from 0% of itself at the bottom of this range
    # to 100% at the top. Fits the same three actuals (Susquehanna 40 -> none,
    # Skidmore 88 and MPG 100 -> full) with no discontinuity in between.
    "vol_add_ramp": [40, 60],           # [no opportunity, full opportunity] % not ranking
    "vol_free_below": 10000,            # normalized: base already covers this
    "volume_add_cap": 450,              # max partner $ from volume: Brendan's quotes
                                        # flex a few hundred for market size, never
                                        # thousands (Waytek: his +$500 total vs the
                                        # formula's former +$1,400-4,500 vol adds)
    # Rates are PARTNER $ per search (scaled from the old basis rates by 0.8775
    # when the constants were converted, 2026-08-13).
    "volume_brackets": [
        [10000, 20000, 0.0702],
        [20000, 35000, 0.0439],
        [35000, 50000, 0.0351],
        [50000, None,  0.0263],         # open-ended top bracket so it keeps escalating
    ],
    # NATIONWIDE service clients (Skidmore Studio datapoint, 2026-07-20):
    # Brendan's national ladder $3,950/$5,450/$6,950 backs out to hard
    # $2,926/$4,037/$5,148 — base = the bare nationwide anchor, steps of 38%.
    #
    # (2026-07-25 REVISION — Brendan meeting) This multiplier was 0.0 on the
    # theory that at national scope the volume add and zero-ranking uplift are
    # tautological ("every nationwide client has >10k volume and ranks for
    # almost nothing"). Brendan says the opposite: volume, competition and
    # CURRENT VISIBILITY are precisely what separate one national client from
    # another — a brand with nothing ranking pays more, an established one
    # pays less. At 0.0 those signals were multiplied out and every national
    # client priced identically. Set to 1.0 (extras live). The Skidmore fit
    # must be re-validated at 1.0 — the original +$1,327 finding was measured
    # against the old inflated-volume lookup and the flat-adder era.
    # If Skidmore comes back high, prefer lowering volume_add_cap over
    # re-zeroing this — the cap is the honest lever, the multiplier is a mute.
    "nationwide_service_extras": 1.0,
    # Brendan steps his ladder in FLAT dollars (~$900-1,000 client per tier),
    # not proportionally — the old 38% ratio made the gap widen with every tier
    # (+15/18/20% on Keller, +13/24/34% on Waytek). Flat $700 hard = ~$950
    # client at 35% markup. step_ratio remains as fallback if flat is nulled.
    # Industry pricing: industries known to carry additional tiered pricing.
    # Matched by substring against the RZ-fed industry text ("DTC ecommerce
    # supplements" matches "ecommerce"). Rule keys:
    #   anchor_add      hard $ added to the base
    #   step_mode       "ratio" (proportional 38% steps) or "flat" (default)
    #   extras_off      skip volume + zero-ranking (org size, not SERPs, prices)
    #   national_demand price on GEO-LESS volume — sets no price of its own
    #
    # (2026-07-25 REVISION — Brendan meeting) The ecommerce family previously
    # carried anchor_add 250 + ratio steps, fit to MPG Gummies. Brendan:
    # ecommerce is "not auto more expensive, but normally in more competitive
    # industries — look at the volumes w/o the geo." So industry no longer
    # moves the price for these; it flips the volume lookup to national and
    # lets volume + CPC adder + zero-ranking do the pricing themselves.
    # NOTE the nationwide anchor ($2,900 hard = $3,915 client) already
    # reproduces Brendan's $3,950 card base on its own — the +$250 was very
    # likely fitting a number the band was going to hit anyway.
    "industry_pricing": {
        "ecommerce":  {"national_demand": True, "note": "Product brand — price on national demand, not a geo-qualified pull. Carries NO price of its own (pricing authority, 2026-07-25). Legacy toggle key."},
        "e-commerce": {"national_demand": True, "note": "Matches RZ “Retail - General / E-commerce”. Price on national demand; no anchor add."},
        # Sibling RZ values an operator would reasonably pick for a product
        # brand (MPG is literally a supplements company) — same volume mode,
        # so the behaviour can't silently vanish on an equally-valid tag.
        "supplements":             {"national_demand": True, "note": "Sibling of e-commerce (MPG is a supplements brand)."},
        "consumer packaged goods": {"national_demand": True, "note": "Sibling of e-commerce — product brand tag."},
        # Brendan's premium/big-org card (Serene Health, 2026-07-20 — one
        # datapoint, provisional): large multi-site / telehealth healthcare
        # orgs price on ORGANIZATION size, not keyword signals — his
        # $3,950/$5,450/$6,950 card. anchor_add lands the base at the card;
        # extras_off skips volume + zero-ranking (size, not SERPs, drives it);
        # ratio steps give the card's $1,500 rungs.
        # Keys must match the RZ industry taxonomy VERBATIM (substring) — the
        # line item ships values like "Health Services - Hospital", not the
        # client's marketing vocabulary. Add each big-org RZ value as Brendan
        # prices one.
        # Insurance carriers (Rockingham, 2026-07-20 — one datapoint,
        # provisional): +$800 with extras ON and default steps lands his
        # $5,450/$6,750/$7,950 within 1% per tier. Note the composition differs
        # from the hospital card: uplift stays (SEO genuinely starts from
        # scratch) and steps run the standard 24%-of-base, not the 38% card.
        # Key "insurance -" matches the RZ "Insurance - *" family only — it
        # deliberately misses "B2B - Insurance Business Solutions". OPEN
        # QUESTION for Brendan: RZ doesn't distinguish carriers from two-agent
        # local agencies; confirm whether small agencies carry the same +$800.
        "insurance -":       {"anchor_add": 395, "note": "Carrier premium — Rockingham re-calibration 2026-07-20 at the CURRENT piecewise CPC adder (which already carries ~$1,000 of insurance click value at a $120 median; the original +$800 was fit against the old +$350-capped adder and double-counted). Contiguous NoVA 9-city scope; lands 5,450/6,750/8,050 vs his 5,450/6,750/7,950. Open: do small agencies carry it too?"},
        # Legal (Ooten Law, 2026-08-22 — one datapoint, provisional). The tool
        # had NO legal rule at all, which is how a personal-injury firm priced
        # at the $2,950 client floor. Fitted the same way insurance was: on top
        # of the piecewise CPC adder, which already carries most of the click
        # value in this vertical. At a $150-160 median top-of-page bid the adder
        # runs $1,250-1,300 and +$700 lands his 5,950/7,250/8,450 within $50 a
        # tier. Note it converges on the same +$700 as the big-org cards rather
        # than needing a number of its own.
        #
        # Matches the RZ "Legal - *" family. Volume stays $0 here: fifty legal
        # terms in one metro total a few thousand a month, well under the
        # 10,000/mo first bracket, so the adder and this card carry the quote.
        #
        # (2026-08-23 DECISION, Kiri) The open question was whether a low-CPC
        # practice area - estate planning, family law - should carry the same
        # card as personal injury, since the substring picks up all nine RZ
        # "Legal - *" values automatically. Answer: yes, one card for all legal.
        # The reasoning that makes it safe is that the two components move
        # independently. The CPC adder is measured per client and already
        # separates a $150 personal-injury click from an $8 estate-planning
        # click; an estate-planning quote comes out roughly $1,000 under a PI
        # quote on the adder alone. The card is the flat premium for the
        # vertical - the sales cycle, the compliance copy, the review posture -
        # and that part does not vary by practice area. So no per-area split,
        # and no second calibration needed before a legal quote goes out.
        "legal -":           {"anchor_add": 700, "note": "Legal family — Ooten Law calibration 2026-08-22. Fitted over the CPC adder on a personal-injury firm. DECIDED 2026-08-23 (Kiri): the card is the legal family, not the practice area — all nine RZ \u201cLegal - *\u201d values carry the same +$700. A cheap-click practice area gets a smaller CPC adder, so the quote still falls; the card does not."},
        "hospital":          {"anchor_add": 700, "step_mode": "ratio", "extras_off": True, "note": "Big-org card ($3,950/$5,450/$6,950 shape) — Serene Health calibration via RZ “Health Services - Hospital”."},
        "telehealth":        {"anchor_add": 700, "step_mode": "ratio", "extras_off": True, "note": "Big-org card — non-RZ vocabulary key, kept for free-text matches."},
        "behavioral health": {"anchor_add": 700, "step_mode": "ratio", "extras_off": True, "note": "Big-org card — non-RZ vocabulary key, kept for free-text matches."},
    },
    # Core SEO + AI Search — GEO PRICING.
    # (2026-07-25 REVISION — Brendan meeting) The $2,950/$4,050/$5,250 card was
    # read off the MPG proposal and hard-coded as universal. It is NOT: MPG had
    # near-zero visibility (little traditional-search presence, almost no AI
    # presence, very few backlinks) which is why its GEO landed ~95% of SEO.
    # Brendan's actual rule: GEO runs 30-50% LESS than SEO on average — i.e.
    # 50-70% of the SEO price — and rises toward parity when nothing ranks.
    # So GEO is a PERCENTAGE of the client's own Core SEO quote, and the
    # percentage is driven by current visibility (the same pct_not_ranking
    # signal the zero-ranking uplift already computes off the top-100 check).
    # Tiers are [min_pct_not_ranking, geo_pct_of_seo], evaluated high-to-low.
    "geo_pricing_mode": "pct",                # "pct" (Brendan rule) or "card" (legacy MPG)
    # CALIBRATION NOTE: MPG's GEO list price is 78% of its SEO price, almost
    # exactly. His intermediate GEO list of $4,250 / SEO intermediate $5,450 =
    # 77.98%; base and advanced back-solve to 78.6% and 79.5% once the 5%
    # bundle discount is removed. So the zero-visibility ceiling is ~78% of
    # SEO -- slightly ABOVE Brendan's stated 50-75% normal band, which is
    # precisely what he said should happen for a client with no visibility.
    # (This assumes MPG's SEO ladder was 3,950/5,450/6,950 -- CONFIRM.)
    # REVISED 2026-07-28: the separate bundle discount is gone and its 5% is
    # folded into these percentages, so this table is now the WHOLE of GEO
    # pricing. Every rate is its old value x 0.95, which leaves the quoted
    # numbers identical — this is a simplification of how the price is
    # expressed, not a change to what anyone pays. Two knobs describing one
    # decision meant the headline rate was never the rate actually charged.
    "geo_pct_tiers": [
        [90, 74],   # <10% of head terms rank  -> the MPG ceiling (was 78 less 5%)
        [70, 66],   # 10-30% rank              -> top of the normal band
        [40, 59],   # 30-60% rank              -> mid of the normal band
        [0,  48],   # 60%+ rank (established)  -> the established-client floor
    ],
    "geo_pct_default": 57,                    # used when no ranking data exists
    # (The GEO bundle discount is GONE, 2026-08-27. It was a real 5% off the AI
    # Search line when sold with Core SEO — MPG, 2026-06-10, "$4,050, discounted
    # from $4,250 in conjunction with the SEO campaign". On 2026-07-28 the 5%
    # was folded INTO geo_pct_tiers above: every rate there is its old value
    # x 0.95, which left the quoted numbers identical. The constant stayed
    # behind at 0 and editable, which meant anyone could re-apply a discount
    # that was already in the rates. Removed outright rather than left as a
    # trap. If a separate discount ever comes back, un-fold geo_pct_tiers in
    # the same edit — divide each rate by 0.95.)
    # Minimum term. Brendan: "we usually do 6 months for both, however where
    # someone has like ZERO visibility sometimes we do 12 because it takes
    # that long to get results." Same trigger as the top geo_pct rung.
    "min_term_months": 6,
    # A SITE BEING REBUILT IS NOT THE SITE BEING PRICED. Three of the tool's
    # inputs -- the rank check, the on-page score and the domain authority --
    # are measured on the CURRENT site. When the client is relaunching, those
    # describe something that is going away, and a client who ranks today would
    # be DISCOUNTED for equity they are about to throw out.
    #
    # Deliberately NOT a price card. The zero-ranking uplift already is the
    # "you own nothing yet" component, and stacking a second premium on it is
    # how the insurance card ended up double-counting the CPC adder (+$800, cut
    # to +$395). A new domain routes through that existing lever instead: read
    # as 100% not ranking, which takes the top uplift band, the full volume add
    # term. (The 12-month "no visibility" term this used to pick up was removed
    # 2026-08-27 — the term is now a judgement call, so a new domain moves the
    # uplift and the volume add and nothing else.)
    #
    # A SAME-DOMAIN rebuild is a different claim. Redirects carry the rankings
    # and the authority, so the measurement still describes the campaign; only
    # the on-page score is stale. It does not touch the price. (2026-08-24)
    "rebuild_new_domain_pct_not_ranking": 100,
    "rebuild_vetoes_performance": True,
    # "up" (default) or "nearest". Governs the figures a CLIENT reads — the tier
    # ladder, the AI Search line and the add-on market rate. Partner cost and
    # the tier step keep round-to-nearest: they describe the shape of the quote
    # rather than the number on it. (2026-08-26)
    "client_round_mode": "up",
    # (The 12-month "no visibility" term is GONE, 2026-08-27. A quote where
    # almost nothing ranks used to be committed for 12 months instead of 6, on
    # the reasoning that results take longer. It is a judgement call per client
    # rather than a rule the tool should enforce, so every quote now carries the
    # one minimum term and the operator sets the dates.)
    # Legacy MPG card, kept for reference / geo_pricing_mode="card" only.
    "geo_card": {"base": 2950, "intermediate": 4050, "advanced": 5250},
    "geo_card_list": {"base": 2950, "intermediate": 4250, "advanced": 5250},
    # (AI Search no longer carries a minimum term of its own, 2026-08-27. It was
    # 12 months against Core SEO's 6, on the reasoning that GEO takes longer to
    # show. Both now run on the quote's single min_term_months — 6, or 12 when
    # >=90% of terms are not ranking — so adding AI Search to a quote can no
    # longer double the commitment on its own.)
    "ai_search_uplift_pct": 75,               # legacy flat-pct mode only
    "ecom_anchor_add": 0,                     # RETIRED 2026-07-25 (Brendan): ecommerce carries no anchor add
    # CALIBRATED ON BE'S SIX FLOOR-BOUND PROPOSALS (2026-08-13). 700 produced a
    # $900 client step; BE steps $1,000 in four of the six quotes whose base sits
    # on the $2,950 floor (Nob Hill, Keller, Visit Central PA, Red Shoes), $900 in
    # one (Junk Bee Gone) and $1,110 in the last (Media Venue). 750 gives $1,000.
    # Scored by pricebench.py: total error on Intermediate + Advanced 1,790 -> 590,
    # exact ladders 1 of 6 -> 4 of 6. The cost is Junk Bee Gone, which was the
    # exact match and is now $100/$200 high — a deliberate trade of one match for
    # four. Only the FLAT path moves: manual overrides and the nationwide card
    # step on step_ratio and are untouched (verified both ways).
    #
    # This is why "too low" got worse up the ladder — NPAIHB read -17% at base and
    # -20% at advanced. It does NOT lift a base off the floor; nothing measurable
    # in BE's proposals separates his $2,950 clients from his $3,550 ones (Nob Hill
    # and Amare Homes are both 20 terms, one city, 80% not ranking, and $600
    # apart), so that gap stays a judgement call on the override.
    "tier_step_flat": 650,                    # partner $ per tier; null -> use step_ratio
    "tier_step_pct_of_base": 0.24,            # step grows past the flat floor on big bases
    "step_ratio": 0.38,                       # fallback: proportional step
    # CALIBRATED ON 12 BE PROPOSALS (2026-08-10). His base has never gone below
    # $2,925 and sits within $25 of $2,950 in six of them, so a quote landing
    # under that is below anything he has ever sent. The floor only LIFTS; every
    # quote already at or above it is untouched, which is why this is the one
    # calibration change that cannot disturb a past quote. Junk Bee Gone is the
    # check: the tool said 2,850 / 3,750 / 4,700 and BE sent 2,950 / 3,850 /
    # 4,750 — the floor makes all three match exactly.
    "client_floor": 2950,
    # Add-on market pricing, per tier. Confirmed against TN Water & Air
    # (2026-03-25): a Knoxville ladder of 2,250 / 2,950 / 3,650 with add-on
    # markets at 950 / 1,250 / 1,750 — 42%, 42%, 48%. The flat 0.42 was right
    # on the first two tiers and 11% light on advanced, because the extra work
    # an advanced campaign does per location scales more than the shared work
    # does. Per-tier ratios reproduce all three exactly.
    # Add-on recommendation thresholds. Both fit on TWO proposals (Skills of
    # Central PA and TN Water & Air) — treat the suggestion as a prompt to
    # think, not a decision, until more actuals confirm them.
    # Markets within this many miles of each other are one market. 25 is a
    # metro radius and separates the two clients we can check — Brent Cogan's
    # towns span 22 miles, TN Water & Air's are 101 and 160 apart — but it is
    # fitted on two proposals, so treat it as a starting point.
    # Straight-line, not drive time: two towns 20 miles apart across a ridge or
    # a state line can still be separate markets.
    "market_radius_miles": 25,
    "addon_free_markets": 3,        # at or below this, always one campaign
    # Need rank data on this share of the ENTERED markets before suggesting.
    # Note the interaction with grid_max_cities: at 5 crossed cities, a client
    # with 8+ markets can never clear 70% and will always be told there isn't
    # enough data. That is the honest answer rather than a bug — the tool
    # genuinely hasn't looked at those markets — but it means wide-footprint
    # clients need Grid max cities raised for one run to decide add-ons, then
    # dropped back for the proposal list.
    "addon_min_measured_share": 0.7,
    "addon_covered_share": 0.6,     # rank in this share of measured markets = one footprint
    # Add-ons are priced off the CORE SEO ladder only — no AI Search component.
    # This is an assumption, not a calibration: TN Water & Air is the only
    # proposal with add-on markets and predates AI Search, while MPG is the
    # only AI Search proposal and has no add-ons. The reasoning is that AI
    # Search work is brand-level — Brendan's GEO tiers are defined by premium
    # content placements, which lift the whole brand rather than a suburb — but
    # it is worth a lot: at Woodstock's 7 add-ons it is ~$5,600/mo on base
    # alone. Confirm before treating it as settled.
    # RETIRED (2026-08-21). An add-on market was priced at a FRACTION of the
    # tier — 0.42 base/intermediate, 0.48 advanced, confirmed against TN Water
    # & Air — on the reasoning that one extra city is a smaller piece of work
    # than the primary campaign. That fraction sat next to the Add-On Market %
    # and read as a second, larger discount, and the decision is that the
    # bracket is the ONLY thing between the tier price and the add-on price.
    #
    # An add-on market is now priced as a full campaign at that tier, less the
    # Add-On Market % for the market count. This is a ~2.4x increase on the
    # add-on leg and a deliberate departure from the TN Water & Air figure;
    # nothing in bench.py measures this leg, so there is no test that would
    # have argued with either number.
    #
    # Both keys are kept so a saved config payload still round-trips, and both
    # are ignored by the price. Use addon_override to negotiate a per-market
    # rate on a single quote.
    "addon_market_ratio": 1.0,                     # RETIRED — ignored by stage4_price
    "addon_market_ratio_tiers": {},                # RETIRED — ignored by stage4_price
    # ---- PERFORMANCE-BASED SEO (pay per ranking term) ---------------------
    # Brendan, 2026-08-22: "For SEO we only offer pay for performance if their
    # keywords are already ranking in the first 5 pages of results. If they're
    # starting from scratch / not ranking we generally don't offer it, because
    # the time frame to results can be 6-12+ months."
    #
    # So this is a GATE, not a strategy the operator can simply select: the
    # client has to already be within reach. Ooten Law clears it at 70% of
    # terms inside the top 50 and is the calibration for everything below.
    "perf_page_depth": 50,               # "the first 5 pages"
    "perf_eligible_min_share": 0.5,      # of the terms actually MEASURED
    "perf_min_measured": 5,              # below this the share means nothing
    "perf_min_monthly_value": 10000,     # BE: minimum potential ranking value
    # Cost Page 1 per term, from the term's own measured top-of-page bid. The
    # bid is what the market pays for one click, which is the same thing this
    # column prices. Fitted on Ooten: his median Page-1 is $280 against a
    # measured median bid of $132 — 2.12x. Floor from his cheapest row ($80).
    "perf_page1_mult": 2.1,
    "perf_page1_floor": 80,
    "perf_page1_round": 5,
    # Top 5 / Top 3 / #1 as multiples of Page 1. Two bands: his cheap rows run
    # a clean 1.6 / 2.4 / 3.6 (measured 1.56-1.62, 2.25-2.44, 3.38-3.69) and his
    # four most expensive are compressed. The #1 column in that top band is
    # irregular (2.10, 1.99, 2.20, 2.66) and looks hand-set, so 2.2 is the
    # middle of it rather than a fit — expect to override those rows.
    "perf_ladder": [1.6, 2.4, 3.6],
    "perf_ladder_high": [1.47, 1.83, 2.2],
    "perf_high_knee": 600,               # Page-1 cost at which the high band starts
    # HOW MANY TERMS THE PERFORMANCE TABLE CARRIES. Brendan's Ooten table is
    # FIFTY, and it is the same fifty as his SEO keyword table — he does not
    # build a separate list. Ours is capped at 20 by grid_max_services, which is
    # his floor rather than his norm (his proposals run 20 to 99; Ooten is 50
    # because it is two full practice areas).
    #
    # The priced grid stays at 20: it is what the campaign commits to WORKING,
    # it is what the option scope lines describe, and every calibrated price in
    # bench.py was built on it. The performance table is a different question —
    # a menu of terms you could be paid for, with no work commitment per term,
    # which is why Brendan calls it "a sample set of potential keywords" — so it
    # draws from a wider pool. Terms past the grid still need a measured rank
    # before they can appear in a client document. (2026-08-22)
    # WHERE THE PROPOSAL SCREENSHOT OPENS. The capture is 1100x1700 and framed
    # 16:9, so it shows roughly the top 640px of a page whose first 700+ are an
    # AI Overview map, the ads and the local pack — none of which is the
    # organic result the client is missing from. Ooten's exhibit was a Google
    # Maps card. Opening 40% down lands on organic results on a typical page;
    # the frame still slides either way. (2026-08-22)
    "serp_frame_offset": 0.40,
    # Pixels of the ORIGINAL capture kept on top of the frame: the Google bar
    # and the query. Without it a slid frame is a result page that never says
    # what was searched. 190 covers the logo row and the search box at 1100px
    # wide; it is joined to the results window with a hairline.
    "serp_head_px": 190,
    "perf_table_terms": 50,
    "perf_initial_term_months": 12,
    "perf_tail_months": 6,
    # ADD-ON MARKET % — the volume break, FLAT BY BRACKET, not graduated.
    # Twelve markets is 15% off all twelve, not nine at 10% and three at 15%.
    # One rate per quote, because that is the number the client argues about;
    # a blended 11.3% is a number nobody can check against a rate card.
    #
    # [min_market_count, percent_off] — read HIGH TO LOW, first match wins,
    # same convention as geo_pct_tiers.
    #
    # The break comes off PARTNER COST, and the client price is then derived
    # from that cost through the margin exactly as every other figure here is.
    # Discounting retail while cost stood still would have walked realised
    # margin down as markets accumulated — which is the specific failure the
    # 2026-08-05 rewrite of add-on pricing was done to end. (2026-08-21)
    "addon_volume_discount_tiers": [[26, 20], [10, 15], [1, 10]],
    # Campaign goals that flip the national-demand switch. See
    # GOAL_NATIONAL_DEMAND — editable here so the list can be widened without a
    # deploy. Matched case-insensitively against the exact order-form option.
    "goal_national_demand": ["Online Sales"],
    # Two clusters count as ONE region if any city in one is within this many
    # miles of any city in the other. Wider than market_radius_miles on purpose:
    # 25 miles is "same market", 60 is "same trade area".
    "scope_join_radius_miles": 60,
    # A service name below this many monthly searches (summed across the
    # client's markets) is treated as a phrase nobody types, and is replaced by a
    # higher-volume term from the operator's own list in the same topic. Set to 0
    # to switch the check off.
    "service_min_volume": 30,
    "service_max_swaps": 3,
    # UPGRADE pass. A service that clears the floor is still replaced when an
    # unused term from the operator's own list, in the SAME topic, measures at
    # least this many times more. 10x is deliberately steep: it only fires when
    # the alternative is in a different league, never on a close call, so a
    # deliberate service choice isn't second-guessed over noise. 0 turns it off.
    "service_upgrade_ratio": 10,
    # How much a STORE-INTENT term ("ski shop") outweighs a product term
    # ("outdoor furniture") of the same size when ordering the tier columns. A
    # multiplier, not a veto: 3 means a product term must carry 3x the demand to
    # take a slot ahead of a shop term. 1 = order on raw volume only.
    "store_intent_tier_boost": 3.0,
    # How many unused seed terms to measure as replacement candidates. Each one
    # adds a keyword to the existing per-city volume calls. Spread round-robin
    # across topics, shortest term first, so every product line gets measured.
    "service_candidate_cap": 21,
    # TIER MIX, measured off eight real BE proposals (303 terms, 2026-08-07)
    # rather than assumed. His splits are PROPORTIONAL, not fixed counts, and
    # they move with the client:
    #
    #   Rockingham Insurance   99 terms   39/40/20   ->  39% / 40% / 20%
    #   Keller Builds          64         16/24/24   ->  25% / 38% / 38%
    #   Waytek                 30          5/12/13   ->  17% / 40% / 43%
    #   Red Shoes              30          7/ 2/21   ->  23% /  7% / 70%
    #   PA Dental Excellence   20          6/10/ 4   ->  30% / 50% / 20%
    #   Nob Hill Dental        20          6/10/ 4   ->  30% / 50% / 20%
    #   Visit Central PA       20          5/ 9/ 6   ->  25% / 45% / 30%
    #   Media Venue            20          6/ 7/ 7   ->  30% / 35% / 35%
    #   -------------------------------------------------------------
    #   TOTAL                 303         90/114/99  ->  30% / 38% / 33%
    #
    # The old fixed "3 ultra / 6 competitive / rest is long tail" produced
    # 15% / 30% / 55% on a 20-term list — ultra half what BE writes, long tail
    # nearly double. Proportions travel across list sizes; counts do not.
    "tier_mix": {"ultra": 0.30, "competitive": 0.38, "long_tail": 0.32},
    # Hard counts kept as an escape hatch. Null -> derive from tier_mix; an
    # integer pins that bucket regardless of list length.
    "ultra_bucket_size": None,
    "competitive_bucket_size": None,
    # BE's lists run 20-99 terms (median 25). A flat cap of 20 was his FLOOR
    # applied as a ceiling — it would have truncated Rockingham by 79 terms.
    "list_cap": 60,
    "rank_check_workers": 8,   # parallel SERP calls — avoids timeout on free Render
    # Long-tail sourcing
    # HOW MANY SEEDS keyword_suggestions is asked about. It is one request each,
    # so this is a real cost — but it was ALREADY capped, by a hardcoded
    # seeds[:6] buried in the function, which is why the call count never looked
    # like the seed count. What changes here is WHICH seeds: the slice took the
    # first six in typing order, and this takes the BROADEST, fewest words first,
    # because a broad seed has far more long-tail children than a specific one.
    # "home for rent" yields pages of them; "rental homes with washer dryer"
    # yields almost nothing and Amare was spending a call on it. The saving is
    # one call; the point is which five. (2026-08-14)
    "suggest_seed_cap": 5,
    # How many variants of one procedure the TOOL may propose — see
    # cap_service_family. Never applies to a term the operator typed.
    "service_family_cap": 3,
    # A token in this many of the operator's OWN terms is the business, not a
    # runaway family, and is exempt from the cap above.
    "family_core_seeds": 2,
    # HOW MANY DIFFERENT SERVICES THE GAP-FINDER IS ASKED FOR. The cap above
    # takes slots BACK, so something has to be there to fill them or the grid
    # only gets smaller — which is the complaint that started this ("they all
    # come back so low now"). This is the one proposal source that returns a
    # DIFFERENT service rather than the same one reworded, so it is the one to
    # widen: proposing more costs no DataForSEO calls (one Anthropic call either
    # way, then one volume request for the whole batch), and every term it
    # proposes is measured and floored before it can reach a chip. (2026-08-17)
    "industry_gap_n": 22,
    # ROOM TO GROW. Slots held for terms the client does NOT already rank for.
    # Ranking a grid purely by measured demand hands every slot to the terms a
    # client with existing SEO already owns — Ski Barn came back ranking in the
    # top 100 for 19 of 19, nine of them in the top four — so the proposal
    # argued for winning what was already won. A reservation, not a quota:
    # nothing is invented to fill it. 0 turns it off. (2026-08-19)
    "grid_headroom_slots": 4,
    # A client ranking for EVERYTHING has no story in the proposal. When the rank
    # check comes back at or near 100%, this many terms they do NOT rank for are
    # hunted down and offered. Measured with the real rank check, not inferred
    # from ranked_keywords — those two disagreed on Ski Barn, where the labs
    # lookup called 19 terms unranked that the SERP found at #1-#4. 0 turns it
    # off. (2026-08-20)
    "min_unranked_terms": 3,
    # FOUR, NOT TEN — THE PROBES WERE STARVING EACH OTHER. The pacing holds the
    # process to dfs_calls_per_minute; a probe that queues ten SERP lookups, then
    # retries the ones that did not answer, then hands over to the cross-market
    # check, puts more than a minute of work into a minute. The later calls sit
    # waiting for a slot until the request times out, and a timed-out probe is
    # reported as "could not be read" — which is how six of ten and then two of
    # two came back unread on a client whose earlier runs answered fine.
    #
    # Four candidates is two batches. With the sieve in front of it that is
    # usually enough to find three misses, and what it does not spend is what
    # the cross-market check needs. (2026-08-21)
    "unranked_probe_max": 4,
    # DataForSEO allows 12 calls a minute. Pacing at 10 leaves headroom for the
    # calls made outside the build (signals, the SERP screenshot poll) without
    # anyone tripping the cap — see _dfs_take_slot(). 0 disables the pacing.
    "dfs_calls_per_minute": 10,
    # A throttled call is retried on its own budget, so a bad minute cannot
    # spend the retries meant for timeouts and 5xx.
    "dfs_rate_limit_retries": 4,
    # HOW THE SITE READER INTRODUCES ITSELF. DataForSEO's default is
    # "Mozilla/5.0 (compatible; RSiteAuditor)", which is a bot string, and a
    # WAF that refuses it costs a whole client's site-condition reading.
    "onpage_user_agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                          "AppleWebKit/537.36 (KHTML, like Gecko) "
                          "Chrome/126.0 Safari/537.36"),
    # Second opinion via Google Lighthouse when the read above comes back empty.
    # A second request per blocked site, never on the happy path.
    "technical_health_fallback": True,
    # How many head terms the back-measure buys a SERP for, per saved quote.
    # Five was enough to characterise all three clients measured by hand.
    "backmeasure_terms": 5,
    # HOW FULL THE GRID HAS TO BE BEFORE THE GROUNDING FILTER STOPS TAKING FROM
    # IT. Its own valve is a ratio against its candidates — drop more than half
    # and stand down — which passed comfortably on MPG Gummies at 7 of 18, and
    # left the grid at 11 of 20. The nine empty slots took total volume under
    # vol_free_below, the volume add went to $0, and the quote came out $1,000
    # under the one Brendan sent. Below this fraction of the slots, the best of
    # what it dropped comes back until the line is reached. 0 turns it off.
    "grounding_min_slot_fill": 0.75,
    # WHO ALREADY HOLDS PAGE ONE, IN DOLLARS. Partner cost added to the anchor,
    # keyed on the median incumbent's backlink authority through
    # _pageone_bucket() — the SAME function the Calibration driver buckets on, so
    # moving a cut moves both and they can never disagree.
    #
    # Four thresholds already in this formula sit above where these clients live
    # (volume pays nothing under 10,000/mo and they run 200-980; the CPC adder
    # needs a $20 click and they measure $2.89-$16.75; difficulty's first break
    # is 30 and they read 0-25), so every quote computed under the floor and the
    # floor was the price. Nob Hill and Amare came out identical at $2,950 and
    # Brendan sent $2,950 and $3,550.
    #
    # REVERTED TO ZERO 2026-08-18, THE DAY IT WAS SET. It was fitted on two
    # quotes — Amare and NPAIHB — and the back-measure across all fifteen killed
    # it on its first run: page-one strength against the price actually sent is
    # Spearman rho -0.09. No relationship. It fired on FOUR OF THE SIX clients
    # Brendan priced at his floor (Media Venue 650, Junk Bee Gone 704, Keller
    # 728, Red Shoes 728) and would have overcharged every one of them; Keller
    # went from $200 over to $750 over. Meanwhile the two he priced highest that
    # it was fitted on, NASSCO at 424 and NPAIHB at 437, sit BELOW four of those
    # floor clients.
    #
    # The mechanism still sounds right — one Zillow does change the job — and it
    # is measured and reported on every quote. It simply is not what moved his
    # prices, and a story that good is exactly the kind that needs a number
    # against it. Kept as a dict at zero rather than deleted: the measurement is
    # still collected, so if a later cut of the data earns a band the lever is
    # here. (2026-08-18)
    "pageone_anchor_add": {
        "page one: local businesses (under 200)": 0,
        "page one: regional or institutional (200-399)": 0,
        "page one: national platforms (400+)": 0,
    },
    "use_suggestions": True,           # pull keyword_suggestions for longer phrases
    "use_site_keywords": True,         # pull keywords_for_site from the client domain (Labs)
    "site_keywords_limit": 200,        # cap rows returned from keywords_for_site
    "longtail_min_words": 4,           # >= this many words qualifies as long-tail
    "longtail_prefixes": ["how","what","why","when","where","which","who","best",
                          "affordable","cheap","near","cost","top","is","can","do"],
    "longtail_target": 10,             # how many long-tails to keep in the list
    "rank_check_cap": 60,              # max keywords sent to the SERP rank check
    # --- GRID MODE (matches Brendan's proposals) -----------------------------
    # His keyword tables are a systematic SERVICE x CITY grid, with the tier
    # assigned to the SERVICE (every city inherits it): e.g. "auto insurance" is
    # Ultra-Competitive in all ten cities, "umbrella insurance" is Long Tail in
    # all ten. He does NOT use question-style long-tails (2 instances across 18
    # proposals), so the long-tail tier is just lower-competition services.
    "grid_mode": True,
    # Brendan targets a keyword COUNT, trading services against cities:
    #   Rockingham  10 cities x 10 services = ~104
    #   Serene       1 metro  x ~14 services = 20
    #   Skidmore     0 cities x ~20 services = 24
    # So services scale INVERSELY with cities to hold the total near target.
    # REVISED 2026-07-28 after live tuning: min services 4 -> 7, max cities
    # 10 -> 5. At the old settings a wide-footprint client hit the service
    # FLOOR — 10 cities forced services down to 4, so a 40-row list showed
    # only four things and read as padding rather than strategy. Trading
    # cities for services keeps the total near target while nearly doubling
    # the variety, and the cities that survive are the highest-demand ones
    # for this client's own service, so the ones dropped cost least.
    "grid_target_keywords": 32,
    "grid_min_services": 7,
    "grid_max_services": 20,
    # AXIS CHOICE (2026-08-10). A market must clear this monthly figure for the
    # crossing to be worth a term slot; 20 sits just above Google's 10/mo floor
    # for thin terms, which is what an empty market reports. Below that the
    # budget is better spent on service breadth — see choose_grid_axis.
    # How many of the lead services also get a "<service> near me" term. Measured
    # like any other; only forms clearing near_me_min_volume are added.
    # Four, not three — Brendan's Ooten list carries four "near me" terms out of
    # fifty (car accident lawyer, personal injury lawyer, dui lawyer, criminal
    # defense lawyer near me), all on the lead services. (2026-08-22)
    "near_me_terms": 4,
    "near_me_probe_cap": 12,
    # Monthly searches a proposed extra service must clear to be offered.
    "expand_min_volume": 20,
    # A word in more than this share of the seeds cannot partition them.
    "topic_token_max_share": 0.5,
    "topic_min_seeds": 2,          # below this it is a term, not a topic
    "ranked_keywords_limit": 80,
    # A MUCH DEEPER READ, FOR ONE PURPOSE. 80 is the right size for SEEDING a
    # list — the client's best terms, most relevant first. It is the wrong size
    # for asking "do they rank for this": Ski Barn ranks for far more than 80
    # terms, so nineteen of twenty grid rows fell outside the sample and were
    # reported as unranked while the live SERP had them at #1-#4.
    #
    # Labs charges the same for a deep page as a shallow one, so the ownership
    # question gets the deep one. It is still ONE call, and it replaces the ten
    # SERP calls the probe was spending to answer the same question.
    # (2026-08-20)
    "ranked_keywords_own_limit": 1000,   # Labs rows pulled per client
    # Above this multiple of the floor a vertical has real demand somewhere, so
    # the sub-floor terms are genuinely the dregs and the floor should hold. Below
    # it the whole market is small and the floor is refusing the only terms that
    # exist — Santa Fe apartment rentals top out at 90/mo. (2026-08-13)
    "expand_thin_market_mult": 10,
    "near_me_min_volume": 30,
    # SINGULAR OR PLURAL — see pick_service_forms. Nothing chose between them,
    # and Amare shipped "home for rent" where Brendan's list leads with "homes
    # for rent". A variant has to clearly win before the operator's phrasing is
    # rewritten: this much volume, and this multiple of the incumbent.
    "service_form_probe_cap": 20,
    # TIER BY SHAPE WHEN VOLUME CANNOT SEPARATE — see service_shape. A token
    # shared by this share of the list is the vertical's core; what a term adds
    # on top of the core is its qualification, and that is what breaks the ties
    # the tier reconciliation cannot break on volume.
    "shape_core_share": 0.4,
    # Below this many distinct stems the keyword pool is too small to prove a
    # qualifier is unused — see drop_ungrounded_services.
    "pool_vocab_min": 60,
    # The absolute floor under the median — see pool_vocabulary. Google Ads
    # reports 10/mo for a phrase it holds no data on, so 11 is "measured at all".
    "pool_min_volume": 11,
    "market_pool_cap": 60,          # local rows carried forward for the refill
    # How far down page one to read the incumbents — see _serp_parse_items.
    "serp_competitor_depth": 10,
    "serp_rival_cap": 12,           # distinct incumbents carried to the panel
    # Above this share of the whole list, dropping dead slots stops being a
    # targeted correction and becomes a cull — see swap_low_volume_services. Same
    # value and same denominator as grounding_max_drop_ratio, for the reason
    # written there.
    "dead_slot_max_drop_ratio": 0.5,
    # OFF until the backfill stops drawing New York and San Antonio terms out of
    # a nationally-ranked pool for a Santa Fe client. The filter half is sound;
    # the refill half is not. (2026-08-16)
    # False | "report" | True. "report" decides everything and applies nothing,
    # so its verdicts can be read against a real pool without moving a quote.
    "pool_qualifier_filter": "report",
    "service_form_min_ratio": 2.0,
    "service_form_min_volume": 50,
    "axis_city_volume_floor": 20,
    "axis_min_seeds_for_services": 8,
    "grid_max_cities": 5,             # cities crossed against each service
    # When a city needs no ", ST" in the keyword. Brendan writes "adhd treatment
    # san diego" but "auto insurance alexandria va" — the test is whether the
    # name is unmistakable on its own. It used to be "is this city in the
    # CITY_STATE lookup table", which is a city->state map, not a list of
    # metros: it holds Farragut (1 ZIP) and Maryville (4), so one grid built
    # "junk removal farragut" beside "junk removal clinton tn" (2026-08-10).
    # Measured instead: a big place that also owns its own name nationally.
    # These two numbers reproduce every example in the sample proposals —
    # Knoxville 31/.82 and San Diego 81/.99 pass, Alexandria VA 23/.57 does not.
    # National monthly searches the bare service terms must clear before the tool
    # will say "this looks national". Below it, zero local volume is just a thin
    # vertical and says nothing about the frame.
    "frame_national_min": 200,
    # Monthly national searches an acronym must clear to be offered as a seed.
    # Below it, it is an internal code rather than something buyers type.
    "acronym_min_volume": 20,
    # A real industry program does not do six figures a month. Anything that big
    # is a common word the miner mistook for an acronym — the volume ceiling is
    # the second line of defence behind the lowercase-word test, because it
    # catches the case where the word never appears in lowercase on the page.
    # PACP, the biggest genuine one seen, is 2,400/mo. (2026-08-12)
    "acronym_max_volume": 50000,
    # A bare acronym carrying this many times its qualified sibling's volume is
    # probably being searched for a different meaning. 8x is well clear of the
    # 1.5x a genuinely owned term shows and well under the 72x/147x of a
    # collision. Only applied to acronyms above acronym_collision_min_volume,
    # since a small number proves nothing either way.
    "acronym_collision_ratio": 8.0,
    "acronym_collision_min_volume": 100,
    # Monthly searches a proposed replacement term must clear to be offered.
    "replacement_min_volume": 20,
    "metro_no_suffix_zips": 25,
    "metro_no_suffix_share": 0.6,
    "grid_state_suffix": "auto",       # auto = suffix only cities that need it
}

def r50(x):
    return int(round(x / 50.0) * 50)


def r50up(x):
    """Round UP to the next $50 — the CLIENT-FACING convention (2026-08-26).

    Everything used to round to NEAREST, which rounds DOWN half the time: a
    59% AI Search share of a $3,900 core computes $2,301 and was quoted at
    $2,300. Kiri: a client price never rounds down. The partner side and the
    tier STEP are untouched — those are internal shape, not the number on the
    proposal — so only the figures a client reads move, and only ever by less
    than $50.

    Switchable at /config: client_round_mode "up" (default) or "nearest",
    because it moves every quote and a calibration run may want the old basis.
    """
    import math
    if str(CFG.get("client_round_mode", "up")).strip().lower() == "nearest":
        return int(round(round(x, 6) / 50.0) * 50)
    return int(math.ceil(round(x, 6) / 50.0) * 50)


def perf_eligibility(rows, site_rebuild=""):
    """Is this client a candidate for pay-for-performance at all?

    Brendan only offers it where the terms are ALREADY within the first five
    pages — a client starting from scratch is 6-12+ months from any of it
    ranking, so there is nothing to bill against and the model does not work.

    Reads the rank table, counts only rows that were actually measured, and
    returns the numbers rather than a verdict alone: "not eligible" and "we
    could not measure enough of it" are different answers and only one of them
    is about the client.
    """
    depth = int(CFG.get("perf_page_depth", 50) or 50)
    # A REBUILD ANSWERS THIS BEFORE THE RANK TABLE DOES. Brendan: "if they're
    # starting from scratch / not ranking we generally don't offer pay for
    # performance SEO because the time frame to results can be 6-12+ months."
    # The gate cannot see it on its own -- it reads positions on the OUTGOING
    # site, so a client whose current site ranks well comes back eligible for a
    # deal billed against rankings that are about to move. True on a same-domain
    # rebuild too: redirects hold most positions, but "most" is not the basis
    # for a per-ranking contract. (2026-08-24)
    _rb = str(site_rebuild or "").strip().lower()
    if _rb in ("new", "same") and CFG.get("rebuild_vetoes_performance", True):
        return {"eligible": False, "measured": 0, "within": 0, "share": 0,
                "depth": depth, "min_share": round(float(
                    CFG.get("perf_eligible_min_share", 0.5) or 0.5) * 100),
                "enough": True, "unmeasured": 0, "site_rebuild": _rb,
                "reason": ("this client is moving to a NEW DOMAIN, so there are "
                           "no rankings to bill against and the first ones are "
                           "6-12+ months out" if _rb == "new" else
                           "this client's site is being rebuilt, so today's "
                           "positions describe a site that is going away — not a "
                           "basis for a contract billed per ranking")}
    measured, within, unmeasured = 0, 0, 0
    # QUEUED IS NOT FAILED. A row still in Google's queue fills in by itself;
    # a row that errored needs Retry pressed. Step 3 has always drawn that
    # distinction and this told the operator to "Retry those 1" about a row
    # with no Retry to press. (2026-08-22)
    queued, failed = 0, 0
    for r in (rows or []):
        if not isinstance(r, dict):
            continue
        if r.get("error") or r.get("queued") or r.get("expired"):
            unmeasured += 1
            if r.get("queued"):
                queued += 1
            else:
                failed += 1
            continue
        pos = r.get("pos", r.get("position"))
        if isinstance(pos, bool):
            continue
        if isinstance(pos, int):
            measured += 1
            if 1 <= pos <= depth:
                within += 1
        elif str(pos).strip().lower() in ("not found", "not ranking"):
            measured += 1
    share = (within / measured) if measured else 0.0
    min_n = int(CFG.get("perf_min_measured", 5) or 5)
    min_share = float(CFG.get("perf_eligible_min_share", 0.5) or 0.5)
    if measured < min_n:
        return {"eligible": False, "measured": measured, "within": within,
                "share": round(share * 100), "depth": depth,
                "min_share": round(min_share * 100), "enough": False,
                "unmeasured": unmeasured,
                "reason": f"only {measured} terms have been measured — "
                          f"the rank check has to land before this can be judged"}
    ok = share >= min_share
    # A FEW TIMEOUTS MUST NOT DECIDE THIS. Ooten came back eligible at 65% on
    # one run and "not a pay-for-performance client" at 47% on the next — same
    # client, same list, three lookups that happened to time out. With
    # seventeen rows measured against a 50% line, one term either way settles
    # it. So ask whether the rows that did NOT answer could still change the
    # verdict, and if they could, do not give one. (2026-08-22)
    total = measured + unmeasured
    if unmeasured and total:
        worst, best = within / total, (within + unmeasured) / total
        if worst < min_share <= best:
            if queued and not failed:
                what = (f"{queued} row{'' if queued == 1 else 's'} still in "
                        f"Google's queue — {'it fills' if queued == 1 else 'they fill'} "
                        f"in without you, and this settles itself when "
                        f"{'it lands' if queued == 1 else 'they land'}")
            elif failed and not queued:
                what = (f"{failed} lookup{'' if failed == 1 else 's'} failed — "
                        f"retry {'it' if failed == 1 else 'them'} and this "
                        f"decides itself")
            else:
                what = (f"{queued} still queued and {failed} failed — the queued "
                        f"{'one fills' if queued == 1 else 'ones fill'} in on "
                        f"{'its' if queued == 1 else 'their'} own, the failed "
                        f"{'one needs' if failed == 1 else 'ones need'} a retry")
            return {"eligible": False, "measured": measured, "within": within,
                    "share": round(share * 100), "depth": depth,
                    "min_share": round(min_share * 100), "enough": False,
                    "unmeasured": unmeasured, "queued": queued, "failed": failed,
                    "reason": f"too close to call — {within} of {total} terms are "
                              f"inside the first {depth // 10} pages, and "
                              f"{'that last one is' if unmeasured == 1 else f'those {unmeasured} are'} "
                              f"enough to settle it either way. {what}"}
    return {"eligible": ok, "measured": measured, "within": within,
            "share": round(share * 100), "depth": depth, "enough": True,
            "min_share": round(min_share * 100), "unmeasured": unmeasured,
            "reason": (f"{within} of {measured} measured terms already rank "
                       f"inside the first {depth // 10} pages"
                       if ok else
                       f"only {within} of {measured} measured terms rank inside "
                       f"the first {depth // 10} pages — this client is starting "
                       f"close enough to scratch that the first rankings are "
                       f"6-12+ months out")}


def grid_suffix(keywords):
    """The market suffix every grid row shares, recovered from the rows.

    _strip_markets removes the ENTERED market, and since the grid started
    choosing its own wording those two can differ: Ooten was entered as "Knox
    County, TN" and quoted on "knoxville tn", so nothing matched and the bare
    form never came out — which is why eleven rows had no bid to price on. The
    grid is the authority on its own suffix; same trick gridSuffix() uses in
    the panel. (2026-08-22)
    """
    parts = [str(k or "").strip().split() for k in (keywords or []) if str(k or "").strip()]
    parts = [p for p in parts if p]
    if len(parts) < 2:
        return ""
    tail, shortest = [], min(len(p) for p in parts)
    for i in range(1, shortest):
        w = parts[0][-i]
        if all(p[-i] == w for p in parts):
            tail.insert(0, w)
        else:
            break
    return " ".join(tail)


def suffix_market(keywords, state=""):
    """The grid's own suffix expressed as a MARKET, or "".

    THE ROWS ARE THE AUTHORITY ON WHERE THIS QUOTE IS MEASURED. The grid picks
    its own wording now, so a client entered as "Whatcom County, WA" is quoted
    on "... bellingham wa" -- and every SERP was still requested for Whatcom
    County, because the market list is what the operator typed. Google then
    answered a county-wide result page for a Bellingham phrase, and the
    zero-ranking uplift keyed off it. The grid form and the rank location have
    to be the same place. (2026-08-24)

    Read off the rows that carry a state abbreviation, so the near-me terms --
    which have no suffix at all -- cannot collapse the common tail to nothing.
    """
    abbrs = set(STATE_ABBREV.values())
    geo = [str(k or "").strip() for k in (keywords or []) if str(k or "").strip()]
    # "near me" ENDS IN MAINE. Same trap as the market list, which is why
    # is_non_place_geo exists: ME is a real abbreviation, so a proximity row
    # passes the suffix test and then shares no tail with anything, collapsing
    # the common suffix to nothing. Strip the proximity rows first -- they are
    # deliberately place-less and have no business naming the SERP location.
    geo = [k for k in geo if not _PROXIMITY_RE.search(" " + k.lower())]
    geo = [k for k in geo if k.split() and k.split()[-1].lower() in abbrs]
    if len(geo) < 2:
        return ""
    sfx = grid_suffix(geo)
    toks = sfx.split()
    if len(toks) < 2 or toks[-1].lower() not in abbrs:
        return ""
    city = " ".join(toks[:-1]).strip()
    if not city:
        return ""
    return f"{city.title()}, {toks[-1].upper()}"


def rank_markets(keywords, markets, state, national=False):
    """`markets` reordered so the grid's own suffix leads, when it names one.

    Only ever PREPENDS -- the entered markets keep their order behind it, so
    the per-keyword lookups that match on a market name are unaffected.
    """
    if national:
        return markets
    m = suffix_market(keywords, state)
    if not m:
        return markets
    bare = parse_market(m, state)[0].strip().lower()
    rest = [x for x in (markets or [])
            if parse_market(x, state)[0].strip().lower() != bare]
    return [m] + rest


def perf_candidates(d, want=None):
    """Terms that could join the performance table but are not in the grid.

    Drawn from what the demand ranking cut — already measured for volume, not
    yet rank-checked — and qualified with the grid's own suffix so they read
    like every other row. Returns the terms only; nothing enters the table
    until a rank has actually been measured for it.
    """
    want = int(want or CFG.get("perf_table_terms", 50) or 50)
    rows = _proposal_rows(d)
    have = {re.sub(r"\s+", " ", r["kw"].strip().lower()) for r in rows}
    short = max(0, want - len(rows))
    if not short:
        return []
    sfx = grid_suffix([r["kw"] for r in rows])
    sr = ((d.get("kw") or {}).get("seed_ranking") or {})
    out = []
    for item in (sr.get("order") or []):
        if not isinstance(item, (list, tuple)) or not item:
            continue
        bare = str(item[0] or "").strip().lower()
        if not bare:
            continue
        kw = clean_kw(f"{bare} {sfx}".strip())
        k = re.sub(r"\s+", " ", kw.lower())
        if not kw or k in have:
            continue
        have.add(k)
        out.append(kw)
        if len(out) >= short:
            break
    return out


_PERF_TOPIC_CACHE = {}


def perf_topics(d):
    """Topics WITH their member terms, for the Practice Area column.

    A quote saved before the topics payload carried its members has labels and
    nothing else, and the column comes out blank — which is where Ooten sat for
    four builds. Rebuilding step 1 fixes it, but a quote that has already been
    priced should not have to be rebuilt to fill a column.

    So fall back to claude_topics(), which is the SAME call the build uses and
    which assigns every term to a topic as part of its answer. One call, the
    answer the build would have given.

    NOT topic_clusters(): it is deterministic and free, and I tried it first.
    It groups by token overlap, so Ooten's list came back as "personal injury"
    and "accident" rather than the two practice areas and put "sex crime
    attorney" under Personal Injury. Wrong is worse than blank here.
    (2026-08-22)
    """
    kw = d.get("kw") or {}
    topics = [t for t in (kw.get("topics") or []) if isinstance(t, dict)]
    if any(t.get("terms") for t in topics):
        return topics
    seeds = list(((kw.get("seed_ranking") or {}).get("was")) or [])
    if not seeds:
        return topics
    # refreshPerf() runs on every step-3 render, so without this the fallback
    # would buy an AI call each time the rank table repainted.
    ck = hashlib.sha1(("|".join(sorted(str(x).lower() for x in seeds))
                       ).encode()).hexdigest()
    if ck in _PERF_TOPIC_CACHE:
        return _PERF_TOPIC_CACHE[ck] or topics
    try:
        rebuilt = claude_topics(seeds, d.get("business_desc") or "",
                                d.get("brand") or "")
    except Exception:                                     # noqa: BLE001
        return topics
    out = [{"label": str(t.get("label") or "").strip(),
            "terms": list(t.get("seeds") or [])}
           for t in (rebuilt or []) if t.get("seeds")]
    out = [t for t in out if t["label"]]
    _PERF_TOPIC_CACHE[ck] = out
    return out or topics


def perf_area(term, topics):
    """The practice area / service line a term belongs to.

    Brendan fills this column by hand — Personal Injury against Criminal
    Defense, 25 each on Ooten. The tool already works the same split out at
    build time: topic_clusters() groups the operator's seeds into the things
    the client actually sells, which for a law firm IS the practice areas.
    Reuse it rather than asking anyone to type it fifty times.

    Matched on stems, so "wrongful death attorney knoxville tn" finds the
    personal-injury topic through "injury" only if the label carries it — where
    the label does not overlap the term at all the cell stays empty rather than
    guessing, because a wrong practice area in a client document is worse than
    a blank one. (2026-08-22)
    """
    if not topics:
        return ""
    stems = _topic_tokens(term)
    # A WORD IN EVERY TOPIC DECIDES NOTHING. "lawyer" and "attorney" are role
    # words a law firm's whole list carries, so they overlap both practice areas
    # and whichever sorted first would have won. Drop the shared vocabulary and
    # score on what actually separates the topics. General, not a legal
    # special case: the same is true of "repair" across a trade's service
    # lines. (2026-08-22)
    _vocab = []
    for t in (topics or []):
        if not isinstance(t, dict):
            continue
        v = set()
        for m in (t.get("terms") or []):
            v |= _topic_tokens(str(m or ""))
        v |= _topic_tokens(str(t.get("label") or ""))
        _vocab.append(v)
    _shared = set()
    for i, a in enumerate(_vocab):
        for b in _vocab[i + 1:]:
            _shared |= (a & b)
    scores = []
    for t in (topics or []):
        if not isinstance(t, dict):
            continue
        lab = str(t.get("label") or "").strip()
        if not lab:
            continue
        # MEMBERSHIP FIRST. The topic carries the seeds that built it, so a term
        # made from one of them is that topic — no guessing.
        members = [str(m or "").strip().lower() for m in (t.get("terms") or [])]
        low = " " + re.sub(r"\s+", " ", term.strip().lower()) + " "
        if any(m and (" " + m + " ") in low for m in members):
            return " ".join(w if w.isupper() else w.capitalize() for w in lab.split())
        # Otherwise the topic's whole vocabulary, which is far wider than its
        # label: the personal-injury topic carries "dog bite" and "slip and
        # fall" even though its label carries neither.
        toks = set()
        for m in members:
            toks |= _topic_tokens(m)
        toks |= _topic_tokens(lab)
        n = len((stems & toks) - _shared)
        scores.append((n, lab))
    scores.sort(reverse=True)
    best, best_n = (scores[0][1], scores[0][0]) if scores else ("", 0)
    # A TIE IS NOT AN ANSWER. "lawyer knoxville tn" overlaps both topics on the
    # one word they share, and whichever sorted first would have been printed as
    # this term's practice area. A blank cell is honest; a coin flip in a client
    # document is not.
    if len(scores) > 1 and scores[0][0] == scores[1][0]:
        return ""
    if not best or not best_n:
        return ""
    # Title Case for a client document: the labels come back lower case.
    return " ".join(w if w.isupper() else w.capitalize() for w in best.split())


def perf_term_price(bid):
    """Cost Page 1 / Top 5 / Top 3 / #1 for one term, from its measured bid."""
    try:
        b = float(bid or 0)
    except (TypeError, ValueError):
        b = 0.0
    step = int(CFG.get("perf_page1_round", 5) or 5)
    p1 = max(int(CFG.get("perf_page1_floor", 80) or 0),
             int(round(b * float(CFG.get("perf_page1_mult", 2.1)) / step) * step))
    ladder = (CFG.get("perf_ladder_high") if p1 >= int(CFG.get("perf_high_knee", 600))
              else CFG.get("perf_ladder")) or [1.6, 2.4, 3.6]
    t5, t3, one = (int(round(p1 * float(x) / step) * step) for x in ladder[:3])
    return {"page1": p1, "top5": t5, "top3": t3, "one": one}


def perf_tier_label(pos):
    """The 'Current Achieved Tier' column — what they are billed at today."""
    if isinstance(pos, int):
        if pos <= 0:
            return "Not ranking"
        if pos == 1:
            return "#1"
        if pos <= 3:
            return "Top 3"
        if pos <= 5:
            return "Top 5"
        if pos <= 10:
            return "Page 1"
        return "Not Page 1"
    return "Not ranking"


def addon_discount_pct(n):
    """The one Add-On Market % that applies to EVERY add-on market on a quote.

    Flat by bracket. A quote with 12 add-on markets is 15% off all twelve, not
    nine at 10% and three at 15% — so there is a single percentage to put on
    the quote and defend, rather than a blended rate that matches no published
    number.
    """
    n = max(0, int(n or 0))
    if n <= 0:
        return 0.0
    for lo, pct in (CFG.get("addon_volume_discount_tiers") or []):
        try:
            if n >= int(lo):
                return float(pct)
        except (TypeError, ValueError):
            continue
    return 0.0

# REQUEST TIMEOUT BUDGET (2026-07-27).
# dfs_post defaulted to a 120s timeout while gunicorn's default worker timeout
# is 30s, so a slow DataForSEO call got the worker killed BEFORE the app's own
# timeout could fire. A killed worker returns no body at all, which the
# frontend can only report as "Server 500 (timeout or non-JSON)" — the app
# never got the chance to say what went wrong. Step 2 made this worse by
# retrying up to three locations in sequence: 3 x 120s against a 30s ceiling.
#
# Two changes: per-call timeouts now fit inside a server window (DFS_TIMEOUT),
# and multi-call stages carry a wall-clock DEADLINE so the chain stops itself
# and returns a readable partial result instead of being killed mid-flight.
#
# IMPORTANT: also raise the server's own timeout, or long calls still die.
# In Render, set the start command to:
#     gunicorn app:app --timeout 120 --workers 1 --threads 4
# Threads, not workers: nearly all of this app's time is spent WAITING on
# DataForSEO and Anthropic, so concurrency should come from threads (one
# process's memory) rather than a second worker process, which doubles memory
# on a 512 MB Starter instance for no gain on I/O-bound work.
# REQUEST_BUDGET_S must stay comfortably BELOW that number.
DFS_TIMEOUT     = int(os.environ.get("DFS_TIMEOUT", "25"))      # per API call
REQUEST_BUDGET_S = int(os.environ.get("REQUEST_BUDGET_S", "90"))  # per route


class BudgetExceeded(Exception):
    """Raised when a multi-call stage runs out of wall clock. Carries a message
    the operator can act on rather than a stack trace."""


def _deadline(budget=None):
    """Start a wall-clock budget for the current request."""
    return _time.time() + float(budget or REQUEST_BUDGET_S)


def _remaining(deadline, minimum=5):
    """Seconds left before the deadline, or None if the budget is spent."""
    left = deadline - _time.time()
    return left if left >= minimum else None


# Seconds to wait out a per-minute rate limit. Long enough for a slot to free up,
# short enough to stay inside the per-route budget.
DFS_RATE_LIMIT_WAIT = 8.0
# The rate-limit family. 40202 is the per-minute cap; the neighbours are the
# per-second and concurrent-task versions of the same refusal.
_DFS_RATE_CODES = {40202, 40203, 40204, 40205}


# ---------------------------------------------------------------------------
# DON'T HIT THE LIMIT IN THE FIRST PLACE.
#
# The 12-per-minute cap was handled entirely by reacting to it: one retry after
# an 8-second wait, and if that attempt landed in the same minute it was refused
# again and the caller got zeros. Ski Barn's rebuild made the whole grid read
# "(no data)": every exact-match volume came back empty, so total volume was 0,
# so the price lost its volume component AND the panel advised switching the
# client to national demand — "31,000/mo bare vs 0/mo with a city attached" is
# what a throttled geo lookup looks like, not what a national business looks
# like. A rate limit turned into a pricing decision and a strategy
# recommendation.
#
# So requests are PACED. A token bucket holds the process under the cap, and
# every caller waits its turn rather than firing and being refused. It is the
# same total wall-clock — the calls were always going to take a minute — but
# they come back with data. (2026-08-20)
_DFS_BUCKET_LOCK = threading.Lock()
_DFS_BUCKET = []          # timestamps of recent sends, newest last


def _dfs_take_slot():
    """Block until sending now keeps us under the per-minute cap."""
    cap = int(CFG.get("dfs_calls_per_minute", 10) or 0)
    if cap <= 0:
        return
    while True:
        with _DFS_BUCKET_LOCK:
            now = _time.monotonic()
            while _DFS_BUCKET and now - _DFS_BUCKET[0] >= 60.0:
                _DFS_BUCKET.pop(0)
            if len(_DFS_BUCKET) < cap:
                _DFS_BUCKET.append(now)
                return
            wait = 60.0 - (now - _DFS_BUCKET[0]) + 0.05
        time.sleep(max(0.05, min(wait, 60.0)))


def _dfs_rate_limited(data):
    """Is this HTTP 200 actually a rate-limit refusal?"""
    if not isinstance(data, dict):
        return False
    if int(data.get("status_code") or 0) in _DFS_RATE_CODES:
        return True
    for t in (data.get("tasks") or []):
        if isinstance(t, dict) and int(t.get("status_code") or 0) in _DFS_RATE_CODES:
            return True
    return False


def dfs_post(path, payload, timeout=None, method="POST", retries=1):
    """One DataForSEO call, retried once on a TRANSIENT failure.

    There was no retry at all, so a single read timeout was fatal to whatever
    depended on it. On a Ski Barn quote the volume lookup timed out once and
    the whole volume component of the price silently became $0 (2026-08-04).

    Only network-level failures and 5xx are retried — a 4xx is a real answer
    about the request and repeating it just wastes the budget. Two attempts at
    the 25s default fit inside the 90s per-route budget.

    RATE LIMITS ARE THE EXCEPTION, and they do not arrive as an HTTP status: a
    "40202: The rates limit per minute has been exceeded: 12 >= 12" comes back
    inside an HTTP 200, in the TASK, exactly like the 40501 that once turned every
    rank check into "Not Found". Ski Barn lost its whole volume component to one
    ($0 volume, every keyword reading "no data") on a build that had simply made
    thirteen calls in a minute. A per-minute limit is transient by definition, so
    it is retried — but after a real pause, not the 1s a timeout gets, because
    retrying immediately just spends another slot on the same refusal.
    (2026-08-12)
    """
    if timeout is None:
        timeout = DFS_TIMEOUT
    login = os.environ.get("DFS_LOGIN", "")
    pw    = os.environ.get("DFS_PASSWORD", "")
    # No credentials means no call can succeed, and pacing a call that cannot
    # be made spends a minute of the per-minute budget on a guaranteed 401.
    # Fail on the first one instead of queueing every later call behind it.
    if not login or not pw:
        raise requests.HTTPError(
            "DataForSEO credentials are not set (DFS_LOGIN / DFS_PASSWORD).")
    token = base64.b64encode(f"{login}:{pw}".encode()).decode()
    hdrs = {"Authorization": f"Basic {token}", "Content-Type": "application/json"}
    last = None
    # A rate limit is not a failure of the request, it is a failure of the
    # PACING, and it is free to wait for. Given its own budget so a throttled
    # minute cannot spend the transient-failure retries.
    rl_left = int(CFG.get("dfs_rate_limit_retries", 4) or 0)
    attempt = 0
    total = int(retries)
    while attempt <= total:
        try:
            _dfs_take_slot()
            if method == "GET":
                resp = requests.get(BASE + path, headers=hdrs, timeout=timeout)
            else:
                resp = requests.post(BASE + path, headers=hdrs,
                                     data=json.dumps(payload), timeout=timeout)
            if resp.status_code >= 500 and attempt < total:
                last = requests.HTTPError(f"HTTP {resp.status_code}")
                attempt += 1
                time.sleep(1.0 + attempt)
                continue
            resp.raise_for_status()
            data = resp.json()
            if _dfs_rate_limited(data):
                if rl_left > 0:
                    rl_left -= 1
                    last = RuntimeError("40202: rate limit per minute exceeded")
                    # Long enough to leave the minute that refused us, and the
                    # bucket above means the next attempt is paced rather than
                    # racing whatever else is in flight.
                    app.logger.warning("dfs %s rate-limited, waiting %.0fs (%d left)",
                                       path, DFS_RATE_LIMIT_WAIT, rl_left)
                    time.sleep(DFS_RATE_LIMIT_WAIT)
                    continue
                # Out of patience: hand the refusal back AS a refusal. It must
                # never reach a caller looking like an empty result.
                raise RuntimeError("40202: rate limit per minute exceeded")
            return data
        except (requests.Timeout, requests.ConnectionError) as e:
            last = e
            if attempt >= total:
                break
            attempt += 1
            time.sleep(1.0 + attempt)
    raise last

def recommend_addons(markets, state, rows, top_n=None, site_locations=None,
                     site_pages_found=None, metro_groups=None, city_volumes=None):
    """Suggest how many markets should be priced as separate campaigns.

    The judgement, per the pricing authority: 2-3 related nearby markets run
    under one campaign; many locations become separate campaigns; what tips it
    is how different the markets are, how competitive they are, and the
    client's budget. The last of those the tool cannot see, so this is a
    suggestion with its reasoning attached, never an answer.

    Count alone doesn't separate the two real proposals — Skills of Central PA
    got ONE campaign across 8 Pennsylvania towns while TN Water & Air was
    scoped to Knoxville with 12 markets offered as add-ons. What does separate
    them is whether the client's footprint is ALREADY multi-market:

      Skills ranks 1st/1st/1st/2nd/2nd/3rd across its towns — one domain
      already serves them, so the work is improving a presence that exists.
      TN ranks nowhere outside Knoxville — every other market is greenfield,
      which is a separate build each.

    "How many markets" is the wrong question. "Are we improving one footprint
    or creating eleven" is the right one, and step 3 already answers it.
    """
    mk = [m for m in (markets or []) if m and m.strip()]
    # Collapse cities Google Ads treats as one place. Counting pills instead of
    # markets inflates everything downstream: fourteen metro-Atlanta suburbs
    # are two or three markets, not fourteen, and an add-on count built on
    # fourteen is four times too big.
    # "Thin" = fewer than half the markets return any volume at all. Below
    # that, grouping has nothing to work with and the count is guesswork.
    _cv = city_volumes or {}
    thin_volume = bool(_cv) and sum(1 for v in _cv.values() if v) < max(2, len(_cv) / 2)
    collapsed = 0
    for g in (metro_groups or []):
        members = [m for m in mk if m in g]
        collapsed += max(0, len(members) - 1)
    n = len(mk) - collapsed
    out_pills = len(mk)
    out = {"markets": n, "pills": out_pills, "collapsed": collapsed,
           "markets_absent": [],
           "metro_groups": [g for g in (metro_groups or []) if len(g) > 1],
           "suggested": 0, "basis": "", "covered": 0,
           "measured": 0, "unmeasured": 0, "states": 0, "confident": False,
           "site_locations": 0}
    if n <= 1:
        # A single market is a CONFIDENT zero, not an absence of information.
        # Leaving confident False made the panel say "not enough data to
        # suggest" and — worse — meant the stepper was never clamped, so a
        # field left at 5 from an earlier run kept charging for five add-on
        # markets that the tool had just determined don't exist (2026-08-03).
        out["basis"] = "single market — nothing to add on."
        out["confident"] = True
        return out

    # Only count states we actually KNOW — a market tagged "City, ST". Inferring
    # them from a name-collision table and then telling the operator the client
    # "spans 2 states" is a confident claim built on a guess.
    states = {(parse_market(m, "")[1] or "").lower() for m in mk if "," in m}
    states.discard("")
    out["states"] = len(states)

    if n <= int(CFG.get("addon_free_markets", 3)):
        out["basis"] = f"{n} markets — three or fewer run as one campaign."
        out["confident"] = True
        return out


    # ABSTAIN when the markets carry almost no search volume. Grouping infers
    # market identity from demand patterns, and a town with no demand has no
    # pattern — three separate attempts at this (geo-modified probe, resolved
    # location name, per-city volume) all failed on Brent Cogan's seven Blair
    # County towns for the same underlying reason: the data does not contain
    # the answer (2026-08-03). Counting them as seven separate markets is a
    # confident wrong answer worth $6,000/mo, which is worse than no answer.
    if thin_volume:
        out["basis"] = (f"{n} markets entered, but they carry almost no search "
                        f"volume individually — there is no demand pattern to tell "
                        f"whether they are one market or several. Decide by hand: "
                        f"towns within a short drive of each other are normally one "
                        f"campaign.")
        return out


    # PRESENCE PER MARKET, not a count of locations. Counting gets both known
    # outcomes backwards: Skills of Central PA has ~8 facilities and got ONE
    # campaign, TN Water & Air has one Knoxville location and was offered 11
    # add-ons. Physical location count is nearly anti-correlated with add-on
    # count, because the question isn't "how many places do they have" — it's
    # "which of the markets we're targeting do they already operate in".
    #
    #   Skills targets 8 markets and has a presence in all 8  -> improving
    #     coverage that exists -> one campaign.
    #   TN targets 12 and has a presence in 1                 -> entering 11
    #     new markets -> a campaign each.
    #
    # A location page or a Google Business Profile answers that per market,
    # and so does a top-100 ranking. All three are the same question asked
    # different ways, which is why they agree where we can check them.
    locs = [str(l).lower() for l in (site_locations or []) if l]
    out["site_locations"] = len(locs)
    # "No location pages" and "we couldn't read the site" look identical from
    # an empty list, and they mean opposite things: the first is evidence, the
    # second is a gap. Woodstock's site refused a TLS handshake, so its crawl
    # returned nothing and the strongest signal was silently absent rather
    # than negative.
    out["site_read"] = (None if site_pages_found is None
                        else bool(site_pages_found))
    if locs:
        with_page = [m for m in mk
                     if any(l in (parse_market(m, state)[0] or m).lower()
                            or (parse_market(m, state)[0] or m).lower() in l
                            for l in locs)]
        # Presence has to be counted in MARKETS too. Counting cities against a
        # collapsed market total gave "present in 5 of 2 markets", and the
        # missing-market arithmetic went negative.
        def _key(m):
            g = next((tuple(sorted(x.lower() for x in grp))
                      for grp in (metro_groups or []) if m in grp), None)
            return g or m.lower()
        present_keys = {_key(m) for m in with_page}
        out["markets_with_location_page"] = len(present_keys)
        # Name the gaps. A count tells the operator a decision was made; the
        # list tells them WHICH markets it rests on, which is the part they can
        # actually check against what they know about the client. When cities
        # have been collapsed into metros, report one name per metro so the
        # list matches the market count rather than the pill count.
        absent, seen_grp = [], set()
        for m in mk:
            if m in with_page:
                continue
            grp = next((tuple(g) for g in (metro_groups or []) if m in g), None)
            if grp:
                if grp in seen_grp or any(x in with_page for x in grp):
                    continue
                seen_grp.add(grp)
                absent.append(f"{m} +{len(grp)-1}" if len(grp) > 1 else m)
            else:
                absent.append(m)
        out["markets_absent"] = absent
        if present_keys:
            missing = n - len(present_keys)
            if missing <= 0:
                out["suggested"] = 0
                out["basis"] = (f"The client has a presence in all {n} markets — one "
                                f"footprint to improve, not {n} to build.")
            elif len(present_keys) >= max(2, int(n * 0.7)):
                out["suggested"] = 0
                out["basis"] = (f"The client has a presence in {len(present_keys)} of "
                                f"{n} markets — operating in most, so one footprint "
                                f"rather than {n} builds.")
            else:
                out["suggested"] = missing
                out["basis"] = (f"The client has a presence in {len(present_keys)} of "
                                f"{n} markets. They aren't in the other {missing} yet, "
                                f"so each of those is a campaign from scratch.")
            out["confident"] = True
            return out

    # Which markets does the client already rank in? A keyword carries its city,
    # so match each market against the terms that mention it.
    top = int(top_n or CFG.get("zero_ranking_top_n", 100))
    covered, measured = set(), set()
    for r in (rows or []):
        kw = (r.get("kw") or "").lower()
        pos = r.get("pos")
        for m in mk:
            city = (parse_market(m, state)[0] or m).strip().lower()
            if city and city in kw:
                measured.add(m)
                if isinstance(pos, (int, float)) and pos <= top:
                    covered.add(m)
    out["covered"], out["measured"] = len(covered), len(measured)

    out["unmeasured"] = n - len(measured)
    if len(measured) < 2:
        out["basis"] = (f"Rankings measured in only {len(measured)} markets — run "
                        f"step 3 first.")
        return out

    # Coverage has to be read against the markets ENTERED, not the ones that
    # happened to be measured. The grid crosses only the top few cities by
    # demand, so a client can show "5 of 5 measured" while nine markets were
    # never looked at — and those nine are precisely the ones most likely to be
    # greenfield, because they were the LOWEST-demand markets. Concluding "one
    # footprint" from the best five is the tool marking its own homework.
    meas_share = len(measured) / n
    if meas_share < float(CFG.get("addon_min_measured_share", 0.7)):
        out["basis"] = (f"Only {len(measured)} of {n} markets have rank data, and "
                        f"the unmeasured ones are the lowest-demand — the ones "
                        f"the client is least likely to already be in. Raise the Grid max "
                        f"cities cap \u2014 it limits markets of any kind, counties "
                        f"included \u2014 and re-run step 3.")
        return out

    share = len(covered) / len(measured)
    if share >= float(CFG.get("addon_covered_share", 0.6)):
        out["suggested"] = 0
        out["basis"] = (f"Already ranking in {len(covered)} of {len(measured)} "
                        f"markets — one footprint to improve, not {n} to build.")
        out["confident"] = True
    else:
        out["suggested"] = max(0, n - 1)
        out["basis"] = (f"Ranking in only {len(covered)} of {len(measured)} markets, "
                        f"so the rest are a campaign from scratch each.")
        out["confident"] = True
    if len(states) > 1:
        out["suggested"] = max(out["suggested"], n - 1)
        out["basis"] += f" They span {len(states)} states — separate territories."
    return out


def primary_first(markets, primary):
    """Put the highest-demand market at the front of the list.

    Bid and rank lookups localise to markets[0], which was whatever the partner
    typed first — usually alphabetical, so Acworth (210/mo) anchored a quote
    whose real market was Marietta (880/mo). Step 1 already ranks the markets
    by actual demand for the client's service; steps 2 and 3 should use that
    answer rather than input order. Two signals that drive price — the CPC
    adder and the zero-ranking uplift — were both being measured in the wrong
    town.
    """
    mk = [m for m in (markets or []) if m and m.strip()]
    p = (primary or "").strip()
    if not p:
        return mk
    rest = [m for m in mk if m.strip().lower() != p.lower()]
    return [p] + rest


def home_state(markets, state=""):
    """The state the client actually operates in — the one most of their
    markets sit in. Not the same question as "which market has most demand"."""
    counts = {}
    for m in (markets or []):
        st = market_state(m, state)
        if st:
            counts[st] = counts.get(st, 0) + 1
    if not counts:
        return (state or "").strip()
    top = max(counts.values())
    # Ties break on input order: the partner types the client's own town first.
    for m in (markets or []):
        st = market_state(m, state)
        if st and counts.get(st) == top:
            return st
    return ""


def measure_first(markets, state="", primary=""):
    """Order markets for LOCALISED MEASUREMENT — rank checks and bid lookups.

    primary_first() answers "where is the most demand", which is the right
    question for scoping a campaign and the wrong one for measuring a client.
    Ski Barn's markets were Wayne / Paramus / Shrewsbury / Lawrenceville NJ plus
    New York City: NYC carries an order of magnitude more demand, so it became
    the primary and the rank check asked whether a New Jersey ski shop outranks
    Manhattan. It does not, and cannot — 0/20, largest zero-ranking uplift
    (2026-08-07). Their own agency's report says as much: "NYC is a tracked
    market but the client's locations are all in NJ."

    So measurement stays inside the HOME state — the state most of the markets
    sit in — and picks the highest-demand market there. A genuinely multi-state
    client (no state holds a majority) falls back to primary_first unchanged.
    """
    mk = primary_first(markets, primary)
    # Measure in a TOWN, not a county. Both are valid Google locations, but a
    # county SERP is a wider net than the client competes in, and the entered
    # list often leads with counties simply because that is the order someone
    # types a service area. Junk Bee Gone's first ten entries were counties, so
    # rankings were measured county-wide. (2026-08-10)
    # A MEASURED PRIMARY OUTRANKS THIS HEURISTIC. The rule below is a guess
    # about which of several typed markets to measure in; `primary` is the
    # market Step 1 actually measured the most demand in, and a guess must not
    # overturn a measurement. Whatcom County was the one market of five carrying
    # demand, so the grid crossed it and the panel named it -- and then this
    # pushed it to the back for being a county, which handed the rank check to
    # Camano Island, eighty miles south in another county. Every row then asked
    # "does this company rank for a WHATCOM COUNTY phrase, as seen from Camano
    # Island". The grid and the rank check have to measure the same place.
    # (2026-08-24)
    _pinned = (mk[0] if (mk and (primary or "").strip()
                         and mk[0].strip().lower() == (primary or "").strip().lower())
               else None)
    if any(county_key(m, state) for m in mk) and any(not county_key(m, state) for m in mk):
        mk = ([m for m in mk if not county_key(m, state)]
              + [m for m in mk if county_key(m, state)])
        if _pinned:
            mk = [_pinned] + [m for m in mk if m is not _pinned]
    hs = home_state(mk, state)
    if not hs:
        return mk
    inside = [m for m in mk if market_state(m, state) == hs]
    if not inside or len(inside) == len(mk):
        return mk
    # Majority test: one outlier market must not be able to move the anchor,
    # but a real 50/50 two-state footprint keeps the demand ordering.
    if len(inside) * 2 <= len(mk):
        return mk
    return inside + [m for m in mk if m not in inside]


# Strings that arrive in the markets list but are not places. They come from
# imported reports: a ranking table's market column carries "near me" rows
# because the agency tracked "junk removal near me" as its own row, and the
# importer added it as a market like any other.
#
# One of these reaching the front of the list is not cosmetic. loc_string takes
# markets[0], so "near me, TN" became the location_name on every SERP and
# volume call — a location Google has never heard of. Junk Bee Gone was measured
# "in near me, Tennessee", scored 0/35, and drew the largest zero-ranking uplift
# in the table. (2026-08-10)
_NON_PLACE_GEOS = {
    "near me", "nearme", "near by", "nearby", "close to me", "around me",
    "local", "locally", "local area", "my area", "your area", "the area",
    "area", "areas", "surrounding", "surrounding area", "surrounding areas",
    "nationwide", "national", "nation", "usa", "u.s.", "u.s.a.", "us",
    "united states", "america", "statewide", "state", "everywhere", "anywhere",
    "n/a", "na", "none", "no", "tbd", "tba", "various", "multiple", "all",
    "other", "others", "unknown", "blank", "total", "totals", "average",
    "overall", "keyword", "keywords", "market", "markets", "city", "cities",
    "<cityname>", "cityname", "city name", "xxx", "test",
}


def is_non_place_geo(m):
    """True if this entered geo is not a place at all.

    Tested on the bare name with any state suffix removed, because the importer
    qualifies everything: the pill reads "near me, TN".
    """
    s = re.sub(r"[‘’“”]", "", str(m or "")).strip().lower()
    s = re.sub(r"\s+", " ", s.strip(" .-—–\t"))
    if not s:
        return True
    if "," in s:
        head, tail = [p.strip() for p in s.rsplit(",", 1)]
        if tail in _abbrev_to_state() or tail in STATE_ABBREV:
            s = head
    return s in _NON_PLACE_GEOS


def usable_markets(markets):
    """The entered markets minus anything that isn't a place.

    Used wherever a market has to survive contact with a search API. Kept
    separate from the pill list on purpose — dropping the operator's own entry
    behind their back is worse than showing it and refusing to measure in it.
    """
    return [str(m).strip() for m in (markets or [])
            if str(m).strip() and not is_non_place_geo(m)]


# WHAT PEOPLE CALL A CITY IS NOT ALWAYS WHAT THE PROVIDER CALLS IT.
#
# DataForSEO's location database holds "New York,New York,United States". A
# market typed the way everyone types it — "New York City, NY" — builds
# "New York City,New York,United States", which is not a place, so the lookup
# comes back 40501 and the row records as a failed check.
#
# It failed silently for as long as it has existed, because NYC has only ever
# been a secondary market here: the grid drops it and nothing else asked. The
# cross-market probe asked, got nothing back three runs in a row, and the
# panel reported it as "could not be read" — a lookup problem, which it was,
# just not a transient one. Any client whose PRIMARY market is New York would
# have had every rank check in the quote fail this way. (2026-08-21)
_CITY_ALIAS = {
    "new york city": "New York",
    "nyc": "New York",
    "new york, ny": "New York",
    "philly": "Philadelphia",
    "st louis": "St. Louis",
    "st. paul": "Saint Paul",
    "ft lauderdale": "Fort Lauderdale",
    "ft. lauderdale": "Fort Lauderdale",
    "ft worth": "Fort Worth",
    "ft. worth": "Fort Worth",
    "washington dc": "Washington",
    "washington, d.c.": "Washington",
    "d.c.": "Washington",
}

# "Washington" alone is a state before it is a city, so the DC aliases cannot
# go through CITY_STATE without rewriting Seattle. They carry their own state.
_CITY_ALIAS_STATE = {
    "washington dc": "District of Columbia",
    "washington, d.c.": "District of Columbia",
    "d.c.": "District of Columbia",
}


def provider_city_state(city):
    """The state the provider files an aliased city under, if it needs one."""
    return _CITY_ALIAS_STATE.get(
        re.sub(r"\s+", " ", str(city or "").strip().lower()), "")


def provider_city(city):
    """The provider's name for a city, or the city unchanged."""
    return _CITY_ALIAS.get(re.sub(r"\s+", " ", str(city or "").strip().lower()),
                           city)


def loc_string(markets, state):
    for m in usable_markets(markets) or []:
        city, st = parse_market(m, state)
        st = provider_city_state(city) or st
        city = provider_city(city)
        # A two-letter fallback state reaches here unexpanded when the market
        # carries no ",ST" of its own, and "New York,NY,United States" is not a
        # place in the provider's database any more than "New York City" is —
        # same silent failure, different half of the string. (2026-08-21)
        st = _abbrev_to_state().get(str(st or "").strip().lower(), st).title() \
            if str(st or "").strip().lower() in _abbrev_to_state() else st
        if city and st:
            return f"{city},{st},United States"
        if city:                      # city without state — still localizes
            return f"{city},United States"
    if state:
        return f"{state},United States"
    return "United States"

# City -> state auto-derivation. Covers major US metros + the cities in the
# sample proposals. Unknown cities fall back to "City,United States", which
# DataForSEO usually resolves to the largest match.
CITY_STATE = {
    "san diego":"California","chula vista":"California","el cajon":"California",
    "oceanside":"California","escondido":"California","bonita":"California","alpine":"California",
    "los angeles":"California","san francisco":"California","sacramento":"California",
    "san jose":"California","fresno":"California","long beach":"California","irvine":"California",
    "knoxville":"Tennessee","nashville":"Tennessee","memphis":"Tennessee",
    "farragut":"Tennessee","alcoa":"Tennessee","maryville":"Tennessee","louisville":"Tennessee",
    "hampton roads":"Virginia","norfolk":"Virginia","virginia beach":"Virginia",
    "chesapeake":"Virginia","newport news":"Virginia","hampton":"Virginia","richmond":"Virginia",
    "wichita":"Kansas","kansas city":"Missouri","topeka":"Kansas",
    "altoona":"Pennsylvania","state college":"Pennsylvania","hanover":"Pennsylvania",
    "harrisburg":"Pennsylvania","lancaster":"Pennsylvania","york":"Pennsylvania",
    "philadelphia":"Pennsylvania","pittsburgh":"Pennsylvania","bedford":"Pennsylvania",
    "lava hot springs":"Idaho","pocatello":"Idaho","boise":"Idaho","idaho falls":"Idaho",
    "anchorage":"Alaska","fairbanks":"Alaska","juneau":"Alaska",
    "new york":"New York","brooklyn":"New York","buffalo":"New York","albany":"New York",
    "chicago":"Illinois","houston":"Texas","dallas":"Texas","austin":"Texas",
    "san antonio":"Texas","phoenix":"Arizona","tucson":"Arizona","denver":"Colorado",
    "seattle":"Washington","portland":"Oregon","miami":"Florida","orlando":"Florida",
    "tampa":"Florida","atlanta":"Georgia","boston":"Massachusetts","detroit":"Michigan",
    "minneapolis":"Minnesota","charlotte":"North Carolina","raleigh":"North Carolina",
    "las vegas":"Nevada","salt lake city":"Utah","columbus":"Ohio","cleveland":"Ohio",
    "cincinnati":"Ohio","indianapolis":"Indiana","milwaukee":"Wisconsin","st louis":"Missouri",
}
_ABBREV_TO_STATE = None   # built lazily — STATE_ABBREV is defined later in the module

def _abbrev_to_state():
    global _ABBREV_TO_STATE
    if _ABBREV_TO_STATE is None:
        _ABBREV_TO_STATE = {v: k for k, v in STATE_ABBREV.items()}   # 'nj' -> 'new jersey'
    return _ABBREV_TO_STATE

def parse_market(m, default_state=""):
    """Split an entered market into (city, state). Accepts 'Cherry Hill, NJ',
    'Cherry Hill, New Jersey', or plain 'Cherry Hill' (state then comes from
    the metro map or the global State field). Multi-state regions — a tri-state
    MSP, say — need per-city suffixes: 'it support cherry hill nj' but
    'it support wilmington de'; one global state would mislabel two-thirds
    of the grid."""
    m = (m or "").strip()
    city, st = m, ""
    if "," in m:
        head, tail = [p.strip() for p in m.rsplit(",", 1)]
        t = tail.lower()
        if t in _abbrev_to_state():              # 'NJ'
            city, st = head, _abbrev_to_state()[t].title()
        elif t in STATE_ABBREV:                  # 'New Jersey'
            city, st = head, tail.title()
    if not st:
        cl = city.strip().lower()
        st = CITY_STATE.get(cl, "")
        if not st and cl.endswith(" county"):
            # "san diego county" -> derive the state from "san diego". Counties
            # are REAL DataForSEO locations ("San Diego County,California,
            # United States") and real search phrasing ("bucks county roofing")
            # — they just need the state attached to resolve.
            st = CITY_STATE.get(cl[:-len(" county")].strip(), "")
        # An EXPLICIT fallback state beats the metro map. The map is a guess
        # from the city name alone, and city names collide: "Cleveland" maps to
        # Ohio, so a Tennessee client with a Cleveland TN branch had its state
        # silently rewritten (2026-07-28). Someone who selected a fallback
        # state has told us where this client operates; a lookup table has not.
        ds = (default_state or "").strip()
        st = ds or st
    return city.strip(), st

def market_city(m, default_state=""):
    return parse_market(m, default_state)[0]

def market_state(m, default_state=""):
    return parse_market(m, default_state)[1]

def derive_state(markets, provided_state=""):
    """Return the client's state as a FULL NAME, or "" if genuinely unknown.

    Three bugs in one function, all the same shape: it only ever recognised a
    full state name, and the tool asks for abbreviations everywhere.

      1. CITY_STATE is keyed on a BARE city name, so a market typed in the
         documented "Knoxville, TN" form matched nothing and this returned "".
      2. provided_state was passed through verbatim, so a state field holding
         "TN" came back as "TN" - and every caller then does
         STATE_ABBREV.get("tn"), which is a miss, because that map goes
         "tennessee" -> "tn".
      3. Both failures are silent. The keyword TEXT was unaffected, because it
         takes the state from parse_market(), which does parse the tag. So the
         grid was built on "junk removal sevierville tn" while pick_grid_cities
         SCORED markets on "junk removal sevierville" - two different strings,
         two different volumes. Sevierville read 0/mo on the probe and 20/mo
         twice in the grid, and the axis recommendation rests on that score.

    Returning a canonical full name fixes every caller at once rather than
    needing a guard at each site. Order: the operator's own value, then any
    market's "City, ST" tag, then the bare-city lookup. (2026-08-11)
    """
    # Built here, not at import: STATE_ABBREV is defined further down the module.
    inv = {a: n.title() for n, a in STATE_ABBREV.items()}

    def canon(v):
        t = (v or "").strip()
        if not t:
            return ""
        low = t.lower()
        if low in STATE_ABBREV:                    # already a full name
            return t.title()
        return inv.get(low, "")                    # "tn" -> "Tennessee"

    got = canon(provided_state)
    if got:
        return got
    for mkt in markets:
        # "Knoxville, TN" / "Knox County, TN" - the format the placeholder asks
        # for. parse_market already handles this; do the same here.
        if "," in (mkt or ""):
            got = canon((mkt or "").rsplit(",", 1)[-1])
            if got:
                return got
    for mkt in markets:
        ml = (mkt or "").strip().lower()
        s = CITY_STATE.get(ml)
        if not s and ml.endswith(" county"):
            s = CITY_STATE.get(ml[:-len(" county")].strip())
        if s:
            return s
    # An unrecognised free-text value is still better than nothing for
    # loc_string, which passes it to the location API as-is.
    return (provided_state or "").strip()

def is_longtail(kw):
    """A keyword qualifies as long-tail if it's long or question/intent-shaped."""
    words = kw.split()
    if len(words) >= CFG["longtail_min_words"]:
        return True
    if words and words[0].lower() in CFG["longtail_prefixes"]:
        return True
    return False

def fetch_suggestions(seeds, markets, state):
    """keyword_suggestions returns queries CONTAINING the seed — structurally
    longer than keyword_ideas. Calls run in parallel; failures are non-fatal.

    Only the broadest few seeds are asked — see suggest_seed_cap. This was
    already capped at six by a hardcoded slice; the cap is now named, config-
    editable, and picks by BREADTH rather than by typing order. Fewest words
    first, ties on the order they were typed, so two builds of the same list ask
    the same questions. (2026-08-14)
    """
    out = []
    if not CFG["use_suggestions"]:
        return out
    cap = int(CFG.get("suggest_seed_cap") or 0)
    if cap > 0 and len(seeds or []) > cap:
        _order = {s: i for i, s in enumerate(seeds)}
        seeds = sorted(seeds, key=lambda s: (len(str(s).split()), _order[s]))[:cap]
    loc = loc_string(markets, state)

    def one(s):
        try:
            payload = [{"keyword": s, "location_name": loc,
                        "language_code": "en", "limit": 150}]
            data = dfs_post("/keywords_data/google_ads/keyword_suggestions/live", payload)
            res = (data["tasks"][0]["result"] or [])
            rows = []
            for block in res:
                for it in (block.get("items") or []):
                    kw = it.get("keyword")
                    if kw:
                        ki = it.get("keyword_info") or {}
                        # Measured AT THE CLIENT'S MARKET — see loc_string above. Tagged,
                        # because the site-keyword pull below is measured across the
                        # whole United States and the two are not comparable numbers.
                        rows.append({"keyword": kw, "scope": "local",
                                     "volume": ki.get("search_volume") or 0})
            return rows
        except Exception:
            return []

    with ThreadPoolExecutor(max_workers=min(len(seeds), CFG["rank_check_workers"]) or 1) as ex:
        for rows in ex.map(one, seeds):
            out.extend(rows)
    return out

def fetch_keywords_for_site(domain, markets, state):
    """Labs 'Keywords For Site' — keywords relevant to the client's domain,
    derived from the site's content/category. Supplements partner seeds for
    established sites; returns little (harmlessly) for brand-new/zero-ranking
    sites, which is why it's additive, not a replacement. One call. Non-fatal."""
    if not CFG["use_site_keywords"] or not domain:
        return []
    dom = domain.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    if not dom:
        return []
    try:
        # Labs endpoint: use numeric location_code (2840 = US), not location_name.
        payload = [{"target": dom, "location_code": 2840,
                    "language_code": "en", "limit": CFG["site_keywords_limit"]}]
        data = dfs_post("/dataforseo_labs/google/keywords_for_site/live", payload)
        res = (data["tasks"][0]["result"] or [])
        rows = []
        for block in res:
            for it in (block.get("items") or []):
                kw = it.get("keyword")
                if kw:
                    ki = it.get("keyword_info") or {}
                    # location_code 2840 is the WHOLE US. These volumes are national
                    # and sit beside Santa Fe figures in the same list.
                    rows.append({"keyword": kw, "scope": "national",
                                 "volume": ki.get("search_volume") or 0})
        return rows
    except Exception:
        return []

def get_client_site(url, timeout=10, headers=None, allow_redirects=True):
    """GET a CLIENT'S OWN marketing site, tolerating a broken TLS chain.

    Plenty of small-business sites serve a valid leaf certificate but omit the
    intermediate, so the chain can't be built. Browsers paper over it — they
    cache intermediates and follow the AIA pointer — and requests does not, so
    the site menu and description auto-fill just fail with a wall of SSL text
    (Woodstock Furniture, 2026-07-27).

    Verification is attempted first and only relaxed on a certificate error,
    for the client's own public pages. That content is read-only, used to
    suggest keywords a human then reviews, and never a credential path. This
    helper is deliberately NOT used for DataForSEO or Anthropic — those carry
    API keys and must always verify.

    Returns (response, insecure_flag). Raises the original error if the retry
    also fails.
    """
    kw = {"timeout": timeout, "headers": headers or {},
          "allow_redirects": allow_redirects}
    try:
        return requests.get(url, **kw), False
    except requests.exceptions.SSLError:
        try:
            import urllib3
            urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
        except Exception:
            pass
        return requests.get(url, verify=False, **kw), True


# A page under one of these paths is a LOCATION page — the thing the industry
# actually counts. Every multi-location pricing model turns on "each location
# needs its own profile, reviews, local content and ideally its own location
# page", and Brendan scoped Skills of Central PA with exactly this test:
# "This campaign would cover all locations mentioned on the website."
_LOCATION_PATH = re.compile(
    r"/(?:our[-_])?(?:locations?|stores?|offices?|branch(?:es)?|showrooms?|"
    r"clinics?|dealers?|service[-_]areas?|areas?[-_]we[-_]serve)(?:/|$)", re.I)
# The index page itself isn't a location — "/locations/" lists them, and
# "/locations/marietta/" is one.
_LOCATION_INDEX = re.compile(
    r"/(?:our[-_])?(?:locations?|stores?|offices?|branch(?:es)?|showrooms?|"
    r"clinics?|dealers?|service[-_]areas?|areas?[-_]we[-_]serve)/?$", re.I)


def google_business_cities(brand, domain):
    """Cities where the client has a Google Business listing.

    The rep tab already pulls these — one Business Listings search, about two
    cents, returning every listing at once regardless of how many markets are
    involved. The SEO tab was ignoring it and inferring presence from location
    pages instead, which fails for anyone using a store-locator widget: it read
    Woodstock Furniture as having zero locations while the rep tab found six
    (2026-07-28).

    A Google Business Profile IS the local presence an SEO campaign optimises,
    so it answers "does the client operate in this market" more directly than
    anything else available. Non-fatal — a failure just leaves the other
    signals to decide.
    """
    try:
        import rep_scan
        res = rep_scan.scan_locations(brand, domain=domain) or {}
    except Exception:
        app.logger.exception("google_business_cities failed")
        return [], None
    cities, seen = [], set()
    for loc in (res.get("locations") or []):
        addr = (loc.get("address") or "")
        # "1234 Main St, Marietta, GA 30060" -> the part before the state
        parts = [p.strip() for p in addr.split(",") if p.strip()]
        city = parts[-2] if len(parts) >= 2 else ""
        city = re.sub(r"\s+\d{5}(-\d{4})?$", "", city).strip()
        if city and city.lower() not in seen:
            seen.add(city.lower())
            cities.append(city)
    return cities, len(res.get("locations") or [])


# Headings that introduce a client's own list of markets.
_SERVICE_AREA_HEADING = re.compile(
    r"(?:where\s+we\s+(?:serve|work)|service\s+areas?|areas?\s+we\s+serve|"
    r"proudly\s+serve|communities\s+we\s+serve|our\s+locations|"
    r"cities\s+we\s+serve)", re.I)

# Words that show up in these lists but aren't markets.
_NOT_A_MARKET = re.compile(
    r"^(?:and\s+)?(?:surrounding|nearby|other|all|more)\b|"
    r"\b(?:areas?|communities|region|counties|county|and\s+beyond)$", re.I)

# Calls to action and nav labels sit immediately after these lists and are
# capitalised exactly like place names, so shape alone can't tell them apart —
# "Book Now" survived every structural test (2026-07-28).
_CTA_WORDS = re.compile(
    r"\b(?:book|call|contact|get|start|started|learn|read|more|free|quote|"
    r"schedule|request|now|today|us|home|about|services?|products?|offers?|"
    r"menu|search|login|sign|apply|shop|buy|order|view|see|click|here|next|"
    r"back|close|submit|send|email|phone|hours|reviews?|blog|news|faq|careers?|"
    r"privacy|terms|sitemap|contact)\b", re.I)


def service_areas_from_html(html):
    """Pull the market list off a client's 'Where We Serve' page.

    TN Water & Air names eighteen service areas on one page under "We proudly
    serve the following areas". Its proposal then prices twelve distinct
    markets — the same list with the metro clusters collapsed. That list is
    the source SSG actually scoped from, and we were reading neither it nor
    anything close: the URL matcher sees no /locations/ paths because the
    areas are body text, and the Google Business pull returns three, because
    three offices serve twelve markets (2026-07-28).

    So read the page. Anchor on a heading, take the short text items that
    follow, and stop at the first stretch of prose — the lists are always
    short link or list-item labels, never sentences.
    """
    if not html:
        return []
    m = _SERVICE_AREA_HEADING.search(html)
    if not m:
        return []
    window = html[m.end():m.end() + 12000]
    items = re.findall(r">([^<>{}]{2,40})<", window)
    out, seen = [], set()
    for raw in items:
        t = re.sub(r"\s+", " ", raw).strip(" \u00b7|,-–—\t")
        if not t or len(t) < 3 or len(t.split()) > 4:
            continue
        if not re.match(r"^[A-Z][A-Za-z.'\-]*(?:\s+[A-Za-z.'\-]+)*$", t):
            continue          # markets are capitalised; nav junk usually isn't
        if _NOT_A_MARKET.search(t) or _CTA_WORDS.search(t):
            continue
        low = t.lower()
        if low in seen:
            continue
        seen.add(low)
        out.append(t)
        if len(out) >= 40:
            break
    return out


def fetch_service_areas(domain):
    """Find and read the client's service-area page. Non-fatal."""
    if not domain:
        return []
    base = domain if domain.startswith("http") else f"https://{domain}"
    base = base.rstrip("/")
    paths = ["/service-areas", "/service-area", "/areas-we-serve", "/where-we-serve",
             "/where-we-work", "/locations", "/service-areas/", ""]
    for p in paths:
        try:
            r, _ = get_client_site(base + p, timeout=8,
                                   headers={"User-Agent": "Mozilla/5.0"})
            if r.status_code != 200 or "<" not in (r.text or ""):
                continue
            found = service_areas_from_html(r.text[:400_000])
            if len(found) >= 3:
                return found
        except Exception:
            continue
    return []


def location_pages_from_urls(urls):
    """Location page slugs found in a set of site URLs, de-duplicated."""
    out, seen = [], set()
    for u in urls or []:
        u = str(u or "")
        # Sitemap files are now collected alongside page URLs (they identify a
        # storefront by name), and "/locations/sitemap.xml" would otherwise be
        # read as a location called "sitemap.xml".
        if u.split("?")[0].lower().endswith(".xml"):
            continue
        if not _LOCATION_PATH.search(u) or _LOCATION_INDEX.search(u):
            continue
        slug = [p for p in u.split("?")[0].rstrip("/").split("/") if p]
        name = (slug[-1] if slug else "").replace("-", " ").replace("_", " ").strip()
        if name and name.lower() not in seen and len(name) < 40:
            seen.add(name.lower())
            out.append(name)
    return out


def fetch_site_pages(domain, limit=30, collect_urls=None):
    """Pull the client's page structure as readable topics — the names of the
    pages they've built, which map directly to their service taxonomy and are
    strong SEO keyword fuel. Tries sitemap.xml first (fast, standard); falls back
    to the DataForSEO On-Page API if there's no usable sitemap. Returns a list of
    short topic strings. Non-fatal: [] on any failure."""
    if not domain:
        return []
    dom = domain.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    if not dom:
        return []

    def slug_to_topic(url):
        path = url.split("//", 1)[-1]
        path = path.split("/", 1)[1] if "/" in path else ""
        path = path.strip("/").split("?")[0].split("#")[0]
        if not path:
            return ""
        seg = [s for s in path.split("/") if s and not s.endswith((".xml", ".jpg", ".png", ".pdf", ".css", ".js"))]
        if not seg:
            return ""
        topic = seg[-1].replace("-", " ").replace("_", " ").replace(".html", "").strip()
        if len(topic) < 3 or topic.isdigit():
            return ""
        if topic.lower() in {"index", "home", "page", "blog", "category", "tag"}:
            return ""
        return topic

    pages, seen = [], set()
    import re
    deadline = time.time() + 8          # hard cap: sitemap work gets <= 8s total
    _UA_B = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36"}
    _UA_T = {"User-Agent": "Mozilla/5.0 (compatible; adtini-seo-quote/1.0)"}
    def _get(url, timeout):
        """Fetch trying both identities — WAFs differ on which they block."""
        last = None
        for hdrs in (_UA_B, _UA_T):
            try:
                r, _insecure = get_client_site(url, timeout=timeout, headers=hdrs)
                if r.status_code == 200 and "<" in (r.text or ""):
                    return r
                last = r
            except Exception:
                pass
        return last

    # Candidate sitemap locations: what robots.txt declares, plus the standard
    # and WordPress-native paths. WP >=5.5 ships /wp-sitemap.xml; Yoast uses
    # /sitemap_index.xml; many themes use /page-sitemap.xml directly.
    candidates = []
    try:
        rr = _get(f"https://{dom}/robots.txt", 4)
        if rr is not None and rr.status_code == 200:
            candidates += re.findall(r"(?im)^sitemap:\s*(\S+)", rr.text)
    except Exception:
        pass
    _dom = dom
    for base_dom in dict.fromkeys([_dom, re.sub(r"^www\.", "", _dom)]):
        candidates += [f"https://{base_dom}/sitemap.xml", f"https://{base_dom}/sitemap_index.xml",
                       f"https://{base_dom}/wp-sitemap.xml", f"https://{base_dom}/page-sitemap.xml"]
    seen_sm = set()

    def _blogish(url):
        u = url.lower()
        return bool(re.search(r"/(blog|news|category|tag|author|20\d\d)/", u))

    for sm in candidates:
        if sm in seen_sm or time.time() > deadline:
            continue
        seen_sm.add(sm)
        try:
            r = _get(sm, 5)
            if r is None or r.status_code != 200 or "<" not in r.text:
                continue
            locs = re.findall(r"<loc>\s*(.*?)\s*</loc>", r.text, re.I)
            index_locs = []
            if locs and all(l.lower().endswith(".xml") for l in locs[:3]):
                # sitemap INDEX — service pages live in "page" sitemaps, so read
                # those first; blog-post sitemaps are last resort
                kids = sorted(locs, key=lambda l: (("page" not in l.lower()),
                                                   ("post" in l.lower())))
                # Keep the index's OWN entries. They are thrown away below when
                # the children are read, and their NAMES are evidence in their
                # own right — "sitemap_products_1.xml" identifies a storefront
                # without costing a fetch. That mattered for Grav: this sort
                # deliberately reads the pages sitemap first, and the 8s budget
                # then ran out before the products sitemap, so every product URL
                # was missed and the store went undetected (2026-08-04).
                index_locs = list(locs)
                child_locs = []
                for child in kids[:4]:
                    if time.time() > deadline:
                        break
                    try:
                        cr = _get(child, 4)
                        if cr is None: continue
                        child_locs += re.findall(r"<loc>\s*(.*?)\s*</loc>", cr.text, re.I)
                    except Exception:
                        pass
                locs = child_locs or locs
            # shallow, non-blog URLs first — service pages are shallow; posts are
            # deep or dated. Service-path hints float to the top.
            def _rank(u):
                depth = u.rstrip("/").count("/") - 2
                hinted = bool(_SERVICE_PATH_HINT.search(u)) if "_SERVICE_PATH_HINT" in globals() else False
                return (_blogish(u), not hinted, depth)
            if isinstance(collect_urls, list):
                collect_urls.extend(locs)
                collect_urls.extend(index_locs)   # names are evidence; see above
            for url in sorted(locs, key=_rank):
                if _blogish(url) and len(pages) >= 5:
                    continue
                t = slug_to_topic(url)
                if t and t.lower() not in seen:
                    seen.add(t.lower()); pages.append(t)
                if len(pages) >= limit:
                    break
            if len(pages) >= 3:
                return pages[:limit]
        except Exception:
            continue
    if pages:
        return pages[:limit]

    # On-Page fallback only if we have time budget left
    if time.time() > deadline:
        return pages[:limit]
    try:
        payload = [{"url": f"https://{dom}", "max_crawl_pages": limit}]
        data = dfs_post("/on_page/instant_pages", payload)
        res = (data["tasks"][0]["result"] or [])
        for block in res:
            for it in (block.get("items") or []):
                u = it.get("url") or ""
                # This fallback never fed collect_urls, so on any site without a
                # readable sitemap the caller got page topics but ZERO urls —
                # silently disabling both storefront and location-page detection.
                if u and isinstance(collect_urls, list):
                    collect_urls.append(u)
                t = slug_to_topic(u)
                if t and t.lower() not in seen:
                    seen.add(t.lower()); pages.append(t)
        return pages[:limit]
    except Exception:
        return pages[:limit]


def _bare_city(m, state=""):
    """The city name alone — the form every per-city lookup uses."""
    try:
        return (parse_market(m, state)[0] or m).strip().lower()
    except Exception:
        return (m or "").strip().lower()


# Cache for canonical_city_name — the nearest-point scan walks 42k ZIP rows, and
# loc_string is called on every lookup.
_CANON_CITY = {}


def canonical_city_name(city, st=""):
    """A DIFFERENT name for the same place, for when Google rejects the first.

    Google Ads has no location called "New York City" — its canonical name is
    "New York". None called "Lawrenceville, NJ" either; it is "Lawrence
    Township". So the name is wrong, not the place, and the old ladder jumped
    straight from a rejected city to its whole STATE — which is how a five-town
    New Jersey retailer ended up priced on New York statewide demand
    (2026-08-07).

    Order: drop a trailing "city", then a known alias, then the nearest name in
    the ZIP data at that exact point (both real cases resolve at 0.0 miles, so
    this is the same place relabelled rather than a neighbour).

    Returns "" when there is no alternative worth trying.
    """
    c = (city or "").strip().lower()
    key = (c, (st or "").strip().lower())
    if key in _CANON_CITY:
        return _CANON_CITY[key]

    cands = []
    if c.endswith(" city") and len(c.split()) > 1:
        cands.append(c[: -len(" city")].strip())
    for a in _CITY_ALIASES.get(c, []):
        if " " in a or len(a) > 3:          # skip 2-3 letter shorthand
            cands.append(a)

    out = ""
    abbr = (STATE_ABBREV.get((st or "").strip().lower(), "") or "").upper()
    idx = _zip_index()
    for cand in cands:
        if cand and cand != c and (not abbr or (cand, abbr) in idx):
            out = cand
            break
    if not out:
        # Nearest ZIP-data name at the same coordinates.
        pt = idx.get((c, abbr)) if abbr else None
        if pt is None:
            hits = [v for (cc, _s), v in idx.items() if cc == c]
            pt = hits[0] if len(hits) == 1 else None
        if pt:
            best = None
            for (cc, ss), p in idx.items():
                if abbr and ss != abbr:
                    continue
                if cc == c:
                    continue
                d = miles_between(pt, p)
                if best is None or d < best[0]:
                    best = (d, cc)
            # Reject an abbreviation of the name we already have. The ZIP data
            # carries "Phila" alongside "Philadelphia"; offering it to Google as
            # an alternative name is worse than the name that already works.
            if best and best[0] <= 2.0:
                cand = best[1]
                # Strict prefix test, no length tolerance. "Phila" is an
                # abbreviation of Philadelphia and "Los Angeles AFB" is a base
                # inside Los Angeles — neither is the city under another name.
                # This only gates the nearest-point scan; the explicit candidate
                # list above (drop trailing "city", alias map) is what carries
                # New York City -> New York, so that case is unaffected.
                prefix_of_each_other = c.startswith(cand) or cand.startswith(c)
                if not prefix_of_each_other:
                    out = cand
    _CANON_CITY[key] = out
    return out


def fetch_local_volume(terms, markets, state, national=False):
    """Search volume for bare service terms across THE CITIES BEING TARGETED.

    A single lookup only covers markets[0], which undercounts a multi-city grid
    by roughly the city count (e.g. 'auto insurance' is ~480/mo in Alexandria but
    the campaign also covers nine other cities). So query each city and sum per
    service — that's the client's real addressable demand.
    Returns ({term_lower: summed_volume}, error_or_None)."""
    if not terms:
        return {}, {}, None
    cities = [c for c in (markets or []) if c and c.strip()]
    if state:
        cities = [c for c in cities if c.strip().lower() != state.strip().lower()]
    if national:
        # Product brands / national scope: the client's cities still build the
        # GRID (the proposal table stays per-city), but pricing demand is the
        # national figure. A geo-qualified pull structurally undercounts a DTC
        # brand — nobody searches "collagen gummies fairfax va".
        cities = [""]
    if not cities:
        cities = [""]                      # nationwide / no city: single lookup
    cities = cities[:CFG.get("grid_max_cities", 10)]
    kws = [t.lower() for t in terms]

    resolved_codes = {}
    renamed = {}

    def one(city):
        # loc_string parses "City, ST" itself; each city localizes to its own state
        loc = loc_string([city], state) if city else loc_string([], state)
        def call(location):
            payload = [{"keywords": dfs_kw_list(kws), "location_name": location,
                        "language_code": "en"}]
            data = dfs_post("/keywords_data/google_ads/search_volume/live", payload,
                            timeout=25)
            task0 = (data.get("tasks") or [{}])[0]
            if task0.get("status_code") not in (20000, None):
                raise RuntimeError(f"{task0.get('status_code')}: {task0.get('status_message')}")
            rows = task0.get("result") or []
            # WHICH LOCATION ANSWERED. DataForSEO resolves a city it doesn't
            # carry up to a metro or state and returns THAT area's volume with
            # no error and no flag — Lawrenceville NJ reported 8,100/mo for
            # "outdoor furniture" against Paramus's 70 (2026-08-07). The only
            # trustworthy signal is the location_code echoed back; captured
            # defensively because it is not guaranteed to be present, and shown
            # rather than acted on until it proves reliable.
            code = None
            try:
                code = (task0.get("data") or {}).get("location_code")
                if code is None:
                    for it in rows:
                        if it.get("location_code") is not None:
                            code = it.get("location_code")
                            break
            except Exception:
                code = None
            return rows, code
        try:
            _rows, _code = call(loc)
            resolved_codes[city] = _code
            return _rows, loc
        except Exception as e:
            # An unrecognized city (misspelling, a regional phrase like "south
            # jersey", or a name DataForSEO doesn't carry) returns 40501. Retry
            # at a broader location so the quote still gets *some* demand signal
            # — but report WHICH location answered, because broad-location
            # volume must never be attributed per-city and summed: three cities
            # falling back to the same national number would count the same
            # searches three times and wildly inflate the volume add.
            if "40501" in str(e) or "not found" in str(e).lower():
                city_st = market_state(city, state)
                # FIRST try the same place under the name Google actually uses.
                # Jumping straight to the state throws away the city entirely,
                # and for New York City that meant pricing on New York STATE.
                alt = canonical_city_name(market_city(city, state), city_st or state)
                if alt:
                    alt_loc = (f"{alt},{city_st},United States" if city_st
                               else (f"{alt},{state},United States" if state
                                     else f"{alt},United States"))
                    try:
                        _rows, _code = call(alt_loc)
                        resolved_codes[city] = _code
                        renamed[city] = alt
                        # Its own place, just spelled Google's way — so this is
                        # NOT a broader-area fallback and it counts in the total.
                        return _rows, loc
                    except Exception:
                        pass
                broader = (f"{city_st},United States" if city_st
                           else (f"{state},United States" if state else "United States"))
                _rows, _code = call(broader)
                resolved_codes[city] = _code
                return _rows, broader
            raise

    totals, per_city, errs, ok = {}, {}, [], 0
    city_locs = {}
    counted_locs, fallback_cities, results = set(), [], []
    try:
        with ThreadPoolExecutor(max_workers=min(len(cities), 8)) as ex:
            futs = {ex.submit(one, c): c for c in cities}
            for fut in futs:
                city = futs[fut]
                try:
                    rows, used_loc = fut.result()
                    was_fallback = (used_loc != (loc_string([city], state) if city
                                                 else loc_string([], state)))
                    if was_fallback:
                        fallback_cities.append(city)
                    results.append((city, rows, used_loc, was_fallback))
                    ok += 1
                except Exception as e:
                    errs.append(str(e))
    except Exception as e:
        return {}, {}, str(e)
    if not ok:
        return {}, {}, (errs[0] if errs else "no volume rows returned")
    # Aggregate deterministically:
    #   1. each effective location counts into the TOTAL exactly once;
    #   2. a location that is NOT the city's own never counts while any city did
    #      return its own figure. Previously only a "United States" fallback was
    #      excluded, so a STATEWIDE one was counted in full: Ski Barn's
    #      "outdoor furniture" total of 20,350/mo was New York statewide (12,100)
    #      plus New Jersey statewide (8,100) plus three real towns (150), for a
    #      five-town New Jersey retailer — and that total drives the volume
    #      component of price and promoted the term to Ultra Competitive
    #      (2026-08-07). A town Google cannot locate must not contribute its
    #      whole state's demand.
    #   3. if NO city returned its own figure, broader locations count once each
    #      rather than pricing the client at zero demand — reported either way.
    own_data = [r for r in results if not r[3]]
    us_skipped = False
    broader_skipped = []
    if own_data:
        # Own-location results first so they claim their location before any
        # fallback that resolved to the same place.
        ordered = sorted(results, key=lambda r: (r[3], r[2] == "United States"))
    else:
        ordered = sorted(results, key=lambda r: r[2] == "United States")
    for city, rows, used_loc, was_fb in ordered:
        # DataForSEO tells us which location it ACTUALLY used for each city —
        # that is exact market identity, not an inference. Two cities resolving
        # to the same effective location are the same market, which is the
        # question the add-on count turns on. Inferring it from volume vectors
        # instead only worked where the geo-modified probe had volume, so it
        # failed silently in exactly the small rural markets where collapsing
        # matters most (2026-08-03).
        city_locs[city] = used_loc
        count_it = used_loc not in counted_locs
        if was_fb and own_data:
            # Not this city's location, and someone else's figures are real.
            count_it = False
            broader_skipped.append(city)
        if used_loc == "United States" and [r for r in results if r[2] != "United States"]:
            count_it = False
            us_skipped = True
        counted_locs.add(used_loc)
        for it in rows:
            k = (it.get("keyword") or "").lower()
            if k:
                v = it.get("search_volume") or 0
                if count_it:
                    totals[k] = totals.get(k, 0) + v
                # Key on the BARE city, which is how every consumer reads it.
                # Writing "altoona, pa" while the grid rows and the metro
                # grouping both look up "altoona" meant every per-city lookup
                # missed, rows silently fell back to the summed service total —
                # printing the same number for every city — and the grouping
                # had nothing to compare (2026-08-03).
                per_city[(_bare_city(city, state), k)] = v
    notes = []
    if broader_skipped:
        notes.append(
            "no city-level volume for " + ", ".join(sorted(set(
                c.strip() for c in broader_skipped)))
            + " — Google doesn't hold those as targetable locations, so a wider "
              "area answered. Shown per keyword but EXCLUDED from the pricing "
              "total, because a town's whole state is not that town's demand")
    if not own_data and fallback_cities:
        notes.append(
            "NO market returned volume of its own — every figure below is a "
            "wider area's, counted once each so the quote isn't priced at zero "
            "demand. Treat the volume component as an estimate")
    if us_skipped:
        notes.append("some geos had no local volume data and fell back to "
                     "national numbers — shown per keyword but EXCLUDED from "
                     "the pricing total to avoid inflating regional demand")
    if ok < len(cities):
        notes.append(f"volume summed over {ok}/{len(cities)} cities (some lookups failed)")
    if fallback_cities:
        notes.append("no city-level volume for "
                     + ", ".join(sorted(set(c.strip() for c in fallback_cities)))
                     + " — used broader-location volume, counted once (not per city)")
    # Hand the resolved locations back so the caller can collapse cities that
    # Google treats as one market. Stashed on the dict rather than widening the
    # signature, since three call sites unpack this tuple.
    per_city["__city_locs__"] = city_locs
    # Two cities sharing a location_code resolved to the SAME place, which means
    # at least one of them is not reporting its own demand.
    per_city["__location_codes__"] = {_bare_city(c, state): v
                                      for c, v in resolved_codes.items() if v is not None}
    # Which cities had NO volume of their own and were answered by a broader
    # location. Previously only mentioned inside the prose note, so the per-row
    # numbers showed a county/state/national figure as though it were the city's
    # — Lawrenceville NJ read 8,100/mo for "outdoor furniture" next to Paramus's
    # 70 (2026-08-07). Callers need this as data to mark those rows.
    # Normalised the same way per_city keys are — bare city, lowercase. Storing
    # the raw pill ("Lawrenceville, NJ") meant the grid rows, which carry the
    # bare city ("lawrenceville"), never matched and the flag never fired
    # (2026-08-07).
    per_city["__fallback_cities__"] = sorted({_bare_city(c, state) for c in fallback_cities})
    # The original pill text too, so a suggestion can be built from it.
    per_city["__fallback_markets__"] = sorted({str(c).strip() for c in fallback_cities})
    # Markets Google accepted only under a different name. Worth showing: the
    # figure is genuinely the city's, and the operator may want the pill to match.
    per_city["__renamed__"] = dict(renamed)
    return totals, per_city, ("; ".join(notes) or None)


def _labs_loc_field(markets, state, national=False):
    """location field for a LABS payload. Always country-level US (2840).

    Two rounds of trying to target Labs at the client's actual market both
    failed against the API (2026-08-04): location_name is rejected outright
    ("40501 Invalid Field: 'location_name'"), and a Google Ads city code is
    rejected too ("Invalid Field: 'location_code'") because Labs keys off its
    own, much smaller location set — Google Ads criteria IDs for suburbs are
    not in it.

    Country level is also the CORRECT granularity here, which is why the
    keyword-difficulty call has always hardcoded 2840 with the note that
    difficulty "is a national-level organic metric". Labs CPC is modelled the
    same way. Nothing is lost: per-market volume comes from fetch_local_volume,
    which calls a Google Ads endpoint that does accept a location_name.

    Returns (payload_fragment, label) — the label is what the UI reports.
    """
    return {"location_code": 2840}, "United States"


def fetch_exact_volume(keywords, markets, state, national=False):
    """Exact-match search volume. The Google Ads keywords_for_keywords endpoint
    we use to GENERATE terms returns GROUPED (broad) volumes that merge similar
    terms — which is why the numbers looked inflated/off. For the FINAL list we
    re-pull volume from the Labs keyword database, which returns per-term exact
    volume. Returns {keyword_lower: volume}. Non-fatal: {} on any failure."""
    if not keywords:
        return {}
    out = {}
    # The comment here used to say "use the city if known, else US" while the
    # code hardcoded 2840 (US) unconditionally, so EVERY client's exact volume
    # was pulled nationally — overstating demand, and therefore price, on every
    # local quote (2026-08-04). Labs takes a numeric location_code ONLY — the
    # first cut of this sent location_name and every local lookup came back
    # "40501 Invalid Field", which this function swallows into {} and prices as
    # zero volume. Resolve to a code instead.
    loc_field, _loc_used = _labs_loc_field(markets, state, national)
    try:
        # batch up to 1000 per call
        for i in range(0, len(keywords), 1000):
            chunk = keywords[i:i+1000]
            payload = [{"keywords": [k.lower() for k in chunk],
                        **loc_field, "language_code": "en"}]
            data = dfs_post("/dataforseo_labs/google/keyword_overview/live", payload)
            res = (data["tasks"][0]["result"] or [])
            for block in res:
                for it in (block.get("items") or []):
                    kw = (it.get("keyword") or "").lower()
                    ki = it.get("keyword_info") or {}
                    if kw:
                        out[kw] = ki.get("search_volume") or 0
        return out
    except Exception:
        return {}

def infer_business(domain, seeds, site_terms, industry="", pages=None):
    """What this client SELLS, in one or two sentences, read off their own site.

    IT NO LONGER GUESSES EXCLUSIONS (2026-08-13). The old prompt asked for what
    the business "does NOT offer", and a website almost never says what it does
    not do — so that clause was invented, and then enforced as a filter, quietly
    removing good keywords whenever the guess was wrong. Exclusions now come from
    the operator's Negative terms; this call supplies VOCABULARY.

    Vocabulary is the job. The text returned here joins the seeds, the site pages
    and the brand as the corpus the grounding filter checks every AI-proposed
    service against, so it wants the client's OWN words and enough of them — a
    thin description makes the filter trigger-happy, which is how NPAIHB's whole
    proposal set got stood down. Returns '' if unavailable; non-fatal.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or not (domain or site_terms):
        return ""
    site_list = [s["keyword"] for s in site_terms][:40]
    page_list = [str(p) for p in (pages or [])][:30]
    prompt = f"""Describe what this business SELLS, in one or two sentences, for SEO targeting.

WEBSITE: {domain or "(none)"}
INDUSTRY: {industry or "not given"}
SERVICES/VERTICAL THE PARTNER ENTERED: {", ".join(seeds)}
PAGES ON THEIR SITE: {json.dumps(page_list, ensure_ascii=False)}
KEYWORDS THEIR SITE SURFACES FOR: {json.dumps(site_list, ensure_ascii=False)}

Rules:
1. Say what they SELL. Do NOT say what they do not sell, do not offer, or do not
   provide — that is supplied separately by a human and guessing it removes real
   keywords.
2. Use THEIR words, the ones on their site, not a synonym you prefer. If their
   pages say "tribal", write "tribal", not "native american".
3. Name the actual service lines. "A home services company" is useless; "a
   build-for-rent community renting detached single-family homes" is the answer.
4. If the evidence is too thin to say anything specific, return an empty string
   rather than a generic sentence.

Return ONLY the description, no preamble."""
    try:
        resp = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({"model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 200, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=20)
        resp.raise_for_status()
        body = resp.json()
        return "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
    except Exception:
        return ""

STATE_ABBREV = {
    "alabama":"al","alaska":"ak","arizona":"az","arkansas":"ar","california":"ca",
    "colorado":"co","connecticut":"ct","delaware":"de","florida":"fl","georgia":"ga",
    "hawaii":"hi","idaho":"id","illinois":"il","indiana":"in","iowa":"ia","kansas":"ks",
    "kentucky":"ky","louisiana":"la","maine":"me","maryland":"md","massachusetts":"ma",
    "michigan":"mi","minnesota":"mn","mississippi":"ms","missouri":"mo","montana":"mt",
    "nebraska":"ne","nevada":"nv","new hampshire":"nh","new jersey":"nj","new mexico":"nm",
    "new york":"ny","north carolina":"nc","north dakota":"nd","ohio":"oh","oklahoma":"ok",
    "oregon":"or","pennsylvania":"pa","rhode island":"ri","south carolina":"sc",
    "south dakota":"sd","tennessee":"tn","texas":"tx","utah":"ut","vermont":"vt",
    "virginia":"va","washington":"wa","west virginia":"wv","wisconsin":"wi","wyoming":"wy",
}

# Words that carry no identifying weight, so they never need grounding.
# Words that carry no evidence about WHOSE service a term is. "can" and "get"
# were missing, so a pinned question was reported as blocked because the client
# "never uses the word can" — technically true, entirely beside the point
# (2026-08-04).
# CAMPAIGN GOAL — verbatim from the adtini order form. Like the RZ industry
# list, this ARRIVES with the order; the demo carries a selector so it can be
# exercised. Added 2026-08-05 after a Ski Barn quote priced a New Jersey ski
# retailer on national demand: the goal would have said in-store sales, and the
# contradiction was sitting there unread.
GOAL_OPTIONS = [
    "In-Store Sales", "B2B Sales", "Information Requests",
    "Branding & Awareness", "Mobile App Download", "Website Traffic",
    "Leads/Form Fillouts", "Online Sales", "Phone Calls",
    "Job Recruitment", "Event Attendance",
]

# What each goal implies about WHERE demand lives. Deliberately three-valued:
# "" means the goal genuinely doesn't say, and guessing would be worse than
# staying quiet. Used to CONTRADICT the operator's scope, never to override it —
# the goal is a statement of intent, not a pricing input, and there is no
# calibration data tying goals to price.
GOAL_SCOPE = {
    "In-Store Sales":       "local",      # footfall has an address
    "Phone Calls":          "local",      # service businesses take local calls
    "Event Attendance":     "local",      # an event happens somewhere
    "Job Recruitment":      "local",      # roles are located
    "Information Requests": "",
    "Leads/Form Fillouts":  "",           # local trade or national B2B
    "Branding & Awareness": "",
    "Website Traffic":      "",
    "B2B Sales":            "national",
    "Online Sales":         "national",   # ecommerce sells everywhere
    "Mobile App Download":  "national",
}

# Goals an SEO keyword campaign is a poor instrument for. Not blocked — flagged,
# because the price would look like a retail campaign for work that isn't one.
GOAL_OFF_PATTERN = {"Mobile App Download", "Job Recruitment"}

# Goals that DO drive the national-demand switch (2026-08-07, operator request).
# Distinct from GOAL_SCOPE, which only ever warns: this set actually flips the
# volume pull. The reasoning is the same one that makes the markets veto correct
# — inferences lose to statements. A detected shopping cart is an inference about
# what a site can do; a goal of "Online Sales" is the client telling the order
# form what they are buying, so it ranks with the manual switch and an explicit
# Nationwide scope rather than with the RZ taxonomy.
#
# Deliberately NOT every GOAL_SCOPE "national" entry. "B2B Sales" maps national
# but a regional IT firm selling to businesses in three counties is a local
# campaign, and "Mobile App Download" is already flagged as off-pattern for SEO.
# Config key so the list can be widened without a deploy.
GOAL_NATIONAL_DEMAND = ["Online Sales"]


def goal_forces_national(goal):
    """Do the selected goals put the quote on national demand?

    Only when NOTHING selected is local. "Online Sales" alone is a national
    campaign; "In-Store Sales + Online Sales" is a shop that also ships, and the
    markets-veto reasoning applies — a client with premises is priced locally
    unless they say so outright (2026-08-09).
    """
    goals = goal_list(goal)
    if not goals:
        return ""
    want = {str(o).strip().lower()
            for o in (CFG.get("goal_national_demand") or GOAL_NATIONAL_DEMAND)}
    hits = [g for g in goals if g.strip().lower() in want]
    if not hits:
        return ""
    if any(GOAL_SCOPE.get(g, "") == "local" for g in goals):
        return ""            # a local goal is present — not a national campaign
    return " + ".join(hits)


def goal_list(goal):
    """Goals arrive as one string, now possibly several joined with ' | '."""
    raw = str(goal or "")
    parts = [p.strip() for p in raw.split("|")] if "|" in raw else [raw.strip()]
    return [p for p in parts if p]


def goal_scope(goal):
    """What the selected goals say about WHERE demand lives.

    With several goals, LOCAL wins: a client who picked "In-Store Sales" as well
    as "Online Sales" has premises, and pricing them on national demand quotes a
    campaign they didn't ask for. Only an all-national selection reads national.
    """
    scopes = {GOAL_SCOPE.get(g, "") for g in goal_list(goal)}
    if "local" in scopes:
        return "local"
    if "national" in scopes:
        return "national"
    return ""


_GROUNDING_STOP = set("""a an and or the of for in on to with your our best top near me
services service company companies agency agencies firm firms group inc llc co
local affordable cheap professional expert experts quality quote quotes free
how what why when where which who whose can could should would will does did
is are was were do has have had am get gets getting rid without into from
about after before during over under you your they them their there here
that this these those not new more most less much many any all out off per
# HOW MUCH / WHAT KIND, not WHO. The filter exists to catch a competitor's company
# name — Turner, Clark — and a proper noun is what it should be testing. Ordinary
# commercial modifiers fail the "did the client ever say this word" test for
# reasons that have nothing to do with competitors: Brendan's own Amare list has
# "luxury rentals", "pet friendly homes for rent", "gated community homes for
# rent" and "homes for rent with garage", and a two-sentence business description
# contains none of those adjectives. Every one was refused. (2026-08-13)
luxury luxurious upscale premium deluxe modern updated renovated remodeled
furnished unfurnished available now open immediate ready vacant
small large big little mini compact spacious oversized
cheap budget discount value low high mid upper lower
gated private secure safe quiet
pet friendly pets dog dogs cat cats family friendly
senior seniors student students corporate
short long term temporary permanent monthly weekly daily nightly
single double triple multi one two three four five bedroom bedrooms bed beds
bath baths bathroom bathrooms studio unit units
emergency urgent same day next 24 hour hours weekend evening
licensed insured bonded certified accredited approved
custom bespoke standard basic full partial complete
indoor outdoor interior exterior residential commercial
""".split())

# Questions, as opposed to services. The service-list prompt has always said
# "never a question", but pinning skips the model entirely and reads straight
# from the volume-ranked candidate pool, where questions are plentiful and
# frequently outrank every real service in the set.
# Interrogatives are conclusive wherever they lead a phrase.
_QUESTION_LEAD = re.compile(r"^(?:how|what|why|when|where|which|who|whose)\b", re.I)
# Auxiliaries are NOT. "can am repair" is a Can-Am dealer, "am radio antenna" is
# a product, "will call service" is a fulfilment term — all led by an auxiliary,
# none of them questions. Real questions built on an auxiliary run longer ("does
# insurance cover chiropractic care"), so require length before treating one as
# a question. A false positive here silently deletes a real service, which is
# worse than a question slipping through to the grounding filter.
_AUX_LEAD = re.compile(r"^(?:is|are|was|were|do|does|did|can|could|should|"
                       r"would|will|has|have|had|am)\b", re.I)
_AUX_MIN_WORDS = 4


# Looking a term UP, not buying it. Used only to warn on the operator's own
# seeds, never to filter: "l/c/f meaning" and "what does lcf stand for" took two
# of PEO Brokers' twenty slots and is_question_kw sees neither, because neither
# leads with a question word. (2026-08-13)
_DEFINITIONAL = re.compile(
    r"\b(?:meaning|meanings|definition|defined|stand[s]? for|abbreviation|"
    r"acronym|explained|explain|means)\b", re.I)

# LOOKING SOMETHING UP, NOT BUYING IT — the support-desk twin of the definition
# case. PEO Brokers came back with "new york state insurance fund workers comp
# phone number" in Competitive: a real search with real volume, made by somebody
# who already deals with that fund and wants its switchboard. Same shape as an
# abbreviation lookup and just as unquotable, and no volume floor catches either.
# (2026-08-14)
# Every pattern here is at least two words or an unambiguous one. A bare "phone"
# or "portal" reads as support intent and is not: "phone systems for business"
# and "portal cranes for sale" are both real service lines, and both were caught
# by the first draft of this.
_SUPPORT_INTENT = re.compile(
    r"\b(?:phone number|telephone number|customer service|customer support|"
    r"contact (?:number|info|information|us)|mailing address|"
    r"hours of operation|opening hours|log ?in page|sign in page)\b", re.I)


def is_question_kw(text):
    """Is this a question phrase rather than a service a client could sell?"""
    t = (text or "").strip().lower()
    if not t:
        return False
    if "?" in t or "how to" in t:
        return True
    if _QUESTION_LEAD.search(t):
        return True
    return bool(_AUX_LEAD.search(t)) and len(t.split()) >= _AUX_MIN_WORDS


def is_lookup_kw(text):
    """Someone reading, not someone buying — a question or a definition.

    The two tests were already here and already agreed; nothing had joined them
    into the one idea they describe, so each caller re-wrote the pair. Never a
    filter: a seed the operator typed is quoted whatever this says. It decides
    what a term is allowed to EARN. (2026-08-13)
    """
    t = (text or "").strip()
    return bool(t) and (is_question_kw(t) or bool(_DEFINITIONAL.search(t))
                        or bool(_SUPPORT_INTENT.search(t)))


def drop_ungrounded_services(services, seeds, business_desc, site_pages, brand, domain):
    """Drop model-invented services containing a word the client never used.

    Three rounds of prompt wording failed to stop competitor names getting in:
    Keller Builds still came back with "turner construction company" and "clark
    construction company" — the two largest national contractors — because the
    keyword-idea pool is ranked by national volume and they sit at the top of
    it looking exactly like services (2026-07-27).

    Detecting "is this a competitor" in the abstract is hard. Detecting "did
    the client ever say this word" is easy and gets the same answer: Keller's
    seeds, description and site say commercial, industrial, agricultural,
    warehouse — they never say Turner. A service has to be GROUNDED in the
    client's own vocabulary. Seeds are exempt — they are the client's words by
    definition.

    Pins are NOT exempt (2026-08-04). They used to be, on the reasoning that a
    pin is backed by real search volume. Volume proves a term is SEARCHED, not
    that it is the client's: Grav, a smoke shop, had "glass water pitcher",
    "glass chess set" and "glass water carafe" pinned — kitchenware and board
    games that match on the word "glass" and sit high in the keyword pool. They
    then bypassed this filter, which would have caught all three, because
    pitcher/chess/carafe appear nowhere in the client's vocabulary.

    Returns (services, dropped, blocked_pins, stood_down). `dropped` is the
    model's own picks that failed — reported WHETHER OR NOT the filter stood
    down, because "here is what it would have cut" is the only thing that tells
    an operator whether standing down was right. `blocked_pins` is separate
    because a pin is forced in for PRICE STABILITY, so a refused one has to be
    seen rather than vanish. `stood_down` used to be signalled by dropped=None,
    which threw away the evidence and left the panel guessing at the cause —
    it told NPAIHB its business description was too short when the description
    was fine and the real problem was vocabulary. (2026-08-12)
    """
    corpus = " ".join([
        " ".join(seeds or []), business_desc or "",
        " ".join(site_pages or []), brand or "", (domain or "").replace(".", " "),
    ]).lower()
    # Matching is singular/plural-insensitive. Site navigation is written in
    # the plural ("Bongs For Sale", "Bubblers", "Nectar Collectors") while a
    # service is written in the singular ("glass bong"), so exact word matching
    # called the client's own catalogue foreign and removed real services as
    # unrecognised (Grav, 2026-08-04). Fold both sides to a bare stem instead.
    def _stem(w):
        w = w.strip("-/")
        if len(w) > 3 and w.endswith("es") and w[-3] in "sxzh":
            return w[:-2]
        if len(w) > 3 and w.endswith("s") and not w.endswith("ss"):
            return w[:-1]
        return w

    # THE CLIENT'S OWN NOUN, SPELLED THE OTHER WAY. Brendan's Amare list has
    # "houses for rent santa fe nm"; the description says "homes", so "houses"
    # read as foreign and the term was refused. These are head nouns for what a
    # business sells, where the two words are the same offer — not a blocklist of
    # bad terms but a short, closed list of commercial synonyms. It will need
    # extending; that is the honest cost of the approach. (2026-08-13)
    _SYN = [
        {"home", "house", "housing", "residence"},
        {"apartment", "apt", "flat", "condo", "condominium"},
        {"rent", "rental", "lease", "leasing", "letting"},
        {"auto", "car", "vehicle", "automotive"},
        {"doctor", "physician", "medical", "clinic"},
        {"dentist", "dental"},
        {"attorney", "lawyer", "legal"},
        {"realtor", "realty", "estate", "broker"},
        {"restaurant", "dining", "eatery", "cafe"},
        {"store", "shop", "retail", "boutique"},
        {"repair", "fix", "service", "servicing"},
        {"removal", "haul", "hauling", "disposal", "pickup"},
        {"cleaning", "cleaner", "janitorial", "maid"},
        {"builder", "contractor", "construction", "building"},
        {"salon", "spa", "barber"},
        {"vet", "veterinary", "veterinarian", "animal"},
    ]
    known = set()
    # HYPHENS SPLIT. "build-for-rent" is one token, so "build for rent homes" was
    # dropped on the word "build" — a term describing the client's own business
    # model, refused because their description hyphenated it. (2026-08-13)
    for w in corpus.replace(",", " ").replace("-", " ").replace("/", " ").split():
        known.add(w.strip("-/"))
        known.add(_stem(w))

    # Fold in every synonym of a word the client DID use.
    for grp in _SYN:
        if known & grp:
            for w in grp:
                # The plural's stem too: "houses" stems to "hous", which is not
                # what _stem("house") returns.
                known.update({w, _stem(w), w + "s", _stem(w + "s")})

    def _alien(svc):
        return [w for w in (svc.get("service") or "").lower().split()
                if w not in _GROUNDING_STOP and len(w) > 2
                and w not in known and _stem(w) not in known]

    out, dropped, blocked_pins = [], [], []
    for svc in services or []:
        if svc.get("from_seed"):
            out.append(svc)
            continue
        alien = _alien(svc)
        if alien and svc.get("pinned"):
            # Always enforced, and deliberately NOT counted toward the
            # stand-down ratio below. That valve asks "is the corpus too thin
            # to judge the MODEL fairly"; pins are a handful of terms whose
            # only job is to hold the price steady, and a junk pin corrupts
            # pricing directly. Refuse it and say so.
            blocked_pins.append((svc.get("service"), alien[0]))
            continue
        if alien:
            dropped.append((svc.get("service"), alien[0]))
            continue
        out.append(svc)
    # Stand down when the corpus is too thin to be a fair test — but measure
    # that against the WHOLE list, not just the model's share of it. Keller
    # contributed only 3 non-seed services and 2 were competitors: a correct
    # 2-of-3 read as a 67% drop and tripped the valve, so Turner and Clark
    # survived (2026-07-27). Against the full list those 2 are 29% — plainly
    # a targeted removal — while MPG's 8 wipe out 73% of its list, which is
    # the runaway this valve exists to catch.
    # Pins sit outside this measurement on both sides of the fraction — they
    # are enforced unconditionally above, so counting them would let three bad
    # pins trip the valve and hand the model's competitors back too.
    unpinned = [s for s in (services or []) if not s.get("pinned")]
    total = len(unpinned)
    max_ratio = float(CFG.get("grounding_max_drop_ratio", 0.5) or 0.5)
    if total and len(dropped) / total > max_ratio:
        # Stood down for the model's picks only. Blocked pins stay blocked.
        blocked_l = {(b[0] or "").lower() for b in blocked_pins}
        return ([s for s in (services or [])
                 if (s.get("service") or "").lower() not in blocked_l],
                dropped, blocked_pins, True)
    return out, dropped, blocked_pins, False


def grounding_gap_words(dropped, limit=6):
    """The client's VOCABULARY HOLE, in their own frequency order.

    NPAIHB's description was a full, accurate sentence and the panel still told
    the operator to go write one. The filter had not failed on length — it failed
    because npaihb.org says "tribal" and never says "native american",
    "indigenous" or "wellness", so every service using those words looked as
    foreign as a competitor's name would. That is a fixable, nameable gap and the
    words are already in hand: they are the ones that tripped the filter.
    (2026-08-12)

    Returns [(word, count)], commonest first.
    """
    tally = {}
    for _term, word in (dropped or []):
        w = str(word or "").strip().lower()
        if len(w) > 2:
            tally[w] = tally.get(w, 0) + 1
    return sorted(tally.items(), key=lambda kv: (-kv[1], kv[0]))[:int(limit)]


def enforce_seed_services(services, seeds, max_services, markets, state, phrase_geos=None):
    """Make the partner's own seed terms the backbone of the service list.

    The prompt asks for this and the model does not reliably comply: a
    Northern Virginia insurance agency supplied eight clean seeds (auto, home,
    renters, landlord, homeowners...) and got back "state farm homeowners
    insurance quote" — a competitor — while its own seeds went unused
    (Rockingham, 2026-07-27). No wording fixes this dependably, because the
    keyword-idea pool is ranked by national volume and the largest competitor
    in any trade sits at the top of it looking like a service.

    So the seeds are enforced here. Someone who knows the account typed them;
    they outrank anything the model invents. Model-chosen services only fill
    the slots the seeds don't.

    Returns (services, used_seed_count, clean_seed_total). The total matters
    because the slice below is clean[:max_services] IN ENTRY ORDER, not by
    volume — hand the tool 80 focus terms against a 20-service grid and 60 of
    them are silently dropped on typing order alone. The caller surfaces that.
    (2026-08-11)
    """
    clean = []
    seen = set()
    for sd in seeds or []:
        name = clean_kw(strip_placeholders(strip_proximity(
            _strip_markets((sd or "").lower(),
                           list(markets or []) + list(phrase_geos or []),
                           state)))).strip()
        if name and name not in seen and len(name.split()) <= 6:
            seen.add(name)
            clean.append(name)
    if not clean:
        return list(services or []), 0, 0

    # Seeds FIRST, then model picks fill what's left. The earlier version
    # appended seeds to the model's list and displaced from the tail, which
    # left the model's competitor pick in place whenever the seeds ran out —
    # exactly the case this exists to prevent. Building seeds-first makes the
    # partner's list the default and the model's contribution the remainder.
    def norm(t):
        return " ".join((t or "").lower().split())

    # KEYED ON MEANING, NOT SPELLING. This deduped on the exact string, so the
    # refinement pass could hand the grid "new york state fund workers comp" AND
    # "new york state fund workers compensation", and "usl&h" alongside "usl&h
    # insurance" — four of PEO Brokers' twenty slots on two services, each pair
    # then splitting its own rankings across two grid rows. Same key the seed fold
    # uses, so the two agree on what "the same service" is, and equality-only, so
    # "auto insurance" and "bundle home and auto insurance" still both belong.
    # A skipped duplicate frees its slot for the next pick rather than shortening
    # the grid. (2026-08-13)
    # THE STAGE THE BACKSTOP KEPT POINTING AT. The alias went into the seed fold
    # and the final pass and not into this one, so "workers compensation nysif"
    # — a machine pick, which the seed fold never sees — survived to be caught
    # at the last step and reported as a bug. It was: this key needed telling.
    # Computed from the seeds AND the model's picks together, because the two
    # spellings are usually one of each. (2026-08-14)
    # From the RAW seeds, not `clean` — clean drops anything over six words, and
    # the long form is nearly always the wordy one. PEO Brokers' expansion lives
    # in "new york state insurance fund workers compensation", seven words, which
    # never reaches `clean` and so left the acronym with nothing to match.
    _alias = acronym_aliases(
        [str(x) for x in (seeds or [])]
        + [str(x.get("service", "")) for x in (services or [])])

    def key(t):
        return _seed_key(t, _alias) or frozenset({norm(t)})

    out, taken = [], set()
    for term in clean[:max_services]:
        if key(term) in taken:
            continue
        taken.add(key(term))
        # Tier assigned below, once the final length is known.
        out.append({"service": term, "tier": "competitive", "from_seed": True})
    used = len(out)
    # Model-chosen services fill any remaining slots, in the order it ranked
    # them. Exact duplicates only — "auto insurance" and "bundle home and auto
    # insurance" are different services and both belong.
    for svc in services or []:
        if len(out) >= max_services:
            break
        n = norm(svc.get("service"))
        if not n or key(n) in taken:
            continue
        taken.add(key(n))
        out.append(dict(svc))
    out = out[:max_services]

    # TIERS FROM THE MEASURED MIX, NOT A FIXED LADDER. This used to assign from a
    # seven-element list by position, so everything from index 5 on was long tail
    # forever: any seed-driven list of 20 landed at 2/3/15 - 10/15/75 against the
    # 30/38/32 measured off twelve of BE's own proposals. It also fought
    # CFG["tier_mix"], which the keyword pool already obeys (see tier_split at the
    # candidate-bucketing call), so the tool held two tier models and the seed
    # path won because it ran last.
    #
    # Scored on every proposal we hold rather than on the client in front of us
    # (bench.py): tier_split wins on 10 of 12 and cuts aggregate error from 1047
    # to 319 percentage points. The two it loses are long-tail-heavy outliers -
    # NASSCO at 11/14/74 and Red Shoes at 23/7/70 - where the old ladder was
    # right by accident rather than by rule.
    #
    # Assignment is by POSITION, which is demand rank: the list arrives ordered
    # by measured volume from rank_seeds(), or by the model's own ranking. Same
    # basis step 1 uses to choose head terms. (2026-08-11)
    n_u, n_c, _n_l = tier_split(len(out))
    for i, svc in enumerate(out):
        svc["tier"] = ("ultra" if i < n_u
                       else "competitive" if i < n_u + n_c
                       else "long_tail")
    return out, used, len(clean)


# Words that describe the SHAPE of a retail term rather than its subject. They
# appear across every topic a client sells, so clustering on them would merge
# "ski shop" with "bbq grill store" and report one topic where there are two.
_TOPIC_STOP = set("""shop shops store stores storefront outlet outlets retailer retailers
service services repair rental rentals sales sale supplier suppliers supply
best top cheap affordable local near me nearby quality premium discount
buy buying sell selling new used online cheapest price prices cost
and or the of for in on to with a an my your our
company companies co inc llc dealer dealers center centre centers
cityname city_name citystate locationname marketname statename clientname tbd
""".split())


def clean_seeds(seeds):
    """Normalise the operator's seed terms once, at the door.

    Seeds arrive from three places — typed, imported from a report, restored
    from a saved quote — and only the report path was being scrubbed. A saved
    Ski Barn quote still carried "bbq grill store <cityname>", so the
    placeholder became a TOPIC ("cityname — 24% of what you typed") and was
    handed its own services (2026-08-07). Cleaning here means every consumer
    downstream is clean: topic clustering, seed enforcement, grounding, grid.

    Order is preserved deliberately — the seed list is a priority list and
    enforce_seed_services fills the grid from the front.
    """
    out, seen = [], set()
    for s in (seeds or []):
        t = re.sub(r"\s+", " ", strip_placeholders(str(s or "").strip())).strip()
        if not t:
            continue
        k = t.lower()
        if k in seen:
            continue
        seen.add(k)
        out.append(t)
    return out


def _topic_stem(t):
    """Crude stem, enough to make two spellings of one subject match.

    Plural AND gerund, because retail seed lists mix them freely: "skis" with
    "ski", "grills" with "grill", and — the one that split Ski Barn's snowboard
    topic in two — "snowboarding" with "snowboard".
    """
    t = (t or "").strip()
    if len(t) > 5 and t.endswith("ing"):
        stem = t[:-3]
        # "snowboarding" -> "snowboard"; "shopping" -> "shop" (doubled letter)
        if len(stem) > 3 and stem[-1] == stem[-2]:
            stem = stem[:-1]
        if len(stem) > 3:
            return stem
    if len(t) > 3 and t.endswith("es") and t[-3] in "sxzh":
        return t[:-2]
    if len(t) > 3 and t.endswith("s") and not t.endswith("ss"):
        return t[:-1]
    return t


def _topic_tokens(text):
    """Subject words in a term, stemmed, with retail-shape words removed."""
    out = set()
    for w in re.split(r"[^a-z0-9]+", (text or "").lower()):
        if not w or w in _TOPIC_STOP:
            continue
        s = _topic_stem(w)
        if s and s not in _TOPIC_STOP and len(s) > 2:
            out.add(s)
    return out


def topic_clusters(seeds):
    """Group the operator's seed terms into TOPICS the client actually sells.

    Ski Barn entered 25 terms covering two businesses — ski/snowboard gear and
    BBQ/patio furniture — and got a 7-service grid that was entirely ski, so half
    the company was missing from its own proposal (2026-08-07). The service
    selector ranks by volume, and ski volume dwarfs patio volume, so the smaller
    topic loses every time no matter how many terms the operator types for it.

    Deliberately NOT an AI call. Topic coverage is a correctness guarantee, and a
    guarantee that depends on a non-deterministic call is not one — the same
    reasoning that put pin_head_services and enforce_seed_services in code. AI
    still gets to NAME the topics for display; this decides membership.

    Single-link on shared subject words after dropping retail-shape words.
    Returns [{"label":..., "seeds":[...], "tokens":set()}] biggest topic first.
    """
    def toks(s):
        return _topic_tokens(s)

    items = [(s, toks(s)) for s in (seeds or []) if str(s).strip()]
    items = [(s, t) for s, t in items if t]
    if not items:
        return []
    n = len(items)
    df = {}
    for _s, t in items:
        for tok in t:
            df[tok] = df.get(tok, 0) + 1

    # WAS SINGLE-LINK, WHICH CHAINS. A seed joined any group sharing one token, so
    # on Amare Homes "pet friendly apartments" and "pet friendly homes for rent"
    # bridged the two halves on the word "pet", "rent" bridged nearly everything
    # else, and 25 of 27 seeds collapsed into one topic — leaving the coverage
    # guarantee with nothing to guarantee on exactly the client that needed it.
    #
    # A seed now belongs to the topic of the most common token that is not in most
    # of the seeds. Two ideas, both about what a topic IS: the thing being sold is
    # the word that recurs (apartment, home), and a word present in most of the
    # list cannot tell any of it apart (rent). Modifiers like "pet friendly" are
    # rare, so they no longer decide membership. Deterministic, same as before.
    # (2026-08-13)
    cap = max(1, int(n * float(CFG.get("topic_token_max_share", 0.5) or 0.5)))
    common = {tok for tok, c in df.items() if c > cap}

    def key(t):
        cands = [tok for tok in t if tok not in common] or sorted(t)
        return max(sorted(cands), key=lambda tok: (df[tok], len(tok)))

    groups = {}
    for s, t in items:
        g = groups.setdefault(key(t), {"seeds": [], "tokens": set()})
        g["seeds"].append(s)
        # High-frequency words are left OUT of the topic's token set too, or
        # service_topic() would match every service to every topic on "rent".
        g["tokens"] |= (t - common)

    out = [{"label": k, "seeds": g["seeds"], "tokens": (g["tokens"] or {k}),
            "size": len(g["seeds"])} for k, g in groups.items()]
    out.sort(key=lambda g: (-g["size"], g["label"]))
    # A TOPIC BACKED BY ONE SEED IS A TERM, NOT A TOPIC. Keying is sharper than the
    # old chaining, which is the point — but sharp enough to turn Ski Barn's two
    # halves into six, and every topic claims a service slot, so a 7-slot grid
    # would have been decided entirely by the guarantee with nothing left for
    # demand. Coverage is for a part of the business the ranking would otherwise
    # eliminate; one term is not that, and the volume ranking can be trusted with
    # it. Kept if it is all there is. (2026-08-13)
    _min = int(CFG.get("topic_min_seeds", 2) or 1)
    real = [g for g in out if g["size"] >= _min]
    return real or out[:1]


def service_topic(service, topics):
    """Which topic a chosen service belongs to, or '' if none claim it."""
    stems = _topic_tokens(service)
    best, best_n = "", 0
    for t in topics:
        n = len(stems & t["tokens"])
        if n > best_n:
            best, best_n = t["label"], n
    return best


def enforce_topic_coverage(services, seeds, max_services, cands=None, topics=None):
    """Every topic the operator typed must appear in the service list.

    Proportional to how much of the input each topic represents: a client whose
    seeds are 19 ski terms and 6 patio terms should not get 7 ski services and
    zero patio ones. Under-represented topics take slots from over-represented
    ones, cheapest slot first (the last service in the biggest topic).

    Returns (services, report) where report lists what was added and why, so the
    operator can see it happened rather than wondering why the list changed.
    """
    topics = topics if topics else topic_clusters(seeds)
    if len(topics) < 2 or not services:
        return services, []

    n_slots = min(int(max_services or len(services)), len(services)) or len(services)
    # SHARE IS OUT OF WHAT WAS TYPED, not out of what clustered. topic_clusters
    # drops a group backed by one seed (it is a term, not a topic), so those seeds
    # left the denominator too and the survivors' shares inflated to fill the gap:
    # PEO Brokers' nine industry terms each keyed on a unique word once "worker"
    # crossed the common-token threshold, leaving three topics holding 8 of 17
    # seeds and claiming 38%, 38% and 25% of a grid they are 18% of. Quotas then
    # reserved every slot in the grid. Counting the seeds themselves leaves the
    # unclustered ones to the volume ranking, which is the correct owner of a term
    # no topic claims. No effect when the model assigns the topics: it assigns
    # every term, so the two counts agree. (2026-08-13)
    total_seeds = max(sum(t["size"] for t in topics),
                      len([x for x in (seeds or []) if str(x).strip()])) or 1

    # A topic only earns a guaranteed slot if the operator's input actually
    # weights it that far. One seed out of 29 is 3% of the input; handing it one
    # of 7 services would be 14% — over-rewarding a stray term at the expense of
    # the business. Topics below the threshold can still be picked on merit,
    # they just aren't protected.
    min_share = 1.0 / max(n_slots, 1)
    _kept_t = [t for t in topics if t["size"] / total_seeds >= min_share * 0.75]
    # WHY A TOPIC WASN'T COVERED, ON THE PANEL. Ski Barn — the client this pass
    # was written for — came back "bbq & grills 9% -> 0 services · patio
    # furniture 9% -> 0 services" with no explanation available anywhere. Run in
    # isolation the pass promotes four of those terms; the live build promoted
    # none, and nothing recorded which branch it took. The protection threshold
    # moves with the slot count (1/7 is 14%, 1/20 is 5%), so a topic that is
    # protected on a one-city grid is unprotected on a five-city one — and this
    # returns silently in that case. Guessing at it from the outside cost most of
    # an afternoon. (2026-08-17)
    _unprot = [{"kind": "unprotected", "added": "", "replaced": "",
                "topic": t["label"],
                "share": round(t["size"] / total_seeds, 4),
                "needed_share": round(min_share * 0.75, 4), "slots": n_slots}
               for t in topics if t not in _kept_t]
    topics = _kept_t
    if len(topics) < 2:
        return services, (_unprot + [{"kind": "stood_down", "added": "",
                                      "replaced": "", "slots": n_slots,
                                      "protected": len(topics),
                                      "total_seeds": total_seeds}])

    quota = {}
    for t in topics:
        quota[t["label"]] = max(1, round(n_slots * t["size"] / total_seeds))
    # THE DECISION ITSELF, RECORDED UNCONDITIONALLY. Two rounds of instrumenting
    # the branches I THOUGHT were firing both came back empty on Ski Barn, which
    # means my model of this function is wrong rather than the branches being
    # rare. So this records the inputs — slot count, quota, and what each topic
    # already holds — whatever happens next. Guessing from the outside has now
    # cost three builds. (2026-08-17)
    _decision = {"kind": "quota", "added": "", "replaced": "",
                 "slots": n_slots, "total_seeds": total_seeds,
                 "quota": dict(quota),
                 "sizes": {t["label"]: t["size"] for t in topics}}
    # Trim quotas back to the slots available, smallest topics protected.
    while sum(quota.values()) > n_slots:
        big = max(quota, key=lambda k: quota[k])
        if quota[big] <= 1:
            break
        quota[big] -= 1

    out = [dict(x) for x in services]
    _decision["quota_after_trim"] = dict(quota)
    for x in out:
        svc = x.get("service", "")
        # A LOOKUP CLAIMS NO TOPIC. Even with the definitional seeds kept out of
        # the clustering, "what does lcf stand for" can still land in a topic on a
        # shared stem and then sit there defended by the quota. It is unclaimed by
        # construction, which also makes it the first slot donated below.
        x["_topic"] = "" if is_lookup_kw(svc) else service_topic(svc, topics)

    vol = {str(r.get("keyword", "")).lower(): (r.get("volume") or 0)
           for r in (cands or [])}
    report = []
    _decision["have"] = {}
    for t in topics:
        lab = t["label"]
        have = [x for x in out if x.get("_topic") == lab]
        need = quota.get(lab, 1) - len(have)
        _decision["have"][lab] = len(have)
        if need <= 0:
            continue
        # Best unused seed from this topic, by measured volume then by order.
        # KEYED, NOT STRING-MATCHED. PEO Brokers came back with "new york state
        # fund workers comp" AND "new york state fund workers compensation" in
        # the same six-slot top tier, both at 880/mo. enforce_seed_services had
        # already folded them — this pass put the second one back, because it
        # only checked the exact spelling of what was already there. A guarantee
        # that fills a topic with a rewording of a service the grid already
        # holds has covered nothing. (2026-08-13)
        used = {str(x.get("service", "")).lower() for x in out}
        _al = acronym_aliases([str(x.get("service", "")) for x in out]
                              + [str(s) for s in (t.get("seeds") or [])])
        used_keys = {_seed_key(str(x.get("service", "")), _al) or
                     frozenset({str(x.get("service", "")).lower()}) for x in out}

        def _dupe(term):
            t2 = str(term).strip()
            if not t2 or t2.lower() in used:
                return True
            return (_seed_key(t2, _al) or frozenset({t2.lower()})) in used_keys

        pool = [s for s in t["seeds"] if not _dupe(s)]
        pool.sort(key=lambda s: (-vol.get(str(s).lower(), 0), t["seeds"].index(s)))
        # Which candidates are the OPERATOR'S OWN — everything from t["seeds"] is,
        # everything the cands fallback below adds is not. The donor rule needs to
        # know, because "seed in, seed out" and "seed out, machine pick in" are
        # completely different trades. (2026-08-17)
        _seed_sourced = {str(x).strip().lower() for x in pool}
        # WHEN THE TOPIC HAS NO SEED LEFT TO PROMOTE. PEO Brokers typed 21 terms
        # for 20 slots, so freeing three of them from abbreviation lookups freed
        # nothing: every remaining seed was already in the grid and the loop had
        # no candidate. Fall back to the MEASURED keyword pool, filtered to this
        # topic — a coverage guarantee that can only reshuffle the operator's own
        # typing is not covering anything the typing already missed. Measured and
        # topic-matched only, never a lookup, so it cannot reintroduce what the
        # slot was just taken back from. (2026-08-13)
        if len(pool) < need:
            extra = []
            for r in (cands or []):
                kw = str(r.get("keyword", "")).strip()
                if not kw or not (r.get("volume") or 0):
                    continue
                lo = kw.lower()
                if _dupe(kw) or lo in {str(x).lower() for x in pool}:
                    continue
                if any((_seed_key(kw) or frozenset({lo})) ==
                       (_seed_key(str(x)) or frozenset({str(x).lower()}))
                       for x in pool):
                    continue
                if is_lookup_kw(kw) or service_topic(kw, topics) != lab:
                    continue
                extra.append((-(r.get("volume") or 0), kw))
            extra.sort()
            seen_x = set()
            for _v, kw in extra:
                if kw.lower() in seen_x:
                    continue
                seen_x.add(kw.lower())
                pool.append(kw)
                if len(pool) >= need:
                    break
        if not pool:
            report.append({"kind": "no_candidate", "added": "", "replaced": "",
                           "topic": lab, "need": need,
                           "slots": n_slots})
        for s in pool[:need]:
            # A SERVICE NO TOPIC CLAIMS IS THE CHEAPEST SLOT IN THE LIST. It used
            # to be untouchable: donors were drawn only from topics that were OVER
            # quota, so five unclaimed services sat protected while the apartment
            # topic — quota 6 — got 2 and the loop stopped. Nothing is defending
            # an unclaimed service, so it goes first. (2026-08-13)
            # A TERM THE OPERATOR TYPED IS NOT A SPARE SLOT. Every other filter
            # in the build exempts a seed; this pass did not, and it is the one
            # pass that REMOVES a service to make room for another. PEO Brokers
            # put five SWIF terms on the focus list, they survived the classifier
            # and the grounding filter, and then none of them reached the grid —
            # three swaps, three seeds gone. Whether a seed is "unclaimed"
            # depends on labels a model wrote that build; it says nothing about
            # whether the operator meant it. Donate machine picks only, and if
            # there are none, leave the topic short rather than overrule them.
            # (2026-08-14)
            _incoming_is_seed = str(s).strip().lower() in _seed_sourced
            orphans = [x for x in out
                       if not x.get("_topic")
                       and (_incoming_is_seed or not x.get("from_seed"))]
            if orphans:
                # Last one first is "cheapest slot"; a lookup is cheaper still,
                # so it goes ahead of any other unclaimed service. Stable, so
                # rank order still decides within each group.
                orphans.sort(key=lambda x: 1 if is_lookup_kw(x.get("service", "")) else 0)
                drop = orphans[-1]
            else:
                # THE SAME ASYMMETRY, ONE LINE EARLIER. This picks which topics
                # are even eligible to donate, and it counted only non-seed rows —
                # so on a grid made entirely of seeds every topic looked to have
                # nothing to give and the loop broke out before the donor test was
                # reached at all. Fixing the donor list without this one changed
                # nothing, which is what the reproduction showed. (2026-08-17)
                _over = [k for k in quota
                         if len([x for x in out if x.get("_topic") == k
                                 and (_incoming_is_seed
                                      or not x.get("from_seed"))]) > 0]
                if not _over:
                    break
                donor_lab = max(_over, key=lambda k: len([x for x in out if x.get("_topic") == k])
                                - quota.get(k, 1))
                # SEED FOR SEED IS ALLOWED; SEED FOR A MACHINE PICK IS NOT.
                #
                # "A term the operator typed is not a spare slot" (2026-08-14)
                # excluded every from_seed row from donating, and that is right
                # when the incoming term is something the tool invented. But
                # enforce_seed_services stamps from_seed on EVERY row it creates,
                # so a client with more focus terms than slots has a grid made
                # entirely of seeds — and then this rule leaves no donor at all
                # and the topic guarantee breaks out having done nothing.
                #
                # Ski Barn is exactly that shape: 59 focus terms, 20 slots, every
                # service from_seed. Slots 20, quota bbq 2 / patio 2 / ski 16,
                # already held bbq 0 / patio 0 / ski 20 — the pass had work to do,
                # a full pool to do it from, and broke out on the donor test. Its
                # BBQ and patio lines have been dropping out of that quote since
                # the seed protection landed.
                #
                # Two guarantees were in direct conflict, and the operator settles
                # it: they typed the BBQ terms too. Refusing to trade a ski seed
                # for a BBQ seed does not protect their intent, it overrules it —
                # silently deleting a whole product line they asked for. So a seed
                # may donate to another SEED, and still never to an invention.
                # (2026-08-17)
                _incoming_is_seed = str(s).strip().lower() in _seed_sourced
                donors = [x for x in out if x.get("_topic") == donor_lab
                          and (_incoming_is_seed or not x.get("from_seed"))]
                if len(donors) <= 1 or len(
                        [x for x in out if x.get("_topic") == donor_lab]) <= quota.get(donor_lab, 1):
                    break
                drop = donors[-1]
            # Scrub on the way in. This function runs AFTER the last
            # scrub_services pass, so a seed added here is the only service that
            # never gets cleaned — which is exactly how "bbq grill store
            # cityname new york city ny" reached a grid on a build that already
            # stripped placeholders everywhere else (2026-08-07).
            name = clean_kw(strip_placeholders(strip_proximity(str(s)))).strip()
            if not name or name in {str(x.get("service", "")).lower() for x in out}:
                continue
            out.remove(drop)
            out.append({"service": name, "tier": drop.get("tier", "competitive"),
                        "_topic": lab})
            report.append({"added": name, "topic": lab,
                           "replaced": drop.get("service", ""),
                           "from_topic": drop.get("_topic") or "unclaimed"})
    for x in out:
        x.pop("_topic", None)
    return out, (_unprot + [_decision] + report)


def tier_split(n_terms):
    """How many terms belong in each tier for a list of this length.

    Proportional, from CFG["tier_mix"] — measured off BE's own proposals, whose
    splits scale with list size rather than sitting at fixed counts. Returns
    (n_ultra, n_competitive, n_long_tail) summing exactly to n_terms, with at
    least one in each tier once there are three terms to spread (the proposal
    renders three columns and an empty one reads as an incomplete strategy).
    """
    n = max(int(n_terms or 0), 0)
    if n <= 0:
        return (0, 0, 0)
    mix = CFG.get("tier_mix") or {"ultra": 0.30, "competitive": 0.38, "long_tail": 0.32}
    order = ("ultra", "competitive", "long_tail")

    hard = {"ultra": CFG.get("ultra_bucket_size"),
            "competitive": CFG.get("competitive_bucket_size")}
    if hard["ultra"] is not None or hard["competitive"] is not None:
        u = int(hard["ultra"]) if hard["ultra"] is not None else round(n * mix["ultra"])
        c = int(hard["competitive"]) if hard["competitive"] is not None else round(n * mix["competitive"])
        u, c = max(0, min(u, n)), max(0, min(c, n - min(u, n)))
        return (u, c, max(0, n - u - c))

    # Largest-remainder so the three always sum to n exactly.
    raw = {t: n * float(mix.get(t, 0)) for t in order}
    base = {t: int(raw[t]) for t in order}
    left = n - sum(base.values())
    for t in sorted(order, key=lambda x: -(raw[x] - base[x]))[:max(0, left)]:
        base[t] += 1
    if n >= 3:
        # Top up any empty tier from the largest one.
        for t in order:
            if base[t] == 0:
                donor = max(order, key=lambda x: base[x])
                if base[donor] > 1:
                    base[donor] -= 1
                    base[t] += 1
    return (base["ultra"], base["competitive"], base["long_tail"])


def rebalance_tiers(services):
    """Guarantee all three tiers are represented.

    Anything that removes a service — the out-of-area filter, seed
    enforcement, pinning — can empty a tier, and the proposal renders three
    columns. An empty one reads as an incomplete strategy rather than a
    deliberate choice. Move from the most-crowded tier into the empty one,
    preferring the longest phrase for long_tail and the shortest for ultra,
    which is how the tiers actually differ.
    """
    order = ["ultra", "competitive", "long_tail"]
    out = [dict(x) for x in (services or [])]
    if len(out) < 3:
        return out                     # too few to fill three tiers honestly
    for _ in range(len(order)):
        counts = {t: [x for x in out if x.get("tier") == t] for t in order}
        empty = [t for t in order if not counts[t]]
        if not empty:
            break
        need = empty[0]
        donor = max(order, key=lambda t: len(counts[t]))
        if len(counts[donor]) < 2:
            break
        pool = counts[donor]
        pick = (max(pool, key=lambda x: len(x["service"].split())) if need == "long_tail"
                else min(pool, key=lambda x: len(x["service"].split())))
        for x in out:
            if x is pick:
                x["tier"] = need
                break
    return out


_NARROW_QUALIFIER = re.compile(
    r"\b(?:with|without|w/|no|near|next to|walking distance|close to|"
    r"that (?:allow|accept)|accepting|allowing|includes?|including)\b", re.I)


def pool_vocabulary(candidates, seeds=None, site_terms=None, min_volume=None):
    """Every word the MARKET uses, from the keyword pool this build already paid
    for, plus the operator's own terms.

    WEIGHTED BY VOLUME, because the pool is not an independent sample. It comes
    from `keywords_for_keywords` and `keyword_suggestions` SEEDED WITH THE
    SERVICE LIST — and keyword_suggestions in particular returns queries
    containing the seed. Ask it about "rental homes no credit check" and it will
    hand back phrasings of that, which then appear in the vocabulary and vouch
    for the term that produced them. Amare judged all seventeen qualifiers and
    grounded all seventeen on exactly that loop.
    Existence is not the test — a phrase can be real, returned by Google, and
    still be searched by nobody. A qualifier earns its place from a keyword
    someone actually runs.

    AND ONLY FROM THE ROWS MEASURED WHERE THE CLIENT SELLS. The pool is two
    populations wearing the same shape: keyword_suggestions is localised
    (loc_string -> "Santa Fe,New Mexico,United States") and keywords_for_site is
    Labs location_code 2840, the whole United States. Santa Fe rentals report
    10-90/mo; the national rows report tens of thousands. An absolute floor
    across both selects AGAINST the client's own market and FOR the national
    noise — which is exactly what happened: "luxury" was judged unused in Santa
    Fe while "new york city apartments for rent" counted as this market's
    vocabulary. So national rows are excluded, and the floor is taken from the
    LOCAL population's own distribution rather than assumed. (2026-08-16)

    `keywords_for_keywords` and `keyword_suggestions` come back with a few
    hundred phrases people really type. The expansion prompt has always been
    handed that pool as "evidence of real demand", and advisory is exactly what
    it stayed: the model reads it and still writes "rental homes with washer
    dryer" whether or not anything resembling it is in there. That is where the
    feature permutations come from — plausible variants, generated rather than
    observed, spending slots Brendan gives to bedroom counts and landmarks.

    Returns the stemmed token set. The operator's seeds are in it because a
    planner who knows the account outranks the pool, and the site's own terms
    because a service a client sells is legitimate whether or not Google's idea
    list happened to surface it. (2026-08-16)
    """
    local = [c for c in (candidates or [])
             if isinstance(c, dict) and c.get("scope") != "national"]
    # A FLOOR READ OFF THIS MARKET, not carried in from another one. Santa Fe's
    # rows sit at 10-90/mo, so 30 is a coin flip; a metro's sit in the
    # thousands, where 30 excludes nothing. The median of what came back is the
    # same question asked in the market's own units.
    # The median, but never below the no-data marker: Google Ads reports 10/mo
    # for a phrase it holds nothing on, so a pool sitting entirely at 10 is not
    # a vocabulary — it is an absence of one, and should stand the filter down
    # rather than ground everything at the floor.
    _abs = int(CFG.get("pool_min_volume", 11))
    if min_volume is not None:
        floor = int(min_volume)
    else:
        vols = sorted(int(c.get("volume") or 0) for c in local)
        vols = [v for v in vols if v > 0]
        floor = max(vols[len(vols) // 2] if vols else 0, _abs)
    out = set()
    for src, weighed in ((local, True), (seeds or [], False),
                         (site_terms or [], False)):
        for c in src:
            if isinstance(c, dict):
                t = c.get("keyword", "")
                # The operator's terms and the client's own pages are not
                # measured against demand — they are statements about the
                # business. Only the keyword POOL has to earn its way in.
                if weighed and int(c.get("volume") or 0) < floor:
                    continue
            else:
                t = c
            for w in re.split(r"[^a-z0-9]+", str(t or "").lower()):
                st = _seed_stem(w)
                if st and len(st) > 1:
                    out.add(st)
    return out


def ungrounded_qualifiers(service, core, vocab):
    """The words this service adds on top of the head term that the market never
    uses. Empty when the service is grounded, or when there is nothing to add."""
    words = [w for w in re.split(r"[^a-z0-9]+", (service or "").lower()) if w]
    bad = []
    for w in words:
        if w in _SEED_SHAPE or w in _FORM_SKIP or w in _PREPOSITIONS:
            continue
        st = _seed_stem(w)
        if not st or len(st) < 2 or st in core or st in vocab:
            continue
        bad.append(w)
    return bad


def backfill_services(services, candidates, want, markets=None, state="",
                      brand="", vocab=None, tier="long_tail"):
    """Refill slots the qualifier filter took back.

    Removing a service without replacing it shortens the grid, and the grid
    length is what the volume total — and therefore the price — is built from.
    A filter that quietly makes a quote cheaper is worse than no filter: the
    operator asked for twenty terms and would get thirteen with nothing on
    screen connecting the two.

    Fills from the keyword pool by measured volume, skipping anything already
    covered, the client's own name, questions, and — the point of the exercise —
    anything whose own qualifiers are ungrounded, so the backfill cannot walk
    straight back into what was just removed. Returns (services, added).
    (2026-08-16)
    """
    want = int(want or 0)
    if want <= 0 or not candidates:
        return services, []
    out = [dict(x) for x in (services or [])]
    have = [(x.get("service") or "").lower() for x in out]

    def covered(term):
        return any(term == h or term in h or h in term for h in have if h)

    # SAME MARKET, SAME UNITS. The first outing of this ranked the whole pool by
    # volume and offered "new york city apartments for rent" (60,500) and
    # "apartments in san antonio tx" (33,100) to a Santa Fe rental community,
    # because the national rows outrank every local one by three orders of
    # magnitude. They are not candidates for a local grid at any volume.
    candidates = [c for c in (candidates or [])
                  if isinstance(c, dict) and c.get("scope") != "national"]
    if not candidates:
        return services, []
    _v = sorted(int(c.get("volume") or 0) for c in candidates)
    _v = [v for v in _v if v > 0]
    floor = max(_v[len(_v) // 2] if _v else 0, int(CFG.get("pool_min_volume", 11)))
    b = (brand or "").strip().lower()
    keys = {n: _seed_key(n) for n in have if n}
    tally = {}
    for k in keys.values():
        for tok in k:
            tally[tok] = tally.get(tok, 0) + 1
    core = {tok for tok, c in tally.items()
            if c >= len(have) * float(CFG.get("shape_core_share", 0.4))} if have else set()

    added = []
    for c in sorted(candidates, key=lambda r: (-(r.get("volume") or 0),
                                               str(r.get("keyword") or ""))):
        if len(added) >= want:
            break
        vol = int(c.get("volume") or 0)
        if vol < floor:
            break                       # sorted, so nothing below here qualifies
        term = clean_kw(strip_placeholders(strip_proximity(
            _strip_markets((c.get("keyword") or "").lower(),
                           list(markets or []), state)))).strip()
        if not term or covered(term) or is_lookup_kw(term):
            continue
        if b and is_brand_term(term, brand):
            continue
        if vocab and ungrounded_qualifiers(term, core, vocab):
            continue
        out.append({"service": term, "tier": tier})
        have.append(term)
        added.append((term, vol))
    return out, added


def drop_ungrounded_qualifiers(services, candidates, seeds=None, site_terms=None,
                               pinned=None, suggested=None, vocab=None):
    """Take back the slots spent on qualifiers nobody searches.

    NOT to be confused with drop_ungrounded_services() above, which asks whether
    the SERVICE is something this client sells, read off their own site and
    description. This asks whether the QUALIFIER is a word this market uses,
    read off the keyword pool. Different question, different evidence.

    A service survives when every word it adds on top of the list's own core
    appears somewhere in the market's vocabulary. "3 bedroom homes for rent"
    survives if the pool anywhere says "bedroom"; "rental homes with washer
    dryer" does not if it never says "washer".

    THREE THINGS IT WILL NOT DO. It never touches a term the operator TYPED or
    the build pinned — same exemption every other filter gives them. It never
    judges the head term itself, only the qualification. And below
    `pool_vocab_min` distinct tokens the pool is too thin to be evidence of
    anything, so it stands down entirely and says so rather than quietly
    narrowing the list on a small sample.

    TYPED, NOT MERELY PRESENT. The first version exempted every seed, and on
    Amare every seed WAS the tool's own suggestion sitting in the focus box with
    a ✦ on it — seventeen of seventeen. So the filter ran, found nothing it was
    allowed to touch, and reported honestly that it had changed nothing, which
    read exactly like the pool having no objection. `suggested` is the same
    record drop_suggested_nonservices uses for the same reason: a term the tool
    proposed does not get the credit for the operator's judgement.
    (2026-08-16)

    Returns (kept, dropped, status). (2026-08-16)
    """
    out = [dict(x) for x in (services or [])]
    # TYPED seeds only, in BOTH places. They join the market's vocabulary because
    # a planner who knows the account outranks the pool — but a seed the tool
    # suggested is the tool's own guess, and letting it into the vocabulary lets
    # every guess vouch for itself. That is the whole filter defeated in one
    # line: on Amare it put "credit", "pool", "yard" and "move" into the market's
    # words, and then found nothing ungrounded. (2026-08-16)
    _sug = {str(t).strip().lower() for t in (suggested or []) if str(t).strip()}
    typed = [t for t in (seeds or []) if str(t).strip().lower() not in _sug]
    # A vocabulary read at stage 1, where provenance still existed, beats one
    # rebuilt here from rows that no longer know where they were measured.
    # `vocab` is THE POOL'S CONTRIBUTION ALONE.
    market = set(vocab) if vocab else pool_vocabulary(candidates)
    # THE STAND-DOWN IS MEASURED ON THE POOL, NOT ON THE UNION. The client's own
    # seeds and site pages join the vocabulary for the membership test — they are
    # statements about the business and outrank the pool — but they must not
    # count towards "is there enough here to judge from". Amare cleared a
    # sixty-word bar almost entirely on its own website's words while the Santa
    # Fe pool contributed nothing above the no-data floor, so a filter with no
    # evidence behind it went ahead and condemned eight terms, including
    # "luxury" — which is Brendan's, and the client's, own framing.
    # (2026-08-16)
    floor = int(CFG.get("pool_vocab_min", 60))
    if len(market) < floor:
        return out, [], f"thin:{len(market)}"
    vocab = market | pool_vocabulary([], typed, site_terms)
    shape = service_shape(out)
    names = [(x.get("service") or "").strip().lower() for x in out]
    keys = {n: _seed_key(n) for n in names}
    tally = {}
    for k in keys.values():
        for tok in k:
            tally[tok] = tally.get(tok, 0) + 1
    core = {tok for tok, c in tally.items()
            if c >= len(names) * float(CFG.get("shape_core_share", 0.4))}
    safe = {seed_norm(str(t)).lower() for t in typed}
    safe |= {str(t).strip().lower() for t in (pinned or [])}
    safe.discard("")
    kept, dropped = [], []
    for x in out:
        n = (x.get("service") or "").strip().lower()
        if n in safe or not shape:
            kept.append(x)
            continue
        bad = ungrounded_qualifiers(n, core, vocab)
        if bad:
            dropped.append((n, bad[0]))
        else:
            kept.append(x)
    # It cannot empty the list, and it cannot cut so deep that the grid has
    # nothing to quote. Same instinct as every other filter here.
    if len(kept) < max(3, int(len(out) * 0.4)):
        return out, [], f"held:{len(dropped)}"
    # How many were actually JUDGED, not how many exist. Silence has cost two
    # builds on this feature already: "ran and found nothing" has to be
    # distinguishable from "never got to run".
    return kept, dropped, f"ok:{len(out) - len(safe & set(names))}"


def service_shape(services):
    """How QUALIFIED each service is, relative to the list it sits in.

    The core of a list is whatever tokens most of it shares — {hom, rent} for a
    rental community, {comp} for a workers-comp broker. What a term adds on top
    of the core is its qualification, and that is the axis Brendan's tiers
    actually run on: bare head terms are Ultra, one qualifier is Competitive, an
    amenity or a landmark is Long Tail.

    Returns {service: (depth, narrow)}. `depth` counts qualifier tokens beyond
    the core; `narrow` marks the constructions that are long tail no matter how
    few tokens they add — "homes for rent WITH a garage", "rentals with NO
    credit check", "homes for rent NEAR meow wolf". Empty when the list has no
    shared core, because relative depth means nothing across unrelated services.
    (2026-08-16)
    """
    names = [(x.get("service") or "").strip().lower() for x in (services or [])]
    names = [n for n in names if n]
    if len(names) < 4:
        return {}
    keys = {n: _seed_key(n) for n in names}
    share = float(CFG.get("shape_core_share", 0.4))
    tally = {}
    for k in keys.values():
        for tok in k:
            tally[tok] = tally.get(tok, 0) + 1
    core = {tok for tok, c in tally.items() if c >= len(names) * share}
    if not core:
        return {}
    return {n: (len(keys[n] - core), bool(_NARROW_QUALIFIER.search(n)))
            for n in names}


def drop_foreign_geo_services(services, markets, state):
    """Remove services naming a US state the client doesn't operate in.

    The keyword-idea pool is ranked by national volume, so it fills with terms
    anchored to whichever state searches most — a Northern Virginia insurer
    came back with "state of california fire insurance" as a service
    (Rockingham, 2026-07-27). No prompt wording reliably suppresses this,
    because the term reads as a plausible insurance product; it is only wrong
    once you know where the client sells. That is a fact the code has and the
    model has to be told, so enforce it here rather than asking.

    The client's own state and any state named in their markets are kept, so
    "virginia farm insurance" survives for a Virginia client.
    """
    # If we cannot establish the client's own state, we cannot judge which
    # states are foreign — and dropping on a guess would delete a legitimate
    # home-state service ("virginia farm insurance" for a Virginia client whose
    # market pills are bare city names). Say nothing rather than guess wrong.
    known_state = bool((state or "").strip()) or any(
        p in STATE_ABBREV or p in set(STATE_ABBREV.values())
        for m in (markets or [])
        for p in (m or "").lower().replace(",", " ").split())
    if not known_state:
        return list(services or []), None      # None = filter could not run

    ours = set()
    for m in list(markets or []) + [state or ""]:
        t = (m or "").strip().lower()
        if not t:
            continue
        for part in t.replace(",", " ").split():
            if part in STATE_ABBREV:
                ours.add(part)
                ours.add(STATE_ABBREV[part])
            elif part in set(STATE_ABBREV.values()):
                ours.add(part)
                for full, ab in STATE_ABBREV.items():
                    if ab == part:
                        ours.add(full)
    out, dropped = [], []
    for svc in services or []:
        name = (svc.get("service") or "").lower()
        toks = set(name.replace(",", " ").split())
        foreign = [st for st in STATE_ABBREV
                   if st not in ours and (f" {st} " in f" {name} ")]
        # Two-word state names ("new york", "north carolina") need a substring
        # test; single-word ones are covered by the token check above.
        if not foreign:
            foreign = [st for st in STATE_ABBREV
                       if " " in st and st not in ours and st in name]
        # AND THE ABBREVIATION, which this only ever tested full names for. Nob
        # Hill Dental — Salem, OREGON — was quoted "dentist salem ct" because the
        # loop above iterates STATE_ABBREV's KEYS, the full names, so it catches
        # "arizona" and never "ct". Abbreviations are how these keywords are
        # actually written.
        #
        # Only the LAST token, and only when the word before it is not a shape
        # word — because half these abbreviations are ordinary English. "me" is
        # Maine and "dentist near me" is a term the tool adds on purpose; "or" is
        # Oregon; "in" is Indiana; "hi", "de", "la", "ok", "pa" are all words or
        # fragments. Requiring the trailing position and a real place-word before
        # it is what separates "salem ct" from "near me". (2026-08-17)
        if not foreign:
            _toks = name.split()
            if len(_toks) >= 3:
                _last, _prev = _toks[-1], _toks[-2]
                if (len(_last) == 2 and _prev not in _SEED_SHAPE
                        and _last not in ours
                        and _last in set(STATE_ABBREV.values())):
                    foreign = [_last]
        if foreign:
            dropped.append((svc.get("service"), foreign[0]))
            continue
        out.append(svc)
    return out, dropped


# Proximity phrases. Google reads these against the SEARCHER's location, so
# they are a substitute for a place name, not a companion to one.
_PROXIMITY_RE = re.compile(
    r"\b(near me|nearby|near by|close to me|around me|in my area|"
    r"near my location|closest|near you)\b")


def strip_proximity(text):
    """Remove 'near me'-style phrases from a SERVICE name.

    The grid appends a city to every service, so a service carrying "near me"
    becomes "mattress store near me acworth ga" (Woodstock Furniture,
    2026-07-27). Nobody types that: "near me" IS the location, so pairing it
    with an explicit city is a contradiction. Those terms report almost no
    volume, and they go into a proposal as keywords the client is quoted to
    rank for. Bare "near me" terms are legitimate — the grid just isn't where
    they belong, because the grid's whole job is to add the place.
    """
    return re.sub(r"\s+", " ", _PROXIMITY_RE.sub(" ", (text or "").lower())).strip()


# Template placeholders. Agency reports print a keyword row once with a token
# standing in for the city — "bbq grill store <cityname>" covers a dozen
# markets in one line. Imported verbatim, the token travels as a literal word:
# "bbq grill store cityname" reached the rank check and the proposal for Ski
# Barn (2026-08-07). Nobody searches it, so it is guaranteed not-ranking, which
# drags the zero-ranking ratio — and the price — up on a phantom. Angle
# brackets are optional because most readers strip them before we see them.
_PLACEHOLDER_RE = re.compile(
    r"[<\[\{\(]?\s*\b(cityname|city_name|city|citystate|location|locationname|"
    r"market|marketname|statename|state|region|geo|area|town|zip|zipcode|"
    r"keyword|kw|service|brand|clientname|client|xxx+|tbd)\b\s*[>\]\}\)]?",
    re.I)


def strip_placeholders(text):
    """Remove template tokens like <cityname> from an imported term.

    Only fires when the token is wrapped (<city>, [City], {city}) or is one of
    the tokens that is never a real search word on its own — "cityname",
    "clientname", "tbd". A BARE "city", "state" or "area" is left alone,
    because "city hall furniture" and "bay area movers" are real terms; those
    only strip inside brackets.
    """
    t = (text or "")

    def keep_or_drop(m):
        whole = m.group(0)
        word = (m.group(1) or "").lower()
        wrapped = whole[:1] in "<[{(" or whole[-1:] in ">]})"
        never_real = word in {"cityname", "city_name", "citystate", "locationname",
                              "marketname", "statename", "clientname", "zipcode", "tbd"} \
                     or word.startswith("xxx")
        return " " if (wrapped or never_real) else whole

    return re.sub(r"\s+", " ", _PLACEHOLDER_RE.sub(keep_or_drop, t)).strip()


def scrub_services(services, markets, state, phrase_geos=None):
    """Strip any market/state text out of the SERVICE names and de-duplicate.

    A service is meant to be a bare offering ("business tech") that the grid
    crosses with cities. When one arrives already carrying a geo — the model
    echoes a candidate back verbatim, or a pinned candidate keeps a market the
    strip list didn't cover — the grid appends a SECOND city on top, producing
    "business tech cherry hill south jersey" (Waytek, 2026-07-27). It also
    burns a slot on a near-duplicate of a service already in the list, which
    on a 6-service grid is a sixth of the proposal.

    Scrubbing here means the grid only ever sees bare services, so the
    crossing is the only thing that adds geography.
    """
    strip_list = list(markets or []) + [g for g in (phrase_geos or []) if g]
    out, seen = [], set()
    for svc in services or []:
        name = clean_kw(strip_placeholders(strip_proximity(
            _strip_markets((svc.get("service") or "").lower(),
                           strip_list, state)))).strip()
        if not name or name in seen:
            continue                      # empty after scrubbing, or a duplicate
        seen.add(name)
        out.append({**svc, "service": name})
    return out


def pin_head_services(services, cands, markets, state, brand, max_services):
    """Force the highest-volume candidate terms into the service list.

    `services` is the model's pick; `cands` are the keyword rows the search API
    returned, each with a real volume. Any top-volume term the model dropped is
    inserted, displacing the lowest-priority model picks so the list length is
    unchanged. Returns (services, pinned_terms).

    Matching is containment-based in both directions so we don't double up:
    if the model already chose "energy gummies", the candidate "energy gummies"
    (or "best energy gummies") is considered covered by it.
    """
    n_pin = int(CFG.get("pin_head_terms", 0) or 0)
    if n_pin <= 0 or not cands:
        return services, []
    min_vol = int(CFG.get("pin_min_volume", 0) or 0)
    b = (brand or "").strip().lower()

    # Bare, brand-free candidate terms ranked by real search volume.
    ranked, seen = [], set()
    for c in sorted(cands, key=lambda r: (-(r.get("volume") or 0),
                                          str(r.get("keyword") or ""))):
        if (c.get("volume") or 0) < min_vol:
            break
        term = clean_kw(strip_placeholders(strip_proximity(
            _strip_markets((c.get("keyword") or "").lower(), markets, state)))).strip()
        if not term or (b and b in term) or term in seen:
            continue
        # A question is never a service, so it should never be pinned. Skipping
        # it HERE rather than letting the grounding filter catch it downstream
        # means the slot goes to the next real head term instead of being spent
        # and then discarded — and the operator stops being told that "how to
        # get rid of a headache" was blocked over the word "get" (2026-08-04).
        if is_question_kw(term):
            continue
        seen.add(term)
        ranked.append((term, c.get("volume") or 0))
        if len(ranked) >= n_pin:
            break
    if not ranked:
        return services, []

    have = [(x.get("service") or "").lower() for x in services]
    def covered(term):
        return any(term == h or term in h or h in term for h in have if h)

    missing = [(t, v) for t, v in ranked if not covered(t)]
    if not missing:
        return services, []

    keep_l = {t for t, _ in ranked}
    out = list(services)

    def _drop_one():
        """Free a slot WITHOUT emptying a tier.

        Displacing from the tail looked reasonable — the model returns its
        strongest picks first — but the model also orders long_tail last, so
        on a short list (6 services for a 6-city client) the pins ate every
        long-tail service and the proposal came back with an empty Long Tail
        column (Waytek, 2026-07-27). Take from the tier that can best afford
        it instead: the most-represented one, from its own tail.
        """
        counts = {}
        for x in out:
            if (x.get("service") or "").lower() not in keep_l:
                counts[x.get("tier")] = counts.get(x.get("tier"), 0) + 1
        # Only tiers that would still have a member afterwards are eligible.
        eligible = {t: n for t, n in counts.items() if n > 1}
        pool = eligible or counts
        if not pool:
            return False
        victim_tier = max(pool.items(), key=lambda kv: kv[1])[0]
        for i in range(len(out) - 1, -1, -1):
            x = out[i]
            if x.get("tier") == victim_tier and (x.get("service") or "").lower() not in keep_l:
                out.pop(i)
                return True
        return False

    for term, vol in missing:
        if len(out) >= max_services and not _drop_one():
            break                          # everything is pinned; nothing to give
        rank = [t for t, _ in ranked].index(term)
        tier = "ultra" if rank < int(CFG.get("pin_as_ultra", 2)) else "competitive"
        out.insert(min(rank, len(out)), {"service": term, "tier": tier, "pinned": True})
    return out, [t for t, _ in missing]


# Term INTENT differs by goal even when the product list is identical. A ski
# retailer chasing footfall wants "ski shop", "ski rental near me", "ski shop
# open sunday"; the same retailer chasing ecommerce wants "buy ski jackets
# online", "ski jackets free shipping". Same catalogue, different keywords, and
# nothing was telling the model which campaign it was building (2026-08-05).
_GOAL_TERM_RULE = {
    "In-Store Sales": "GOAL IS IN-STORE SALES. Favour terms with shopping-trip "
        "intent — shop/store/rental/hire forms, and the categories someone "
        "drives to a shop for. Avoid pure ecommerce phrasing (\"free shipping\", "
        "\"online only\").",
    "Online Sales": "GOAL IS ONLINE SALES. Favour transactional ecommerce "
        "phrasing — buy/order/online/shipping forms and specific product "
        "categories a basket is built from. Avoid store-visit phrasing.",
    "Phone Calls": "GOAL IS PHONE CALLS. Favour terms someone searches when "
        "they want to speak to a person now — service + urgency forms, "
        "emergency/same-day/consultation variants.",
    "Leads/Form Fillouts": "GOAL IS LEAD FORMS. Favour consideration-stage "
        "service terms — quote/estimate/consultation/pricing variants.",
    "B2B Sales": "GOAL IS B2B SALES. Favour commercial and trade phrasing — "
        "wholesale/supplier/commercial/bulk/for-business variants, not consumer "
        "retail terms.",
    "Information Requests": "GOAL IS INFORMATION REQUESTS. Favour service and "
        "offering terms that lead to an enquiry, still never questions.",
    "Branding & Awareness": "GOAL IS BRANDING AND AWARENESS. Favour the broad "
        "category head terms the client wants to be known for, over narrow "
        "long-tail variants.",
    "Website Traffic": "GOAL IS WEBSITE TRAFFIC. Favour the highest-demand "
        "category terms the client can credibly rank for.",
    "Event Attendance": "GOAL IS EVENT ATTENDANCE. Favour event, tickets, "
        "schedule and things-to-do phrasing tied to the client's offering.",
    "Job Recruitment": "GOAL IS JOB RECRUITMENT. Favour jobs/careers/hiring "
        "phrasing for the roles this employer fills.",
    "Mobile App Download": "GOAL IS APP DOWNLOADS. Favour app, download and "
        "platform phrasing alongside the core category terms.",
}


def claude_expand_services(seeds, business_desc, site_pages, brand, domain,
                           candidates, max_services, n_cities=1, national=False,
                           goal=""):
    _goal_rule = _GOAL_TERM_RULE.get((goal or "").strip(),
                                     "No campaign goal supplied \u2014 pick the "
                                     "terms that best describe what the business "
                                     "sells, without assuming a channel.")
    """Expand the partner's seed terms into the SERVICE list a proposal would
    target, assigning a competitiveness TIER to each service (not to each
    keyword). This mirrors how the real proposals are built: 'auto insurance' is
    Ultra-Competitive in every city, 'umbrella insurance' is Long Tail in every
    city. Returns [{"service":..., "tier": "ultra"|"competitive"|"long_tail"}]
    or None on failure (caller falls back to the seeds themselves)."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None
    pages = [p for p in (site_pages or [])][:40]
    cands = [c.get("keyword", c) if isinstance(c, dict) else c for c in (candidates or [])][:80]
    prompt = f"""You are an SEO strategist choosing which SERVICES a local business should target in a proposal.

BUSINESS: {business_desc or "(infer from the vertical, website and pages below)"}
SEED TERMS FROM THE PARTNER: {", ".join(seeds)}
WEBSITE: {domain or "(none)"}
BRAND (never include this in a service): {brand or "(none)"}

THEIR ACTUAL WEBSITE PAGES (their real service taxonomy):
{json.dumps(pages, ensure_ascii=False) if pages else "(none available)"}

KEYWORDS THE SEARCH API RETURNED FOR THIS BUSINESS (evidence of real demand):
{json.dumps(cands, ensure_ascii=False)}

TASK: choose exactly {max_services} SERVICES this business should target, and assign each a competitiveness tier.

RULES:
1. A SERVICE is a short, generic phrase with NO city and NO brand — e.g. "auto insurance", "home insurance", "insurance agency", "umbrella insurance". {"This is a NATIONAL product brand: these terms are the final keyword list and will NOT be crossed with cities. Qualify the long-tail entries by AUDIENCE or USE CASE instead of location (e.g. \'electrolyte gummies for athletes\', \'energy gummies for teen athletes\'), never by place." if national else "It will be crossed with city names later, so do NOT include any location."}
2. Only services this business actually offers. Exclude anything they don't do.
2g. {_goal_rule}
2p. NEVER include "near me", "nearby", "closest" or any other RELATIVE proximity phrase. Every service is
   crossed with a city later, and "mattress store near me acworth ga" is not a phrase any human
   types — "near me" IS the location. Write the bare service; the grid adds the place.
2n. A NAMED LANDMARK IS DIFFERENT, and is allowed. "homes for rent near meow wolf", "apartments near
   the university", "storage near the airport" — a specific place a local would name is real local
   phrasing and reads as knowledge of the market, which is why Brendan puts two of them in a
   twenty-term list. Only use a landmark that appears in the keyword-idea list below; a landmark
   nobody searches is worse than a plain term, and inventing one is obvious to the reader.
2s. THE PARTNER'S SEED TERMS COME FIRST. They were typed by someone who knows the account, so treat
   them as the client's own service list. Any seed that is already a clean, bare service belongs in the
   output essentially as written. Only when you have fewer seeds than {max_services} should you add
   services of your own — and then from the business description and site pages before the keyword-idea
   list, which is ranked by national volume and is full of competitor names and out-of-area terms.
   Never spend a slot on an invented term while a usable seed goes unused.
2a. NEVER include a COMPETITOR'S company name as a service. The keyword-idea list WILL contain them
   (a search for "commercial construction company" surfaces the big national contractors) and they look
   like plausible services because they end in a trade word. They are not. Someone searching a specific
   firm's name wants that firm — the intent is navigational, the client will not outrank the brand for
   its own name, and because those terms come back "not ranking" every time they also inflate the quoted
   price on a client who actually ranks well.
2b. NARROW EXCEPTION: a brand name is allowed ONLY when it is a PHYSICAL PRODUCT LINE the client
   resells or installs — "butler building contractor", "trane furnace installation", "andersen window
   replacement". The client stocks or fits that manufacturer's goods, so customers really do search the
   brand plus the service.
   This exception does NOT cover a firm that PROVIDES THE SAME SERVICE as the client. An insurance
   agency does not "carry" State Farm; a builder does not "carry" Turner; an agency does not "carry" a
   rival agency. If the name belongs to a company a customer could hire INSTEAD of the client, it is a
   competitor — exclude it, no matter how much search volume it has. The keyword-idea list is ranked by
   volume, so the largest national competitor in the trade will almost always appear near the top and
   will look tempting. It is still wrong.
   When you cannot tell which kind of name it is, leave it out.
2h. WHEN THE BUSINESS SELLS ONE THING, THE HEAD TERM'S SYNONYMS ARE SEPARATE SERVICES.
   A rental community is not a dental practice with eight departments: there is one thing for
   sale, so the variety has to come from the words customers use for it. "homes for rent",
   "houses for rent", "apartments for rent" and "rentals" are four different keywords with four
   different volumes and four different result pages — a proposal buys two or three of them, not
   one. Lead the list with the BARE, UNQUALIFIED head term in each wording; those are the money
   terms, and a list made entirely of qualified variants has bought none of them.
   This is also the rule that decides the ultra tier for such a business: bare head terms are
   ultra, one-qualifier forms are competitive, and narrow forms — a single amenity, a micro-segment
   — are long tail.
   Rule 2b's cap on variants-per-family does NOT apply here. It exists to stop thirteen implant
   variants crowding out a general practice's other departments; a business with no other
   departments has nothing to crowd out.
2b. BALANCE ACROSS SERVICE LINES — this is the rule that most often gets missed.
   Cover the business's WHOLE service range the way their own website menu does:
   no more than 2-3 variants of any one service family unless the business
   description explicitly says that family is the focus. A general dental
   practice gets family dentistry, cleanings, crowns, invisalign, veneers,
   emergency — NOT thirteen implant variants because one seed said "implants".
   Bread-and-butter services beat exotic variants: they carry the demand and
   the client's existing rankings.
3. Spread across tiers so the proposal has all three. Aim for roughly:
   - 2 "ultra"        (the biggest, most competitive money terms)
   - 1 "competitive"  (solid mid-competition terms)
   - 1 "long_tail"    (a genuine but lower-competition service, e.g. a niche product line)
   Adjust the mix if {max_services} differs, but ALWAYS include at least one long_tail and at least one ultra —
   even when {max_services} is small (a 6-service list still needs a long_tail; the proposal shows three columns
   and an empty one reads as an incomplete strategy).
4. long_tail means a LOWER-COMPETITION SERVICE — never a question. Do NOT produce phrases starting with how/what/why/when/where.
5. Prefer the phrasing a customer would actually search.
6. TIER GUIDANCE — how these tiers are actually assigned in practice (insurance example):
   - ultra: the core high-demand money terms — "auto insurance", "car insurance", "home insurance", "insurance quotes"
   - competitive: solid mid-demand services — "homeowners insurance", "renters insurance", "insurance agency", "insurance company"
   - long_tail: niche or compound product lines with genuinely lower demand — "umbrella insurance", "home and auto insurance"
   Note that a mainstream service like "renters insurance" is COMPETITIVE, not long tail. Reserve long_tail for genuinely niche lines.
   LONG-TAIL PHRASING: prefer COMPOUND or QUALIFIED service phrases over bare two-word niches, so the long-tail tier reads as
   genuinely longer than the head terms. Good: "home and auto insurance", "commercial umbrella insurance", "business auto insurance",
   "classic car insurance". Weaker (still valid, but use sparingly): "umbrella insurance", "boat insurance".
   Aim for at least one multi-word compound in the long_tail tier. These must still be real services the business offers —
   never invent a service, and never turn it into a question.
7. VARIETY: these will be crossed with {n_cities} cit{"y" if n_cities == 1 else "ies"}, so you must supply {max_services} DISTINCT services.
   {"Because there are few or no cities to cross against, the variety has to come from the services themselves. Include close variants and qualified forms the way a real proposal does — but ONLY WITH QUALIFIERS THE KEYWORD-IDEA LIST ABOVE ACTUALLY USES. That list is what people really type in this market, and it is the difference between a real long-tail term and an invented one: '3 bedroom homes for rent' and 'pet friendly rentals' are in it, 'rental homes with washer dryer' and 'move in ready rentals' are phrases nobody searches that read as padding. If you cannot find a qualifier in that list, use a barer form of the head term rather than making one up — e.g. for a supplement brand: 'energy gummies', 'electrolyte gummies', 'hydration gummies', 'energy gummies for athletes', 'electrolyte gummies for kids sports', 'best energy gummies'. For a clinic: 'adhd treatment', 'anxiety treatment', 'depression counseling', 'couples therapy', 'family therapy', 'mental health clinic', 'behavioral health services'. Synonyms, sub-services, audience qualifiers and 'best X' forms all count as distinct services." if n_cities <= 2 else "With several cities to cross against, keep the services broad and distinct rather than near-duplicates."}

Return ONLY valid JSON, no prose:
{{"services": [{{"service": "auto insurance", "tier": "ultra"}}, {{"service": "umbrella insurance", "tier": "long_tail"}}]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 1000, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=30)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", []) if b.get("type") == "text").strip()
        if text.startswith("```"):
            text = text.split("```", 2)[1]
            if text.startswith("json"):
                text = text[4:]
        parsed = json.loads(text.strip())
        out = []
        for s in parsed.get("services", []):
            svc = (s.get("service") or "").strip().lower()
            tier = (s.get("tier") or "competitive").strip().lower()
            if tier not in ("ultra", "competitive", "long_tail"):
                tier = "competitive"
            if svc:
                out.append({"service": svc, "tier": tier})
        return out[:max_services] or None
    except Exception:
        return None


def choose_grid_axis(city_scores, n_seeds, forced=""):
    """Spend the term budget on SERVICES or on GEOGRAPHY.

    A 32-term list is a fixed budget and there are two ways to spend it: a few
    services across many cities, or many services in one city. The tool always
    chose geography. Brendan chooses whichever axis the opportunity is on, and
    the two Junk Bee Gone lists show it plainly (2026-08-10):

      tool  7 services x 5 cities  — junk removal / roll off dumpster / junk /
            remove junk / rent a dumpster ... x Knoxville, Morristown, Oak Ridge,
            Sevierville, Tellico Village
      BE    ~17 services x 1 city — junk removal, hauling, commercial,
            residential, hoarding cleanup, construction debris, dumpster rental,
            roll off, furniture, appliance, pickup, same-day, house cleanout,
            estate cleanout, demolition, shed demolition, paper shredding

    BE covered demolition, hoarding cleanup and paper shredding — three service
    lines the client sells and the tool's list never mentioned — precisely
    because he was not paying for four extra cities. And the cities he skipped
    had nothing to buy: Knoxville measured 170/mo for the lead service while
    every other market sat at Google's 10/mo floor.

    That is the test. Geography only earns the budget when more than one market
    has measurable demand. Otherwise the crossing buys near-duplicates of one
    city and the service breadth is what is missing.

    On NASSCO he did the reverse — many jurisdictions, few services — because
    there the ordinance towns WERE the opportunity. Same rule, opposite answer.

    Returns (axis, reason, evidence).
    """
    floor = int(CFG.get("axis_city_volume_floor", 20))
    min_seeds = int(CFG.get("axis_min_seeds_for_services", 8))
    scored = [(c, int(v or 0)) for c, v in (city_scores or [])]
    real = [c for c, v in scored if v >= floor]
    ev = {"cities_with_demand": len(real), "cities_scored": len(scored),
          "floor": floor, "seeds": int(n_seeds or 0),
          "top": scored[:5]}
    if forced in ("services", "geography"):
        # WHAT THE MEASUREMENT SAID, EVEN WHEN OVERRIDDEN. The panel used to
        # print the chosen axis beside the raw evidence, which on a hand-forced
        # build read "Budget spent on geography - 1 of 5 markets carry demand"
        # - the evidence for the OPPOSITE choice, presented as its
        # justification. Recompute the unforced verdict and carry it. (2026-08-11)
        m_axis, m_reason, _m_ev = choose_grid_axis(city_scores, n_seeds, "")
        ev["by_hand"] = True
        ev["measured_axis"] = m_axis
        ev["measured_reason"] = m_reason
        if m_axis == forced:
            return forced, f"set by hand, and the measurement agrees: {m_reason}", ev
        return forced, f"set by hand, overriding the measurement: {m_reason}", ev
    if len(scored) <= 1:
        return "services", "only one market, so the budget can only buy services", ev
    if len(real) <= 1 and (n_seeds or 0) >= min_seeds:
        return ("services",
                (f"only {len(real)} of {len(scored)} markets carry demand above "
                 f"{floor}/mo, and you supplied {n_seeds} focus terms — the budget "
                 "buys more by covering services in the market that has demand "
                 "than by copying it across markets that have none"), ev)
    if len(real) <= 1:
        return ("geography",
                (f"only {len(real)} of {len(scored)} markets carry measurable "
                 f"demand, but there are too few focus terms ({n_seeds}) to fill a "
                 "service-led list — add more services to switch axis"), ev)
    return ("geography",
            f"{len(real)} of {len(scored)} markets carry demand above {floor}/mo, "
            "so crossing them is buying real reach", ev)


def services_needed(n_cities):
    """How many services to generate so services x cities lands near the target
    keyword count. Few cities -> many services (a one-metro client needs service
    variety); many cities -> fewer services (the crossing supplies the volume)."""
    import math
    target = CFG.get("grid_target_keywords", 32)
    lo, hi = CFG.get("grid_min_services", 4), CFG.get("grid_max_services", 20)
    n = max(int(n_cities), 1)
    return max(lo, min(hi, math.ceil(target / n)))



# --- ranking the partner's own seed list ------------------------------------
# enforce_seed_services() fills the grid from clean[:max_services] IN ENTRY
# ORDER, so a 69-term focus list against a 20-service grid is decided by typing
# order and nothing else — and the first two seeds typed become the "ultra
# competitive" tier whatever their demand. Junk Bee Gone's list spent eight of
# its twenty slots on synonyms of one service (haul away / haul away junk / haul
# away service / junk haulers / hauling services / junk remover / remove junk /
# junk) while demolition, hoarding cleanup and paper shredding — all on the
# client's own site — never reached the grid.
#
# Two problems, two fixes: FOLD the near-duplicates so one service line cannot
# take eight slots, then RANK what survives by measured demand. Deliberately NOT
# wired into the build: reordering seeds moves tier assignment and therefore
# price, so this runs only when the operator asks and shows every fold and drop
# before anything changes. (2026-08-11)

# Words that describe the SHAPE of a service phrase, not its subject. Dropped
# before folding, so "haul away" and "haul away service" collapse together.
_SEED_SHAPE = frozenset("""
mens men womens women kids kid child children childrens boys girls toddler baby
adult adults junior juniors youth unisex family
service services company companies contractor contractors business businesses
pro pros professional professionals expert experts specialist specialists
near me nearby local best top cheap affordable cost price prices pricing quote
quotes estimate estimates free same day emergency 24 7 hour hours a an the and
or for of in my your our
""".split())


# COMPENSATION IS COMP. The stemmer takes "workers compensation" to `compens`
# and "workers comp" to `comp`, so PEO Brokers carried four near-duplicate pairs
# into a twenty-slot grid — swif workers compensation / swif workers comp,
# workers comp insurance for staffing agencies / workers compensation for
# staffing agencies, new york state fund workers comp / ...compensation. Four
# slots, and each pair splits its own rankings in the grid. (2026-08-13)
_SEED_SYN = {"compens": "comp"}

# Words that QUALIFY the subject without changing it. Dropped from the key only
# when something else survives, so "insurance" on its own is still a key — this
# folds "usl&h insurance" into "usl&h", not every insurance term into one.
_SEED_QUALIFY = frozenset("insurance insurances coverage coverages".split())


# IRREGULAR PLURALS THE SUFFIX RULES CANNOT REACH. Nob Hill Dental came back with
# "tooth overlay" AND "overlays for teeth" in the same thirteen-slot grid — one
# procedure, two slots — because no amount of suffix-stripping turns "teeth" into
# "tooth". Six of that client's thirteen terms were overlay variants while
# Brendan's list for the same practice covers cleanings, crowns, whitening, root
# canals, extractions and dentures. A short closed list, not a lexicon: these are
# the ones that actually show up in service names. (2026-08-17)
_IRREGULAR = {
    "teeth": "tooth", "feet": "foot", "children": "child", "men": "man",
    "women": "woman", "people": "person", "mice": "mouse", "geese": "goose",
    "lice": "louse", "oxen": "ox", "leaves": "leaf", "knives": "knife",
    "wives": "wife", "lives": "life", "shelves": "shelf", "halves": "half",
    "wolves": "wolf", "loaves": "loaf", "thieves": "thief", "calves": "calf",
    "hooves": "hoof", "roofs": "roof",
}


def _seed_stem(word):
    """Crude, deterministic stemmer. A real one is not worth the dependency:
    the only job is making 'hauling', 'haulers' and 'haul' the same token."""
    w = re.sub(r"[^a-z0-9]+", "", (word or "").lower())
    if w in _IRREGULAR:
        w = _IRREGULAR[w]
    # Trailing "e" is last on purpose: without it "remove" and "removal" stem to
    # different tokens and "remove junk" survives as a second slot for the same
    # service as "junk removal".
    for suf in ("ations", "ation", "ings", "ing", "ers", "er", "ors", "or",
                "als", "al", "ies", "es", "s", "e"):
        # One-letter suffixes get a shorter guard so three-letter plurals stem:
        # without it "tvs" and "tv" are different tokens.
        floor = 2 if len(suf) == 1 else len(suf) + 2
        if len(w) > floor and w.endswith(suf):
            return w[: -len(suf)]
    return w


# THE VERB IS NOT THE SERVICE, THE OBJECT IS. "yard waste removal" and "yard
# waste disposal" both took a slot on Junk Bee Gone, as did "furniture removal"
# and "furniture pick up", and "mattress removal" was kept while "mattress
# disposal" fell below the cut — the same service bought twice, or split across
# the cut so neither read as important. Canonicalising the action verbs merges
# them; it can only ever merge phrases that already share their object noun, so
# it stays safe across verticals ("tooth extraction" / "tooth removal" are also
# one service). (2026-08-11)
_SEED_ACTION = {
    # get-rid-of family
    "remov": "remov", "remo": "remov", "dispos": "remov", "disposit": "remov",
    "pickup": "remov", "pick": "remov", "haul": "remov", "hauling": "remov",
    "junk": "junk",          # 'junk' is an object here, never an action
    "takeaway": "remov", "discard": "remov", "dump": "remov", "toss": "remov",
    "recycl": "remov", "scrap": "remov",
    "extract": "remov", "extraction": "remov",
    # clear-the-space family
    "cleanout": "clean", "cleanup": "clean", "clean": "clean", "clear": "clean",
    "clearanc": "clean", "clearout": "clean", "declutt": "clean",
    # equipment-hire family
    "rental": "rent", "rent": "rent", "leas": "rent", "hir": "rent",
}


def acronym_aliases(terms):
    """{acronym: [words it stands for]} — worked out from the list itself.

    "new york state fund workers comp" and "workers compensation nysif" are the
    same fund and share no token, so every fold in the build looks straight past
    them and PEO Brokers spends two of six Ultra slots on one thing. Nothing here
    is a lookup table: if some term in the list contains consecutive words whose
    initials spell a single word in another term, that is the expansion.

    Deliberately strict — four letters or more, three consecutive words or more,
    exact initials. "peo" is three letters and would otherwise swallow anything
    starting p-e-o; "usl&h" flattens to "uslh" and matches nothing. False
    positives here merge two real services, so the bar is high. (2026-08-14)
    """
    words, singles = [], set()
    for t in terms or []:
        ws = [w for w in re.sub(r"[^a-z0-9 ]+", " ", str(t).lower()).split() if w]
        words.append(ws)
        for w in ws:
            if len(w) >= 4 and w.isalpha():
                singles.add(w)
    out = {}
    for acr in singles:
        n = len(acr)
        for ws in words:
            if acr in ws:
                continue                      # its own term proves nothing
            for i in range(len(ws) - n + 1):
                run = ws[i:i + n]
                if len(run) >= 3 and "".join(w[0] for w in run) == acr:
                    out[acr] = run
                    break
            if acr in out:
                break
    return out


def _seed_key(term, alias=None):
    """Stemmed token set with shape words removed and action verbs canonicalised
    — the near-duplicate key. `alias` expands acronyms first (see
    acronym_aliases), so "nysif" and "new york state insurance fund" key alike."""
    toks, quals = set(), set()
    _words = []
    for w in (term or "").lower().split():
        # An acronym stands in for the words it abbreviates, so both spellings
        # reach the same key. Everything downstream — the fold, the grid dedupe,
        # the topic guarantee — inherits this for free.
        _words.extend((alias or {}).get(w, [w]))
    for w in _words:
        if w in _SEED_SHAPE:
            continue
        st = _seed_stem(w)
        st = _SEED_ACTION.get(st, st)
        st = _SEED_SYN.get(st, st)
        if not st or st in _SEED_SHAPE:
            continue
        if w in _SEED_QUALIFY:
            quals.add(st)
        else:
            toks.add(st)
    # Qualifiers only count when they are all there is.
    return frozenset(toks or quals)


def seed_norm(term, markets=None, state=""):
    """One definition of "the same seed", shared by rank_seeds() and the build.

    Both need to answer "have I already got this term?" and they have to answer
    it identically, or reordering the list in the build silently duplicates
    entries whose raw text differs only in case or a stripped market name.
    """
    return clean_kw(strip_placeholders(strip_proximity(
        _strip_markets((term or "").lower(), list(markets or []), state)))).strip()


def cap_service_family(terms, seeds=None, cap=None, markets=None, state=""):
    """No more than a few variants of any one procedure, enforced in code.

    The expansion prompt has carried a rule about this since July — "no more than
    2-3 variants of any one service family... NOT thirteen implant variants
    because one seed said implants" — and a prompt cannot be checked. Nob Hill
    Dental came back with nine: onlay overlay, overlay onlay, crown overlay,
    overlay crown, overlays on teeth, overlays for teeth, tooth overlay, dental
    onlays and overlays, composite tooth restoration. Six of them reached a
    thirteen-slot grid. Brendan's list for the same practice covers cleanings,
    crowns, whitening, root canals, extractions, dentures and emergency care.

    The damage compounds: those nine became SEEDS, which made overlays half the
    seed list, which earned them half the grid through the topic quota — working
    exactly as designed, on an input that was already lopsided.

    ONLY WHAT THE TOOL PROPOSED. A family cap is lethal in the wrong place: a
    rental community's list is ALL "homes for rent" variants and every one of
    them is the business. The operator's own terms are what say which is which,
    so they are exempt and they are not counted towards the cap — thirteen typed
    homes-for-rent terms can never make "home" look over-represented.

    Groups on the shared significant token, keeps the earliest few of each group
    (the ranking that produced them already put the best first), and returns
    (kept, dropped) with dropped as (term, token). (2026-08-17)
    """
    cap = int(cap if cap is not None else CFG.get("service_family_cap", 3))
    typed = {seed_norm(str(t), markets, state).lower() for t in (seeds or [])}
    typed |= {str(t).strip().lower() for t in (seeds or [])}
    if cap <= 0:
        return list(terms or []), []

    def toks(t):
        return {w for w in _seed_key(seed_norm(str(t), markets, state))
                if w and len(w) > 2}

    # WHICH FAMILY IS THE BUSINESS, AND WHICH RAN AWAY. The operator's own list
    # is what says so. Amare's typed terms are "apartment for rent", "home for
    # rent", "single family homes for rent" — "home" and "rent" are in most of
    # them, so a suggestion pool full of homes-for-rent variants is the client's
    # whole offer and capping it would gut the quote. Nob Hill typed five terms
    # and exactly one mentions overlays, so nine overlay suggestions are a
    # runaway. A token carried by this many of the TYPED seeds is core and is
    # never capped. (2026-08-17)
    core_n = int(CFG.get("family_core_seeds", 2))
    typed_freq = {}
    for t in (seeds or []):
        for w in toks(t):
            typed_freq[w] = typed_freq.get(w, 0) + 1
    core = {w for w, n in typed_freq.items() if n >= core_n}

    # A CORE TOKEN EXEMPTS THE WHOLE TERM, and keeps it out of everyone else's
    # count. Exempting only the core token itself sent the term to its NEXT most
    # crowded token instead, and that token is usually a category word rather
    # than a service: with "overlay" core, the five terms carrying "tooth"
    # regrouped onto tooth and the cap dropped `teeth whitening` — a different
    # procedure entirely, thrown out for sharing a body part with overlays.
    # A term the operator's own vocabulary vouches for is the business; it is
    # not capped, and it does not make anything else look crowded either.
    # (2026-08-17)
    def _exempt(t):
        return str(t).strip().lower() in typed or bool(toks(t) & core)

    # Counted over the PROPOSALS only. A token the operator uses constantly is
    # their vocabulary, not a family that has run away with the list.
    freq = {}
    for t in (terms or []):
        if _exempt(t):
            continue
        for w in toks(t):
            freq[w] = freq.get(w, 0) + 1
    crowded = {w for w, n in freq.items() if n > cap and w not in core}
    if not crowded:
        return list(terms or []), []

    seen, kept, dropped = {}, [], []
    for t in (terms or []):
        if _exempt(t):
            kept.append(t)
            continue
        # The most crowded token this term belongs to decides its group, so a
        # term in two runaway families is counted once, against the worse one.
        mine = sorted(toks(t) & crowded, key=lambda w: (-freq[w], w))
        if not mine:
            kept.append(t)
            continue
        w = mine[0]
        seen[w] = seen.get(w, 0) + 1
        if seen[w] <= cap:
            kept.append(t)
        else:
            dropped.append((str(t), w))
    return kept, dropped


def fold_proposals(terms, seeds=None, markets=None, state="", limit=None):
    """Collapse a list of PROPOSED service terms the way rank_seeds folds seeds.

    Ski Barn's menu produced 39 chips: ski jackets, boys ski jackets, girls ski
    jackets, ski pants, boys ski pants, girls ski pants, ski socks, ski hats and
    accessories, kids ski gear, toddler ski clothing... A grid holding 20 services
    cannot use 39 near-duplicates, and "Add all 39" would have pushed the seed
    list past 50 for the build to then rank back down again. Same containment rule
    as the seed folder, so the two agree on what "the same service" means.

    Keeps the FIRST form of each group rather than the longest: these arrive in
    the client's own menu order, and "ski jackets" is the label they lead with.
    Anything already covered by an existing seed drops out entirely — proposing a
    term the operator has already got is not a suggestion.

    Returns (kept, folded_away) preserving input order.
    """
    have = set()
    for sd in seeds or []:
        k = _seed_key(seed_norm(sd, markets, state))
        if k:
            have.add(k)
    kept, folded, groups = [], [], []
    for t in terms or []:
        name = seed_norm(t if isinstance(t, str) else (t or {}).get("term", ""),
                         markets, state)
        k = _seed_key(name) or frozenset({name})
        if not name:
            continue
        # A HEAD TERM IS NOT A DUPLICATE OF ITS OWN QUALIFIER. "homes for rent"
        # keys to {hom, rent}, a subset of "single family homes for rent"
        # {hom, rent, singl}, so containment folded the BROAD term away as
        # already-covered — and the broad term is the one with the volume: 40/mo
        # against 10 for every qualified variant, straight into Ultra when the
        # operator finally typed it by hand. Brendan's Amare list leads with it.
        #
        # Kept only when it carries two or more meaningful tokens. That is the
        # existing bare-head guard, reused: "junk" ({junk}, one token) is still
        # folded into "junk removal", which is the case this fold was written for.
        # (2026-08-13)
        _broader = any(k < h for h in have)
        if _broader and len(k) >= 2:
            pass
        elif any(k == h or k <= h or h <= k for h in have):
            folded.append(t)
            continue
        hit = next((g for g in groups if k == g or k <= g or g <= k), None)
        if hit is not None:
            folded.append(t)
            continue
        groups.append(k)
        kept.append(t)
        if limit and len(kept) >= int(limit):
            break
    return kept, folded


# A company name wearing its own suffix. Only the LAST token counts, so "llc
# formation services" — a real thing a lawyer sells — is untouched.
_CORP_SUFFIX = frozenset("""llc l.l.c. inc inc. incorporated corp corp. corporation
ltd ltd. limited llp lllp pllc lp gmbh""".split())


def _name_tokens(text):
    """Significant stemmed tokens of a name, shape words removed."""
    raw = re.sub(r"[^a-z0-9 ]+", " ", (text or "").lower())
    out = set()
    for w in raw.split():
        if w in _SEED_SHAPE:
            continue
        st = _seed_stem(w)
        if st and len(st) > 1 and st not in _SEED_SHAPE:
            out.add(st)
    return out


def is_brand_term(term, brand):
    """Is this the client's OWN name?

    PEO Brokers' focus list came back holding "peo insurance brokers network" —
    their own company — proposed by the ranked-keywords pass, which of course
    found it: they rank #1 for their name and always will. There is no work to
    sell against a term you already own, and it takes a slot from one you don't.
    Needs two significant words in the brand, so a one-word brand that doubles as
    a service ("Amare", "Prime") cannot swallow the list. (2026-08-13)
    """
    bt = _name_tokens(brand)
    if len(bt) < 2:
        return False
    tt = _name_tokens(term)

    # PREFIX, NOT EQUALITY. The stemmer is crude by design — it strips a short
    # list of suffixes — so "dental" reduces to `dent` and "dentistry" reduces to
    # nothing at all. That made {dent, hill, nob} not a subset of
    # {dentistry, hill, nob}, and Nob Hill Dental was quoted to rank for "nob hill
    # dentistry" at 140/mo: their own name, the biggest number in their grid, and
    # work nobody can sell because they already own it.
    #
    # A brand token matches a term token when either is a prefix of the other and
    # the shorter is at least four characters. That covers dental/dentistry/
    # dentist, plumb/plumbing/plumber and the rest of the trade-word family
    # without a lexicon. EVERY brand token must still match, so "dentures" — one
    # loose prefix hit and nothing else — is untouched. (2026-08-17)
    def _hit(b):
        for t in tt:
            if b == t:
                return True
            if len(b) >= 4 and t.startswith(b):
                return True
            if len(t) >= 4 and b.startswith(t):
                return True
        return False

    # RUN TOGETHER, THE WAY DOMAINS AND HANDLES WRITE IT. Red Shoes Inc came back
    # quoted on "redshoe appleton wi" — their own name, one word, because the
    # per-token matcher compares {red, shoes} against {redshoe, appleton, wi} and
    # neither brand token is a prefix of "redshoe": "red" is three characters and
    # falls under the four-character floor, and "shoes" starts a different
    # letter. Same failure as "nob hill dentistry", arriving by spelling instead
    # of by suffix.
    #
    # The squashed form is how people actually type a brand they know from a URL
    # — redshoesinc.com, mpgxtreme — so it is checked as one token against each
    # token of the term, with the same prefix rule. Two or more brand words are
    # required: squashing a one-word brand just restates it, and a one-word
    # brand has never been allowed to swallow a list. (2026-08-18)
    # AND THE TERM TOKEN HAS TO BE LONGER THAN ANY SINGLE BRAND WORD, or the
    # squashed form matches one of its own components: "gummies" is a prefix of
    # "gummiesmpg", so a first pass at this read "energy gummies" as MPG Gummies'
    # own name and would have deleted the biggest term in their grid. A run-
    # together brand is by definition longer than the words it is made of.
    # Built from the RAW brand, not from bt: _name_tokens returns a SET (so
    # "".join() on it is in arbitrary order) and it STEMS, so "Nob Hill Dental"
    # arrives as {nob, hill, dent} and could never reproduce "nobhilldental".
    _raw = [w for w in re.split(r"[^a-z0-9]+", (brand or "").lower()) if w]
    if len(_raw) >= 2:
        _squash = "".join(_raw)
        _longest = max(len(w) for w in _raw)
        for t in tt:
            if len(t) <= _longest:
                continue
            if _squash.startswith(t) or t.startswith(_squash):
                return True

    return all(_hit(b) for b in bt)


def drop_suggested_nonservices(seeds, suggested, kinds, brand,
                               markets=None, state="", ranked=None, rivals=None):
    """Terms the TOOL proposed do not get the operator's exemption.

    "You typed it, it wins" exempts a focus term from every filter in the build.
    It was written to protect the judgement of someone who knows the account —
    and it was covering the tool's own guesses, because nothing recorded which
    terms those were. Now something does.

    Only touches terms in `suggested`. Anything typed is returned untouched, and
    the list is never emptied. Returns (keep, dropped).
    """
    sug = {str(x).strip().lower() for x in (suggested or []) if str(x).strip()}
    _rank = {str(x).strip().lower() for x in (ranked or []) if str(x).strip()}
    _rival = {str(x).strip().lower() for x in (rivals or []) if str(x).strip()}
    if not sug:
        return list(seeds or []), []
    kinds = kinds or {}
    keep, dropped = [], []
    for t in seeds or []:
        raw = str(t or "").strip().lower()
        norm = seed_norm(t, markets, state)
        if raw not in sug and norm not in sug:
            keep.append(t)
            continue
        k = kinds.get(raw) or kinds.get(norm) or {}
        why = ""
        # A POSITION IS PROOF OF OWNERSHIP, and it beats a name-shaped guess.
        # PEO Brokers lost five SWIF terms in one build: SWIF is Pennsylvania's
        # State Workers' Insurance Fund, a program this broker places business
        # into, and the classifier read the capitalised name and called it another
        # business. The evidence against that was already in the same build —
        # those terms surfaced from the ranked-keywords pass, which means the
        # client RANKS for them. The two deterministic rules below still apply:
        # a corporate suffix and the client's own brand name do not guess.
        # (2026-08-13)
        if k.get("kind") == "other_business" and not (
                (raw in _rank or norm in _rank) and raw not in _rival):
            why = k.get("why") or "another company, not a service"
        # RANKING IS NOT AN EXEMPTION HERE — IT IS THE SYMPTOM. A reference term
        # is one the client's own site ranks for BECAUSE it publishes about it,
        # so "they rank for it" is exactly how it got proposed and cannot also be
        # what saves it. (2026-08-16)
        elif k.get("kind") == "reference":
            why = k.get("why") or "something they publish about, not something they sell"
        elif raw.split() and raw.split()[-1] in _CORP_SUFFIX:
            why = "a company name"
        elif is_brand_term(t, brand):
            why = "their own name — they already rank for it"
        if why:
            dropped.append({"term": t, "why": why})
        else:
            keep.append(t)
    if not keep:
        return list(seeds or []), []
    return keep, dropped


def _neg_tokens(text):
    """Stemmed tokens of a term, punctuation flattened, order kept.

    One extra fold on top of the shared stemmer, local to this matcher: a
    trailing 'y' is dropped, because _seed_stem takes "policies" to `polic` and
    "policy" to `policy` and a negative has to catch both spellings of its own
    word. Kept out of _seed_stem itself so the duplicate fold, which every grid
    slot depends on, is not moved by a change made for this.
    """
    raw = re.sub(r"[^a-z0-9]+", " ", (text or "").lower())
    out = []
    for w in raw.split():
        # "policies" -> "policy" BEFORE stemming. Folding the trailing y off the
        # stem instead was the obvious fix and it collided "policy" with
        # "policing", which is the exact class of over-match this whole function
        # exists to avoid. Undoing the plural first keeps them apart:
        # policies -> policy -> policy, policing -> polic.
        if len(w) > 4 and w.endswith("ies"):
            w = w[:-3] + "y"
        out.append(_seed_stem(w))
    return out


def negative_hit(term, negatives):
    """Which negative this term trips, or ''.

    WHOLE WORDS, NEVER FRAGMENTS. A substring test is the obvious implementation
    and it is wrong in a way that is hard to see afterwards: negating "carrier"
    would also take out "career", and "auto" would take "automotive" and
    "automation" with it. Everything is compared on the same stems the duplicate
    fold uses, so "policy" catches "policies" without catching "policing".

    A multi-word negative matches as a PHRASE — consecutive tokens, in order —
    so "workers comp" does not fire on every term containing "workers".
    (2026-08-13)
    """
    toks = _neg_tokens(term)
    if not toks:
        return ""
    for neg in negatives or []:
        nt = _neg_tokens(neg)
        if not nt:
            continue
        n = len(nt)
        if any(toks[i:i + n] == nt for i in range(len(toks) - n + 1)):
            return str(neg).strip()
    return ""


def drop_negative_services(services, negatives, seeds=None):
    """Remove proposed services that trip a negative. Returns (kept, dropped).

    Deterministic and in code, not in a prompt — the same class of rule as the
    out-of-area filter. A model can be talked out of an instruction; this cannot.

    A FOCUS TERM IS EXEMPT. It reaches this list through enforce_seed_services
    like anything else, so without `seeds` the filter deletes the operator's own
    term and the panel reports it as a conflict that was "kept" — which it then
    would not be. Both the term and the negative are the operator speaking; the
    term is the more specific instruction, so it stays and the disagreement is
    reported instead. Never empties the list either: if every service trips a
    negative, the negatives are what is wrong, not the build. (2026-08-13)
    """
    if not negatives:
        return list(services or []), []
    mine = {str(t).strip().lower() for t in (seeds or []) if str(t).strip()}
    kept, dropped = [], []
    for x in services or []:
        name = x.get("service", "") if isinstance(x, dict) else str(x)
        if str(name).strip().lower() in mine:
            kept.append(x)
            continue
        hit = negative_hit(name, negatives)
        if hit:
            dropped.append([name, hit])
        else:
            kept.append(x)
    if not kept:
        return list(services or []), []
    return kept, dropped


def negative_seed_conflicts(seeds, negatives):
    """Focus terms that trip a negative. REPORTED, never removed.

    Both of these are the operator speaking, so the tool does not pick a winner
    quietly. The focus term is the more specific instruction and stays in the
    quote; the panel says so, because a planner deleting their own focus term by
    typing a negative is the worse failure. (2026-08-13)
    """
    out = []
    for t in seeds or []:
        hit = negative_hit(t, negatives)
        if hit:
            out.append([str(t), hit])
    return out


def claude_competitor_check(terms, brand="", domain="", industry="", business_desc=""):
    """Of the named companies in this list, which are RIVALS and which are things
    the client sells or places business into?

    The distinction the build needs and the kinds classifier does not draw. It
    answers "is this somebody else's company", which is true of both a competitor
    and a carrier — and those two want opposite treatment. PEO Brokers rank for
    "swif pa" because SWIF is the Pennsylvania state fund they place business
    into; they rank for "prime peo brokers" because a rival is a rival. Ranking
    for a term earned both of them an exemption, and one of them put a
    competitor's name in front of a client.

    Returns {term: {"verdict": "competitor"|"offering"|"neither", "why": str}}.
    Empty on any failure, and every caller treats empty as "no opinion" — a dead
    API must not start quoting rivals OR start deleting the client's own
    products. (2026-08-14)
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    tl = [str(t).strip() for t in (terms or []) if str(t).strip()]
    if not api_key or not tl:
        return {}
    prompt = f"""Each term below names a company, brand, programme or fund. For THIS client,
say whether each one is a COMPETITOR or something they OFFER.

CLIENT: {brand or '(not given)'}
WEBSITE: {domain or '(not given)'}
INDUSTRY: {industry or '(not given)'}
WHAT THEY DO: {(business_desc or '').strip()[:500] or '(not given)'}

TERMS:
{chr(10).join('- ' + t for t in tl[:25])}

Verdicts:
  "competitor" - a firm selling the SAME service to the SAME buyer. This client
                 competes with them for the customer. Never worth quoting: the
                 searcher is looking for that firm by name.
  "offering"   - a third-party product, carrier, fund, programme, platform or
                 network that this client SELLS, PLACES BUSINESS INTO, RESELLS,
                 IS APPOINTED WITH or otherwise works through. Their own
                 customers search for it, and this client can legitimately rank
                 for and win it.
  "neither"    - not a company or programme at all.

Worked example. An insurance broker who places workers' compensation:
  "prime peo brokers"  -> competitor (another broker chasing the same client)
  "swif pa"            -> offering  (the state fund they place business into)
  "the hartford"       -> offering  (a carrier they are appointed with)
  "workers comp peo"   -> neither   (a service, not a company)

Be careful in both directions. Calling an offering a competitor deletes a real
service line. Calling a competitor an offering puts a rival's name in a client
proposal. If you genuinely cannot tell, answer "neither".

Return ONLY JSON, no prose:
{{"verdicts": [{{"term": "swif pa", "verdict": "offering", "why": "PA state fund"}}]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 1200, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=25)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        rows = json.loads(text).get("verdicts") or []
    except Exception:
        app.logger.exception("claude_competitor_check failed")
        return {}
    out = {}
    for r in rows:
        t = str(r.get("term", "")).strip().lower()
        v = str(r.get("verdict", "")).strip().lower()
        if t and v in ("competitor", "offering", "neither"):
            out[t] = {"verdict": v, "why": str(r.get("why", "")).strip()}
    return out


def rival_terms(terms, kinds, brand="", domain="", industry="", business_desc=""):
    """The subset of `terms` that are COMPETITORS. Everything else is spared.

    Only asks about terms the kinds classifier already flagged as another
    business — at most a handful — so this is one small call, not a pass over
    the list.
    """
    ask = [t for t in (terms or [])
           if ((kinds or {}).get(str(t).strip().lower()) or {}).get("kind")
           == "other_business"]
    if not ask:
        return set()
    v = claude_competitor_check(ask, brand, domain, industry, business_desc)
    return {t for t in ask
            if (v.get(str(t).strip().lower()) or {}).get("verdict") == "competitor"}


def demote_nonservices(seeds, kinds, markets=None, state="", ranked=None,
                       rivals=None):
    """Split seeds into what this client actually sells and what belongs to
    somebody else, per claude_seed_kinds().

    VOLUME-FREE ON PURPOSE. This has to run on every build, and the check used to
    live inside rank_seeds() — which needs a DataForSEO call, is capped at twelve
    a minute, and is therefore gated on the ranking mattering. NPAIHB had 8 seeds
    for 20 slots, so the ranking was skipped, so nothing ever asked whether
    "confederated tribes warm springs" is a service NPAIHB sells. It is a member
    tribe the site-heading miner scraped off their own pages. Classifying is an
    Anthropic call and costs nothing against the DataForSEO allowance, so it is
    split out and always runs. (2026-08-12)

    Matches on the raw seed and on its seed_norm() form, because the classifier
    is keyed on what the operator typed and the build works in normalised terms.

    Returns (keep, demoted) with `keep` in the order given. Never demotes
    everything: if the classifier disliked every seed, the classifier is what is
    wrong, not the seed list.
    """
    kinds = kinds or {}
    if not kinds:
        return list(seeds or []), []
    _rank = {str(x).strip().lower() for x in (ranked or []) if str(x).strip()}
    _rival = {str(x).strip().lower() for x in (rivals or []) if str(x).strip()}
    keep, demoted = [], []
    for t in seeds or []:
        raw = str(t or "").strip().lower()
        k = kinds.get(raw)
        if not k:
            n = seed_norm(t, markets, state)
            if n:
                k = kinds.get(n) or kinds.get(n.strip().lower())
        # A POSITION OUTRANKS THE CLASSIFIER HERE TOO. The same exemption went
        # into drop_suggested_nonservices and stopped there, so PEO Brokers' SWIF
        # terms were spared by that filter and then removed by this one, three
        # builds running — silently, because this path's panel line was taken out
        # on request and a demoted term keeps its pill. It survived every filter
        # the operator could see and never reached the grid. (2026-08-14)
        # The exemption covers "other_business" ONLY. Ranking for a term says
        # somebody at this domain owns it, which answers "is this somebody
        # else's company" and answers nothing about "is this a thing rather
        # than a service" — a shop ranks for "old tvs" precisely because it is
        # an object it sells, and that verdict should still stand.
        _kind = (k or {}).get("kind")
        # ...and ranking does not save a RIVAL. Both a competitor and a carrier
        # read as "another business" and both can be ranked for; only one of them
        # is a service line. (2026-08-14)
        _exempt = (_kind == "other_business"
                   and (raw in _rank or seed_norm(t, markets, state) in _rank)
                   and raw not in _rival)
        if _kind in ("item", "other_business", "reference") and not _exempt:
            demoted.append({"term": t, "kind": k.get("kind", "item"),
                            "why": k.get("why", "")})
        else:
            keep.append(t)
    if not keep:
        return list(seeds or []), []
    return keep, demoted


def fold_seed_duplicates(seeds, markets=None, state=""):
    """Collapse near-duplicate seeds. NO volume call — pure string work.

    The fold lived inside rank_seeds(), which needs a DataForSEO lookup, so the
    build only ran it when there were more seeds than slots. At exactly twenty
    seeds for twenty slots the gate is false and nothing folded: Amare quoted
    "home for rent santa fe nm" and "homes for rent santa fe nm" as two of its
    twenty terms. Folding costs nothing; only the RANKING needs the API, so they
    are separated. Same containment rule as everywhere else. (2026-08-13)

    EQUALITY, NOT CONTAINMENT. Containment is what rank_seeds uses, and it works
    there because measured volume picks the survivor. With no volumes it is far
    too blunt: "home for rent" keys to {home, rent}, a subset of every qualified
    variant, so the bare head term swallowed "homes for rent no deposit",
    "month to month home rental" and the rest of the list. Equal keys is the case
    this function exists for — singular/plural and word order — and the
    broader/narrower question stays with the passes that can measure it.

    Keeps the first form the operator typed. Returns (kept, [(dropped, kept_as)]).
    """
    kept, folded, groups = [], [], []      # groups: [(key, term)]
    # Acronyms first, worked out from this list alone, so "nysif" and the words
    # it stands for reach the same key here as everywhere else. (2026-08-14)
    _alias = acronym_aliases([seed_norm(x, markets, state) for x in (seeds or [])])
    for raw in seeds or []:
        t = seed_norm(raw, markets, state)
        if not t:
            continue
        k = _seed_key(t, _alias) or frozenset({t})
        hit = next((g for g in groups if k == g[0]), None)
        if hit is not None:
            folded.append((raw, hit[1]))
            continue
        groups.append((k, raw))
        kept.append(raw)
    return kept, folded


def geo_qualified_volume(terms, markets, state):
    """Volume for the terms AS THEY WILL BE QUOTED — "ski shop wayne nj", not
    "ski shop".

    THE RANKING AND THE GRID WERE MEASURING DIFFERENT STRINGS. Seeds are ordered
    by fetch_local_volume, which sends the BARE phrase with location_name set to
    the city: "how many people in Wayne search *ski shop*" — 2,420/mo, a real
    number and a fair measure of local demand. But the row that reaches the
    proposal is "ski shop wayne nj", which draws 340/mo. Roughly 7x apart, and
    consistently so, which means the slot order was set by numbers that do not
    describe the terms being sold. A term with strong bare local demand and no
    geo-qualified search took a slot and delivered very little.

    National on purpose: the geography is IN the string, so a Wayne-localised
    lookup would be asking how many people in Wayne search for a phrase that
    already says Wayne. One call for the whole candidate list.

    Returns ({term_lower: volume}, error) — {} when there is nothing to ask.
    """
    ts = [clean_kw(str(t).lower()).strip() for t in (terms or [])]
    ts = [t for t in ts if t]
    mk = [m for m in (markets or []) if m and str(m).strip()]
    if not ts or not mk:
        return {}, ""
    forms = geo_form_candidates(mk[0], state)
    sfx = forms[0] if forms else ""
    if not sfx:
        return {}, ""
    probe, back = [], {}
    for t in ts:
        kw = clean_kw(f"{t} {sfx}")
        if kw:
            probe.append(kw)
            back[kw] = t
    probe = dfs_kw_list(probe)
    if not probe:
        return {}, ""
    out = {}
    try:
        data = dfs_post("/keywords_data/google_ads/search_volume/live",
                        [{"keywords": probe[:700], "location_name": "United States",
                          "language_code": "en"}], timeout=25)
        for it in (data["tasks"][0]["result"] or []):
            k = str(it.get("keyword", "")).lower()
            if k in back:
                out[back[k]] = int(it.get("search_volume") or 0)
    except Exception as e:                            # noqa: BLE001
        return {}, str(e)[:140]
    return out, ""


def rank_seeds(seeds, markets, state, national=False, limit=None, kinds=None):
    """Fold near-duplicate seeds, rank the survivors by measured demand, and
    say which ones fit the grid.

    `kinds` is the claude_seed_kinds() map. Anything it calls an item or another
    trade is set aside BEFORE the ranking rather than filtered after it, so a
    high-volume object like "old tvs" cannot displace a real service line. It is
    reported, not deleted.

    Returns a dict the panel renders verbatim. Never mutates anything.
    """
    limit = int(limit or CFG.get("grid_max_services", 20))
    kinds = kinds or {}
    clean, order = [], {}
    for s in seeds or []:
        t = seed_norm(s, markets, state)
        if t and t not in order:
            order[t] = len(clean)
            clean.append(t)
    if not clean:
        return {"kept": [], "folded": [], "dropped": [], "demoted": [],
                "limit": limit, "total": 0, "measured": False, "basis": "",
                "error": ""}

    vols, err = {}, ""
    try:
        vols, _pc, err = fetch_local_volume(clean, [] if national else markets,
                                           state, national=national)
    except Exception as e:                       # noqa: BLE001
        err = str(e)[:140]
    vols = vols or {}
    vol = lambda t: int(vols.get(t, 0) or 0)     # noqa: E731
    measured = any(vol(t) for t in clean)

    # ---- ORDER on the quoted form, PRICE on local demand -------------------
    # Two jobs, two numbers, and they were sharing one. `vol` (bare, measured in
    # the client's cities) is the addressable demand the campaign captures and
    # is what total_volume and the volume component are built from — untouched.
    # `gvol` is the volume of the string that actually reaches the proposal, and
    # that is what decides which terms are worth a slot.
    #
    # gvol LEADS, vol BREAKS TIES. Small markets return 0/mo for most qualified
    # phrases; ordering on gvol alone would flatten those lists into entry order
    # and lose the signal the tool does have. Where the qualified form is
    # measurable it decides, and where it is not the old ordering stands.
    gvols, gerr = ({}, "")
    if not national:
        gvols, gerr = geo_qualified_volume(clean, markets, state)
    gvol = lambda t: int(gvols.get(t, 0) or 0)   # noqa: E731
    geo_measured = any(gvol(t) for t in clean)

    # Set aside anything that is not a service for this client. Done here, not
    # after the cut, because the whole point is that these must not take a slot.
    demoted = []
    if kinds:
        clean, _dem = demote_nonservices(clean, kinds, markets, state)
        demoted = [dict(r, volume=vol(r["term"])) for r in _dem]
    demoted.sort(key=lambda r: -r["volume"])

    # ---- fold near-duplicates -------------------------------------------
    # Highest volume wins the group; ties fall back to entry order so the
    # result is stable across runs. Containment counts as duplication: "junk"
    # and "junk removal" are one service line, not two, and a 20-slot grid
    # cannot afford to spend two slots saying it twice.
    # A BARE HEAD NOUN NEVER REPRESENTS THE GROUP. "junk" outsearches "junk
    # removal" 2,400 to 1,900 and is not a service anybody sells — same
    # bare-versus-qualified trap the acronym check exists for. So single-token
    # seeds sort last inside their group and only win when nothing else is left.
    # Counted on the RAW phrase, not the stemmed key: "services" is a shape word,
    # so keying "demolition services" gives one token and the penalty below would
    # hand its group to "shed demolition" (90/mo) over the real label (880/mo).
    ntok = lambda t: len((t or "").split())                # noqa: E731
    # THE CLIENT'S OWN NOUN OUTRANKS A NEIGHBOUR'S BIGGER NUMBER. Amare typed
    # thirteen "homes for rent" seeds and got a quote made of apartments, because
    # "apartments for rent" measures 90/mo and "3 bedroom homes for rent" measures
    # 10 — so volume alone spent every slot on a product they do not rent. Not a
    # filter: an adjacent term still gets quoted when slots remain, which is how
    # Brendan's own list carries two apartment terms among eighteen homes.
    # (2026-08-13)
    adj = lambda t: 1 if (kinds.get(t) or {}).get("kind") == "adjacent" else 0
    ranked = sorted(clean, key=lambda t: (adj(t), ntok(t) < 2,
                                          -gvol(t), -vol(t), order[t]))
    groups = []                                   # [{"keep":t, "fold":[t,...], "key":set}]
    for t in ranked:
        # A seed made entirely of shape words ("services", "near me") keys to the
        # empty set. Skipping it dropped the term from every bucket, so applying
        # the trim deleted a term the operator was never shown. Give it a key of
        # its own instead: it appears, and it folds with nothing. (2026-08-11)
        k = _seed_key(t) or frozenset({t})
        hit = None
        for g in groups:
            if k == g["key"] or k <= g["key"] or g["key"] <= k:
                hit = g
                break
        if hit is None:
            groups.append({"keep": t, "fold": [], "key": k})
        else:
            hit["fold"].append(t)
            hit["key"] = hit["key"] | k

    survivors = [g["keep"] for g in groups]
    kept_terms = survivors[:limit]
    dropped_terms = survivors[limit:]
    keptset = set(kept_terms)
    # Both numbers on every row: the panel has to be able to show WHY a term
    # outranked one with a bigger headline figure.
    row = lambda t: {"term": t, "volume": vol(t),        # noqa: E731
                     "geo_volume": gvol(t)}
    return {
        "kept": [row(t) for t in kept_terms],
        "dropped": [row(t) for t in dropped_terms],
        # Only report folds on seeds that actually made the cut — a fold under a
        # term that was dropped anyway is noise.
        "folded": [{"keep": g["keep"], "volume": vol(g["keep"]),
                    "fold": [row(f) for f in g["fold"]]}
                   for g in groups if g["fold"] and g["keep"] in keptset],
        "folded_total": sum(len(g["fold"]) for g in groups),
        "demoted": demoted,
        "adjacent": [t for t in clean if adj(t)],
        "limit": limit, "total": len(clean) + len(demoted), "measured": measured,
        "basis": "US national" if national or not markets else "targeted cities",
        # Which number actually did the ordering, said out loud. A probe that
        # came back empty leaves the old bare ordering in place, and that has to
        # be visible rather than looking like the new behaviour.
        "order_basis": ("national demand — no market to qualify with" if national
                        else "geo-qualified volume, ties broken on local demand"
                        if geo_measured
                        else "local demand — the qualified forms did not measure"),
        "geo_measured": geo_measured,
        "geo_ordered": int(sum(1 for t in clean if gvol(t))),
        "geo_error": gerr or "",
        "error": err or "",
    }

# --- geographic market grouping -------------------------------------------
# Three attempts at inferring market identity from SEARCH DEMAND all failed on
# the same client: geo-modified probes, resolved location names, and per-city
# volumes. The reason was the same each time — Blair County towns have no
# measurable individual demand, so no demand-based method can say whether they
# are one market or seven. Distance always can, and it is the thing the rule
# actually means: "2-3 related NEARBY markets normally run under one campaign".
#
# Brent Cogan's seven towns span 22 miles end to end. TN Water & Air's markets
# are 101 and 160 miles apart. One threshold separates them.
_ZIP_INDEX = None


def _zip_index():
    """city+state -> mean lat/long, built once from the bundled ZIP dataset."""
    global _ZIP_INDEX
    if _ZIP_INDEX is not None:
        return _ZIP_INDEX
    idx = {}
    try:
        import zipcodes
        acc = {}
        for r in zipcodes.list_all():
            if r.get("country") != "US" or not r.get("lat"):
                continue
            la, lo = float(r["lat"]), float(r["long"])
            # The dataset carries placeholder rows at 0,0 — Kennesaw GA has one,
            # and averaging it in put the city in the Atlantic, 1,100 miles from
            # its neighbours (2026-08-03).
            if not la or not lo:
                continue
            st_ = str(r.get("state", "")).upper()
            acc.setdefault((str(r.get("city", "")).lower(), st_), []).append((la, lo))
            # Some cities have no ZIPs of their own and appear only as an
            # alternate name on a neighbour's — Milton GA shares Alpharetta's
            # 30004/30009 — so they were reported as unplaceable and counted
            # as separate markets (2026-08-03). Index the alternates too, at
            # the host ZIP's coordinates, which is where they actually are.
            for alt in (r.get("acceptable_cities") or []):
                acc.setdefault((str(alt).lower(), st_), []).append((la, lo))

        def _med(xs):
            xs = sorted(xs)
            n_ = len(xs)
            return xs[n_ // 2] if n_ % 2 else (xs[n_ // 2 - 1] + xs[n_ // 2]) / 2

        # Median, not mean: one stray record can't drag a city across the map.
        idx = {k: (_med([a for a, _ in v]), _med([b for _, b in v]))
               for k, v in acc.items()}
    except Exception:
        app.logger.warning("zipcodes not available — geographic grouping is off")
    _ZIP_INDEX = idx
    return idx


# ---- counties -------------------------------------------------------------
# A county is a legitimate thing to type: it is how a service business describes
# its footprint ("we cover Knox, Anderson and Blount"), it is a real DataForSEO
# location, and "bucks county roofing" is real search phrasing. But the ZIP
# index is keyed by CITY, so every county came back unplaceable — and
# api_markets counts each unplaceable entry as its own market.
#
# Junk Bee Gone: ten counties plus thirteen towns inside those same counties.
# The towns clustered correctly into four markets; the ten counties were counted
# separately on top, so 23 entries covering about six trade areas were reported
# as 16 markets. Everything downstream — grid shape, coverage percentages,
# add-on recommendation — is built on that number. (2026-08-10)
_COUNTY_INDEX = None
_CITY_COUNTY = None
_COUNTY_SUFFIX = re.compile(
    r"\s+(county|counties|parish|borough|census area|municipality)$")


def _county_indexes():
    """Build (county -> coords/zips/cities, city -> county) from the ZIP data."""
    global _COUNTY_INDEX, _CITY_COUNTY
    if _COUNTY_INDEX is not None:
        return _COUNTY_INDEX, _CITY_COUNTY
    cidx, ccidx = {}, {}
    try:
        import zipcodes
        acc = {}
        for r in zipcodes.list_all():
            if r.get("country") != "US":
                continue
            co = str(r.get("county") or "").strip().lower()
            st = str(r.get("state") or "").upper()
            if not co or not st:
                continue
            city = str(r.get("city") or "").strip().lower()
            e = acc.setdefault((co, st), {"pts": [], "cities": {}})
            e["cities"][city] = e["cities"].get(city, 0) + 1
            try:
                la, lo = float(r.get("lat") or 0), float(r.get("long") or 0)
            except (TypeError, ValueError):
                la = lo = 0
            if la and lo:
                e["pts"].append((la, lo))
            # Alternate ZIP names live in the same county as their host.
            for alt in (r.get("acceptable_cities") or []):
                a = str(alt).strip().lower()
                if a:
                    ccidx.setdefault((a, st), co)
            if city:
                ccidx[(city, st)] = co

        def _med(xs):
            xs = sorted(xs)
            n_ = len(xs)
            return xs[n_ // 2] if n_ % 2 else (xs[n_ // 2 - 1] + xs[n_ // 2]) / 2

        for k, e in acc.items():
            pts = e["pts"]
            cities = e["cities"]
            cidx[k] = {
                "coords": ((_med([a for a, _ in pts]), _med([b for _, b in pts]))
                           if pts else None),
                "zips": sum(cities.values()),
                "cities": cities,
                # The county's recognisable centre, for naming and for the one
                # case where a county has to stand in for a city.
                "principal": (max(cities, key=lambda c: (cities[c], -len(c)))
                              if cities else ""),
            }
    except Exception:
        app.logger.warning("zipcodes not available — county grouping is off")
    _COUNTY_INDEX, _CITY_COUNTY = cidx, ccidx
    return cidx, ccidx


def county_key(market, state=""):
    """('knox county', 'TN') for a market that names a county, else None."""
    city, st = parse_market(market, state)
    name = (city or "").strip().lower()
    if not _COUNTY_SUFFIX.search(name):
        return None
    name = _COUNTY_SUFFIX.sub(" county", name)
    abbr = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").upper()
    cidx, _cc = _county_indexes()
    if abbr and (name, abbr) in cidx:
        return (name, abbr)
    # No state given: accept a unique national match, never a guess. There are
    # Jefferson Counties in twenty-five states.
    hits = [k for k in cidx if k[0] == name]
    return hits[0] if len(hits) == 1 else None


def county_of(market, state=""):
    """The county a city sits in, as ('knox county','TN'), or None."""
    city, st = parse_market(market, state)
    c = (city or "").strip().lower()
    abbr = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").upper()
    _ci, ccidx = _county_indexes()
    if abbr:
        co = ccidx.get((c, abbr))
        return (co, abbr) if co else None
    hits = [(co, s) for (cc, s), co in ccidx.items() if cc == c]
    return hits[0] if len(hits) == 1 else None


def city_size(market, state=""):
    """Rough size proxy — how many ZIP codes a city has.

    Used to name a market after its recognisable centre. Sorting by latitude
    made a seven-town Blair County market read as "Claysburg +6" when anyone
    would call it Altoona.
    """
    city, st = parse_market(market, state)
    city = (city or "").strip().lower()
    abbr = STATE_ABBREV.get((st or state or "").strip().lower(), "").upper()
    # A county's size is its ZIP count, but it must never out-rank its own seat
    # when a group is being named: "Knoxville +4" is the market, "Knox County +4"
    # is a filing cabinet. Knox County has 35 ZIPs to Knoxville's 31, so scored
    # honestly the county would win. Halve it — enough to lose to its principal
    # city, enough to still beat a hamlet.
    ck = county_key(market, state)
    if ck:
        return (_county_indexes()[0].get(ck) or {}).get("zips", 0) // 2
    # The ZIP data uses its own names: "New York" not "New York City", "Lawrence
    # Township" not "Lawrenceville". An exact match alone scored New York City at
    # ZERO — the same as a town that doesn't exist — which made it read as the
    # smallest market in a five-market grid (2026-08-07).
    variants = [city]
    if city.endswith(" city") and len(city.split()) > 1:
        variants.append(city[: -len(" city")].strip())
    for alias in _CITY_ALIASES.get(city, []):
        variants.append(alias)
    try:
        import zipcodes
        rows = zipcodes.list_all()
        for v in variants:
            n = sum(1 for r in rows
                    if str(r.get("city", "")).lower() == v
                    and (not abbr or str(r.get("state", "")).upper() == abbr))
            if n:
                return n
        # Last resort: a township or borough carrying the same stem.
        stem = variants[-1]
        n = sum(1 for r in rows
                if str(r.get("city", "")).lower().startswith(stem + " ")
                and (not abbr or str(r.get("state", "")).upper() == abbr))
        return n
    except Exception:
        return 0


_NAME_SHARE = {}


def name_is_unmistakable(market, state=""):
    """True if this city name needs no state suffix to mean what it means.

    Two conditions, both measured off the bundled ZIP data: the place is big
    (ZIP count), and it is the dominant bearer of its name nationally. Clinton
    TN has 2 ZIPs out of 33 Clintons — a bare "clinton" is not this client's
    market. Knoxville has 31 of 38.
    """
    if county_key(market, state):
        return False                       # "knox county" always keeps its state
    city, st = parse_market(market, state)
    c = (city or "").strip().lower()
    abbr = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").upper()
    if not c or not abbr:
        return False
    if _NAME_SHARE.get("__built__") is None:
        # Built ONCE, not scanned per city: Render's free tier gives 0.1 vCPU
        # and a 16-market grid would otherwise walk 42,000 ZIP rows sixteen
        # times inside a request.
        counts = {}
        try:
            import zipcodes
            for r in zipcodes.list_all():
                cc = str(r.get("city", "")).lower()
                st_ = str(r.get("state", "")).upper()
                if not cc:
                    continue
                counts[(cc, st_)] = counts.get((cc, st_), 0) + 1
                counts[cc] = counts.get(cc, 0) + 1
        except Exception:
            pass
        _NAME_SHARE.clear()
        _NAME_SHARE.update(counts)
        _NAME_SHARE["__built__"] = True
    n_self = _NAME_SHARE.get((c, abbr), 0)
    n_all = _NAME_SHARE.get(c, 0)
    if not n_all:
        return False
    return (n_self >= int(CFG.get("metro_no_suffix_zips", 25))
            and (n_self / n_all) >= float(CFG.get("metro_no_suffix_share", 0.6)))


def name_share(market, state=""):
    """What fraction of US ZIPs carrying this city name are in THIS state.

    A distinctiveness score. Clinton TN holds 2 of 33 Clintons; Oak Ridge TN
    holds 2 of 7. Both are the same size in ZIP count, and this is what
    separates them when one has to stand for the market.
    """
    name_is_unmistakable(market, state)          # builds the index on first call
    city, st = parse_market(market, state)
    c = (city or "").strip().lower()
    abbr = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").upper()
    n_self = _NAME_SHARE.get((c, abbr), 0)
    n_all = _NAME_SHARE.get(c, 0)
    return (n_self / n_all) if n_all else 0.0


def market_anchor(group, state=""):
    """The one name that stands for a clustered market.

    Used by BOTH the market summary and the keyword grid, so the panel that says
    "Sevierville +3" and the keywords that get built cannot disagree — they did,
    and the grid crossed its terms with Pigeon Forge while the operator read
    Sevierville on screen. (2026-08-10)

    Towns before counties (a county describes coverage, not search), then size,
    then distinctiveness, then the shorter name, then alphabetical so the same
    input always gives the same answer.
    """
    g = [m for m in (group or []) if str(m).strip()]
    if not g:
        return ""
    towns = [m for m in g if not county_key(m, state)] or g
    return sorted(towns, key=lambda m: (-city_size(m, state), -name_share(m, state),
                                        len(str(m)), str(m).lower()))[0]


def city_coords(market, state=""):
    """Latitude/longitude for an entered market, or None."""
    city, st = parse_market(market, state)
    city = (city or "").strip().lower()
    if not city:
        return None
    # A county has coordinates too — the median of its ZIPs. Without this every
    # county entered was "couldn't place" and counted as a market of its own.
    ck = county_key(market, state)
    if ck:
        e = _county_indexes()[0].get(ck) or {}
        if e.get("coords"):
            return e["coords"]
    abbr = STATE_ABBREV.get((st or state or "").strip().lower(), "")
    idx = _zip_index()
    if abbr:
        hit = idx.get((city, abbr.upper()))
        if hit:
            return hit
    # No usable state: accept a unique national match, never a guess between
    # several — there are Springfields in thirty states.
    hits = [v for (c, _s), v in idx.items() if c == city]
    return hits[0] if len(hits) == 1 else None


def miles_between(a, b):
    from math import radians, sin, cos, asin, sqrt
    (la1, lo1), (la2, lo2) = a, b
    la1, lo1, la2, lo2 = map(radians, [la1, lo1, la2, lo2])
    h = sin((la2 - la1) / 2) ** 2 + cos(la1) * cos(la2) * sin((lo2 - lo1) / 2) ** 2
    return 2 * 3958.8 * asin(sqrt(h))


def is_state_geo(m):
    """True if this targeting area is a whole STATE, not a market in it.

    "california" as a targeting area is unplaceable, so it was counted as its own
    market and reported as "couldn't place california" — which reads like a data
    gap rather than what it is: a state entered where a city belongs. Its cities
    are usually in the list already, so it is also double coverage. (2026-08-10)
    """
    t = re.sub(r"\s+", " ", str(m or "").strip().lower())
    if not t:
        return False
    if "," in t:
        head, tail = [p.strip() for p in t.rsplit(",", 1)]
        # "Austin, TX" is a city; only a bare state name or "Texas, TX" counts.
        if tail in _abbrev_to_state() or tail in STATE_ABBREV:
            t = head
    return t in STATE_ABBREV or t in _abbrev_to_state()


def acronym_collisions(rows, min_ratio=None, min_volume=None):
    """Which bare acronyms in the list are probably measuring someone else.

    A short acronym rarely belongs to one industry, and Google Ads volume is the
    SUM ACROSS EVERY MEANING. NASSCO's list came back with LACP at 4,400/mo and
    MACP at 2,900 — against PACP, its flagship and by far its best-known
    programme, at 480. LACP is also the Link Aggregation Control Protocol and
    MACP a Master of Arts in Counselling Psychology, and between them those two
    rows were 87% of the volume setting the price.

    The test needs no extra call, because the list already carries both forms.
    Compare the bare acronym with its qualified sibling ("lacp" vs "lacp
    certification"). Where the client genuinely owns the term the two track each
    other — PACP 480 against 320, a ratio of 1.5. Where the bare form is
    dwarfing its qualified form by two orders of magnitude, the traffic is
    arriving for a different meaning: 147x for LACP, 72x for MACP.

    Reported, not removed. An acronym can legitimately outrun its qualified form,
    and only a human looking at the SERP can settle it. (2026-08-10)
    """
    ratio_cap = float(min_ratio if min_ratio is not None
                      else CFG.get("acronym_collision_ratio", 8.0))
    vol_floor = int(min_volume if min_volume is not None
                    else CFG.get("acronym_collision_min_volume", 100))
    vols = {}
    for r in (rows or []):
        kw = str((r or {}).get("keyword") or (r or {}).get("kw") or "").strip().lower()
        if kw:
            vols[kw] = int((r or {}).get("volume") or (r or {}).get("vol") or 0)
    out = []
    for kw, v in vols.items():
        if " " in kw or not kw.isalpha() or not (3 <= len(kw) <= 6):
            continue
        if v < vol_floor:
            continue
        # Every longer phrase in the list that starts with this acronym.
        quals = {k: n for k, n in vols.items() if k.startswith(kw + " ")}
        if not quals:
            # NOTHING TO COMPARE AGAINST is not the same as passing. The ratio
            # test silently skipped any bare acronym with no qualified sibling in
            # the list — so removing "itcp certification" stopped "itcp" being
            # checked at all, and a big untested number kept setting the price.
            # Flag it as untestable when it is material, and let the SERP decide.
            # (2026-08-10)
            if v >= vol_floor * 2:
                out.append({"acronym": kw, "volume": v,
                            "qualified_volume": None, "qualified": "",
                            "ratio": None, "untestable": True})
            continue
        best_q = max(quals.values())
        ratio = (v / best_q) if best_q else float("inf")
        if ratio >= ratio_cap:
            out.append({"acronym": kw, "volume": v, "qualified_volume": best_q,
                        "qualified": max(quals, key=lambda k: quals[k]),
                        "ratio": (round(ratio, 1) if best_q else None)})
    out.sort(key=lambda x: -x["volume"])
    total = sum(vols.values()) or 1
    flagged = sum(x["volume"] for x in out)
    return {"items": out, "flagged_volume": flagged, "total_volume": total,
            "share": round(flagged / total * 100)}


def geo_overlaps(markets, state=""):
    """Which entered geos are already covered by another entered geo.

    Distance clustering answers "are these the same market". It does not answer
    "is this one INSIDE that one", and those are different questions with
    different consequences. Knox County and Knoxville are one footprint however
    far apart their centroids happen to compute; so are New York City and New
    York. Counting both is counting the same ground twice, and the market count
    is what the grid, the coverage percentages and the add-on recommendation are
    all built on.

    Three kinds, all read off bundled data rather than guessed:
      county    — an entered county contains an entered city
      duplicate — two entries name the same place ("nyc" / "new york city")
      non_place — not a place at all ("near me"), so it covers nothing

    Returns a list of {kind, container, contained[]} plus the non-place list.
    Detection only: nothing is removed from the operator's pills here.
    """
    mk = [m for m in (markets or []) if str(m).strip()]
    junk = [m for m in mk if is_non_place_geo(m)]
    real = [m for m in mk if m not in junk]

    # --- counties containing entered cities
    out = []
    counties = [(m, county_key(m, state)) for m in real]
    counties = [(m, k) for m, k in counties if k]
    have_county = {m for m, _k in counties}
    for m, key in counties:
        inside = [c for c in real
                  if c not in have_county and county_of(c, state) == key]
        if inside:
            out.append({"kind": "county", "container": m, "contained": inside})

    # --- same place entered twice under two names
    seen = {}
    for m in real:
        if m in have_county:
            continue
        city, st = parse_market(m, state)
        cl = (city or "").strip().lower()
        abbr = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").upper()
        # Canonicalise through the same alias ladder the volume lookup uses, so
        # "New York City, NY" and "New York, NY" land on one key.
        canon = canonical_city_name(cl, st or state) or cl
        k = (canon, abbr)
        if k in seen:
            hit = next((o for o in out
                        if o["kind"] == "duplicate" and o["container"] == seen[k]), None)
            if hit:
                hit["contained"].append(m)
            else:
                out.append({"kind": "duplicate", "container": seen[k],
                            "contained": [m]})
        else:
            seen[k] = m

    if junk:
        out.append({"kind": "non_place", "container": "", "contained": junk})
    return out


def group_by_distance(markets, state="", radius=None):
    """Cluster markets that sit within `radius` miles of each other.

    Single-link: A joins B's cluster if it is within the radius of ANY member,
    so a chain of neighbouring towns stays one market rather than splitting on
    the arithmetic of where the centre happens to fall.

    Returns (groups, located, unlocated).
    """
    r = float(radius if radius is not None else CFG.get("market_radius_miles", 25))
    mk = [m for m in (markets or []) if str(m).strip()]

    # Counties are CONTAINERS, not peers, and they must not cluster like towns.
    # Clustered as ordinary points they act as bridges: Sevier County's ZIP
    # median lands near Knoxville, so letting it chain pulled Sevierville and
    # Pigeon Forge — 30 miles out and plainly their own market — into metro
    # Knoxville and collapsed 22 entries to 3. Hold them back, cluster the towns
    # with the diameter guarantee intact, then place each county by what it
    # actually contains. (2026-08-10)
    county_of_entry = {m: county_key(m, state) for m in mk}
    counties = [m for m in mk if county_of_entry[m]]
    towns = [m for m in mk if not county_of_entry[m]]

    pts, unlocated = {}, []
    for m in towns:
        c = city_coords(m, state)
        (pts.__setitem__(m, c) if c else unlocated.append(m))
    # Complete-link: a city joins only if it is within the radius of EVERY
    # member, which caps the cluster's diameter at the radius. Single-link
    # chained Jasper into metro Atlanta through Canton even though Jasper is
    # 36 miles from Marietta — a market has a size, not just neighbours.
    groups = []
    for m in sorted(pts, key=lambda x: (pts[x][0], pts[x][1])):
        for g in groups:
            if all(miles_between(pts[m], pts[o]) <= r for o in g):
                g.append(m)
                break
        else:
            groups.append([m])

    # Now the counties. A county holding an entered town is redundant coverage —
    # it joins that town's cluster and adds no market. It joins ONE cluster even
    # if it spans several, so it can never merge two markets that the distance
    # test kept apart.
    for m in counties:
        key = county_of_entry[m]
        best, best_n = None, 0
        for g in groups:
            n = sum(1 for t in g if county_of(t, state) == key)
            if n > best_n:
                best, best_n = g, n
        if best is not None:
            best.append(m)
            pts.setdefault(m, city_coords(m, state) or (0.0, 0.0))
            continue
        # A county with no entered town inside it IS a market of its own —
        # Roane County here. Place it at its ZIP median and let it join a
        # neighbouring cluster only if it passes the same diameter test.
        c = city_coords(m, state)
        if not c:
            unlocated.append(m)
            continue
        for g in groups:
            if all(miles_between(c, pts[o]) <= r for o in g if o in pts):
                g.append(m)
                pts[m] = c
                break
        else:
            groups.append([m])
            pts[m] = c

    # Same place under two names — "New York City, NY" and "New York, NY" — is
    # one market however the clustering fell out.
    for ov in geo_overlaps(towns, state):
        if ov["kind"] != "duplicate":
            continue
        fam = [ov["container"]] + list(ov["contained"])
        idxs = sorted({i for i, g in enumerate(groups) if any(m in g for m in fam)})
        if len(idxs) < 2:
            continue
        keep = groups[idxs[0]]
        for i in reversed(idxs[1:]):
            keep.extend(groups[i])
            del groups[i]
    return groups, list(pts), unlocated


def suggest_geo_scope(markets, state="", national_demand=False,
                     national_reason=""):
    """Read the geo scope OFF the entered markets instead of asking for it.

    The operator picks a band from a dropdown, and the band chooses the pricing
    anchor — so a wrong pick is a wrong price. But the markets themselves answer
    the question: how many states, how far apart, and do the clusters touch.

    Contiguity is single-link at a JOIN radius (default 60 miles) rather than
    "is every city near every other city". Wayne and Shrewsbury are 55 miles
    apart and Shrewsbury to Lawrenceville is 35 — nobody would call that two
    separate regions, but a complete-link test at 25 miles calls it three. One
    connected chain = one region; two chains that never touch = non-contiguous.

    Returns a suggestion plus the evidence, never a decision. The operator keeps
    the dropdown: "Philadelphia + South Jersey" is one trade area to a human and
    two states to a distance function. (2026-08-07)
    """
    mk = [m for m in (markets or []) if str(m).strip()]
    out = {"suggested": "", "confidence": "", "reason": "", "evidence": {}}

    # NATIONAL DEMAND AND A REGIONAL BAND ARE INCOMPATIBLE, and the band is not
    # cosmetic — it picks the pricing anchor. NASSCO sat on non_contiguous_region
    # while its volume, keywords and competition were all national: $2,300 hard
    # against $2,000, a $400/mo difference decided by a dropdown describing
    # geography the quote had stopped using. Nothing anywhere told the operator
    # to change it. (2026-08-10)
    if national_demand:
        why = f" ({national_reason})" if national_reason else ""
        out.update(suggested="nationwide", confidence="high",
                   reason=(f"Priced on national demand{why}, so the band should "
                           "be Nationwide. The band sets the pricing anchor, and "
                           "a regional one charges for a footprint this quote is "
                           "not measuring — the keywords are bare and the volume "
                           "is a US figure."),
                   evidence={"cities": len(mk), "national_demand": True,
                             "national_reason": national_reason})
        return out
    if not mk:
        return out

    pts, unlocated = {}, []
    for m in mk:
        c = city_coords(m, state)
        (pts.__setitem__(m, c) if c else unlocated.append(m))

    states = sorted({market_state(m, state) for m in mk if market_state(m, state)})
    join_r = float(CFG.get("scope_join_radius_miles", 60))
    out["evidence"] = {"cities": len(mk), "states": states,
                       "located": len(pts), "unlocated": unlocated,
                       "join_radius": int(join_r)}

    if len(pts) < 2:
        # Nothing to measure. One city is one city; anything else can't be read
        # without coordinates, and guessing a band that sets the anchor is worse
        # than saying so.
        if len(mk) == 1:
            out.update(suggested="single_city", confidence="high",
                       reason="One market entered.")
        elif unlocated:
            out.update(reason=f"Could not place {len(unlocated)} of {len(mk)} "
                              "areas on the map, so the scope can't be read "
                              "from them.")
        return out

    # Single-link chain at the join radius: which clusters actually touch.
    names = list(pts)
    chains = []
    for m in names:
        joined = [c for c in chains
                  if any(miles_between(pts[m], pts[o]) <= join_r for o in c)]
        if not joined:
            chains.append([m])
        else:
            merged = [m]
            for c in joined:
                merged += c
                chains.remove(c)
            chains.append(merged)

    span = max((miles_between(pts[a], pts[b])
                for i, a in enumerate(names) for b in names[i + 1:]), default=0.0)
    out["evidence"].update(chains=len(chains), span_miles=round(span),
                           chain_sizes=sorted((len(c) for c in chains), reverse=True))

    if len(states) > 1:
        # Multi-state is only non-contiguous if the clusters are also apart.
        # Philadelphia + Cherry Hill is two states and one trade area.
        if len(chains) == 1:
            out.update(suggested="contiguous_region", confidence="medium",
                       reason=f"{len(mk)} entered areas across {len(states)} states "
                              f"({', '.join(states)}), but they form ONE "
                              f"connected area — {round(span)} miles end to end, "
                              f"no gap wider than {int(join_r)} miles.")
        else:
            out.update(suggested="non_contiguous_region", confidence="high",
                       reason=f"{len(chains)} separate clusters across "
                              f"{len(states)} states, {round(span)} miles end to "
                              "end — they do not touch.")
        return out

    if len(chains) > 1:
        out.update(suggested="non_contiguous_region", confidence="high",
                   reason=f"{len(chains)} separate clusters within "
                          f"{states[0] if states else 'one state'} — "
                          f"{round(span)} miles end to end, with gaps wider than "
                          f"{int(join_r)} miles between them.")
    elif span <= float(CFG.get("market_radius_miles", 25)):
        out.update(suggested="single_city", confidence="medium",
                   reason=f"All {len(mk)} entered areas sit within "
                          f"{round(span)} miles — one metro, not a region.")
    else:
        out.update(suggested="contiguous_region", confidence="high",
                   reason=f"{len(mk)} entered areas form one connected region, "
                          f"{round(span)} miles end to end"
                          + (f", all in {states[0]}" if states else "") + ".")
    return out


def group_by_metro(vectors, min_terms=2):
    """Group cities that Google Ads treats as the same place.

    Google Ads doesn't hold volume for every town — it resolves a city to the
    nearest TARGETABLE location, which for most suburbs is the metro. So two
    cities in one metro return the identical figure for every term, because
    they are literally the same location as far as the data is concerned.

    That makes "same market" free to detect: cities whose volume VECTOR across
    several distinct terms matches exactly are one market. Woodstock entered
    fourteen Georgia towns and seven returned exactly 10/mo on every probe —
    not a floor artifact, one metro answering seven times (2026-07-28).

    It matters because everything downstream counts markets: a client with
    fourteen pills that are really three markets gets an add-on count, a grid
    shape and a coverage percentage all built on a number that is four times
    too big.

    `vectors` maps city -> list of volumes, in a consistent term order.
    Returns a list of groups, biggest first.
    """
    cities = [c for c, v in (vectors or {}).items() if v]
    if len(cities) < 2:
        return [[c] for c in cities]
    # A single shared term proves nothing — thin terms report 0 or 10 almost
    # everywhere. Require several, and require at least one to be non-zero, so
    # cities aren't grouped on a shared absence of data.
    groups, seen = [], set()
    for c in cities:
        if c in seen:
            continue
        vec = tuple(vectors[c])
        if len(vec) < min_terms or not any(vec):
            groups.append([c])
            seen.add(c)
            continue
        same = [d for d in cities
                if d not in seen and tuple(vectors.get(d, ())) == vec]
        for d in same:
            seen.add(d)
        groups.append(same or [c])
    groups.sort(key=len, reverse=True)
    return groups


def pick_grid_cities(markets, state, limit, probe_term="", explain=None,
                    home_hint=""):
    """`explain` is an optional dict that gets filled with WHY these cities won.

    Selection quietly discards markets the partner entered, which is the kind
    of decision that should never be invisible — an operator seeing four of
    thirteen cities in the grid has no way to tell whether the other nine were
    judged or just dropped. The dict records the probe used, every city's
    score, and what was cut.
    """
    """Choose WHICH cities go in the grid when more are supplied than the cap.
    Taking the first N by input order picks alphabetically-early villages over
    real metros (e.g. 'Augusta Springs' before 'Fairfax'). Instead, rank the
    supplied cities by how much search demand they actually carry, using a
    generic '<city>' population-proxy query, and keep the biggest.
    Falls back to input order if the lookup fails."""
    exp = explain if isinstance(explain, dict) else {}
    exp.update({"limit": limit, "method": "", "probe": "", "kept": [], "dropped": []})
    cities = [m.strip() for m in markets if m.strip()]
    # drop a market that is actually the state name — it isn't a city
    if state:
        cities = [c for c in cities if c.lower() != state.strip().lower()]
    # NOTE: we do NOT return early when the cities fit under the cap. The same
    # probe that ranks them is what reveals which ones Google treats as one
    # place, and the market count matters whether or not any city was dropped —
    # four towns in one metro are one market, cap or no cap. Returning early
    # here meant grouping only ever ran for clients OVER the limit
    # (2026-08-03).
    under_cap = len(cities) <= limit
    try:
        # Probe with the state suffix so ambiguous names resolve to the RIGHT
        # place: bare "insurance washington" matches Washington State/DC, and
        # "insurance jersey" matches New Jersey — which would rank tiny Virginia
        # towns above real metros. "insurance washington va" scores correctly.
        abbr = STATE_ABBREV.get((state or "").strip().lower(), "")
        sfx = f" {abbr}" if abbr else ""
        # Probe with the CLIENT'S OWN service where we have one. "insurance"
        # was a population proxy — it ranks cities by how big they are, not by
        # how much demand this client has in them, and those differ: a commuter
        # town can out-search a bigger neighbour for furniture and lose badly on
        # insurance. Measuring the actual service picks the markets that matter
        # to this quote. Falls back to the proxy when no seed is available.
        # Probe with EVERY seed and sum, not just the first one. seeds[0] is
        # whatever the partner happened to type first — for Woodstock that was
        # "furniture design consultation", a niche term with no volume in any
        # market, so the real signal never got a chance and the generic proxy
        # took over (2026-07-28). Summing across the seeds means one thin term
        # can't blind the ranking.
        terms = []
        for t in (probe_term if isinstance(probe_term, (list, tuple)) else [probe_term]):
            t = clean_kw(strip_placeholders(strip_proximity((t or "").lower()))).strip()
            if t and t not in terms:
                terms.append(t)
        terms = terms[:4] or ["insurance"]
        # Build the probe key the SAME way it will be sent. dfs_kw_list strips
        # punctuation, so a market tagged "Hollidaysburg, PA" was looked up as
        # "...hollidaysburg, pa" while the API was asked for "...hollidaysburg
        # pa" — every volume read as zero, the proxy fallback fired, the metro
        # vectors were all zeros so nothing grouped, and the city ranking fell
        # back to alphabetical (2026-08-03). It only bit clients whose markets
        # carry a state tag, which is the format we ask for.
        # PER CITY, FROM THE MARKET ITSELF. The single `sfx` above comes from
        # derive_state(), which reads CITY_STATE keyed on a BARE city name — so
        # for a market typed in the documented "Knoxville, TN" form it returns
        # nothing and the suffix is empty. The grid keyword text does not have
        # that problem: it takes the state from parse_market(), which does parse
        # the tag. Result: markets were SCORED on "junk removal sevierville"
        # while the grid was BUILT on "junk removal sevierville tn" — two
        # different strings, two different volumes. Sevierville scored 0 on the
        # probe and 20/mo twice in the grid, and the axis recommendation rests on
        # that score. parse_market already returns the state here; use it.
        # (2026-08-11)
        # SCORE A COUNTY ON ITS PRINCIPAL CITY. "personal injury lawyer knox
        # county tn" measures nothing, and so does every other county, so the
        # market ranking collapsed to a tie and the primary market fell out of
        # sort order — Ooten Law was priced on Blount County (Maryville) when
        # the firm is in Knoxville, and the whole quote followed that choice.
        # The county is what the operator entered and what coverage means; the
        # city is what carries the demand that ranks it. (2026-08-22)
        # BOTH NAMINGS, AND THE BEST ONE COUNTS. The invariant this probe was
        # built to protect is that a market is SCORED on the same string the
        # grid will be BUILT from — score "junk removal sevierville" and quote
        # "junk removal sevierville tn" and the axis recommendation rests on a
        # number for a different keyword (2026-08-11). A county has two honest
        # namings and the grid picks between them later on measured volume, so
        # scoring only one would break that invariant whichever one was chosen.
        # Ask for both and take the better: it is the same question the grid
        # form probe asks, off the same data, so the two cannot disagree.
        def _ckeys(c):
            cty, st = parse_market(c, state)
            ab = STATE_ABBREV.get((st or "").strip().lower(), "") or (
                abbr if not st else "")
            sfx = f" {ab}" if ab else ""
            names = [clean_kw(cty or c).lower()]
            for seat in county_cities(c, state, limit=1):
                if seat and seat not in names:
                    names.append(seat)
            return [(n, sfx) for n in names]
        _ck = {c: _ckeys(c) for c in cities}
        keys = lambda t, c: [clean_kw(f"{t} {n}{sf}") for n, sf in _ck[c]]
        key = lambda t, c: keys(t, c)[0]
        # The best naming of the market is its score for that term.
        kvol = lambda vmap, t, c: max((vmap.get(k, 0) or 0) for k in keys(t, c))
        probe = [k for c in cities for t in terms for k in keys(t, c)][:700]
        payload = [{"keywords": dfs_kw_list(probe),
                    "location_name": loc_string(cities, state),
                    "language_code": "en"}]
        data = dfs_post("/keywords_data/google_ads/search_volume/live", payload)
        items = (data.get("tasks") or [{}])[0].get("result") or []
        vol = {(it.get("keyword") or "").lower(): (it.get("search_volume") or 0)
               for it in items}
        # Show the form actually sent, suffix included — the label read
        # "junk removal <city>" while the probe carried the state.
        _lbl_sfx = (_ck[cities[0]][0][1] if cities else sfx) or sfx
        exp["probe"] = " / ".join(f"{t} <city>{_lbl_sfx}" for t in terms)
        exp["method"] = "client term"
        scored = {c: sum(kvol(vol, t, c) for t in terms) for c in cities}
        # The same lookup that ranks the cities also reveals which of them
        # Google Ads treats as one place — no extra call.
        vectors = {c: [kvol(vol, t, c) for t in terms] for c in cities}
        exp["metro_groups"] = [g for g in group_by_metro(vectors, min_terms=len(terms))
                               if len(g) > 1]
        term = terms[0]
        # If the client's own term returns nothing anywhere, the ranking is
        # noise — fall back to the population proxy rather than picking cities
        # by accident of ordering.
        if term != "insurance" and not any(scored.values()):
            probe2 = [k for c in cities for k in keys("insurance", c)][:700]
            data2 = dfs_post("/keywords_data/google_ads/search_volume/live",
                             [{"keywords": dfs_kw_list(probe2),
                               "location_name": loc_string(cities, state),
                               "language_code": "en"}])
            v2 = {(it.get("keyword") or "").lower(): (it.get("search_volume") or 0)
                  for it in ((data2.get("tasks") or [{}])[0].get("result") or [])}
            scored = {c: kvol(v2, "insurance", c) for c in cities}
            # Regroup on the proxy too. Woodstock's seeds were niche enough to
            # return nothing anywhere, so the vectors were all zeros and no two
            # cities could be matched — the grouping went silent exactly when
            # the fallback fired, which is the case it is most needed in
            # (2026-07-28). The proxy resolves to the same Google Ads location
            # as any other term, so it groups just as well.
            vectors = {c: [kvol(v2, "insurance", c)] for c in cities}
            exp["metro_groups"] = [g for g in group_by_metro(vectors, min_terms=1)
                                   if len(g) > 1]
            exp["probe"] = f"insurance <city>{sfx}"
            exp["method"] = "population proxy"
            _vmap, _vterms = v2, ["insurance"]
        else:
            _vmap, _vterms = vol, terms
        # WHICH NAMING WON, not just how much the winner scored.
        # kvol() takes the max across a county's two namings and throws the
        # answer away -- yet that answer is exactly what the grid suffix needs,
        # and it is the only LOCALLY measured comparison of the two. The
        # national wording probe in pick_geo_forms asks the same question of
        # geo-qualified phrases, which mostly return nothing at country level:
        # "water damage restoration bellingham wa" and "... whatcom county wa"
        # both score 0, the probe reports "no data", and the default wins --
        # and the default is the string the operator typed, which is the
        # county. That is how a Bellingham restoration company was quoted on
        # "whatcom county wa" on a build that already knew better. Recorded
        # here, used there, no extra call. (2026-08-24)
        exp["market_forms"] = {}
        for _c in cities:
            _best, _bv = None, 0
            for _n, _sf in _ck[_c]:
                _tot = sum((_vmap.get(clean_kw(f"{_t} {_n}{_sf}"), 0) or 0)
                           for _t in _vterms)
                if _tot > _bv:
                    _best, _bv = f"{_n}{_sf}".strip(), _tot
            if _best:
                exp["market_forms"][_c] = _best
        # Ties broken by name so the same input always gives the same cities.
        # Ties are common in small markets — DataForSEO floors thin terms at
        # 10/mo, so seven towns can score identically and an alphabetical
        # tiebreak silently drops the biggest one. Where the client's own name
        # or domain contains a market, that market is almost always their
        # flagship, so it wins a tie.
        # Matching the WHOLE market string ("clarendon hills, il") against the
        # hint could never succeed: a domain is "clarendonchiro.com", with no
        # space, no state and usually only the first word of the town. So
        # compare on the bare city, on its de-spaced form, and on its first
        # token — which is what a domain almost always carries (2026-08-04).
        hint_raw = (home_hint or "").lower()
        hint_squashed = re.sub(r"[^a-z0-9]", "", hint_raw)

        def home_rank(c):
            bare = (parse_market(c, state)[0] or c).strip().lower()
            if not bare:
                return 1
            squashed = re.sub(r"[^a-z0-9]", "", bare)
            if len(squashed) > 3 and squashed in hint_squashed:
                return 0                       # "clarendonhills" in the hint
            if len(bare) > 3 and bare in hint_raw:
                return 0                       # "clarendon hills" spelled out
            first = bare.split()[0]
            if len(first) >= 5 and first in hint_squashed:
                return 0                       # "clarendon" in clarendonchiro
            return 1
        # A county is coverage, not a search target: "junk removal jefferson
        # county tn" is not a phrase anyone types, and a county only earns a grid
        # slot when no town of its market is available to stand for it.
        cty_rank = lambda c: 1 if county_key(c, state) else 0
        # NOTHING MEASURED IS NOT A RANKING. When every market scores zero the
        # order below is sort order wearing a measurement's clothes, and the
        # primary market it hands back sets the grid suffix, the rank-check
        # location and the price. Say so rather than letting it read as a
        # finding. The client's own name or domain breaks the tie first, which
        # is what home_rank is for. (2026-08-22)
        exp["nothing_measured"] = not any(scored.values())
        ranked = sorted(cities, key=lambda c: (-scored.get(c, 0), cty_rank(c),
                                               home_rank(c), c.lower()))
        if under_cap:
            exp["method"] = "all"
            exp["kept"] = [(c, scored.get(c, 0)) for c in ranked]
            exp["dropped"] = []
            return ranked

        # ---- one city per MARKET before a second city from any market --------
        # Ranking cities on volume alone is market-blind, and in a thin vertical
        # every city ties at Google's 10/mo floor, so the cut fell to
        # alphabetical order. Junk Bee Gone's five slots went to Knoxville,
        # Maryville, Clinton, Farragut and Jefferson County — and the first four
        # are ONE market. Three slots bought near-duplicates of Knoxville while
        # Sevierville, Oak Ridge and Morristown, three whole markets the tool had
        # just identified, got no keywords at all.
        #
        # So: round-robin across markets, best city first, and only start a
        # second lap once every market has a representative. The grouping is
        # already computed for the market count — this just stops the grid from
        # ignoring it. (2026-08-10)
        try:
            groups, _loc, _un = group_by_distance(cities, state)
            buckets = [list(g) for g in groups] + [[u] for u in _un]
        except Exception:
            buckets = [[c] for c in cities]
        order = {c: i for i, c in enumerate(ranked)}
        for bk in buckets:
            # The market's ANCHOR represents it, unless another city in it
            # actually carries more measured demand for this client's service.
            anc = market_anchor(bk, state)
            bk.sort(key=lambda c: (-scored.get(c, 0), c != anc, order.get(c, 10**6)))
        # Markets ordered by their strongest city, so the biggest market leads.
        buckets.sort(key=lambda bk: order.get(bk[0], 10**6))
        picked, lap = [], 0
        while len(picked) < limit and any(len(bk) > lap for bk in buckets):
            for bk in buckets:
                if len(bk) > lap and len(picked) < limit:
                    picked.append(bk[lap])
            lap += 1
        chosen = picked or ranked[:limit]
        exp["per_market"] = True
        exp["markets_covered"] = sum(1 for bk in buckets if any(c in chosen for c in bk))
        exp["market_count"] = len(buckets)
        ranked = chosen + [c for c in ranked if c not in chosen]
        exp["kept"] = [(c, scored.get(c, 0)) for c in ranked[:limit]]
        exp["dropped"] = [(c, scored.get(c, 0)) for c in ranked[limit:]]
        # If the cut line falls inside a tie, the choice between those markets
        # was arbitrary and the operator needs to know rather than assume the
        # tool measured something.
        if exp["dropped"]:
            edge = scored.get(ranked[limit - 1], 0)
            tied = [c for c in cities if scored.get(c, 0) == edge]
            exp["tied_at_cut"] = tied if len(tied) > 1 else []
            exp["tie_value"] = edge
        return ranked[:limit]
    except Exception:
        exp.update({"method": "input order",
                    "kept": [(c, None) for c in cities[:limit]],
                    "dropped": [(c, None) for c in cities[limit:]]})
        return cities[:limit]


# DataForSEO's Google Ads keyword endpoints reject a specific punctuation set
# outright: the whole BATCH 40501s, so one dirty term zeroes every volume in
# the run and the volume component of price silently becomes $0. Seen
# 2026-07-25 on a partner seed that carried the client's own city with a comma
# ("corner dental salem,"), which also defeated the grid's "service already
# contains this market" check and produced "corner dental salem, salem or".
# Apostrophes and hyphens are legal and meaningful ("kid's dentist",
# "same-day crowns"), so they are deliberately NOT in this set.
DFS_BAD_CHARS = re.compile(r"[,!@%^()={};~`<>?\\|]")

def clean_kw(text):
    """Strip characters DataForSEO rejects and normalise whitespace."""
    t = DFS_BAD_CHARS.sub(" ", (text or ""))
    return re.sub(r"\s+", " ", t).strip()

def dfs_kw_ok(text):
    """DFS also caps keywords at 80 chars and 10 words. Terms past either
    limit are dropped from LOOKUPS only — they stay in the proposal list."""
    t = (text or "").strip()
    return bool(t) and len(t) <= 80 and len(t.split()) <= 10

def dfs_kw_list(terms):
    """Sanitised, de-duplicated, API-safe keyword list."""
    out, seen = [], set()
    for t in terms or []:
        c = clean_kw(t).lower()
        if c and c not in seen and dfs_kw_ok(c):
            seen.add(c)
            out.append(c)
    return out


def claude_region_names(markets, state, brand="", business_desc=""):
    """Ask for the vernacular names locals use for these markets.

    Returns a list of candidate names — Appleton -> "fox cities", Cherry Hill
    -> "south jersey", Lewisburg -> "central pa". These are proposals only;
    nothing reaches the quote until search volume backs it (see
    validate_region_names). Many markets have no such name and an empty list
    is the correct answer for them.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    cities = [c for c in (markets or []) if c and c.strip()]
    if not api_key or not cities:
        return []
    prompt = f"""Name the informal REGIONAL terms that local people use for these markets, and that they
would plausibly type into Google along with a service.

MARKETS: {", ".join(cities)}
STATE: {state or "unknown"}

Rules:
1. Return only vernacular region names — the ones on local radio and in local business names.
   "fox cities" (Appleton WI), "south jersey" (Cherry Hill NJ), "central pa" (Lewisburg PA),
   "lehigh valley" (Bethlehem PA), "the triangle" (Raleigh NC), "inland empire" (Riverside CA).
2. NEVER return: the city itself, the state on its own, a county name, a neighbourhood, a metro
   area label nobody says out loud ("Allentown-Bethlehem-Easton MSA"), or a compass phrase you
   invented ("north-central Wisconsin") unless it is genuinely in common use.
3. Most markets have NO such term. Returning an empty list is a correct and expected answer —
   a wrong region name is far worse than a missing one, because it becomes a keyword the client
   is quoted to rank for.
4. At most 3, lowercase, no state suffix.

Return ONLY a JSON object, no prose, no markdown:
{{"regions": ["fox cities"]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 300, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=25)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        out = []
        for r in (json.loads(text).get("regions") or []):
            name = clean_kw(str(r).lower()).strip()
            if name and name not in out and len(name.split()) <= 4:
                out.append(name)
        return out[:3]
    except Exception:
        app.logger.exception("claude_region_names failed")
        return []


def claude_topics(seeds, business_desc="", brand=""):
    """Ask which PRODUCT LINES the operator's terms cover, and assign each term.

    Token clustering gets the big split right and the granularity wrong. Ski
    Barn's terms are three lines — ski/snowboard gear, BBQ and grills, and patio
    furniture — but "outdoor grill" shares a word with "outdoor furniture", so
    single-link merged all three into one topic labelled "furniture", under which
    the tool chose two grill services and no furniture at all (2026-08-07). A
    human reads those as obviously different aisles of the store.

    So the MODEL partitions and names; the CODE still enforces the quota, so the
    guarantee never depends on a non-deterministic call. Returns [] on any
    failure and topic_clusters() takes over.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    sd = [s for s in (seeds or []) if s and str(s).strip()]
    if not api_key or len(sd) < 4:
        return []
    listing = "\n".join("- " + str(s) for s in sd[:120])
    prompt = f"""Group these search terms into the PRODUCT LINES a customer would shop separately.

TERMS ({len(sd)}):
{listing}

BUSINESS: {(business_desc or '').strip()[:400] or 'not given'}
BRAND: {brand or 'not given'}

Rules:
1. A topic is an AISLE OF THE STORE — something a customer shops on its own trip. A retailer selling
   ski gear, barbecues and patio furniture has THREE topics, not one "outdoor" topic: nobody shopping
   for a grill is also shopping for skis, and grills and patio furniture are browsed separately.
2. Do NOT group by a shared adjective. "outdoor grill" and "outdoor furniture" are DIFFERENT topics
   even though both say "outdoor". Group by the THING BEING BOUGHT.
3. Do NOT split one thing into synonyms. "ski shop", "ski store", "ski outfitters", "ski gear" and
   "snowboard shop" are ONE topic. Words like shop/store/stores/service/rental/best/near me describe
   the shape of a search, not a different product.
4. Label each topic in 1-3 plain words, the way a person says it: "ski & snowboard gear",
   "bbq & grills", "patio furniture", "auto insurance", "dental implants".
5. Assign EVERY term to exactly one topic. Do not invent terms and do not drop any.
6. Most businesses have ONE topic. One topic covering everything is a correct answer — only split
   when the lines really are shopped separately.

Return ONLY JSON, no prose, no markdown:
{{"topics": [{{"label": "ski & snowboard gear", "terms": ["ski shop", "snowboard rentals"]}}]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 2000, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=30)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        raw = json.loads(text).get("topics") or []
    except Exception:
        app.logger.exception("claude_topics failed")
        return []

    # Map back onto the ACTUAL seeds. The model can reword a term, so only exact
    # (case-insensitive) matches count, and anything it dropped joins the biggest
    # topic rather than silently losing its claim on a slot.
    by_lower = {str(s).strip().lower(): s for s in sd}
    out, claimed = [], set()
    for t in raw:
        label = " ".join(str(t.get("label") or "").strip().lower().split())
        if len(label) > 60:
            # Trim on a word boundary. A hard [:40] produced "pipeline
            # inspection certification traini". (2026-08-10)
            label = label[:60].rsplit(" ", 1)[0]
        terms = []
        for term in (t.get("terms") or []):
            k = str(term).strip().lower()
            if k in by_lower and k not in claimed:
                claimed.add(k)
                terms.append(by_lower[k])
        if label and terms:
            toks = set()
            for x in terms:
                toks |= _topic_tokens(x)
            out.append({"label": label, "seeds": terms, "tokens": toks,
                        "size": len(terms), "source": "ai"})
    if not out:
        return []
    leftover = [by_lower[k] for k in by_lower if k not in claimed]
    if leftover:
        big = max(out, key=lambda g: g["size"])
        big["seeds"] += leftover
        big["size"] = len(big["seeds"])
        for x in leftover:
            big["tokens"] |= _topic_tokens(x)
    out.sort(key=lambda g: (-g["size"], g["label"]))
    return out


def validate_region_names(candidates, service_term, markets, state):
    """Keep only region names with REAL search demand for this client's service.

    A name being locally recognised doesn't mean anyone searches it with a
    service attached, and an unsearched region name in the grid is a keyword
    the client is being quoted to rank for. So each candidate is measured the
    way it will actually be used — service + region — against the same phrase
    built from the client's own city, which is the honest benchmark.

    Testing the bare region name would not work: "fox cities" has volume as a
    navigational query whether or not anyone searches "roofing fox cities".

    Returns (kept, rejected) where each entry is (name, volume).
    """
    svc = clean_kw((service_term or "").lower()).strip()
    cands = [c for c in (candidates or []) if c]
    if not svc or not cands:
        return [], []
    city = next((c for c in (markets or []) if c and c.strip()), "")
    if city:
        _c, _st = parse_market(city, state)
        _sfx = f"{_c} {STATE_ABBREV.get((_st or '').lower(), '')}".strip()
        city_kw = clean_kw(f"{svc} {_sfx}")
    else:
        city_kw = svc
    probe = dfs_kw_list([f"{svc} {c}" for c in cands] + ([city_kw] if city_kw else []))
    if not probe:
        return [], []
    vols = {}
    try:
        payload = [{"keywords": probe,
                    "location_name": loc_string(markets, state) or "United States",
                    "language_code": "en"}]
        data = dfs_post("/keywords_data/google_ads/search_volume/live", payload)
        for row in ((data.get("tasks") or [{}])[0].get("result") or []):
            vols[(row.get("keyword") or "").lower()] = row.get("search_volume") or 0
    except Exception:
        app.logger.exception("validate_region_names lookup failed")
        return [], [(c, None) for c in cands]

    floor = int(CFG.get("region_min_volume", 10) or 10)
    kept, rejected = [], []
    for c in cands:
        v = vols.get(f"{svc} {c}".lower(), 0)
        (kept if v >= floor else rejected).append((c, v))
    kept.sort(key=lambda t: -t[1])
    return kept, rejected


# Aliases people actually type instead of the full city name. Only entries that
# are genuinely more common in search than the formal name — this is not a list
# of every nickname a city has.
_CITY_ALIASES = {
    "new york city": ["nyc", "new york"],
    "new york": ["nyc"],
    "los angeles": ["la"],
    "san francisco": ["sf"],
    "philadelphia": ["philly"],
    "washington": ["dc", "washington dc"],
    "las vegas": ["vegas"],
    "atlanta": ["atl"],
    "new orleans": ["nola"],
    "saint louis": ["st louis"],
    "saint petersburg": ["st pete", "st petersburg"],
    "fort lauderdale": ["ft lauderdale"],
    "minneapolis": ["twin cities"],
}


_COUNTY_SEATS = None


def _county_seats():
    """county+state -> the cities in it, most ZIPs first.

    Built once from the bundled ZIP dataset. ZIP count is a proxy for size and
    it is a good one where it matters: Knox County TN is 31 ZIPs of Knoxville
    against 1 for the next town, San Diego County 81 against 8. Where no city
    dominates — Roane County TN is four towns with one ZIP each — the list is
    still returned and measured volume decides, which is the point.
    """
    global _COUNTY_SEATS
    if _COUNTY_SEATS is not None:
        return _COUNTY_SEATS
    idx = {}
    try:
        import zipcodes
        acc = {}
        for r in zipcodes.list_all():
            if r.get("country") != "US":
                continue
            cty = (r.get("county") or "").strip().lower()
            st_ = (r.get("state") or "").strip().upper()
            city = (r.get("city") or "").strip().lower()
            if cty and st_ and city:
                acc.setdefault((cty, st_), {}).setdefault(city, 0)
                acc[(cty, st_)][city] += 1
        for k, cities in acc.items():
            idx[k] = [c for c, _n in sorted(cities.items(), key=lambda kv: (-kv[1], kv[0]))]
    except Exception:                                     # noqa: BLE001
        idx = {}
    _COUNTY_SEATS = idx
    return idx


def county_cities(market, state, limit=2):
    """The principal cities of a county market, or [] if it is not a county."""
    city, st = parse_market(market, state)
    c = re.sub(r"\s+", " ", (city or "").strip().lower())
    if not c.endswith(" county"):
        return []
    # THE STATE ARRIVES IN BOTH SHAPES AND ONLY ONE OF THEM WORKED.
    # parse_market() expands a state written INSIDE the market string
    # ("Whatcom County, WA" -> "Washington"), but a state typed in the separate
    # Geo field is passed through as the operator wrote it, which is almost
    # always the two-letter form. STATE_ABBREV is name -> abbreviation, so
    # .get("wa") missed, ab came back empty, and this returned [] -- for EVERY
    # county on EVERY quote whose state field held an abbreviation. The
    # conversion built for Ooten was therefore dead in the common case, which is
    # how a Bellingham restoration company was quoted on "water damage
    # restoration whatcom county wa" at 30/mo. (2026-08-24)
    _st = (st or "").strip()
    ab = (STATE_ABBREV.get(_st.lower(), "") or "").upper()
    if not ab and len(_st) == 2 and _st.lower() in set(STATE_ABBREV.values()):
        ab = _st.upper()
    if not ab:
        return []
    return (_county_seats().get((c, ab)) or [])[:max(1, int(limit))]


def geo_form_candidates(market, state):
    """The ways a searcher might write this market, most formal first.

    "new york city ny" is what the grid produced for Ski Barn and nobody types
    it — the state is redundant on a city that famous, and "nyc" outsearches the
    full name several times over (2026-08-07). So generate the plausible forms
    and let measured volume choose, rather than guessing from a suffix rule.
    """
    city, st = parse_market(market, state)
    c = (city or "").strip().lower()
    if not c:
        return []
    ab = (STATE_ABBREV.get((st or "").strip().lower(), "") or "").lower()
    forms = []

    def add(f):
        f = clean_kw(re.sub(r"\s+", " ", (f or "")).strip().lower())
        if f and f not in forms:
            forms.append(f)

    if ab:
        add(f"{c} {ab}")          # the current default
    add(c)                        # bare city
    for a in _CITY_ALIASES.get(c, []):
        add(a)
        if ab and " " not in a and len(a) > 3:
            add(f"{a} {ab}")
    # "new york city" -> "new york": a trailing "city" is usually dropped.
    if c.endswith(" city") and len(c.split()) > 2:
        add(c[: -len(" city")].strip())
        if ab:
            add(f"{c[: -len(' city')].strip()} {ab}")
    # A COUNTY IS NOT HOW PEOPLE SEARCH, USUALLY. Ooten Law was entered as
    # thirteen East Tennessee counties and the grid quoted "personal injury
    # lawyer blount county tn" at 30/mo. Brendan's own list for the same client
    # is 42 of 50 terms on "knoxville tn" and 2 on "knox county" — the county
    # form is real, it is just long tail. County-qualified terms also carry
    # almost no advertisers, so the median bid read $25 against a vertical that
    # runs $80-200, and the competitive adder collapsed with it. (2026-08-22)
    #
    # Not a rewrite: the county stays a candidate and the principal cities join
    # it, so the volume probe picks between them the way it already picks
    # between "new york city ny" and "nyc". Where the county genuinely is the
    # phrase — "bucks county roofing" — it wins on its own numbers.
    seats = county_cities(market, state)
    for seat in seats:
        if ab:
            add(f"{seat} {ab}")
        add(seat)
    # AND THE PRINCIPAL CITY LEADS. forms[0] is the DEFAULT -- what
    # pick_geo_forms keeps when the wording probe reads nothing, which for
    # geo-qualified phrases at national level is most of the time. Leaving the
    # county there meant the tie always fell to the string the operator typed,
    # so the whole conversion only ever fired when the city won outright. The
    # prior is the other way round: a county is long tail and starves the CPC
    # adder, so where nothing separates them the city is the better guess. The
    # county keeps its place in the list and still wins on its own numbers.
    # (2026-08-24)
    if seats:
        lead = clean_kw(f"{seats[0]} {ab}".strip() if ab else seats[0])
        if lead in forms:
            forms = [lead] + [f for f in forms if f != lead]
    return forms


def pick_geo_forms(markets, state, service_terms):
    """Choose each market's grid form by MEASURED search volume.

    Two things had to change after the first attempt reported nothing at all
    (2026-08-07):

    1. PROBE WITH GENERIC TERMS. It used the grid's lead service, which was
       "alpine ski shop" — "alpine ski shop nyc" has no measurable volume in any
       wording, so all five candidates tied at zero and the default won by
       default. Now it probes with the SHORTEST few client terms, because short
       means generic means measurable, and sums across them.

    2. MEASURE NATIONALLY. The probe was localised to the primary market, so
       "ski shop nyc" was being counted only among searchers in Wayne, New
       Jersey. The question here is not "how much demand is there" — it is
       "which spelling do people type", and that is a national property of the
       language. Absolute demand is measured elsewhere, per city, as before.

    Returns (forms, report). `forms` maps market -> chosen form. `report` has one
    entry per market ALWAYS, with a status, so a run that changed nothing is
    visibly different from a run that could not measure.
    """
    terms = [clean_kw(str(t).lower()).strip() for t in (service_terms or [])]
    terms = [t for t in terms if t]
    # Shortest first: the most generic phrasing carries the volume that makes a
    # comparison possible. Cap at 3 to keep one API call small.
    terms = sorted(dict.fromkeys(terms), key=lambda t: (len(t.split()), len(t)))[:3]
    mk = [m for m in (markets or []) if m and str(m).strip()]
    if not terms or not mk:
        return {}, []

    cand = {m: geo_form_candidates(m, state) for m in mk}
    probes = []
    for m, forms in cand.items():
        for f in forms:
            for t in terms:
                kw = clean_kw(f"{t} {f}")
                if kw:
                    probes.append(kw)
    probes = dfs_kw_list(probes)
    if not probes:
        return {}, []

    vols = {}
    try:
        payload = [{"keywords": probes[:700],
                    # NATIONAL on purpose — see the docstring. This is a question
                    # about wording, not about local demand.
                    "location_name": "United States",
                    "language_code": "en"}]
        data = dfs_post("/keywords_data/google_ads/search_volume/live", payload)
        for it in (data["tasks"][0]["result"] or []):
            vols[str(it.get("keyword", "")).lower()] = it.get("search_volume") or 0
    except Exception as e:
        return {}, [{"market": m, "status": "error", "detail": str(e)[:120]} for m in mk]

    forms_out, report = {}, []
    for m, flist in cand.items():
        scored = []
        for f in flist:
            total = sum(vols.get(clean_kw(f"{t} {f}").lower(), 0) for t in terms)
            scored.append((f, total))
        default = flist[0] if flist else ""
        best = max(scored, key=lambda x: x[1]) if scored else None
        table = [{"form": f, "volume": v} for f, v in
                 sorted(scored, key=lambda x: -x[1])]
        if not best or best[1] <= 0:
            report.append({"market": m, "status": "no data", "kept": default,
                           "tested": table, "probed_with": terms})
            continue
        forms_out[m] = best[0]
        report.append({"market": m,
                       "status": "changed" if best[0] != default else "confirmed",
                       "chose": best[0], "instead_of": default,
                       "volume": best[1],
                       "default_volume": dict(scored).get(default, 0),
                       "tested": table, "probed_with": terms})
    return forms_out, report


# Words whose number must not be touched. A service already says which of these
# it means, and "1 bedrooms homes for rent" is not a phrase.
_NUMBER_FIXED = frozenset("""
bedroom bedrooms bathroom bathrooms bath baths br ba story stories storey
class classes size sizes series
north south east west northwest northeast southwest southeast midwest
""".split())


_PREPOSITIONS = frozenset("for to on by in with near at from of".split())
_SIBILANT = ("s", "x", "z", "ch", "sh")
# Function words only. NOT _SEED_SHAPE, which is a KEYING set and holds real
# nouns on purpose — "company", "services", "contractor" are shape for the
# purpose of "is this the same service", but they are the head noun of
# "property management company" and pluralise like one.
_FORM_SKIP = frozenset("a an the and or my your our".split())


def _flip_number(w):
    """The other number of one English noun. "" when it has no sensible other."""
    if not w.isalpha() or len(w) < 3:
        return ""
    if w.endswith("ies") and len(w) > 4:
        return w[:-3] + "y"
    if w.endswith("es") and w[:-2].endswith(_SIBILANT) and len(w) > 4:
        return w[:-2]
    # "business" and "class" end in s and are singular. They pluralise with -es
    # like any other sibilant; stripping the s gives "busines".
    if w.endswith("ss"):
        return w + "es"
    if w.endswith("s") and len(w) > 3:
        return w[:-1]
    if w.endswith("s"):
        return ""
    if w.endswith("y") and len(w) > 3 and w[-2] not in "aeiou":
        return w[:-1] + "ies"
    if w.endswith(_SIBILANT):
        return w + "es"
    return w + "s"


def service_number_forms(svc):
    """The same service spelled singular and plural.

    NOTHING IN THE BUILD CHOOSES BETWEEN THEM, and the choice is worth two
    orders of magnitude. Amare shipped "home for rent santa fe nm" at 10/mo in
    Competitive; Brendan's list leads with "homes for rent santa fe nm" as an
    Ultra head term. The stemmer folds the pair to one key everywhere — the
    fold, the grid dedupe, the topic guarantee — so the tool has always known
    they are the same service. It just kept whichever spelling arrived first,
    which is the model's or the operator's typing, not the market's.

    ONLY THE HEAD NOUN MOVES: the last content word before the first
    preposition, or the last content word when there is none. That is the thing
    being sold — "homes" in "homes for rent", "apartments" in "luxury
    apartments". Moving anything else invents a phrase rather than re-spelling
    one: pluralise every word and "homes for rent" produces "homes for rents",
    "luxury apartment" produces "luxurys apartment". One variant out, or none.
    (2026-08-16)
    """
    words = clean_kw((svc or "").lower()).split()
    if not words:
        return []
    stop = next((i for i, w in enumerate(words) if w in _PREPOSITIONS), len(words))
    head = next((i for i in range(stop - 1, -1, -1)
                 if words[i].isalpha() and words[i] not in _FORM_SKIP), None)
    if head is None or words[head] in _NUMBER_FIXED:
        return []
    other = _flip_number(words[head])
    if not other:
        return []
    cand = " ".join(words[:head] + [other] + words[head + 1:])
    return [cand] if cand != " ".join(words) else []


def service_form_probes(service_terms, cap=None):
    """The alternate spellings to MEASURE, for services already chosen.

    Handed to the volume call the build already makes, so this costs extra
    keywords and not an extra request — the same trade the "<service> near me"
    probe makes a few lines below it. The first version of this made its own
    national call, which was both a twelfth request against a 12/min limit and
    a silent failure: a rate-limited response comes back as HTTP 200 with an
    empty result, every spelling measured zero, and "no data on both" is
    indistinguishable from "nothing to do". (2026-08-16)
    """
    terms = [clean_kw(str(t).lower()).strip() for t in (service_terms or [])]
    terms = [t for t in dict.fromkeys(terms) if t]
    cap = int(cap if cap is not None else CFG.get("service_form_probe_cap", 20))
    out = []
    for t in terms[:cap]:
        for v in service_number_forms(t):
            if v not in out and v not in terms:
                out.append(v)
    return out


def choose_service_forms(service_terms, vols, err="", cap=None, pool=None):
    """Choose each service's spelling — by the MARKET'S OWN PHRASING first, and
    by measured volume only where that is silent.

    THE VOLUME TEST CANNOT ANSWER THIS QUESTION. Google Ads groups close
    variants, so it returns one figure for both spellings: Amare came back
    "apartment for rent 90 vs apartments for rent 90 · home for rent 10 vs homes
    for rent 10 · 3 bedroom homes for rent 10 vs 3 bedroom home for rent 10" —
    six pairs, six exact ties, including the zeros. A test that ties every time
    is not a strict test, it is no test, and it had quietly decided nothing on
    two consecutive builds.

    So ask the keyword pool instead. `keywords_for_keywords` and
    `keyword_suggestions` return phrases people really type, carrying their
    natural number: if the pool says "homes for rent" and never "home for rent",
    that IS the market's spelling, and it costs nothing to read because the
    build already fetched it. It also generalises where a rule could not —
    "roof repair" stays singular and "homes for rent" goes plural without
    anyone writing down which nouns pluralise.

    Volume remains the tie-break for the case the pool has no opinion on, under
    the same bar as before: `service_form_min_ratio` times the incumbent and
    above `service_form_min_volume`.

    Returns (forms, report). `forms` maps the original service -> the spelling
    to use. `report` carries one entry per service that was TESTED, with a
    status, so a run that changed nothing is visibly different from a run that
    could not measure. (2026-08-16)
    """
    terms = [clean_kw(str(t).lower()).strip() for t in (service_terms or [])]
    terms = [t for t in dict.fromkeys(terms) if t]
    if not terms:
        return {}, []
    cap = int(cap if cap is not None else CFG.get("service_form_probe_cap", 20))
    terms = terms[:cap]
    variants = {t: service_number_forms(t) for t in terms}
    if err:
        return {}, [{"service": t, "status": "error", "detail": str(err)[:120]}
                    for t in terms if variants[t]]
    vols = {str(k).lower(): v for k, v in (vols or {}).items()}
    seen = set()
    for c in (pool or []):
        t = c.get("keyword", "") if isinstance(c, dict) else c
        t = clean_kw(str(t or "").lower()).strip()
        if t:
            seen.add(t)

    def _in_pool(phrase):
        """The pool holds whole keywords, usually with a place on the end, so a
        containment test on word boundaries is the honest one."""
        return any(p == phrase or p.startswith(phrase + " ")
                   or p.endswith(" " + phrase) or (" " + phrase + " ") in p
                   for p in seen)

    ratio = float(CFG.get("service_form_min_ratio", 2.0))
    floor = int(CFG.get("service_form_min_volume", 50))
    forms_out, report = {}, []
    taken = {t.lower() for t in terms}
    for t in terms:
        if not variants[t]:
            continue
        base = int(vols.get(t, 0) or 0)
        scored = sorted(((v, int(vols.get(v, 0) or 0)) for v in variants[t]),
                        key=lambda x: -x[1])
        table = [{"form": f, "volume": v} for f, v in [(t, base)] + scored]
        best, best_v = scored[0]
        # THE POOL GETS THE FIRST WORD. Only when it names one spelling and not
        # the other has it actually said something.
        pool_says = ""
        if seen:
            _has_t, _has_v = _in_pool(t), _in_pool(best)
            if _has_v and not _has_t:
                pool_says = best
            elif _has_t and not _has_v:
                pool_says = t
        if pool_says and pool_says != t and best.lower() not in taken:
            forms_out[t] = pool_says
            taken.discard(t.lower())
            taken.add(pool_says.lower())
            report.append({"service": t, "status": "changed", "chose": pool_says,
                           "instead_of": t, "volume": best_v,
                           "default_volume": base, "tested": table,
                           "basis": "the market's own phrasing"})
            continue
        if pool_says == t:
            report.append({"service": t, "status": "confirmed", "chose": t,
                           "volume": base, "tested": table,
                           "basis": "the market's own phrasing"})
            continue
        if best_v <= 0 and base <= 0:
            report.append({"service": t, "status": "no data", "tested": table})
            continue
        # Re-spelling onto a phrase the grid already holds would merge two
        # services into one and leave the list a slot short.
        if best.lower() in taken:
            report.append({"service": t, "status": "confirmed", "chose": t,
                           "volume": base, "tested": table,
                           "detail": "the bigger spelling is already in the list"})
            continue
        if best_v >= floor and best_v >= base * ratio:
            forms_out[t] = best
            taken.discard(t.lower())
            taken.add(best.lower())
            report.append({"service": t, "status": "changed", "chose": best,
                           "instead_of": t, "volume": best_v,
                           "default_volume": base, "tested": table})
        else:
            report.append({"service": t, "status": "confirmed", "chose": t,
                           "volume": base, "tested": table})
    return forms_out, report


def suggest_market_name(market, state=""):
    """The name Google is likely to recognise for a market it rejected.

    "New York City" and "Lawrenceville" are not Google Ads locations — it holds
    "New York" and "Lawrence Township". The ZIP index knows both, and the market
    already resolves to coordinates, so the nearest ZIP-data city name at that
    exact point IS the recognised name. Both resolve at 0.0 miles, which is the
    same place under a different label rather than a guess (2026-08-07).

    Returns "City, ST" or "" when nothing better than the input can be found.
    """
    coords = city_coords(market, state)
    if not coords:
        return ""
    city, st = parse_market(market, state)
    have = (city or "").strip().lower()
    try:
        import zipcodes
        rows = zipcodes.list_all()
    except Exception:
        return ""
    best = None
    for r in rows:
        try:
            pt = (float(r["lat"]), float(r["long"]))
        except Exception:
            continue
        d = miles_between(coords, pt)
        if best is None or d < best[0]:
            best = (d, str(r.get("city") or ""), str(r.get("state") or ""))
    # Only worth suggesting if it is the SAME place under another name and the
    # name actually differs from what the operator typed.
    if not best or best[0] > 2.0:
        return ""
    name, abbr = best[1], best[2]
    if not name or name.strip().lower() == have:
        return ""
    return f"{name}, {abbr}" if abbr else name


# Words that make a search a SHOP search rather than a product search. A local
# retailer ranks and converts on these; "ski jackets" is a product query owned by
# Amazon, REI and the manufacturers, and "weber grill" is the maker's own name.
_STORE_INTENT = set("""shop shops store stores storefront outlet outlets dealer dealers
retailer retailers rental rentals service services repair shopping showroom
supplier suppliers supply center centre""".split())

# Manufacturer / product brands a retailer resells. They are legitimate keywords
# (the client stocks them) but they are not the client's own storefront demand.
_RESELLER_BRANDS = set("""weber traeger bigreenegg kamado napoleon blackstone
rossignol salomon atomic burton dakine k2 volkl nordica head fischer elan
patagonia northface columbia arcteryx oakley smith giro
trane carrier lennox rheem andersen pella marvin kohler moen""".split())


def is_store_intent(term):
    """Does this phrase describe a PLACE TO BUY rather than a thing to buy?"""
    words = set(re.split(r"[^a-z0-9]+", (term or "").lower()))
    return bool(words & _STORE_INTENT)


def is_reseller_brand(term):
    """Does this phrase lead with a manufacturer's name?"""
    words = set(re.split(r"[^a-z0-9]+", (term or "").lower()))
    return bool(words & _RESELLER_BRANDS)


def swap_low_volume_services(services, vols, seeds, topics, min_volume=None,
                            max_swaps=None, upgrade_ratio=None, typed=None):
    """Replace service names nobody searches with ones the client's own list has.

    The geo WORDING is measured against search volume; the SERVICE NAMES were
    not. Ski Barn's grid spent three of seven slots on "nearest snowboard shop"
    (10/mo total, four of five markets no-data), "best snowboard store" (20) and
    "bbq grill store" (50), while the operator's own seeds held "ski shop",
    "ski store", "snowboard shop" and "ski equipment" — all unused and all much
    larger. Nobody types "nearest snowboard shop" (2026-08-08).

    A dead service slot costs twice: it fills the proposal with terms the client
    cannot win traffic on, and it drags the volume total that drives price.

    Swaps stay INSIDE the topic, so topic coverage survives — a patio-furniture
    slot is replaced by a better patio-furniture term, never by a ski term. The
    tier is inherited from the service being replaced.

    Returns (services, report).
    """
    floor = int(CFG.get("service_min_volume", 30) if min_volume is None else min_volume)
    cap = int(CFG.get("service_max_swaps", 3) if max_swaps is None else max_swaps)
    ratio = float(CFG.get("service_upgrade_ratio", 0)
                  if upgrade_ratio is None else upgrade_ratio)
    if not services or not vols:
        return services, []

    def vol_of(name):
        return int(vols.get(str(name).lower(), 0) or 0)

    def _intent_ok(current, candidate):
        """A replacement may not be weaker in INTENT than what it replaces.

        Volume alone picked "weber grill" (4,560) over "bbq grill store" (50) and
        "ski jackets" (2,440) over a shop term, so the proposal led with a
        manufacturer's name and a product query for a client whose campaign is
        in-store sales (2026-08-08). Those searches are owned by Amazon, REI and
        the makers; a five-store New Jersey retailer neither ranks for nor
        converts on them.

        So: a store term can be replaced only by another store term, and a
        manufacturer's brand can never be swapped IN. A product term may still
        replace a product term on volume.
        """
        if is_reseller_brand(candidate):
            return False
        if is_store_intent(current) and not is_store_intent(candidate):
            return False
        return True

    out = [dict(x) for x in services]
    used = {str(x.get("service", "")).lower() for x in out}
    report = []
    # Weakest slots first, so a limited number of swaps is spent where it counts.
    order = sorted(range(len(out)), key=lambda i: vol_of(out[i].get("service")))
    for i in order:
        if len(report) >= cap:
            break
        cur = out[i].get("service", "")
        cur_v = vol_of(cur)
        if cur_v >= floor:
            continue
        topic = service_topic(cur, topics) if topics else ""
        # Candidates: unused seeds from the SAME topic, best measured first.
        pool = []
        for t in (topics or []):
            if topic and t.get("label") != topic:
                continue
            pool += list(t.get("seeds") or [])
        if not topic or not pool:
            pool = list(seeds or [])
        cands = []
        for sd in pool:
            k = str(sd).strip().lower()
            if not k or k in used:
                continue
            v = vol_of(k)
            if v <= max(cur_v, floor - 1):
                continue
            if not _intent_ok(cur, k):
                continue
            cands.append((v, k))
        if not cands:
            continue
        cands.sort(reverse=True)
        best_v, best = cands[0]
        report.append({"out": cur, "out_volume": cur_v,
                       "in": best, "in_volume": best_v, "kind": "dead",
                       "topic": topic, "tier": out[i].get("tier", "")})
        used.discard(str(cur).lower())
        used.add(best)
        out[i] = {"service": best, "tier": out[i].get("tier", "competitive")}

    # UPGRADE PASS. The rule above only rescues DEAD slots, so a service that
    # scrapes past the floor keeps its place even when the operator's own list
    # holds something in a different league — Ski Barn kept "alpine ski shop"
    # (110/mo) while "ski store" (880) and "snowboard shop" (720) sat unused
    # (2026-08-08). Same-topic only, so coverage survives.
    if ratio > 0:
        order2 = sorted(range(len(out)), key=lambda i: vol_of(out[i].get("service")))
        for i in order2:
            if len(report) >= cap:
                break
            cur = out[i].get("service", "")
            cur_v = vol_of(cur)
            if cur_v <= 0:
                continue                      # handled by the floor pass
            topic = service_topic(cur, topics) if topics else ""
            pool = []
            for t in (topics or []):
                if topic and t.get("label") != topic:
                    continue
                pool += list(t.get("seeds") or [])
            if not topic or not pool:
                pool = list(seeds or [])
            cands = []
            for sd in pool:
                k = str(sd).strip().lower()
                if not k or k in used:
                    continue
                v = vol_of(k)
                if v < cur_v * ratio:
                    continue
                if not _intent_ok(cur, k):
                    continue
                cands.append((v, k))
            if not cands:
                continue
            cands.sort(reverse=True)
            best_v, best = cands[0]
            report.append({"out": cur, "out_volume": cur_v,
                           "in": best, "in_volume": best_v, "kind": "upgrade",
                           "topic": topic, "tier": out[i].get("tier", "")})
            used.discard(str(cur).lower())
            used.add(best)
            out[i] = {"service": best, "tier": out[i].get("tier", "competitive")}

    # ---- AND WHEN THERE IS NOTHING TO SWAP IN, THE SLOT GOES ---------------
    # Everything above replaces a dead service with one of the operator's unused
    # terms. NPAIHB had none left — eight typed terms, all either used or set
    # aside — so four invented names survived at no measured demand at all:
    # "tribal health resources", "native american epidemiology services",
    # "tribal health advisory services", "native american health board". Plausible
    # English, nobody's search, and they went in front of a client as things this
    # health board would be ranked for.
    #
    # A slot is worth less than nothing when it holds a phrase that does not
    # exist. Dropped rather than kept, which shortens the list and lowers the
    # price — the amber warning above the panel already says so, and a seven-term
    # quote that is true beats an eleven-term one that is padded.
    #
    # ZERO, NOT MERELY BELOW THE FLOOR. Google Ads reports 10/mo for a phrase it
    # holds no data on, so 10 is "small or unmeasured" and 0 is "the tool has
    # nothing at all". Only the second is dropped.
    #
    # Never a term the operator typed, never a pin, and never below
    # grid_min_services — a list can be honest and short without being empty.
    # (2026-08-17)
    # THE OPERATOR'S LIST AS THEY TYPED IT, NOT AS THE BUILD LEFT IT. `seeds` is
    # the right pool to swap FROM — a term the classifier set aside should not be
    # swapped in — and the wrong list to grant the exemption from. NPAIHB's
    # "federally recognized tribes pacific northwest" was set aside by the kinds
    # classifier, left `seeds`, and this pass then dropped it as an invention: the
    # operator watched their own typing removed under a line that says "nothing
    # measurable behind the phrase".
    #
    # This is the SECOND time the same narrowing has caused the same class of bug
    # — the grounding corpus had it two builds earlier and was fixed by passing
    # the typed list separately. Any pass that asks "did the operator type this?"
    # needs `typed`, never `seeds`. (2026-08-17)
    _typed = {str(x).strip().lower() for x in (typed if typed is not None else seeds) or []}
    _min = int(CFG.get("grid_min_services", 7) or 7)
    _dead = []
    for x in out:
        n = str(x.get("service") or "").strip().lower()
        if not n or x.get("pinned") or x.get("from_seed") or n in _typed:
            continue
        if vol_of(n) == 0:
            _dead.append(n)
    # THE SAME VALVE EVERY OTHER FILTER HERE HAS. This one had a floor on the
    # FINAL COUNT and no cap on how much it removed, so it took ten of seventeen
    # services in one build — 59%. The competitor-name filter, looking at the same
    # list on the same build, stood down at exactly that ratio and said so in
    # amber. Two filters, one list, opposite behaviour, and the only difference
    # was that this one did not copy the valve.
    #
    # MEASURED AGAINST THE WHOLE LIST, and the grounding filter already worked out
    # why. Its comment: "measure that against the WHOLE list, not just the model's
    # share of it. Keller contributed only 3 non-seed services and 2 were
    # competitors: a correct 2-of-3 read as a 67% drop and tripped the valve."
    # The first draft of this valve used the model's share and did exactly that —
    # NPAIHB's four inventions were four of four, so a correct removal read as
    # 100% and the padding came back. Against the whole list those four are 36%,
    # plainly targeted, while Ski Barn's ten of seventeen are 59%, which is the
    # cull this valve exists to catch. Above the ratio it reports what it WOULD
    # have cut and removes nothing. (2026-08-17)
    _judged = [x for x in out if not x.get("pinned")]
    _ratio = float(CFG.get("dead_slot_max_drop_ratio", 0.5) or 0.5)
    if _dead and _judged and len(_dead) / len(_judged) > _ratio:
        return out, report + [{"out": n, "out_volume": 0, "in": "", "in_volume": 0,
                               "kind": "would_drop", "topic": "",
                               "tier": ""} for n in _dead]
    if _dead and len(out) - len(_dead) < _min:
        # THE CHEAPEST SLOT GOES FIRST — the same ordering the topic guarantee
        # uses when it has to take one. Lowest tier first, and inside a tier the
        # longest phrase, because a long qualified phrase is the more likely
        # invention and the barer term is the one worth keeping. This was
        # alphabetical, which dropped whichever two terms happened to sort first
        # and left an invention sitting in the list. (2026-08-17)
        _rank = {"long_tail": 0, "competitive": 1, "ultra": 2}
        _tier_of = {str(x.get("service") or "").strip().lower():
                    str(x.get("tier") or "competitive") for x in out}
        _dead.sort(key=lambda n: (_rank.get(_tier_of.get(n), 1),
                                  -len(n.split()), n))
        _dead = _dead[:max(0, len(out) - _min)]
    if _dead:
        _gone = set(_dead)
        kept_out = []
        for x in out:
            n = str(x.get("service") or "").strip().lower()
            if n in _gone:
                report.append({"out": x.get("service"), "out_volume": 0,
                               "in": "", "in_volume": 0, "kind": "dropped",
                               "topic": service_topic(n, topics) if topics else "",
                               "tier": x.get("tier", "")})
                _gone.discard(n)          # one row per name
                continue
            kept_out.append(x)
        out = kept_out
    return out, report


def build_grid(services, markets, state, prepicked=False, geo_forms=None):
    """Cross each SERVICE with each CITY, in the proposal format
    ('auto insurance fairfax va'). The tier comes from the service, so every
    city inherits it. Returns {ultra:[], competitive:[], long_tail:[]}."""
    cities = list(markets) if prepicked else pick_grid_cities(markets, state, CFG["grid_max_cities"])
    suffix_mode = CFG.get("grid_state_suffix", "auto")
    buckets = {"ultra": [], "competitive": [], "long_tail": []}

    # ONE rule per grid. Brendan suffixes small or ambiguous cities but not
    # well-known metros — "auto insurance alexandria va" and "adult autism
    # services hyde pa", but "adhd treatment san diego" and "deck repair
    # knoxville". Applied city-by-city, that produced a table reading "junk
    # removal farragut" on one row and "junk removal clinton tn" on the next,
    # which reads as a mistake to whoever receives the proposal. The cities in a
    # grid are one footprint, so they get one convention: no suffix only if
    # EVERY city in the grid stands on its own name. (2026-08-10)
    _all_unmistakable = bool(cities) and all(
        name_is_unmistakable(c, state) for c in cities)

    def city_suffix(city_lower, city_state):
        """Each city uses ITS OWN state — a tri-state footprint gets
        'cherry hill nj' and 'wilmington de' in the same grid."""
        ab = STATE_ABBREV.get((city_state or "").strip().lower(), "")
        if not ab:
            return ""
        if suffix_mode is False or suffix_mode == 0:
            return ""
        if suffix_mode is True or suffix_mode == 1:
            return f" {ab}"
        return "" if _all_unmistakable else f" {ab}"          # auto
    for s in services:
        svc, tier = clean_kw(s["service"]).lower(), s["tier"]
        if not svc:
            continue
        if not cities:                     # nationwide: no crossing
            buckets[tier].append({"keyword": svc, "volume": 0,
                                  "src": "grid", "origin": "added", "service": svc})
            continue
        for city in cities:
            c_name, c_state = parse_market(city, state)
            c = c_name.strip().lower()
            svc_l = f" {svc.lower()} "
            # DMO-style seeds carry the destination INSIDE the service ("things
            # to do in central pa") — appending the market again produces
            # "central pa pennsylvania". If the service already contains this
            # market, its state name, or ends with the state abbr, keep the
            # service as-is for this crossing.
            st_of_market = (c_state or "").strip().lower() or (c if c in STATE_ABBREV else "")
            ab = STATE_ABBREV.get(st_of_market, "")
            already = (f" {c} " in svc_l
                       or (st_of_market and f" {st_of_market}" in svc_l.rstrip())
                       or (ab and svc.lower().rstrip().endswith(" " + ab)))
            if already:
                kw = svc
            else:
                # A measured form wins over the suffix rule — see pick_geo_forms.
                chosen = (geo_forms or {}).get(city) or (geo_forms or {}).get(c)
                if chosen:
                    kw = clean_kw(f"{svc} {chosen}")
                else:
                    # don't append the state if the "city" IS the state
                    sfx = "" if (c_state and c == c_state.strip().lower()) else city_suffix(c, c_state)
                    kw = clean_kw(f"{svc} {c}{sfx}")
            if any(r["keyword"] == kw for r in buckets[tier]):
                continue                      # same term from two crossings
            buckets[tier].append({"keyword": kw,
                                  "volume": 0, "src": "grid",
                                  "origin": "added", "service": svc, "city": c})
    return buckets


def claude_refine_keywords(seeds, markets, brand, domain, candidates,
                           site_terms, business_desc="", site_pages=None,
                           state=""):
    """Claude pass over the API-generated candidates: removes junk/garbled/off-topic
    terms (using the business description to exclude irrelevant services), folds in
    site-related opportunities, buckets by difficulty, and tags each term's origin
    ('kept' from the candidates, or 'added' by Claude) so the UI can show what AI did.
    Non-fatal: returns None on no key / failure, so the caller falls back to rules."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return None
    cand_terms = [c["keyword"] for c in candidates][:120]
    site_list  = [s["keyword"] for s in site_terms][:40]
    pages_list = [p for p in (site_pages or [])][:60]
    cand_lower = {c.lower() for c in cand_terms}
    mkt = ", ".join(markets) if markets else "national (no specific city)"
    biz = business_desc or "(NOT PROVIDED — infer it yourself from the vertical, website, pages, and site keywords below, and return it in the 'business' field)"
    pages_block = (json.dumps(pages_list, ensure_ascii=False) if pages_list
                   else "(no page structure available)")
    prompt = f"""You are an SEO strategist refining a keyword list for a client proposal. Be strict about relevance to THIS specific business.

WHAT THE BUSINESS DOES (and does not do): {biz}
CLIENT VERTICAL / SERVICES: {", ".join(seeds)}
TARGET MARKET(S): {mkt}
CLIENT BRAND (exclude any keyword containing this): {brand or "(none given)"}
CLIENT WEBSITE: {domain or "(none given)"}

THE CLIENT'S ACTUAL WEBSITE PAGES (their real service taxonomy — each page is a topic they offer and should rank for):
{pages_block}

CANDIDATE KEYWORDS (from a keyword API — contain junk, garbled terms, near-duplicates, and OFF-TARGET terms for services this business does not offer):
{json.dumps(cand_terms, ensure_ascii=False)}

KEYWORDS THE SITE ALREADY RANKS FOR:
{json.dumps(site_list, ensure_ascii=False)}

RULES:
1. EXCLUDE terms for services the business does NOT offer. (Example: a therapy practice that does not prescribe drugs should NOT have "medication", "prescription", or "over the counter" keywords.)
2. EXCLUDE garbled/nonsensical terms ("adhd and therapy", "add therapy" when the vertical is "adhd treatment"), near-duplicates, and brand terms.
3. KEEP real searches a prospective customer of THIS business would type.
4. USE THE WEBSITE PAGES as your primary guide to what this business actually offers. For each real service page, ensure there is a strong head keyword targeting it (geo-modified where local). ADD any the candidate list missed — these are high-priority SEO opportunities.
5. ADD other high-value keywords this business should target, consistent with their pages and services.
6. Keep the city modifier on local-intent terms where the market is local.
7. BALANCE THE VOCABULARY: the ultra/competitive buckets must carry the everyday words customers actually type (for a therapy practice: "therapist [city]", "therapy [city]", "counseling [city]", "mental health services [city]") — these hold the search volume. Clinical, technical, or page-template phrasings ("[condition] treatment [city]") belong in long_tail, and no single template word should dominate the list. If the seeds themselves are templated, FIX the vocabulary rather than propagating it.
8. Bucket by ranking difficulty: "ultra" (hardest/highest value), "competitive" (moderate), "long_tail" (longer/question-style).
9. Do NOT invent search volumes. Only real, searchable terms.

Return ONLY valid JSON in exactly this shape. Each keyword item is [keyword, origin] where origin is "kept" or "added". The "business" field is your one-sentence read of what the business does and does not offer:
{{"business": "one sentence", "ultra": [["keyword","kept"], ...], "competitive": [["keyword","added"], ...], "long_tail": [["keyword","kept"], ...]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key,
                     "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 2500,
                "temperature": 0,
                "messages": [{"role": "user", "content": prompt}],
            }), timeout=30)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", []) if b.get("type") == "text")
        text = text.strip()
        if text.startswith("```"):
            text = text.split("```", 2)[1]
            if text.startswith("json"):
                text = text[4:]
        parsed = json.loads(text.strip())
        def rows(key):
            out = []
            for item in parsed.get(key, []):
                if isinstance(item, list) and item:
                    kw = str(item[0]).strip()
                    origin = item[1] if len(item) > 1 else "kept"
                elif isinstance(item, str):
                    kw = item.strip(); origin = "kept"
                else:
                    continue
                if not kw:
                    continue
                # trust the model's tag but sanity-check against the candidate set
                if origin not in ("kept", "added"):
                    origin = "added" if kw.lower() not in cand_lower else "kept"
                out.append({"keyword": kw, "volume": 0, "src": "claude", "origin": origin})
            return out
        return {"ultra": rows("ultra"), "competitive": rows("competitive"),
                "long_tail": rows("long_tail"),
                "business": (parsed.get("business") or "").strip()}
    except Exception:
        return None



# ---------------------------------------------------------------------------
# STAGE 1 — keyword list
# ---------------------------------------------------------------------------
def stage1_keyword_list(seeds, markets, state, brand, domain="", business_desc=""):
    crossed = []
    for s in seeds:
        crossed.append(s)
        for m in markets:
            crossed.append(f"{s} {m}")
        if state:
            crossed.append(f"{s} {state}")

    payload = [{"keywords": crossed[:200],
                "location_name": loc_string(markets, state),
                "language_code": "en"}]
    data = dfs_post("/keywords_data/google_ads/keywords_for_keywords/live", payload)
    items = (data["tasks"][0]["result"] or [])
    raw = [{"keyword": it["keyword"], "volume": it.get("search_volume") or 0, "src": "ideas"}
           for it in items]

    # Add keyword_suggestions (longer, seed-containing phrases) into the pool
    for r in fetch_suggestions(seeds, markets, state):
        r["src"] = "suggest"; raw.append(r)
    # Add keywords_for_site (terms relevant to the client's domain) into the pool
    for r in fetch_keywords_for_site(domain, markets, state):
        r["src"] = "site"; raw.append(r)

    seed_tokens = {t.lower() for s in seeds for t in s.split()}
    brand_l = (brand or "").lower()
    # Connector words that signal a stitched-together / garbled phrase rather
    # than a real search query ("adhd and therapy", "treatment or counseling").
    CONNECTORS = {"and", "or", "&", "vs", "with"}
    def is_junk(kw):
        toks = kw.split()
        for i, t in enumerate(toks):
            if 0 < i < len(toks) - 1 and t in CONNECTORS:
                return True
        return False
    kept = []
    seen = set()
    for r in raw:
        kw = r["keyword"].lower()
        if kw in seen:
            continue
        seen.add(kw)
        if brand_l and brand_l in kw:
            continue
        if is_junk(kw):
            continue
        # Seed-token relevance filter applies to seed-derived sources only.
        # Site keywords come from the client's own domain and are on-topic by
        # construction, so they bypass it (but still drop the brand name above).
        if r.get("src") != "site" and seed_tokens and not (seed_tokens & set(kw.split())):
            continue
        kept.append(r)

    # Tie-break on the term itself: volumes tie constantly at the thin end
    # (10, 40, 70/mo) and the API's own ordering is not stable between runs,
    # so without this the candidate pool — and therefore the keyword list —
    # reshuffles for reasons that have nothing to do with the client.
    kept.sort(key=lambda r: (-(r["volume"] or 0), r["keyword"]))
    with_vol = [r for r in kept if r["volume"] > 0]

    # THE MARKET'S OWN WORDS, READ HERE BECAUSE HERE IS THE ONLY PLACE THEY CAN
    # BE READ. The refine step rebuilds its candidate rows from what the browser
    # posts back — {kw, vol} and nothing else — so `src` and every other trace of
    # where a row came from is gone by the time the qualifier check runs. It was
    # therefore reading "apartments in san antonio tx" and "apartments for rent
    # los angeles" as this market's vocabulary: those come from
    # keywords_for_site, which is a Labs call at location_code 2840 — the whole
    # United States — while everything else here is measured at the client's own
    # city. Three builds of this filter were judging Santa Fe by Los Angeles
    # numbers. Computed once, carried forward as a token list. (2026-08-16)
    _local = [r for r in kept if r.get("src") != "site"]
    market_vocab = sorted(pool_vocabulary(_local))
    market_pool = [{"keyword": r["keyword"], "volume": r["volume"]}
                   for r in _local[:int(CFG.get("market_pool_cap", 60))]]

    # Bucket sizes come from the measured tier mix, not fixed counts — see
    # CFG["tier_mix"] for the eight-proposal sample behind the proportions.
    _target = min(int(CFG.get("list_cap", 60) or 60), max(len(kept), 20))
    u, c, _lt = tier_split(_target)
    n_head = u + c

    if markets:
        # GEO-SCOPED: head terms are seed × market combinations ("adhd treatment
        # san diego") — the form the proposals actually use. We build these
        # directly from the crossing rather than relying on the API to return
        # them (it strips geo and inflates bare national terms). Volume is looked
        # up where available but NOT required, since local terms often report low
        # or zero volume in keyword tools yet are exactly what we rank/quote on.
        vol_lookup = {r["keyword"].lower(): r["volume"] for r in kept}
        geo_heads, seen_h = [], set()
        # (a) direct seed × market crossings
        seed_phrases = list(seeds)
        # (b) plus the API's related head terms, geo-modified — this is what gives
        # the proposal its variety ("adhd therapy san diego", "couples therapy
        # san diego") beyond the literal seeds. Drawn from the FULL candidate pool
        # (not volume-filtered) so sparse/niche verticals — where local terms
        # report little or no volume — still build a full list instead of
        # collapsing to the bare seeds. (This is the Versability case.)
        # Related expansion terms must share a SUBSTANTIVE seed token (length >= 4)
        # with the seeds. This drops loose API associations and garbled near-words
        # like "add therapy" (the seed was "adhd treatment" — "add" is only 3 chars
        # and isn't a seed word) while keeping real expansions ("adhd therapy").
        seed_long_tokens = {t.lower() for s in seeds for t in s.split() if len(t) >= 4}
        def shares_substantive_seed(kw):
            return bool(seed_long_tokens & set(kw.lower().split()))
        related = [r["keyword"] for r in kept
                   if not is_longtail(r["keyword"])
                   and not any(m.lower() in r["keyword"].lower() for m in markets)
                   and shares_substantive_seed(r["keyword"])]
        seed_phrases += related[:25]
        for s in seed_phrases:
            for m in markets:
                kw = f"{s} {m}".strip()
                kl = kw.lower()
                if kl in seen_h or (brand_l and brand_l in kl):
                    continue
                seen_h.add(kl)
                # rank these by the volume of their BARE form (local volume is
                # usually unreported, but bare volume signals term importance)
                bare_vol = vol_lookup.get(s.lower(), 0)
                geo_heads.append({"keyword": kw, "volume": bare_vol, "src": "geo"})
        # strongest terms first (by bare-form volume)
        geo_heads.sort(key=lambda r: (-(r["volume"] or 0), r["keyword"]))
        # backfill with any remaining bare terms (volume or not) if still short
        bare_backfill = [r for r in kept if not is_longtail(r["keyword"])
                         and r["keyword"].lower() not in seen_h]
        head_ordered = geo_heads + bare_backfill
    else:
        # NATIONAL: no geo modifier; rank bare head terms by volume.
        head_ordered = [r for r in with_vol if not is_longtail(r["keyword"])]

    ultra       = head_ordered[:u]
    competitive = head_ordered[u:u + c]
    head_kws    = {r["keyword"] for r in ultra + competitive}

    # LONG-TAIL bucket: explicitly long / question-shaped phrases, deduped,
    # not already used as a head term. Longer phrases preferred.
    lt_candidates = [r for r in kept
                     if is_longtail(r["keyword"]) and r["keyword"] not in head_kws]
    # prefer more words, then higher volume
    lt_candidates.sort(key=lambda r: (-len(r["keyword"].split()),
                                      -(r["volume"] or 0), r["keyword"]))
    long_tail = lt_candidates[:CFG["longtail_target"]]

    # Backfill: if the API returned few real long-tails (common in local/niche
    # verticals), generate question-form long-tails from the seeds + market so
    # the bucket is never empty at Step 1. PAA harvested in Step 3 will add more.
    if len(long_tail) < CFG["longtail_target"]:
        seen_lt = {r["keyword"].lower() for r in long_tail} | {k.lower() for k in head_kws}
        mkt = markets[0] if markets else ""
        templates = ["how much does {s} cost{inm}", "best {s} near me",
                     "what to look for in {s}{inm}", "affordable {s} for adults{inm}",
                     "is {s} covered by insurance{inm}", "how to find a good {s}{inm}"]
        for s in seeds:
            for t in templates:
                if len(long_tail) >= CFG["longtail_target"]:
                    break
                kw = t.format(s=s, inm=(f" in {mkt}" if mkt else "")).strip()
                kl = kw.lower()
                if kl in seen_lt or (brand_l and brand_l in kl):
                    continue
                seen_lt.add(kl)
                long_tail.append({"keyword": kw, "volume": 0, "src": "gen"})

    # ---- Claude refinement pass (Option 2: API generates, Claude refines) ----
    site_terms = [r for r in raw if r.get("src") == "site"]
    # BUILD stops here (fast: keyword API + rules only). The Claude refinement and
    # exact-match volume run in a SEPARATE request (stage1b_refine) so neither
    # half can exceed the platform request timeout on heavy verticals.
    full = (ultra + competitive + long_tail)[:CFG["list_cap"]]
    fs = {r["keyword"] for r in full}
    return {
        "ultra":       [r for r in ultra if r["keyword"] in fs],
        "competitive": [r for r in competitive if r["keyword"] in fs],
        "long_tail":   [r for r in long_tail if r["keyword"] in fs],
        "head":        [r for r in (ultra + competitive) if r["keyword"] in fs],
        "all":         full,
        "refined_by_ai": False,
        "business_desc": "",
        "site_pages_found": 0,
        "site_terms":  [r["keyword"] for r in site_terms],   # passed to refine step
        "market_vocab": market_vocab,        # both passed to refine step, because
        "market_pool":  market_pool,         # the pool itself cannot survive the trip
    }

def stage1b_refine(seeds, markets, state, brand, domain, business_desc,
                   ultra, competitive, long_tail, site_terms_kw, phrase_geos=None,
                   national_demand=False, goal="", band="",
                   national_reason="", grid_axis="", industry="",
                   product_demand=False, suggested=None, negatives=None,
                   ranked=None, market_vocab=None, market_pool=None):
    """Second half of Step 1, run as its own request: reads the sitemap, runs the
    Claude refinement pass, and re-pulls exact-match volume. Takes the raw buckets
    from stage1_keyword_list. Kept separate so a heavy Claude call can't time out
    the list build."""
    site_terms = [{"keyword": k} for k in (site_terms_kw or [])]
    scope_note = ""
    _site_urls = []
    site_pages = fetch_site_pages(domain, collect_urls=_site_urls)
    site_locations = location_pages_from_urls(_site_urls)
    # A storefront on the client's own site used to flip national demand on
    # unconditionally. This has to happen HERE — before the volume pull below —
    # or the flip would only take effect on a second run, and the quote in front
    # of the operator would still be priced on per-city volume.
    #
    # But the flip itself was too eager. BE on the Ski Barn quote: "They have
    # specific locations, the quote didn't consider the local impact, it just
    # treated them as an ecommerce/nationwide option." Almost every retailer
    # sells online now, so "has a cart" stopped being evidence of a national
    # campaign — and the consequences all point the same way: national volume,
    # geo-less keywords, and a rank check the client cannot pass.
    #
    # Named markets outrank a detected cart, because the operator typed them on
    # purpose. So the flip only fires when there is nothing local to price
    # against — no markets, or an explicitly nationwide scope. Otherwise the
    # storefront is REPORTED and the pull stays local: geo-qualified terms,
    # per-city volume, rankings in the client's own market. (2026-08-07)
    ecom_found, ecom_reason = detect_ecommerce(_site_urls)
    ecom_suppressed = ""
    national_demand_reason = ""
    if ecom_found and not national_demand:
        if markets and (band or "") != "nationwide":
            ecom_suppressed = (
                f"Storefront detected ({ecom_reason}) — but this client has "
                f"{len(markets)} market{'' if len(markets) == 1 else 's'} entered"
                + (f", which read as a {band.replace('_', ' ')} footprint"
                   if band else "")
                + ", so demand is still being pulled LOCALLY. A store that also "
                  "ships is not a national campaign. Turn on Price on national "
                  "demand only if this really is a product brand selling "
                  "everywhere with no local trade to win.")
        else:
            national_demand = True
            national_demand_reason = f"storefront detected — {ecom_reason}"
    # A Google Business listing is stronger evidence of operating in a market
    # than a page on the website, and it is the only one that works when the
    # site uses a store-locator widget. Merge, don't replace — a client can
    # have a location page without a listing and vice versa.
    # The client's own service-area list is the closest thing to the market
    # list a proposal is scoped from — better than listings for anyone whose
    # markets outnumber their premises.
    service_areas = fetch_service_areas(domain)
    for a in service_areas:
        if a.lower() not in {l.lower() for l in site_locations}:
            site_locations.append(a)
    gbp_cities, gbp_count = google_business_cities(brand, domain)
    for c in gbp_cities:
        if c.lower() not in {l.lower() for l in site_locations}:
            site_locations.append(c)
    # THE OPERATOR NO LONGER TYPES THIS (2026-08-13). The field came off the form
    # and the build fills it in, because the description is not really an input —
    # it is the vocabulary sample the grounding filter needs, and it was being
    # asked of a human who could only copy it off the site the tool already reads.
    #
    # It has to happen HERE, before the filter runs. infer_business existed but
    # nothing called it: the only inference lived inside claude_refine_keywords
    # and came back in its 'business' field, which never reached the corpus. So a
    # blank description meant the filter judged every proposal against seeds and
    # page titles alone — the thin-corpus case that stood the whole thing down on
    # NPAIHB. Two sources, cheapest first: their own site, then a SERP read for
    # the clients whose site refuses to be crawled at all.
    biz = business_desc.strip() if business_desc else ""
    biz_inferred = ""
    if not biz:
        try:
            biz = (infer_business(domain, seeds, site_terms, industry,
                                  site_pages) or "").strip()
            biz_inferred = biz
        except Exception:
            app.logger.exception("infer_business failed")
        if not biz and (brand or domain):
            # Site unreadable (403, JS-only, dead host). Their name and market are
            # on the order and the answer is on page one of Google.
            try:
                _q = " ".join(x for x in [brand, (markets or [""])[0]] if x).strip()
                _sn = serp_snippets(_q, loc_string(markets, state)) if _q else []
                _bd = claude_business_desc(brand, markets, industry, _sn, domain) or {}
                biz = (_bd.get("text") or "").strip()
                biz_inferred = biz
            except Exception:
                app.logger.exception("claude_business_desc failed during build")

    # ---- GRID MODE: build a service x city grid like the real proposals -----
    if CFG.get("grid_mode"):
        cands = ultra + competitive + long_tail
        seed_ranking = {}
        seeds_demoted = []
        seeds_folded = []
        seeds_dropped_suggested = []
        negatives_dropped = []
        services_deduped = []
        negative_conflicts = negative_seed_conflicts(seeds, negatives)
        # Decide the city set FIRST so the service count can scale to it.
        city_pick = {}
        cities = pick_grid_cities(markets, state, CFG["grid_max_cities"],
                                  probe_term=list(seeds or []),
                                  explain=city_pick,
                                  # Location pages and named service areas are
                                  # direct evidence of where the client
                                  # actually operates — better than the domain
                                  # alone, which only ever names one town.
                                  home_hint=" ".join(
                                      [brand or "", domain or ""]
                                      + [str(x) for x in (site_locations or [])]
                                      + [str(x) for x in (service_areas or [])]))
        # Search-phrase geos ("south jersey", "fox cities") cross into keyword
        # TEXT exactly like cities, but never touch a location API — no volume
        # lookup, no validation, no rank-check location. Keeps Brendan-style
        # regional phrasing without the invalid-location fallout.
        # ---- WHICH AXIS GETS THE TERM BUDGET --------------------------------
        # pick_grid_cities has already measured per-market demand for this
        # client's own service, so the evidence is in hand: if only one market
        # has anything to buy, crossing the others spends slots on near-
        # duplicates instead of on the services the client actually sells.
        # (2026-08-10)
        axis, axis_reason, axis_ev = choose_grid_axis(
            city_pick.get("kept") or [], len(seeds or []),
            forced=(grid_axis or ""))
        if axis == "services" and cities:
            axis_ev["dropped_cities"] = [c for c in cities[1:]]
            cities = cities[:1]
            city_pick["axis_trimmed"] = True
        phrases = [p.strip() for p in (phrase_geos or []) if p and p.strip()]
        seen_c = {c.strip().lower() for c in cities}
        grid_cities = cities + [p for p in phrases if p.lower() not in seen_c]
        # NATIONAL DEMAND (2026-07-25): a product brand's keywords carry no
        # city. "energy gummies texas" is not a search anyone runs for a DTC
        # supplement, and pairing a national volume figure with a geo-suffixed
        # term misrepresents what the number counts. So the grid stops crossing
        # and becomes a flat national service list — build_grid already has
        # that path (it emits the bare service when cities is empty). The
        # client's geos stay on the order as their targeting area; they just
        # don't enter the keyword text or the volume lookup.
        if national_demand:
            grid_cities = []
        # ---- SEEDS IN DEMAND ORDER, before anything reads them ----------
        # ONE FEWER CALL WHEN IT CANNOT MATTER. Ranking the seeds costs a
        # search_volume call, and DataForSEO allows twelve a minute — adding one to
        # every build is what pushed Ski Barn over the limit. When there are no
        # more seeds than slots, every seed is quoted whatever the order, so the
        # measurement buys nothing. Measure only when the ranking DECIDES
        # something. (2026-08-12)
        _slots = services_needed(len(grid_cities))
        # ---- IS EVERY SEED A SERVICE THIS CLIENT SELLS? ------------------
        # Runs on EVERY build, unlike the ranking below it. The check used to be
        # reachable only through rank_seeds(), which is gated on the DataForSEO
        # allowance — so NPAIHB (8 seeds, 20 slots) skipped it, and nine member
        # tribes the heading miner scraped off npaihb.org became the quote.
        # claude_seed_kinds is an Anthropic call; the gate never applied to it.
        # (2026-08-12)
        _kinds = {}
        if seeds:
            try:
                _kinds = claude_seed_kinds(
                    seeds, brand, domain, industry, biz, site_pages,
                    # NOT national_demand — that is true whenever no geos are
                    # entered, which would switch the item verdict off for a
                    # local hauler mid-setup. Only the operator's own product
                    # tick, or a storefront in the URLs, means "they sell things".
                    sells_products=bool(product_demand or ecom_found))
            except Exception:                             # noqa: BLE001
                app.logger.exception("claude_seed_kinds failed during build")
        # WHICH OF THE FLAGGED NAMES ARE ACTUALLY RIVALS. Asked once, only about
        # terms the classifier already called another business, and only when a
        # ranking would otherwise exempt them — so on most builds this costs
        # nothing. Without it the ranked exemption spares a competitor for the
        # same reason it spares a carrier, and PEO Brokers got "prime peo
        # brokers" quoted at 30/mo. (2026-08-14)
        _rivals = set()
        if _kinds and ranked:
            try:
                _rivals = rival_terms(
                    [t for t in seeds
                     if str(t).strip().lower() in
                     {str(x).strip().lower() for x in (ranked or [])}],
                    _kinds, brand, domain, industry, biz)
            except Exception:
                app.logger.exception("rival_terms failed")
        # THE TOOL'S OWN PROPOSALS GO FIRST, and they get no exemption. Order
        # matters and it was wrong: demote_nonservices ran ahead of this and
        # removed ANY seed the classifier called another business — including a
        # suggested one — into seeds_demoted, whose panel line was removed on
        # request. So PEO Brokers' two company names split between a silent path
        # and a reported one, the ✂ line said "1 taken back off" when two had
        # gone, and the demoted pill stayed in the box to be re-quoted on the
        # next build. A suggested term is the tool's to retract, so it is
        # retracted here, named on the panel, and its pill goes with it. What
        # the operator typed still only ever gets DEMOTED below. (2026-08-13)
        # WHAT THE OPERATOR TYPED, KEPT SEPARATELY FROM WHAT SURVIVED. `seeds` is
        # narrowed below — set aside by the classifier, folded, ranked out — and
        # it is also the corpus the grounding filter judges vocabulary against.
        # So a focus term removed by one filter stopped counting as the client's
        # own words, and any service still using those words then read as
        # foreign: NPAIHB had "federally recognized tribes pacific northwest"
        # set aside as not-a-service, and the same phrase came back a second time
        # under "6 unrecognised terms removed" — one term, two removals, two
        # different explanations, and the second one wrong. Somebody who knows
        # the account typed those words; that is true whatever any filter later
        # decides about the term. (2026-08-16)
        seeds_typed = list(seeds or [])
        if seeds and suggested:
            _sk, _sd = drop_suggested_nonservices(seeds, suggested, _kinds,
                                                  brand, markets, state, ranked,
                                                  _rivals)
            if _sd:
                seeds_dropped_suggested = _sd
                seeds = _sk
        if _kinds and seeds:
            _keep, _dem = demote_nonservices(seeds, _kinds, markets, state,
                                             ranked, _rivals)
            if _dem:
                seeds_demoted = _dem
                seeds = _keep
        # FOLD ALWAYS; RANK ONLY WHEN IT DECIDES SOMETHING. (2026-08-13)
        if seeds:
            _fk, _ff = fold_seed_duplicates(seeds, markets, state)
            if _ff and _fk:
                seeds_folded = [[a, b] for a, b in _ff]
                seeds = _fk
        if seeds and len(seeds) > _slots:
            _sr = rank_seeds(seeds, markets, state, national=national_demand,
                             limit=len(seeds), kinds=_kinds)
            if _sr.get("measured"):
                _ordered = [r["term"] for r in _sr["kept"]]
                _folded = [f["term"] for g in _sr["folded"] for f in g["fold"]]
                # Folded synonyms go to the BACK rather than out: the fold is a
                # ranking device, and a client who really does want both wordings
                # keeps them if the grid has room.
                _have = set(_ordered) | set(_folded)
                _tail = []
                for _t in seeds:
                    _k = seed_norm(_t, markets, state)
                    if _k and _k not in _have:
                        _have.add(_k)
                        _tail.append(_k)
                if _ordered:
                    seed_ranking = {
                        "basis": _sr.get("basis", ""),
                        "adjacent": _sr.get("adjacent") or [],
                        # [term, local demand, demand for the QUOTED form]. The
                        # third figure is what did the ordering; the second is
                        # what prices the volume component. They disagree by
                        # about 7x and the panel has to be able to show why a
                        # term outranked one with a bigger headline number.
                        "order": [[r["term"], r["volume"], r.get("geo_volume", 0)]
                                  for r in _sr["kept"]],
                        "order_basis": _sr.get("order_basis", ""),
                        "geo_measured": bool(_sr.get("geo_measured")),
                        "geo_ordered": int(_sr.get("geo_ordered") or 0),
                        "geo_error": _sr.get("geo_error", ""),
                        "folded": _sr.get("folded") or [],
                        "was": list(seeds),
                    }
                    seeds = _ordered + _folded + _tail
                    # ---------------------------------------------------------
                    # ROOM TO GROW. Ski Barn's quote came back ranking in the
                    # top 100 for 19 of 19 terms — nine of them in the top four
                    # — and Brendan's note was the right one: "a lot of these
                    # keywords they are ranking for, it might be good to have
                    # the tool find a few more that they don't rank for."
                    #
                    # Ranking purely by measured demand does this on any client
                    # with existing SEO: the terms they already own are the ones
                    # with volume attached, so they win every slot, and the
                    # proposal argues for a campaign to win what is already won.
                    #
                    # So a few slots are reserved for terms they do NOT rank
                    # for, taken in demand order from the ones the ranking cut.
                    # They are the campaign's actual upside and the honest
                    # answer to "what am I buying". A RESERVATION, NOT A QUOTA:
                    # nothing is invented to fill it, and if every candidate is
                    # already ranked the grid is unchanged. (2026-08-19)
                    _room = int(CFG.get("grid_headroom_slots", 4) or 0)
                    if _room and _slots:
                        _vol = {r["term"]: r.get("volume") or 0
                                for r in _sr["kept"]}
                        # WHAT THEY ALREADY RANK FOR, FROM SOMETHING THAT KNOWS.
                        # `ranked` only carries terms the operator ADDED from the
                        # rankings panel during expansion — Ski Barn's 59 terms
                        # came off a product line and a report, so it was empty,
                        # every term looked like headroom, the reservation was
                        # already satisfied and the whole thing silently did
                        # nothing. The client's positions are a fact about the
                        # client, not about how the list was assembled, so ask
                        # for them. One Labs call, and only when a domain is
                        # known and the panel gave us nothing. (2026-08-19)
                        _own, _own_src = set(), "the rankings panel"
                        for _r in (ranked or []):
                            _rt = (_r.get("bare") or _r.get("term") or ""
                                   if isinstance(_r, dict) else str(_r))
                            _rk = seed_norm(_rt, markets, state)
                            if _rk:
                                _own.add(_rk)
                        if not _own and domain:
                            try:
                                _rk_rows = fetch_ranked_keywords(
                                    domain, markets, state,
                                    CFG.get("ranked_keywords_own_limit", 1000))
                                for _r in _rk_rows:
                                    if _r.get("bare"):
                                        _own.add(_r["bare"])
                                _own_src = (f"{len(_rk_rows)} ranked keywords "
                                            f"read for {domain}")
                            except Exception as _e:           # noqa: BLE001
                                app.logger.warning(
                                    "headroom: ranked_keywords failed (%s)", _e)
                                _own_src = f"could not be read ({str(_e)[:80]})"
                        if not _own:
                            # NOTHING KNOWN IS NOT THE SAME AS NOTHING RANKED.
                            # With an empty set every term reads as headroom and
                            # the reservation is trivially satisfied, which is
                            # exactly how this shipped doing nothing. Say so.
                            seed_ranking["headroom_skipped"] = _own_src
                        _head, _tailseeds = seeds[:_slots], seeds[_slots:]
                        _fresh = [t for t in _head if t not in _own]
                        _want = min(_room, _slots)
                        # ALWAYS REPORT. This has now shipped twice looking like
                        # it worked, because every outcome except a swap was
                        # silent: no rankings to compare against, reservation
                        # already met, nothing left to promote. Each of those is
                        # a different answer and the panel has to distinguish
                        # them or the next silent no-op costs another day.
                        seed_ranking["headroom_seen"] = {
                            "want": _want, "fresh": len(_fresh),
                            "own": len(_own), "pool": len(_tailseeds),
                            "source": _own_src}
                        if len(_fresh) < _want:
                            # best unranked candidates the cut left behind
                            _cand = [t for t in _tailseeds if t not in _own]
                            _cand.sort(key=lambda t: -_vol.get(t, 0))
                            _need = _want - len(_fresh)
                            _promote = _cand[:_need]
                            if not _cand:
                                seed_ranking["headroom_dry"] = (
                                    "every candidate the ranking cut is also a "
                                    "term they already rank for")
                            if _promote:
                                # displace the LOWEST-demand already-ranked
                                # terms, never a term the ranking put on top
                                _drop = [t for t in reversed(_head)
                                         if t in _own][:len(_promote)]
                                _dropset = set(_drop)
                                _head = [t for t in _head if t not in _dropset]
                                _head = _head + _promote
                                seeds = _head + [t for t in _tailseeds
                                                 if t not in set(_promote)] + _drop
                                seed_ranking["headroom"] = [
                                    [t, _vol.get(t, 0)] for t in _promote]
                                seed_ranking["headroom_displaced"] = [
                                    [t, _vol.get(t, 0)] for t in _drop]
                                seed_ranking["headroom_basis"] = _own_src
                        elif _own:
                            seed_ranking["headroom_met"] = len(_fresh)
            else:
                seed_ranking = {"failed": _sr.get("error") or "no volume data",
                                "was": list(seeds)}

        # WHAT THE OPERATOR'S OWN SEEDS ARE WORTH. Seeds are exempt from every
        # filter — someone who knows the account typed them, and that rule stays.
        # But nothing told them what they were spending slots on: PEO Brokers put
        # five "... workers comp peo" terms that measure nothing and four LCF
        # abbreviation lookups into a twenty-term quote, and the panel reported
        # neither. A warning, not a filter. (2026-08-13)
        seed_quality = {"zero": [], "question": []}
        try:
            _sv = {}
            if seeds:
                _sv, _pc2, _ = fetch_local_volume(
                    [seed_norm(x, markets, state) for x in seeds if x],
                    [] if national_demand else markets, state,
                    national=national_demand)
            for _s in (seeds or []):
                _n = seed_norm(_s, markets, state)
                if not _n:
                    continue
                if is_lookup_kw(_n):
                    seed_quality["question"].append(_s)
                elif not int((_sv or {}).get(_n, 0) or 0):
                    seed_quality["zero"].append(_s)
        except Exception:                                 # noqa: BLE001
            app.logger.exception("seed quality read failed")
            seed_quality = {"zero": [], "question": []}

        n_services = services_needed(len(grid_cities))
        services = claude_expand_services(seeds, biz, site_pages, brand, domain,
                                          cands, n_services,
                                          0 if national_demand else len(cities),
                                          national=national_demand, goal=goal)
        if not services:
            # fall back to the partner's seeds, spread across tiers
            tiers = ["ultra", "ultra", "competitive", "long_tail"]
            services = [{"service": s.strip().lower(), "tier": tiers[min(i, 3)]}
                        for i, s in enumerate(seeds[:n_services])]
        # Pricing must not swing on a non-deterministic model call: force the
        # highest-volume terms the search API returned into the list.
        # Bare the services BEFORE pinning so the containment check compares
        # like with like, and before the grid so nothing carries a geo into
        # the crossing.
        services = scrub_services(services, markets, state, phrase_geos)
        services, geo_dropped = drop_foreign_geo_services(services, markets, state)
        # QUALIFIERS THE MARKET DOES NOT USE. The prompt has always been handed
        # the keyword pool as "evidence of real demand" and advisory is what it
        # stayed — Amare spent five of twenty slots on "rental homes with washer
        # dryer", "move in ready rentals", "rental homes no credit check",
        # "with pool" and "with yard", none of which anyone types, while
        # Brendan's list for the same client spends them on bedroom counts and
        # two named landmarks. Enforced here rather than asked for, because a
        # prompt cannot be checked and this can. (2026-08-16)
        pool_dropped, pool_added, pool_status = [], [], ""
        services, pinned = pin_head_services(services, cands, markets, state,
                                             brand, n_services)
        # A PIN IS THE HIGHEST-LEVERAGE SLOT AND NOTHING WAS CHECKING IT. Pins
        # are pulled straight from the keyword pool by volume, and the pool is
        # ranked nationally, so the biggest names in the trade sit at the top of
        # it looking exactly like head terms. PEO Brokers got "the hartford
        # workers compensation insurance" pinned into Ultra at 2,900/mo — a
        # carrier's brand, in front of a client. The grounding filter could not
        # catch it either, because the client's own site names the carriers they
        # place with, so "hartford" reads as their vocabulary. A pin's whole job
        # is to hold the price steady; a rival's name doing that job is worse
        # than an empty slot. (2026-08-14)
        pins_refused = []
        if pinned:
            try:
                _pk = claude_seed_kinds(list(pinned), brand, domain, industry,
                                        biz, site_pages)
                _pr = rival_terms(list(pinned), _pk, brand, domain, industry, biz)
                if _pr:
                    _lost = {str(x).strip().lower() for x in _pr}
                    pins_refused = [[t, ((_pk.get(str(t).strip().lower()) or {})
                                         .get("why") or "a competitor")]
                                    for t in pinned
                                    if str(t).strip().lower() in _lost]
                    services = [x for x in services
                                if str(x.get("service", "")).strip().lower()
                                not in _lost]
                    pinned = [t for t in pinned
                              if str(t).strip().lower() not in _lost]
            except Exception:
                app.logger.exception("pin rival check failed")
        services = scrub_services(services, markets, state, phrase_geos)
        # Pinning pulls straight from the keyword-idea pool, which is exactly
        # where out-of-area terms live — so a term filtered out above can be
        # re-inserted below it. Filter again AFTER pinning and fold the two
        # result sets together; a pin is not a licence to sell in a state the
        # client doesn't operate in.
        services, seed_used, seed_total = (enforce_seed_services(
                                        services, seeds, n_services,
                                        markets, state, phrase_geos)
                               if seeds else (services, 0, 0))
        services, geo_dropped2 = drop_foreign_geo_services(services, markets, state)
        # QUALIFIERS THE MARKET DOES NOT USE — and it has to run HERE, after
        # enforce_seed_services, not before it. Every term this filter took out
        # was a seed, and enforce_seed_services rebuilds the list FROM the seeds:
        # it put all seven straight back. The panel reported seven drops and
        # seven refills and the quote shipped with neither, which is worse than
        # the filter not existing — a report nobody can act on, about work that
        # did not happen. (2026-08-16)
        #
        # OFF BY DEFAULT until the backfill is trustworthy. Its first outing
        # offered "new york city apartments for rent", "apartments in san
        # antonio tx" and "budget rental santa fe" — the car hire company — for
        # a Santa Fe rental community, because it ranked the keyword pool by
        # NATIONAL volume, which is the exact trap drop_foreign_geo_services and
        # the competitor rules were written for.
        _pq = CFG.get("pool_qualifier_filter")
        if _pq:
            # REPORT MODE. Four builds were spent on this feature, and the one
            # thing that would have caught every failure was seeing its verdicts
            # against the real pool before they touched a quote. In "report" it
            # decides everything and applies nothing. (2026-08-16)
            _dry = (str(_pq).lower() == "report")
            _before = [dict(x) for x in services]
            _typed = [t for t in (seeds or [])
                      if str(t).strip().lower() not in
                      {str(x).strip().lower() for x in (suggested or [])}]
            # The vocabulary and the refill pool BOTH come from stage 1, where
            # the rows still knew where they were measured. `cands` here is the
            # browser's round-trip — keyword and volume, no provenance — so it
            # cannot be used for either. (2026-08-16)
            services, pool_dropped, pool_status = drop_ungrounded_qualifiers(
                services, cands, seeds, site_terms_kw, pinned=pinned,
                suggested=suggested, vocab=set(market_vocab or []))
            if pool_dropped and market_pool:
                services, pool_added = backfill_services(
                    services, market_pool, len(pool_dropped), markets, state,
                    brand, vocab=set(market_vocab or []))
                # Anything the backfill brought in goes through the out-of-area
                # filter like everything else. It is drawing from the same
                # nationally-ranked pool that produced "state of california fire
                # insurance" for a Virginia insurer.
                services, _bf_geo = drop_foreign_geo_services(services, markets,
                                                              state)
                if _bf_geo:
                    _gone = {d[0] for d in _bf_geo}
                    pool_added = [a for a in pool_added if a[0] not in _gone]
                    if not _dry:
                        geo_dropped2 = list(geo_dropped2 or []) + list(_bf_geo)
            if _dry:
                services = _before          # decided everything, applied nothing
                pool_status = "dry:" + str(pool_status or "")
        # Counted HERE, not off the final payload: several more filters run below
        # and "11 of 15" has to mean 15 as the grounding filter saw it.
        _gtotal = len([x for x in (services or []) if not x.get("pinned")])
        _pre_ground = list(services or [])
        services, ungrounded, blocked_pins, grounding_off = drop_ungrounded_services(
            services, seeds_typed, biz,
            [p.get("title", "") if isinstance(p, dict) else str(p)
             for p in (site_pages or [])], brand, domain)

        # THE VALVE WAS WATCHING THE WRONG NUMBER. drop_ungrounded_services stands
        # down when it drops more than half of what it was handed — a ratio
        # against its own candidates. It never asks what the grid LOOKED LIKE
        # afterwards. MPG Gummies came back with 7 dropped out of 18, which is
        # 39% and passes that test comfortably, and left the grid at 11 of 20.
        #
        # The nine empty slots are not cosmetic. Total volume fell to ~8,020/mo,
        # under vol_free_below, so the volume add went to $0 and the quote came
        # out $1,000 under the one Brendan sent for the same client. List length
        # is a pricing input, and a filter that quietly shortens the list is
        # quietly moving the price.
        #
        # So there is a second valve, on the outcome rather than the ratio: if
        # the grid is left materially short and this filter is holding services
        # that would fill it, the best of them come back — highest volume first,
        # in the order the ranking already put them — and only as many as it
        # takes to reach the fill line. Everything past that stays dropped, so a
        # full list still gets the whole protection: Keller's "turner
        # construction company" is only ever restored if the alternative is an
        # empty slot. (2026-08-18)
        _restored = []
        _fill = float(CFG.get("grounding_min_slot_fill", 0.75) or 0)
        if ungrounded and _fill and n_services:
            _want = int(n_services * _fill)
            _have = len([x for x in (services or []) if not x.get("pinned")])
            if _have < _want:
                _gone = {str(d[0]).lower() for d in ungrounded}
                _kept = {(x.get("service") or "").lower() for x in (services or [])}
                for x in _pre_ground:
                    if _have >= _want:
                        break
                    _n = (x.get("service") or "").lower()
                    if _n in _gone and _n not in _kept:
                        services.append(x)
                        _restored.append(x.get("service"))
                        _kept.add(_n)
                        _have += 1
                if _restored:
                    _low = {r.lower() for r in _restored}
                    ungrounded = [d for d in ungrounded
                                  if str(d[0]).lower() not in _low]
        grounding_restored = list(_restored)
        # LAST, after every filter has had its say: make sure each topic the
        # operator typed is still represented. Everything above ranks by volume,
        # and the biggest topic wins every one of those contests — Ski Barn's
        # patio/BBQ half was eliminated seven times over by ski volume before
        # anyone saw the list (2026-08-07). Runs before rebalance_tiers so a
        # swapped-in service can still have its tier corrected.
        # The model partitions and names; token clustering is the fallback so a
        # dead API can't remove the guarantee, only its granularity.
        # WHAT THE OPERATOR SAID NOT TO SELL. Runs after the grounding filter and
        # before the topic guarantee, so a freed slot gets refilled rather than
        # left empty. Only reaches what the TOOL proposed — a focus term that
        # trips a negative is reported on the panel and kept. (2026-08-13)
        if negatives:
            services, negatives_dropped = drop_negative_services(
                services, negatives, seeds)
        # SHARE IS COMPUTED OVER TERMS SOMEONE COULD BUY. PEO Brokers typed four
        # abbreviation lookups; one line of the panel flagged them as lookups and
        # two lines down the coverage guarantee reserved three of twenty slots for
        # them, because 17% of the seed list earns 17% of the grid no matter what
        # the seeds are. The two mechanisms were reading the same terms and
        # disagreeing. Lookups are still QUOTED — seeds are never filtered — they
        # just stop reserving slots against the topics that sell something.
        # Falls back to the whole list rather than leaving nothing to cluster.
        # (2026-08-13)
        _buyable = [s for s in (seeds or []) if not is_lookup_kw(str(s))]
        topic_seeds = _buyable if len(_buyable) >= 2 else list(seeds or [])
        topics = (claude_topics(topic_seeds, biz, brand)
                  or topic_clusters(topic_seeds))
        topic_source = ("ai" if topics and topics[0].get("source") == "ai"
                        else "words")
        services, topic_fixes = enforce_topic_coverage(services, topic_seeds,
                                                      n_services, cands,
                                                      topics=topics)
        # THE LAST PASS TO ADD A SERVICE HAS TO FACE THE OUT-OF-AREA FILTER TOO.
        # The topic guarantee fills an under-represented topic from `cands` when
        # the operator's own terms cannot, and `cands` is ranked by NATIONAL
        # volume — so it reached past two earlier drop_foreign_geo_services passes
        # and put "arizona native peoples" and "arizona native tribes" into a
        # Portland, Oregon quote for a Pacific Northwest health board. They also
        # arrived after the grid had chosen its wording, which is why they read
        # "portland or" while every other row reads "portland".
        #
        # Exactly the hole the qualifier backfill had. Any pass that can add a
        # service from the pool needs this after it, not before it. (2026-08-17)
        if topic_fixes:
            services, _tc_geo = drop_foreign_geo_services(services, markets, state)
            if _tc_geo:
                _seen_tc = {d[0] for d in (geo_dropped2 or [])}
                geo_dropped2 = list(geo_dropped2 or []) + [
                    d for d in _tc_geo if d[0] not in _seen_tc]
                topic_fixes = [f for f in topic_fixes
                               if str((f or {}).get("added") or f) not in
                               {d[0] for d in _tc_geo}]
        services = rebalance_tiers(services)
        # LAST WORD ON DUPLICATES. enforce_seed_services folds on the shared key
        # and so does the topic guarantee, and PEO Brokers still came back with
        # "workers compensation nysif" AND "nysif workers compensation" — the
        # same three words, the same key, three of six Ultra slots on one state
        # fund. Something between those stages and the grid reintroduces it and
        # the isolated repro does not show which. This is the backstop: one pass
        # over the FINAL list, same key as everywhere else, first occurrence
        # wins so pins and seeds keep their place. Reported, not silent — if it
        # ever catches anything the stage above it is still wrong. (2026-08-14)
        _alias = acronym_aliases([(x.get("service") or "") for x in services])
        # SPACING IS NOT A DIFFERENT SERVICE. MPG Gummies' grid carried both
        # "mpgxtreme" (210/mo) and "mpg xtreme" (30/mo) — one product, two slots,
        # and their volumes counted twice into the total the price is computed
        # from. _seed_key compares stemmed token SETS, so {mpgxtreme} and
        # {mpg, xtreme} are simply different and no amount of stemming closes
        # that. Running the stems together gives one key for both, and it is
        # checked alongside the token set rather than replacing it: a brand
        # written solid is a real spelling that a keyword tool will return.
        # (2026-08-18)
        _seen_final, _final, services_deduped = set(), [], []
        _seen_squash = set()
        for _x in services:
            _n = (_x.get("service") or "").strip()
            _k = _seed_key(_n, _alias) or frozenset({_n.lower()})
            _sq = "".join(sorted(_k))
            if _k in _seen_final or (len(_sq) > 5 and _sq in _seen_squash):
                services_deduped.append(_n)
                continue
            _seen_final.add(_k)
            _seen_squash.add(_sq)
            _final.append(_x)
        if services_deduped:
            services = _final
        if geo_dropped is None and geo_dropped2 is None:
            geo_dropped = None
        else:
            seen_d = set()
            geo_dropped = [d for d in (list(geo_dropped or []) + list(geo_dropped2 or []))
                           if not (d[0] in seen_d or seen_d.add(d[0]))]
        pinned = [t for t in pinned
                  if any((x.get("service") or "") == t for x in services)]
        service_forms, service_form_report = ({}, [])
        # Which WORDING of each market to cross with. Measured, not assumed —
        # "new york city ny" is nobody's search (2026-08-07).
        # Probe with the client's own SHORTEST terms, not the grid's lead
        # service: "alpine ski shop nyc" is unmeasurable in every spelling, so a
        # narrow probe makes every candidate tie at zero (2026-08-07).
        probe_terms = list(seeds) + [x.get("service") for x in services if x.get("service")]
        geo_forms, geo_form_report = ({}, [])
        if grid_cities and not national_demand:
            try:
                geo_forms, geo_form_report = pick_geo_forms(grid_cities, state,
                                                           probe_terms)
            except Exception:
                geo_forms, geo_form_report = {}, []
            # A market the wording probe could not read falls back to the FIRST
            # candidate, which is the string the operator typed. For a county
            # that is the wrong default -- see market_forms above. The city
            # ranking measured both namings locally, so use its answer here
            # rather than the typed one.
            _mf = (city_pick or {}).get("market_forms") or {}
            for _m in grid_cities:
                if _m not in geo_forms and _mf.get(_m):
                    geo_forms[_m] = _mf[_m]
                    geo_form_report.append(
                        {"market": _m, "status": "from city ranking",
                         "kept": _mf[_m],
                         "detail": "the wording probe read nothing at national "
                                   "level, so the form is the one that measured "
                                   "when the markets were ranked"})
        g = build_grid(services, grid_cities, state, prepicked=True,
                       geo_forms=geo_forms)
        full = g["ultra"] + g["competitive"] + g["long_tail"]
        # Volume: look up the BARE service term AT THE CLIENT'S MARKET (the
        # geo-modified forms report ~0). The same figure is shown on each city
        # row for that service, so pricing must count it ONCE PER SERVICE — not
        # once per row — or a 10-city grid would inflate volume 10x.
        svc_names = list(dict.fromkeys([s["service"] for s in services]))
        # Measure the UNUSED seeds alongside the chosen services, in the same
        # call, so a dead service name can be swapped for one the client's own
        # list already has. Costs extra keywords, not extra requests.
        _chosen = {x.lower() for x in svc_names}
        _pool = [str(sd).strip().lower() for sd in (seeds or [])
                 if str(sd).strip() and str(sd).strip().lower() not in _chosen]
        _pool = list(dict.fromkeys(_pool))
        _cand_cap = int(CFG.get("service_candidate_cap", 14))
        # ROUND-ROBIN ACROSS TOPICS, not the first N in input order. The pill
        # list arrives alphabetically, so a flat slice took "ski a-s" and cut off
        # every bbq and patio candidate before it was measured — which made an
        # upgrade impossible for any topic that sorts late, silently
        # (2026-08-08).
        if topics and len(topics) > 1:
            by_topic = {}
            for sd in _pool:
                by_topic.setdefault(service_topic(sd, topics) or "", []).append(sd)
            # SHORTEST FIRST inside each topic. The budget can't cover 30 ski
            # terms, and alphabetical order is meaningless — but short means
            # generic means high volume, the same proxy the geo-wording probe
            # uses. Alphabetical would have spent the ski budget on "ski apparel"
            # ... "ski gear stores" and never measured "ski shop" or "ski store".
            for lab in by_topic:
                by_topic[lab].sort(key=lambda t: (len(t.split()), len(t)))
            _alts, _i = [], 0
            while len(_alts) < _cand_cap:
                took = False
                for lab in list(by_topic):
                    bucket = by_topic.get(lab) or []
                    if _i < len(bucket) and len(_alts) < _cand_cap:
                        _alts.append(bucket[_i])
                        took = True
                if not took:
                    break
                _i += 1
        else:
            _alts = _pool[:_cand_cap]
        # ---- "<service> near me" -------------------------------------------
        # BE's Junk Bee Gone list carries "junk removal near me" at rank 2 and
        # "dumpster rental near me" at 6 — real terms with real positions. The
        # tool could not produce them at all: "near me" was correctly banned as a
        # MARKET (it poisoned the rank location) but nothing ever made it a
        # KEYWORD. Measured in the same call as everything else, so a form that
        # nobody searches costs nothing and never reaches the list. (2026-08-10)
        # Probe near-me for EVERY service, then keep the best few. Taking the
        # first three by list order gave Junk Bee Gone "rent a dumpster near me"
        # and "dumpster rentals near me" while missing "junk removal near me" —
        # the one BE put in his proposal at rank 2, and the biggest of the three.
        # Extra keywords in a call already being made; no extra request.
        # SINGULAR OR PLURAL, decided by measurement rather than by whoever typed
        # it first — see service_form_probes. The near-me form of each variant is
        # probed too, so a service that gets re-spelled still has a measured
        # near-me term waiting rather than losing one to the swap.
        form_probes = service_form_probes(svc_names)
        near_n = int(CFG.get("near_me_terms", 3))
        near_forms = []
        if near_n > 0 and not national_demand:
            for nm in (svc_names + form_probes)[:2 * int(CFG.get("near_me_probe_cap", 12))]:
                f = clean_kw(f"{nm} near me")
                if f and f not in near_forms:
                    near_forms.append(f)
        vols, per_city, vol_err = fetch_local_volume(
            svc_names + _alts + near_forms + form_probes,
            [] if national_demand else cities, state,
            national=national_demand)
        # Apply the spellings BEFORE anything reads the grid: rebuilding it is
        # pure string work, so the swap costs nothing beyond the keywords already
        # measured above, and every dedupe downstream sees the chosen form.
        if form_probes:
            try:
                service_forms, service_form_report = choose_service_forms(
                    svc_names, vols, vol_err, pool=cands)
            except Exception:                                 # noqa: BLE001
                app.logger.exception("service form choice failed")
                service_forms, service_form_report = {}, []
        if service_forms:
            for _x in services:
                _new = service_forms.get(_x.get("service") or "")
                if _new:
                    _x["service"] = _new
            pinned = [service_forms.get(t, t) for t in pinned]
            svc_names = list(dict.fromkeys([s["service"] for s in services]))
        if service_forms:
            g = build_grid(services, grid_cities, state, prepicked=True,
                           geo_forms=geo_forms)
            full = g["ultra"] + g["competitive"] + g["long_tail"]
        # Add the near-me forms that earned their place, into the same tier as the
        # service they came from.
        def _attach_near_me():
            """(Re)attach the near-me rows to the grid, returning what stuck.

            THIS RUNS AGAIN AFTER EVERY REBUILD, and that is the whole point.
            build_grid() regenerates the grid from services x cities, so a row
            that is not a crossed service does not survive it -- and every
            near-me term is exactly that. Two passes below rebuild: the
            service-swap pass and the tier-reconciliation pass. Whatcom County
            (2026-08-24) hit the first one: the panel reported "3 near me terms
            added -- water damage restoration near me (210/mo) ..." and the
            keyword list held none of them, because three services were dropped
            against measured demand and the rebuild that followed wiped the
            rows. The note was reading `near_added`, which is a record of what
            was appended, not of what is still there.

            Idempotent by construction: the tier lookup and the duplicate check
            are both read fresh from the CURRENT grid, and the returned list
            replaces the old one rather than extending it. Re-running it after a
            drop is also the correct behaviour on its own terms -- a service
            that just lost its slot should lose its near-me term with it.
            """
            added = []
            if not near_forms or vol_err:
                return added
            _nfloor = int(CFG.get("near_me_min_volume", 30))
            _tier_of = {x["service"]: x["tier"] for x in services}
            # Highest measured demand wins the slots, not list order.
            _ranked = sorted(
                ((clean_kw(f"{nm} near me"), nm) for nm in svc_names),
                key=lambda fn: -int((vols or {}).get(fn[0], 0) or 0))
            for f, nm in _ranked[:near_n]:
                if f not in near_forms:
                    continue
                v = int((vols or {}).get(f, 0) or 0)
                if v < _nfloor:
                    continue
                t = _tier_of.get(nm, "competitive")
                if any(r["keyword"] == f for r in g[t]):
                    continue
                row = {"keyword": f, "volume": v, "src": "grid",
                       "origin": "added", "service": nm, "city": ""}
                g[t].append(row)
                full.append(row)
                added.append((f, v))
            return added

        near_added = _attach_near_me()

        # ---- IS THE LOCAL FRAME RIGHT? ---------------------------------------
        # "Should this be nationwide" was left to the operator's judgement, and
        # the obvious test — does the client have locations — is the wrong one.
        # NASSCO listed seven real cities in a document it sent us, and every one
        # of the 35 city-attached terms returned zero: a contractor looking for
        # ITCP certification types "itcp certification", not "itcp certification
        # san mateo ca". Junk Bee Gone's "junk removal knoxville tn" returns
        # 170/mo, because you need someone to drive to your house.
        #
        # So measure it instead of guessing. Same seeds, once with cities and
        # once bare. If the bare terms carry real demand and the city-attached
        # ones carry none, the buyer does not search with a place attached and
        # the local frame is producing a quote out of nothing. One extra call,
        # only on local quotes, and only reported — never applied, because
        # switching the demand basis changes the price. (2026-08-10)
        frame = {}
        # A NATIONAL build needs the same sanity check. Switching the basis fixes
        # the frame but not the wording, and NASSCO went national and still
        # returned 15 zeros out of 16 — at which point the answer is "these
        # phrases are not what anyone types", and nothing on screen said so
        # because the check only ran on local quotes. Here `vols` already IS the
        # national figure, so no extra call. (2026-08-10)
        if national_demand and svc_names:
            nat_tot = sum(int(v or 0) for v in (vols or {}).values())
            frame = {"local_total": None, "national_total": nat_tot,
                     "terms": svc_names[:6], "error": vol_err}
            if not vol_err and nat_tot < int(CFG.get("frame_national_min", 200)):
                frame["verdict"] = "no_demand"
                frame["reason"] = (
                    f"Priced on national demand, and these terms still return "
                    f"only {nat_tot:,}/mo between them. Geography was not the "
                    "problem — the wording is. These read like descriptions of "
                    "the client rather than phrases anyone types into Google.")
            elif not vol_err and cities:
                # CONFIRM the national choice rather than only questioning it.
                # The scope check reads a footprint off the site's structure and
                # says "maybe this should be local"; measuring the same terms
                # with a city attached answers that outright. NASSCO's site
                # showed 49 Google Business listings — its member directory, not
                # its offices — and the operator got two panels disagreeing with
                # each other while the demand data had already settled it.
                # (2026-08-10)
                try:
                    _loc, _pc2, _le = fetch_local_volume(svc_names, cities, state,
                                                         national=False)
                    loc_tot = sum(int(v or 0) for v in (_loc or {}).values())
                    frame["local_total"] = loc_tot
                    if not _le and loc_tot * 10 <= nat_tot:
                        frame["verdict"] = "national_ok"
                        frame["reason"] = (
                            f"{nat_tot:,}/mo bare against {loc_tot:,}/mo with a "
                            "city attached — national is the right basis, and a "
                            "footprint read off the site's page structure does "
                            "not change that.")
                except Exception:
                    pass
        if not national_demand and svc_names:
            try:
                _nat, _pc, _ne = fetch_local_volume(svc_names, [], state,
                                                    national=True)
                loc_tot = sum(int(v or 0) for v in (vols or {}).values())
                nat_tot = sum(int(v or 0) for v in (_nat or {}).values())
                frame = {"local_total": loc_tot, "national_total": nat_tot,
                         "terms": svc_names[:6], "error": _ne or vol_err,
                         "local_error": vol_err}
                min_nat = int(CFG.get("frame_national_min", 200))
                # ZERO LOCAL DEMAND AND NO LOCAL MEASUREMENT LOOK IDENTICAL, AND
                # ONLY ONE OF THEM IS AN ARGUMENT FOR ANYTHING. Ski Barn's
                # exact-match pull was refused by the 12/min cap, so every term
                # carried no volume, loc_tot was 0, and this told the operator
                # "31,000/mo bare vs 0/mo with a city attached — tick Price on
                # national demand and rebuild". That is a rate limit wearing the
                # shape of a market finding, and acting on it would have moved
                # the client onto the national anchor. The local side has to
                # have been READ before its zero means anything. (2026-08-20)
                if vol_err or _ne:
                    frame["verdict"] = "unmeasured"
                    frame["reason"] = (
                        f"The city-attached volumes could not be read "
                        f"({str(vol_err)[:90]}), so there is nothing to compare "
                        f"the national figure against. This is not a finding "
                        f"about the market — rebuild once the lookup succeeds.")
                elif nat_tot >= min_nat and loc_tot == 0:
                    frame["verdict"] = "national"
                    frame["reason"] = (
                        f"These services draw {nat_tot:,}/mo searches nationally "
                        f"and {loc_tot}/mo with a city attached. Nobody searches "
                        "them with a place, so the city grid is measuring "
                        "something that isn't there.")
                elif not _ne and not vol_err and loc_tot > 0:
                    frame["verdict"] = "local"
                    frame["reason"] = (
                        f"City-attached terms carry {loc_tot:,}/mo against "
                        f"{nat_tot:,}/mo nationally — people do search these with "
                        "a place, so the local frame is measuring real demand.")
                elif not _ne and not vol_err and nat_tot == 0:
                    frame["verdict"] = "no_demand"
                    frame["reason"] = (
                        "These services return no volume either locally OR "
                        "nationally, so the wording is the problem, not the "
                        "geography. Check the Product / Vertical Focus terms "
                        "against what people actually type.")
            except Exception as e:
                frame = {"error": str(e)[:120]}
        # LAST RESORT for a national quote: Labs carries per-term volume and is
        # a different service, so a Google Ads outage doesn't take it down too.
        # Restricted to national_demand deliberately — Labs answers at country
        # level, which for a national quote is the SAME figure, but for a local
        # one it would substitute national demand for local and inflate the
        # price. A local quote keeps the honest error instead.
        volume_source = "google_ads"
        if vol_err and national_demand and svc_names:
            _lv = fetch_exact_volume(svc_names, [], "", national=True)
            if _lv and any(_lv.values()):
                vols = _lv
                volume_source = "labs"
                vol_err = None
        # Collapse cities that resolved to the SAME Google Ads location. This
        # is the exact answer — DataForSEO reports the location it used — and
        # it replaces the volume-vector inference, which needed the geo-modified
        # probe to have volume and therefore gave up in small rural markets.
        # Group on the PER-CITY VOLUMES that were just fetched. Google Ads holds
        # demand at the nearest targetable location, so cities in one market
        # return the identical figure for every service — Brent Cogan's five
        # towns all came back 30 / 10 / 10 (2026-08-03).
        #
        # Two earlier attempts missed this. Matching on the geo-modified probe
        # ("electrical panel upgrade hollidaysburg pa") needed volume the term
        # doesn't have in a rural county, so every vector was zero. Matching on
        # the location NAME failed because the name recorded is the one we
        # SENT — "Bellwood,Pennsylvania" — not the one Google resolved it to.
        # These volumes are the resolution, observed rather than inferred.
        # Distance first — it is the only signal that works when the markets
        # have no measurable demand, which is exactly when grouping matters.
        if not national_demand:
            _g, _loc, _un = group_by_distance(markets, state)
            _real = [x for x in _g if len(x) > 1]
            if _real:
                city_pick["metro_groups"] = _real
                city_pick["grouped_by"] = (
                    f"distance — markets within {int(CFG.get('market_radius_miles', 25))} "
                    f"miles of each other")
                city_pick["unlocated_markets"] = _un
        if not city_pick.get("metro_groups") and per_city and not national_demand:
            _svcs = [x.lower() for x in svc_names]
            _vecs = {}
            for _c in cities:
                _cl2 = _bare_city(_c, state)
                _v = [per_city.get((_cl2, _s)) for _s in _svcs]
                if any(x is not None for x in _v):
                    _vecs[_c] = [(0 if x is None else x) for x in _v]
            _groups = [g for g in group_by_metro(_vecs, min_terms=1) if len(g) > 1]
            if _groups:
                city_pick["metro_groups"] = _groups
                city_pick["grouped_by"] = "identical per-market search volume"
            elif not city_pick.get("metro_groups"):
                city_pick["metro_groups"] = []
        def _apply_volumes(rows):
            """Give every crossed row its own city's measured volume.

            Extracted because the tier-reconciliation pass REBUILDS the grid,
            and the rebuilt rows were left at volume 0 — the old code re-looked
            them up by r["keyword"] ("outdoor furniture nyc") against a map keyed
            by SERVICE ("outdoor furniture"), which never matches. Result: any
            quote whose tiers moved displayed no volume against any keyword,
            which is exactly the number an operator checks the tiering against
            (2026-08-07).
            """
            fb = set()
            for x in ((per_city or {}).get("__fallback_cities__") or []):
                fb.add(str(x).strip().lower())
                fb.add(_bare_city(str(x), state))
            # Cities that resolved to the SAME location as another city are not
            # each reporting their own demand — one of them is borrowing the
            # other's area. Keep the largest as the genuine one and mark the
            # rest, because the big city is the one the shared location is
            # named after. Hard signal from the API, not a size guess.
            # Which location actually answered, per city, in friendly form.
            # "wider area" told the operator a number was wrong without saying
            # what it was — "New York statewide" is actionable (2026-08-07).
            loc_by_city = {}
            for c, l in ((per_city or {}).get("__city_locs__") or {}).items():
                bare = _bare_city(c, state)
                txt = str(l or "").replace(",United States", "").strip()
                if not txt or txt.lower() == "united states":
                    loc_by_city[bare] = "nationwide"
                elif "," in txt:
                    loc_by_city[bare] = txt.split(",")[0].strip()
                elif txt.lower() in STATE_ABBREV:
                    loc_by_city[bare] = f"{txt} statewide"
                else:
                    loc_by_city[bare] = txt
            # city_size needs the market WITH its own state. Passing the bare
            # name plus the fallback state looked up "new york, NJ" and scored
            # New York City at zero, so the smaller town was kept as the genuine
            # figure and the big one got flagged instead.
            market_of = {_bare_city(m, state): m for m in (cities or [])}
            codes = (per_city or {}).get("__location_codes__") or {}
            if codes:
                groups = {}
                for c, code in codes.items():
                    groups.setdefault(code, []).append(c)
                for code, members in groups.items():
                    if len(members) < 2:
                        continue
                    keep = max(members,
                               key=lambda c: city_size(market_of.get(c, c), state))
                    for c in members:
                        if c != keep:
                            fb.add(c)
            for r in rows:
                svc_l = (r.get("service") or "").lower()
                city_l = (r.get("city") or "").lower()
                # A row with NO CITY was measured as itself, not as a
                # service crossed with a place -- the near-me terms are the only
                # ones shaped like this. Looking one up as ("", service) misses,
                # and the miss branch below writes volume 0 / vol_scope
                # "unknown" over a figure that was read from the API. That is
                # how "water damage restoration near me (210/mo)" would have
                # reached the proposal carrying no data even on a build where it
                # survived the rebuilds. (2026-08-24)
                if not city_l:
                    continue
                # This city's own volume for the SERVICE. Not the volume of the
                # geo-modified phrase — local phrases mostly report zero — and
                # not the cross-city total, which would print the same number
                # against every city.
                v = per_city.get((city_l, svc_l))
                if v is not None:
                    r["volume"] = v
                    # A city with no data of its own was answered by its county,
                    # state or the whole US. That number is real but it is NOT
                    # this city's, so say so rather than letting it read as a
                    # small town out-searching Manhattan.
                    if city_l in fb:
                        r["vol_scope"] = "broader"
                        r["vol_area"] = loc_by_city.get(city_l, "")
                    continue
                # No per-city figure at all: leave the row blank rather than
                # substituting the summed total. An unknown is an unknown.
                r["volume"] = 0
                r["vol_scope"] = "unknown"
            return rows

        _apply_volumes(full)
        service_volume = {s: vols.get(s.lower(), 0) for s in svc_names}

        # SERVICE NAMES RECONCILED AGAINST MEASURED DEMAND (2026-08-08).
        # A slot holding a phrase nobody searches is worth less than the client's
        # own unused term. Runs before the tier pass so the tiers are assigned to
        # the final set.
        service_swaps = []
        _svc_check_on = (int(CFG.get("service_min_volume", 0) or 0) > 0
                         or float(CFG.get("service_upgrade_ratio", 0) or 0) > 0)
        if not national_demand and _svc_check_on:
            try:
                services, service_swaps = swap_low_volume_services(
                    services, vols, seeds, topics, typed=seeds_typed)
                # A PASS THAT REMOVES A SERVICE HAS TO RESTORE THE TIER
                # GUARANTEE. rebalance_tiers runs earlier, so dropping the dead
                # slots after it emptied a column outright: NPAIHB came back
                # "Long Tail 0" with a bare em dash where the third column of the
                # proposal goes, which reads as a mistake rather than a thin
                # account. Same shape as every other late-pass bug this week —
                # something added after the guarantee, without re-asking it.
                # (2026-08-17)
                if any(r.get("kind") == "dropped" for r in (service_swaps or [])):
                    services = rebalance_tiers(services)
                if service_swaps:
                    svc_names = list(dict.fromkeys([x["service"] for x in services]))
                    g = build_grid(services, grid_cities, state, prepicked=True,
                                   geo_forms=geo_forms)
                    full = g["ultra"] + g["competitive"] + g["long_tail"]
                    _apply_volumes(full)
                    # service_volume was keyed by the OLD names; the tier pass
                    # below looks services up in it, so a stale map would make
                    # every swapped-in service read as unmeasured.
                    service_volume = {x: vols.get(x.lower(), 0) for x in svc_names}
                    near_added = _attach_near_me()
            except Exception:
                app.logger.exception("swap_low_volume_services failed")
                service_swaps = []

        # TIERS RECONCILED AGAINST MEASURED DEMAND (2026-08-05).
        # Tiers are assigned by the model on judgement, BEFORE any volume is
        # known, and nothing reconciled them afterwards — rebalance_tiers only
        # guarantees no column is empty. On Ski Barn that put "ski jeans" at
        # 60,500/mo in Long Tail while "ski equipment rental" at 1,600/mo sat in
        # Ultra Competitive. A proposal whose Long Tail column carries the
        # biggest numbers reads as though the tool doesn't understand the market.
        #
        # Re-sorts by volume while PRESERVING the tier COUNTS the model chose,
        # so the proposal keeps its shape and only the assignment changes. Terms
        # with no volume data keep their original position rather than being
        # dumped into Long Tail on missing data.
        tier_moves = []
        try:
            _order = ["ultra", "competitive", "long_tail"]
            _counts = {t: sum(1 for x in services if x.get("tier") == t)
                       for t in _order}
            _measured = [x for x in services
                         if (service_volume.get(x["service"]) or 0) > 0]
            if len(_measured) >= 3 and sum(_counts.values()) == len(services):
                # Store intent gets a THUMB ON THE SCALE, not a veto. Absolute
                # precedence pushed "outdoor furniture" (4,580/mo, the largest
                # line in the quote) into Long Tail beneath terms at 80/mo,
                # because it isn't a shop phrase — and the columns stopped
                # running high-to-low, which is the one thing this pass exists
                # to guarantee (2026-08-09). A multiplier keeps a shop term
                # ahead of a comparable product term while letting a term with
                # several times the demand take the slot it has earned.
                # STORE INTENT BREAKS TIES; IT DOES NOT OVERTAKE. It was a 3x
                # multiplier, and its own note said the point was to keep a shop
                # term ahead of a COMPARABLE product term while a term with
                # several times the demand still takes the slot it earned. On a
                # list where everything sits between 0 and 90/mo there is no
                # "several times", so the multiplier stopped being a thumb and
                # became the whole hand: "rental homes with pool" (10/mo, boosted
                # to 30) outranked "luxury apartments" (20/mo), and three 4-to-5
                # word qualifier phrases took Ultra Competitive on Amare Homes
                # while Brendan put them in Long Tail. Comparable now means
                # EQUAL, which is what the sentence always meant.
                #
                # BREADTH IS THE LAST WORD, because in a thin market volume
                # cannot rank anything: fifteen of Amare's twenty terms measured
                # 0-10/mo, so arrival order was deciding the tiers. Brendan tiers
                # the same client on breadth — "homes for rent santa fe nm" Ultra
                # at three words, "homes for rent with garage santa fe nm" Long
                # Tail at six. (2026-08-13)
                # STORE INTENT WAS STILL THE WHOLE HAND, one layer down. It
                # was demoted from a 3x multiplier to a tie-break, but it sat
                # ABOVE breadth in the key — and on a list where every term ties
                # at 10/mo, the first tie-break is the ranking. `_STORE_INTENT`
                # holds "rental", which is a shop word for a ski hire but the
                # PRODUCT NOUN for a rental community, so it fired on "rental
                # homes with pool", "rental homes no credit check" and "short
                # term rental homes" and put all three in Ultra Competitive
                # ahead of "houses for rent" and "home for rent". Brendan's list
                # for the same client leads with the bare head terms and puts
                # every one of those three in Long Tail.
                #
                # Breadth now outranks it, measured as SHAPE rather than word
                # count: what a term adds on top of the list's own core, so
                # "houses for rent" reads as one step off the head term rather
                # than as three words. An amenity or a landmark — "with a
                # garage", "no credit check", "near meow wolf" — is last
                # whatever its length. Volume still leads; this only decides the
                # ties, which in a thin market is most of the list.
                # (2026-08-16)
                _shape = service_shape(services)

                def _rank_key(x):
                    _n = (x.get("service") or "").strip().lower()
                    _depth, _narrow = _shape.get(_n, (0, False))
                    return (-(service_volume.get(x["service"]) or 0),
                            1 if _narrow else 0,
                            _depth,
                            len(_n.split()),
                            0 if is_store_intent(x["service"]) else 1)
                _ranked = sorted(_measured, key=_rank_key)
                # UNMEASURED TERMS SINK; THEY DO NOT SQUAT. They used to keep
                # whatever tier the model gave them, and the per-tier capacity was
                # reduced to make room — so on PEO Brokers five "no data" terms
                # held five of six Ultra slots, capacity fell to one, and
                # "professional employer organization" at 49,500/mo was the only
                # real term that fitted while "peo services" at 4,400 sat in
                # Competitive. A column whose top rows have no numbers is exactly
                # what this pass exists to prevent.
                #
                # Measured terms are placed first, best to worst, across the tier
                # counts the mix asked for. Whatever is left over is unmeasured,
                # and it fills from the bottom up. No term is dropped and the
                # counts do not change. (2026-08-13)
                _unmeasured = [x for x in services if x not in _measured]
                _fill = _ranked + _unmeasured
                _i = 0
                for t in _order:
                    for _ in range(_counts[t]):
                        if _i >= len(_fill):
                            break
                        _svc = _fill[_i]
                        if _svc.get("tier") != t:
                            tier_moves.append({"service": _svc["service"],
                                               "from": _svc.get("tier"), "to": t,
                                               "volume": service_volume.get(_svc["service"], 0)})
                            _svc["tier"] = t
                        _i += 1
                if tier_moves:
                    g = build_grid(services, grid_cities, state, prepicked=True,
                                   geo_forms=geo_forms)
                    full = g["ultra"] + g["competitive"] + g["long_tail"]
                    # Same assignment as the first pass — by service AND city.
                    _apply_volumes(full)
                    near_added = _attach_near_me()
        except Exception:
            app.logger.exception("tier reconciliation failed")
            tier_moves = []

        # National demand on a client with PHYSICAL PREMISES is usually a
        # mis-scope, not a product brand. The signals to catch it are already
        # collected; nothing was checking them (Ski Barn: NJ stores, priced
        # nationwide, so every term came back as national head demand).
        scope_warning = ""
        scope_why = ""
        # "note" = this is the tool working as intended and saying so; "warn" =
        # something needs a decision. They were rendered identically.
        scope_kind = "warn"
        goals_all = goal_list(goal)
        _gs = goal_scope(goal)
        _gforce = goal_forces_national(goal)
        if _gforce and (markets or gbp_count or site_locations):
            # Goal-driven national is intentional, so this is not a warning that
            # something is wrong — it's a statement of what the goal did, and of
            # the one thing it deliberately did NOT do.
            _where = []
            if markets:
                _where.append(f"{len(markets)} market"
                              f"{'' if len(markets) == 1 else 's'} entered")
            if gbp_count:
                _where.append(f"{gbp_count} Google Business listing"
                              f"{'' if gbp_count == 1 else 's'}")
            if site_locations:
                _where.append(f"{len(site_locations)} location page"
                              f"{'' if len(site_locations) == 1 else 's'}")
            # ONE LINE, AND NOT IN AN AMBER BOX. This branch fires when the
            # goal did exactly what it is for, and it was five sentences of
            # prose inside a warning panel — so the loudest note on the build
            # was the one saying nothing is wrong. The reasoning is worth
            # keeping and belongs behind the tooltip with the rest of it.
            # (2026-08-18)
            scope_kind = "note"
            scope_warning = (
                f"Demand pulled nationally — the goal is “{_gforce}”, "
                "with " + " and ".join(_where) + ".")
            scope_why = (
                "The client asked to be sold online sales, so the volumes "
                "describe the whole addressable market rather than one city. "
                "Rankings are still measured in the client's own market, "
                "because whether THIS client is visible is a local question. "
                "Change the goal if the campaign is really about the stores.")
        elif goal and _gs == "local" and national_demand:
            scope_warning = (
                f"Goal is \u201c{goal}\u201d, which happens somewhere \u2014 but this "
                "quote is priced on NATIONAL demand, so every keyword, volume "
                "and ranking here describes a national campaign. Set Geo scope "
                "to the client's region and enter their markets, or change the "
                "goal if this really is a nationwide play.")
        elif goal and _gs == "national" and not national_demand:
            scope_warning = (
                f"Goal is \u201c{goal}\u201d, which usually sells everywhere \u2014 but "
                "this quote is priced on LOCAL demand, so the volumes are "
                "geo-limited and will understate the opportunity. Consider "
                "Nationwide scope or the national-demand switch.")
        elif goals_all and all(g in GOAL_OFF_PATTERN for g in goals_all):
            scope_warning = (
                f"Goal is \u201c{' + '.join(goals_all)}\u201d. An SEO keyword campaign is a weak "
                "instrument for this \u2014 the ladder below prices category-term "
                "ranking work, which may not be what the client is buying. "
                "Confirm the objective before quoting.")
        # A DELIBERATE choice does not need a warning on every rebuild. The scope
        # check exists to catch national demand INFERRED from an industry tag or
        # a storefront; when the operator ticked the switch themselves it fires
        # forever on a decision already made, and NASSCO's count is its member
        # directory rather than its offices anyway. Suppressed for a manual
        # override and for a frame the demand data has confirmed — the reasoning
        # is kept as a quiet note either way. (2026-08-10)
        _explicit = ("manual" in (national_reason or "").lower()
                     or (frame or {}).get("verdict") == "national_ok")
        if national_demand and (gbp_count or site_locations) and _explicit:
            scope_note = (
                f"National demand is set deliberately"
                + (" and confirmed by the volume data"
                   if (frame or {}).get("verdict") == "national_ok" else "")
                + f". The site shows {gbp_count} Google Business listing"
                + ("" if gbp_count == 1 else "s")
                + f" and {len(site_locations)} location page"
                + ("" if len(site_locations) == 1 else "s")
                + " — member, chapter or directory pages if the client is an "
                  "association, their own premises if not.")
        elif (national_demand and (gbp_count or site_locations)
                and (frame or {}).get("verdict") != "national_ok"):
            _bits = []
            if gbp_count:
                _bits.append(f"{gbp_count} Google Business listing"
                             f"{'' if gbp_count == 1 else 's'}")
            if site_locations:
                _bits.append(f"{len(site_locations)} location page"
                             f"{'' if len(site_locations) == 1 else 's'} on the site")
            # "A retailer with premises… if it is a store" reads as a mistake on
            # a nonprofit standards body, and a directory-style site inflates the
            # listing count with OTHER companies' locations — NASSCO scored 49
            # Google Business listings, which it does not have. So say what was
            # counted and let the operator judge it, rather than asserting the
            # client is a shop. (2026-08-10)
            scope_warning = (
                "Priced on NATIONAL demand, but this client's site suggests a "
                "physical footprint \u2014 " + " and ".join(_bits)
                + ". The keywords, volumes and rankings here are all national, "
                  "which prices a different campaign from a local one. If those "
                  "premises are the client's own, set Geo scope to their region "
                  "and enter their markets. If they belong to members, chapters "
                  "or a directory the client publishes, the count is not theirs "
                  "and national is right \u2014 check before trusting it.")

        return {
            "ultra": g["ultra"], "competitive": g["competitive"],
            "long_tail": g["long_tail"],
            "head": g["ultra"] + g["competitive"],
            "all": full,
            "refined_by_ai": True,
            "business_desc": biz,
            "site_pages_found": len(site_pages),
            "grid": True,
            "services": services,
            "pinned_head_terms": pinned,
            # [(term, offending_word)] — pins refused for being ungrounded.
            "blocked_pins": [[b[0], b[1]] for b in (blocked_pins or [])],
            "city_selection": city_pick,
            # per_city is keyed by (city, keyword) tuples but also carries the
            # "__city_locs__" side-channel, so every key has to be checked
            # before it is unpacked — destructuring in the comprehension blew
            # up the whole refine pass with "too many values to unpack"
            # (2026-08-03).
            "city_volumes": {c: max([v for k, v in (per_city or {}).items()
                                     if isinstance(k, tuple) and len(k) == 2
                                     and k[0] == _bare_city(c, state)] or [0])
                             for c in cities},
            "city_locs": {c: l for c, l in
                          ((per_city or {}).get("__city_locs__") or {}).items()},
            # Which location ID actually answered per city. Two cities sharing
            # one ID means only one of them reported its own demand.
            "city_location_codes": (per_city or {}).get("__location_codes__") or {},
            # Markets that contributed NO volume of their own, with the name
            # Google would probably accept. Surfaced per MARKET because reading
            # it off twenty tagged keyword rows is work the tool should do.
            "service_swaps": service_swaps,
            "service_upgrade_ratio": CFG.get("service_upgrade_ratio", 0),
            "market_renames": [{"market": k, "used": v}
                               for k, v in ((per_city or {}).get("__renamed__") or {}).items()],
            "market_volume_gaps": [
                {"market": m,
                 "suggestion": suggest_market_name(m, state)}
                for m in ((per_city or {}).get("__fallback_markets__") or [])],
            "site_locations": site_locations,
            "service_areas": service_areas,
            "gbp_locations": gbp_count,
            "tier_moves": tier_moves,
            "scope_warning": scope_warning,
            "scope_why": scope_why,
            "scope_kind": scope_kind,
            "scope_note": scope_note,
            "gbp_cities": gbp_cities,
            "dropped_out_of_area": [d[0] for d in (geo_dropped or [])],
            "seed_ranking": seed_ranking,
            "business_desc_inferred": biz_inferred,
            "pins_refused": pins_refused,
            "services_deduped": services_deduped,
            "negatives_dropped": negatives_dropped,
            "negative_conflicts": negative_conflicts,
            "seeds_demoted": seeds_demoted,
            "seeds_dropped_suggested": seeds_dropped_suggested,
            "seeds_folded": seeds_folded,
            "seed_quality": seed_quality,
            # HOW FULL THE GRID IS. NPAIHB had room for 20 services and got 9 — a
            # 9-keyword quote against BE's 20 for the same client — and the panel
            # said nothing, because every check was about whether the terms were
            # WRONG rather than whether there were enough of them. The slot count
            # is known; the shortfall is worth naming. (2026-08-12)
            "service_slots": n_services,
            # The reserve list the "find terms they don't rank for" probe draws
            # from, and how many it is trying to find. Carried on the payload so
            # the panel never has to guess either. (2026-08-20)
            "min_unranked_terms": int(CFG.get("min_unranked_terms", 3) or 0),
            "unranked_probe_max": int(CFG.get("unranked_probe_max", 10) or 0),
            "seed_services_used": seed_used,
            "seed_services_total": seed_total,
            "seed_services_dropped": max(0, seed_total - seed_used),
            # Removed for real (empty when the filter stood down) versus what it
            # WOULD have removed — the panel needs both to say anything true.
            "dropped_ungrounded": ([] if grounding_off
                                   else [d[0] for d in (ungrounded or [])]),
            "grounding_stood_down": grounding_off,
            # Put back because the grid would otherwise have gone out short —
            # see the slot-fill valve. Named, because this is the filter
            # deliberately not doing its job and the operator should see which
            # terms it let through and why.
            "grounding_restored": grounding_restored,
            "grounding_slot_fill": CFG.get("grounding_min_slot_fill"),
            "grounding_would_drop": ([list(d) for d in (ungrounded or [])]
                                     if grounding_off else []),
            "grounding_gap_words": (grounding_gap_words(ungrounded)
                                    if grounding_off else []),
            "grounding_total": _gtotal,
            # No foreign states exist on a nationwide quote, so the warning was
            # telling the operator to fix something that is not broken.
            "geo_filter_off": (geo_dropped is None) and not national_demand,
            "service_volume": service_volume,
            "volume_error": vol_err,
            "demand_frame": frame,
            "grid_axis": {"axis": axis, "reason": axis_reason, "evidence": axis_ev},
            "near_me_added": [[f, v] for f, v in (near_added or [])],
            "acronym_collisions": acronym_collisions(full),
            "volume_location": "United States" if national_demand else loc_string(markets, state),
            "volume_source": volume_source,
            "national_demand": bool(national_demand),
            "national_demand_reason": national_demand_reason,
            "ecommerce_detected": bool(ecom_found),
            "ecommerce_reason": ecom_reason,
            "ecommerce_suppressed": ecom_suppressed,
            "state_missing": bool(cities) and not state
                             and not any(market_state(c)
                                         or c.strip().lower() in STATE_ABBREV
                                         for c in cities),
            "grid_cities": [] if national_demand else cities,
            "total_volume": sum(service_volume.values()),   # unique, not per-row
            # Topic coverage: what the operator's terms are ABOUT, how many
            # services each topic got, and any swap made to keep a topic alive.
            "topic_source": topic_source,
            "topics": [{"label": t["label"], "seeds": t["size"],
                        "share": round(t["size"] / max(len(seeds), 1) * 100),
                        # THE MEMBER TERMS, not just the count. The performance
                        # table's Practice Area column is this same split —
                        # Personal Injury against Criminal Defense — and
                        # matching a keyword to a topic LABEL by shared words
                        # catches 2 of 13 real terms, because "dog bite lawyer"
                        # does not contain "personal injury". Membership is
                        # exact and the build already knows it. (2026-08-22)
                        "terms": list(t.get("seeds") or []),
                        "services": len([x for x in services
                                         if service_topic(x.get("service", ""),
                                                          topics) == t["label"]])}
                       for t in topics],
            "topic_fixes": topic_fixes,
            "geo_forms": geo_form_report,
            "service_forms": service_form_report,
            "pool_dropped": [{"service": d[0], "word": d[1]} for d in (pool_dropped or [])],
            "pool_added": [{"service": a[0], "volume": a[1]} for a in (pool_added or [])],
            "pool_status": pool_status,
        }

    refined = claude_refine_keywords(seeds, markets, brand, domain,
                                     ultra + competitive + long_tail, site_terms,
                                     business_desc=biz, site_pages=site_pages)
    used_claude = False
    biz_out = biz
    if refined and (refined["ultra"] or refined["competitive"]):
        ultra       = refined["ultra"][:CFG["ultra_bucket_size"]] or ultra
        competitive = refined["competitive"][:CFG["competitive_bucket_size"]] or competitive
        if refined["long_tail"]:
            long_tail = refined["long_tail"][:CFG["longtail_target"]]
        used_claude = True
        biz_out = biz or refined.get("business", "")

    full = (ultra + competitive + long_tail)[:CFG["list_cap"]]

    exact = fetch_exact_volume([r["keyword"] for r in full], markets, state,
                               national=national_demand)
    if exact:
        for r in full:
            v = exact.get(r["keyword"].lower())
            if v is not None:
                r["volume"] = v

    fs = {r["keyword"] for r in full}
    return {
        "ultra":       [r for r in ultra if r["keyword"] in fs],
        "competitive": [r for r in competitive if r["keyword"] in fs],
        "long_tail":   [r for r in long_tail if r["keyword"] in fs],
        "head":        [r for r in (ultra + competitive) if r["keyword"] in fs],
        "all":         full,
        "refined_by_ai": used_claude,
        "business_desc": biz_out if used_claude else "",
        "site_pages_found": len(site_pages),
        # Carried on this path too — the storefront read is independent of
        # whether the grid build ran, and dropping it here would silently undo
        # the flip for any client that falls through to the non-grid list.
        "national_demand": bool(national_demand),
        "national_demand_reason": national_demand_reason,
        "ecommerce_detected": bool(ecom_found),
        "ecommerce_reason": ecom_reason,
        "ecommerce_suppressed": ecom_suppressed,
    }

# ---------------------------------------------------------------------------
# STAGE 3a — metrics -> competitive adder
# ---------------------------------------------------------------------------
def fetch_keyword_difficulty(kws, markets, state):
    """Labs bulk keyword difficulty (1-100 organic ranking difficulty). Separate
    call from the Google Ads bid data. Returns (kd_map, error_or_None) so the
    caller can surface why it's empty instead of silently failing."""
    if not kws:
        return {}, None
    try:
        # Labs endpoints want a numeric location_code, not location_name (which
        # the Google Ads endpoints use). 2840 = United States. Keyword difficulty
        # is a national-level organic metric, so country-level is appropriate.
        payload = [{"keywords": kws[:1000],
                    "location_code": 2840,
                    "language_code": "en"}]
        data = dfs_post("/dataforseo_labs/google/bulk_keyword_difficulty/live", payload)
        task = (data.get("tasks") or [{}])[0]
        # surface API-level errors (auth, plan, balance) explicitly
        if task.get("status_code") not in (20000, None) and not task.get("result"):
            return {}, f"{task.get('status_code')}: {task.get('status_message')}"
        res = task.get("result") or []
        kd = {}
        for block in res:
            for it in (block.get("items") or []):
                k = it.get("keyword")
                if k is None:
                    continue
                # difficulty can appear as a top-level field or nested
                v = it.get("keyword_difficulty")
                if v is None:
                    v = (it.get("keyword_properties") or {}).get("keyword_difficulty")
                if v is not None:
                    kd[k] = v
        return kd, None
    except requests.HTTPError as e:
        return {}, f"HTTP {e.response.status_code if e.response else '?'}"
    except Exception as e:
        return {}, str(e)[:80]

def _strip_markets(kw, markets, state=None):
    """Remove the trailing geo modifier so we can look up bid/difficulty data,
    which the APIs key to the bare term ('adhd treatment'), not the geo form
    ('adhd treatment san diego'). Grid keywords may also carry a state suffix
    ('commercial contractor kaukauna wi'), so strip that FIRST — otherwise the
    city never matches the end of the string and nothing gets stripped, which
    silently kills the bid lookup."""
    k = kw
    # Strip whichever state abbr this keyword carries — in a multi-state grid
    # different keywords end in different abbrs (nj / pa / de).
    abbrs = set()
    if state:
        a = STATE_ABBREV.get(state.strip().lower(), "")
        if a: abbrs.add(a)
    for m in markets:
        a = STATE_ABBREV.get((market_state(m, state) or "").lower(), "")
        if a: abbrs.add(a)
    for a in abbrs:
        if k.lower().endswith(" " + a):
            k = k[: -(len(a) + 1)].strip()
            break
    # Then strip the city — match on the parsed city name, not the raw
    # "Cherry Hill, NJ" pill text.
    city_names = sorted({market_city(m, state) for m in markets}, key=len, reverse=True)
    for c in city_names:
        if c and k.lower().endswith(" " + c.lower()):
            k = k[: -(len(c) + 1)].strip()
            break
    return k

# Verticals Google Ads restricts. The Keyword Planner "metrics for keywords you
# provide" endpoint (search_volume) policy-filters these and returns NO rows,
# while the keyword-IDEAS endpoint answers normally — which is why a cannabis
# client can show real volumes in Step 1 and a $0 adder in Step 2 off the same
# API key (Grav, 2026-08-04). Matched by substring against the RZ industry text.
# Substrings, so they must not appear inside innocent categories: bare "gun"
# hits "Burgundy" and bare "adult" hits "Adult Day Care" / "Adult Education",
# both real RZ categories that are not restricted at all.
RESTRICTED_VERTICALS = (
    "cannabis", "marijuana", "cbd", "hemp", "kratom", "vape", "smoke shop",
    "dispensary", "tobacco", "firearm", "ammunition", "gun shop", "gun range",
    "gambling", "casino", "sportsbook", "adult entertainment",
)


def restricted_vertical(industry=""):
    """Which restricted term (if any) this industry matches. '' when none."""
    ind = (industry or "").strip().lower()
    for k in RESTRICTED_VERTICALS:
        if k in ind:
            return k
    return ""


def fetch_bids_via_ideas(terms, location_name):
    """Top-of-page bids from the keyword-IDEAS endpoint.

    Used only when search_volume returns nothing. Same Google Ads account, same
    payload shape — but keywords_for_keywords is not policy-filtered, so it is
    the one endpoint that answers for restricted verticals. It returns cpc and
    bid fields alongside search_volume; the candidate-pool parse throws them
    away, so they were being paid for and discarded.

    Ideas responses are seeded BY the terms and come back with extra keywords,
    so only exact matches on what we asked for are kept. Returns
    ({keyword: {bid, cpc, volume}}, error_or_None).
    """
    kws = dfs_kw_list(terms)
    if not kws:
        return {}, None
    try:
        payload = [{"keywords": kws[:200], "location_name": location_name,
                    "language_code": "en"}]
        data = dfs_post("/keywords_data/google_ads/keywords_for_keywords/live", payload)
        task0 = (data.get("tasks") or [{}])[0]
        if task0.get("status_code") not in (20000, None):
            return {}, f"{task0.get('status_code')}: {task0.get('status_message')}"
        want = set(kws)
        out = {}
        for it in (task0.get("result") or []):
            k = (it.get("keyword") or "").lower()
            if k not in want:
                continue                      # ideas the seed pulled in, not ours
            out[k] = {"bid": it.get("high_top_of_page_bid") or 0,
                      "cpc": it.get("cpc") or it.get("high_top_of_page_bid") or 0,
                      "volume": it.get("search_volume") or 0}
        return out, None
    except Exception as e:
        return {}, str(e)


def fetch_bids_via_labs(terms, markets, state, national=False):
    """Bids and CPC from the DataForSEO LABS keyword database.

    The last resort, and the one most likely to answer in a restricted vertical.
    Labs is DataForSEO's own aggregated database rather than a live Google Ads
    call, so it is not subject to the ad-policy filtering that empties
    search_volume for cannabis and friends — which is exactly why keyword
    DIFFICULTY has been coming back for these clients all along.

    The keyword_info block this reads already reaches us on every Step 1 run:
    fetch_exact_volume calls the same endpoint and keeps search_volume alone,
    discarding cpc and the bid fields sitting beside it.

    NOTE: Labs CPC is modelled, not a live auction reading, so it can differ
    from Google Ads. The caller reports which source supplied the number so an
    adder built on Labs is never mistaken for one built on Google Ads.
    Returns ({keyword: {bid, cpc, volume}}, error_or_None).
    """
    kws = dfs_kw_list(terms)
    if not kws:
        return {}, None
    loc_field, _loc_used = _labs_loc_field(markets, state, national)
    try:
        payload = [{"keywords": [k.lower() for k in kws[:1000]],
                    **loc_field, "language_code": "en"}]
        data = dfs_post("/dataforseo_labs/google/keyword_overview/live", payload)
        task0 = (data.get("tasks") or [{}])[0]
        if task0.get("status_code") not in (20000, None) and not task0.get("result"):
            return {}, f"{task0.get('status_code')}: {task0.get('status_message')}"
        out = {}
        for block in (task0.get("result") or []):
            for it in (block.get("items") or []):
                k = (it.get("keyword") or "").lower()
                ki = it.get("keyword_info") or {}
                if not k:
                    continue
                bid = ki.get("high_top_of_page_bid") or 0
                out[k] = {"bid": bid,
                          "cpc": ki.get("cpc") or bid or 0,
                          "volume": ki.get("search_volume") or 0}
        return out, None
    except Exception as e:
        return {}, str(e)


def stage3_metrics(head, markets, state, national=False, industry=""):
    geo_kws = [r["keyword"] for r in head]
    if not geo_kws:
        return {"adder": 0, "median_score": 0, "bids": {}, "cpc": {}, "kd": {}}
    # Map each geo head term -> its bare form; query metrics on the bare forms
    # (which have real bid/difficulty data), then attribute results to both keys.
    bare_of = {g: _strip_markets(g, markets, state) for g in geo_kws}
    bare_unique = list(dict.fromkeys(bare_of.values()))

    # Google Ads bid data is sparse at small-city granularity (e.g. Kaukauna, WI
    # returns no rows even for real terms). Advertiser demand for the adder
    # doesn't need city precision, so fall back city -> state -> US and report
    # which level actually supplied the data.
    # Step 1 prices national demand on geo-less volume; Step 2 has to measure
    # competition on the SAME basis or the two halves of one quote describe
    # different markets. This was never plumbed through, so a nationwide client
    # with cities entered had its bids read from those cities (2026-08-04).
    primary_loc = "United States" if national else loc_string(markets, state)
    loc_chain = [primary_loc]
    if state and f"{state},United States" not in loc_chain:
        loc_chain.append(f"{state},United States")
    if "United States" not in loc_chain:
        loc_chain.append("United States")
    bid_err = None
    items = []
    bid_loc_used = primary_loc
    _dl = _deadline()
    for _loc in loc_chain:
        _left = _remaining(_dl)
        if _left is None:
            break                       # budget spent: use what we have
        payload = [{"keywords": dfs_kw_list(bare_unique),
                    "location_name": _loc,
                    "language_code": "en"}]
        try:
            data = dfs_post("/keywords_data/google_ads/search_volume/live", payload)
            task0 = (data.get("tasks") or [{}])[0]
            # DataForSEO reports per-task problems in status_code/status_message
            # even on an HTTP 200, so surface those rather than returning nothing.
            if task0.get("status_code") not in (20000, None):
                bid_err = f"{task0.get('status_code')}: {task0.get('status_message')}"
                continue
            got = (task0.get("result") or [])
            if got and not items:
                items = got            # keep the first non-empty result set
                bid_loc_used = _loc
            # only stop early if this level actually carries bid values
            if got and any((it.get("high_top_of_page_bid") or 0) for it in got):
                items = got
                bid_loc_used = _loc
                bid_err = None
                break
        except Exception as e:
            bid_err = str(e)
    bare_bid = {it["keyword"]: (it.get("high_top_of_page_bid") or 0) for it in items}
    bare_cpc = {it["keyword"]: (it.get("cpc") or it.get("high_top_of_page_bid") or 0) for it in items}

    # FALLBACK. search_volume gave us no usable bid — either it returned no rows
    # at all, or rows with every bid blank. Ask the ideas endpoint, which is not
    # policy-filtered. Costs one extra call and only on the failing path.
    bid_source = "search_volume"
    ideas_err = labs_err = None

    def _adopt(src, got):
        """Take bids from a fallback source. Returns True if it supplied any."""
        nonlocal bare_bid, bid_source, bid_err
        if not any((v or {}).get("bid") for v in got.values()):
            return False
        bare_bid = {k: v["bid"] for k, v in got.items() if v.get("bid")}
        for k, v in got.items():
            if v.get("cpc"):
                bare_cpc[k] = v["cpc"]
        bid_source = src
        bid_err = None
        return True

    if not any(bare_bid.values()):
        # 2nd: Google Ads keyword IDEAS — same account, not policy-filtered.
        ideas, ideas_err = fetch_bids_via_ideas(bare_unique, bid_loc_used)
        if not _adopt("keyword_ideas", ideas):
            # 3rd: LABS. Its own database rather than a live Ads call, so it is
            # the source that still answers when advertising in the vertical is
            # banned outright. Modelled CPC, hence last.
            labs, labs_err = fetch_bids_via_labs(bare_unique, markets, state,
                                                 national=national)
            _adopt("labs", labs)
    bare_kd, kd_err = fetch_keyword_difficulty(bare_unique, markets, state)

    # Attribute to both the geo key (for the table) and the bare key.
    bids, cpc, kd = {}, {}, {}
    for g in geo_kws:
        b = bare_of[g]
        if bare_bid.get(b):  bids[g] = bare_bid[b]; bids[b] = bare_bid[b]
        if bare_cpc.get(b):  cpc[g]  = bare_cpc[b]; cpc[b]  = bare_cpc[b]
        if bare_kd.get(b) is not None: kd[g] = bare_kd[b]; kd[b] = bare_kd[b]

    kd_vals = [v for v in {bare_of[g]: kd.get(g) for g in geo_kws}.values()
               if isinstance(v, (int, float))]
    median_kd = int(statistics.median(kd_vals)) if kd_vals else None

    lo, hi = CFG["bid_score_breaks"]
    # Score only on head terms that returned bid data (don't let missing data
    # count as 0 and drag the median down).
    have_bid = [bids.get(g, 0) for g in geo_kws if bids.get(g, 0)]
    scores = [2 if b >= hi else 1 if b >= lo else 0 for b in have_bid]
    median_score = int(statistics.median(scores)) if scores else 0
    # Bid distribution so the panel can show what the score is derived from.
    # Use unique bare-term bids (the actual data points the score is built on).
    bid_vals = [v for v in bare_bid.values() if v]
    bid_stats = None
    if bid_vals:
        bid_stats = {"median": round(statistics.median(bid_vals), 2),
                     "min": round(min(bid_vals), 2),
                     "max": round(max(bid_vals), 2),
                     "n": len(bid_vals), "n_total": len(bare_unique)}
    # Competitive adder: prefer CPC-scaled (adder tracks median bid = click value),
    # fall back to the flat score buckets when there's no bid data to scale on.
    flat_adder = CFG["competitive_adder"][median_score]
    adder = flat_adder
    adder_basis = "flat"
    cpc_used = None
    n_bids = (bid_stats or {}).get("n", 0)
    min_n = int(CFG.get("cpc_adder_min_samples", 1) or 1)
    cpc_low_conf = bool(bid_stats) and n_bids < int(CFG.get("cpc_adder_low_confidence_n", 3))
    if (CFG.get("cpc_adder_enabled") and bid_stats and bid_stats["median"]
            and n_bids >= min_n):
        med_cpc = bid_stats["median"]
        cpc_used = med_cpc
        free = CFG.get("cpc_adder_free_below", 5.0)
        if med_cpc > free:
            # Piecewise: $/CPC at the normal rate up to the knee, then a much
            # steeper rate above it. Brendan's premium grows super-linearly with
            # CPC — dental ($18) +$400 over card, Waytek ($60) +$500, Rockingham
            # ($121, insurance carrier) +$2,500. A single multiplier can't fit
            # both ends; the knee can.
            knee = CFG.get("cpc_adder_knee", 50.0)
            raw = (min(med_cpc, knee) * CFG.get("cpc_adder_mult", 3.0)
                   + max(0.0, med_cpc - knee) * CFG.get("cpc_adder_mult_high", 14.0))
            capped = min(raw, CFG.get("cpc_adder_cap", 1500))
            adder = int(round(capped / 50.0) * 50)
            adder_basis = "cpc"
        else:
            adder = 0
            adder_basis = "cpc"
    # No bid data from EITHER endpoint. The adder is then not a measurement of
    # $0 competition, it is an absence of evidence — and this quote is priced as
    # if the vertical were free. Say so and make the operator decide.
    restricted = restricted_vertical(industry)
    no_bids = not bid_vals
    # When there are no bids anywhere, organic difficulty is the only competition
    # signal left. Score it on the SAME 0/1/2 ladder as bids and APPLY it: an
    # adder the tool can derive from evidence should not need to be typed in by
    # hand. The operator can still override, but the default is now a reasoned
    # number rather than a blank field (2026-08-04).
    # ALWAYS SCORED, EVEN WHEN BIDS EXIST. Organic difficulty was computed on
    # every quote and consulted only when bid data was completely missing, so on
    # a quote WITH bids it was measured, printed on the panel, and thrown away.
    # Amare came back at 7/100 and NPAIHB at 25/100 with the price identical, and
    # difficulty is the one competition signal the industry actually scopes on:
    # ranking effort is a function of how hard the page-one incumbents are, not
    # of how many people search. Scored here for every quote so the two numbers
    # can be COMPARED before either is trusted to move money.
    #
    # Reported, not applied, while bids exist — the bid-based adder is what every
    # calibrated price in the bench was built on, and swapping the basis would
    # move every quote already written. When there are no bids it still decides,
    # exactly as before. (2026-08-17)
    kd_suggested_adder = kd_score = None
    if median_kd is not None:
        klo, khi = CFG.get("kd_score_breaks", [30, 60])
        kd_score = 2 if median_kd > khi else 1 if median_kd >= klo else 0
        kd_suggested_adder = CFG["competitive_adder"][kd_score]
    if no_bids and kd_suggested_adder is not None:
        adder = kd_suggested_adder
        adder_basis = "kd"
    # Only a total absence of evidence still stops the quote: no bids from any
    # of the three sources AND no organic difficulty either. Then there really
    # is nothing to reason from and a human has to supply the number.
    adder_blocked = no_bids and kd_suggested_adder is None
    # The disagreement is the interesting number: a quote priced at $0 on bids
    # while organic difficulty says the page is hard is exactly the case the
    # formula currently cannot see.
    kd_vs_cpc = None
    if adder_basis == "cpc" and kd_suggested_adder is not None:
        kd_vs_cpc = {"cpc_adder": int(adder or 0),
                     "kd_adder": int(kd_suggested_adder),
                     "median_kd": median_kd, "kd_score": kd_score,
                     "delta": int(kd_suggested_adder) - int(adder or 0)}
    return {"adder": adder, "adder_basis": adder_basis, "cpc_used": cpc_used,
            "cpc_low_confidence": cpc_low_conf, "cpc_n_bids": n_bids,
            "flat_adder": flat_adder,
            "bid_source": bid_source,
            "bid_ideas_error": ideas_err,
            "bid_labs_error": labs_err,
            "adder_blocked": adder_blocked,
            "kd_suggested_adder": kd_suggested_adder,
            "kd_score": kd_score,
            "kd_vs_cpc": kd_vs_cpc,
            "restricted_vertical": restricted,
            "bid_error": bid_err,
            "bid_location": bid_loc_used,
            "bid_location_fallback": (bid_loc_used != primary_loc),
            "bid_terms_queried": bare_unique[:8],
            "n_markets": len(markets),
            "median_score": median_score, "bids": bids, "cpc": cpc,
            "bid_stats": bid_stats, "breaks": [lo, hi],
            "kd": kd, "median_kd": median_kd, "kd_error": kd_err}

# ---------------------------------------------------------------------------
# STAGE 3b — rank check -> table + zero-ranking + PAA
# ---------------------------------------------------------------------------
def market_for_keyword(kw, markets, state=""):
    """Which entered market a keyword names, read off the keyword itself.

    The grid tags every row with its city, but a keyword list restored from a
    saved quote carries only the text — and that is the common case, because an
    operator re-checks rankings without rebuilding step 1. Depending on the tag
    meant the per-market rank check silently did nothing on every existing quote
    (2026-08-10).

    The keyword ends with the market: "junk removal morristown tn". Match on the
    END, longest first, so "oak ridge" wins over a market called "ridge" and
    "kansas city ks" is never mistaken for "kansas".
    """
    k = " " + re.sub(r"\s+", " ", clean_kw((kw or "").lower())).strip()
    if not k.strip():
        return ""
    cands = []
    for m in (markets or []):
        city, st = parse_market(m, state)
        c = clean_kw((city or "").lower()).strip()
        if not c:
            continue
        ab = (STATE_ABBREV.get((st or state or "").strip().lower(), "") or "").lower()
        forms = [f"{c} {ab}", c] if ab else [c]
        # A county reads either way round: "roane county tn" / "roane tn".
        ck = county_key(m, state)
        if ck:
            bare = _COUNTY_SUFFIX.sub("", c).strip()
            if bare and ab:
                forms.append(f"{bare} {ab}")
        for f in forms:
            if f:
                cands.append((len(f), f, m))
    for _n, form, m in sorted(cands, reverse=True):
        if k.endswith(" " + form):
            return m
    return ""


def _serp_one(kw, domain_dom, markets, state, brand, top_n, deadline=None,
              loc_override=""):
    """One keyword's SERP call. Returns (position, [paa questions], [rival domains]).
    Depth tracks top_n (<=100 is one DataForSEO unit either way). Works within a shared batch DEADLINE: the
    platform kills any request near ~30s, so retrying past the budget doesn't
    save this keyword — it kills the WHOLE batch, failing keywords that had
    already finished. Better to fail one fast and let the retry pass get it."""
    depth = max(top_n, 10)
    payload = [{"keyword": kw,
                # loc_override is an already-built DataForSEO location string —
                # step 3 measures each grid row in the market it names, so the
                # caller resolves the location per keyword rather than per batch.
                "location_name": loc_override or loc_string(markets, state),
                "language_code": "en", "depth": depth}]
    last_err = None
    for attempt in range(2):
        remaining = (deadline - time.time()) if deadline else 20
        if remaining < 4:
            raise last_err or TimeoutError("rank-check batch budget exhausted")
        tmo = min(14 if attempt == 0 else remaining - 1, remaining, 20)
        try:
            # /regular, not /advanced: organic-only, ~10x smaller JSON. Depth-100
            # advanced responses are megabyte-scale and parsing 20 of them
            # serializes on Render free tier's 0.1 vCPU; regular is also cheaper.
            # Cost: no PAA items — only ever used for the non-grid long-tail
            # top-up, an acceptable trade.
            data = dfs_post("/serp/google/organic/live/regular", payload, timeout=tmo)
            break
        except Exception as e:
            last_err = e
            if attempt == 0:
                time.sleep(1)
    else:
        raise last_err
    # DataForSEO answers HTTP 200 and reports per-TASK problems in status_code.
    # Every other caller in this module checks it; _serp_one did not — the one
    # whose result sets the price. An unresolvable location_name ("near me,
    # Tennessee") returns 40501 with result=None, which fell straight through to
    # items=[] -> pos=None -> "Not Found" for EVERY keyword in the batch. 0/35
    # ranked then drew the largest zero-ranking uplift off a lookup that never
    # ran, and the result was cached for six hours, so the re-check returned
    # instantly and agreed with itself. Raise instead: the caller already renders
    # a failed lookup as "—", keeps it out of the ranked fraction, and does not
    # cache it. (2026-08-10)
    task0 = ((data or {}).get("tasks") or [{}])[0] or {}
    if task0.get("status_code") not in (20000, None):
        raise RuntimeError(f"{task0.get('status_code')}: {task0.get('status_message')}")
    res = (task0.get("result") or [{}])[0] or {}
    items = res.get("items", []) or []
    # THE SAME PARSE AS THE TASK PATH. This was a hand-rolled copy of
    # _serp_parse_items, so adding the incumbent domains to that one left this
    # one alone — and this is the path that runs when Google answers fast enough
    # to skip the queue. Nob Hill Dental resolved entirely live and the market
    # signals block came back "0 incumbents measured" for a client whose page one
    # was sitting in the response. One parser now, both paths. (2026-08-17)
    return _serp_parse_items(items, domain_dom, brand)

def stage3_rankcheck(all_kws, domain, markets, state, brand):
    top_n = CFG["zero_ranking_top_n"]
    dom = (domain or "").replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    # Cap the number of SERP calls to stay under the platform timeout.
    capped = all_kws[:CFG["rank_check_cap"]]
    kws = [r["keyword"] for r in capped]

    # Fire SERP calls in parallel; keep results aligned to input order.
    results = [None] * len(kws)
    with ThreadPoolExecutor(max_workers=CFG["rank_check_workers"]) as ex:
        futs = {ex.submit(_serp_one, kw, dom, markets, state, brand, top_n): i
                for i, kw in enumerate(kws)}
        for fut in futs:
            i = futs[fut]
            try:
                results[i] = fut.result() + (False,)   # (pos, paa, rivals, err)
            except Exception:
                # One bad keyword shouldn't sink the quote — but it must not be
                # counted as "not ranking" either. A failed lookup measured
                # nothing, and folding it into the denominator inflates the
                # zero-ranking percentage and therefore the price. Same rule the
                # batched /api/rankings path already follows. (2026-08-10)
                results[i] = (None, [], [], True)

    table, paa, ranked, errors = [], [], 0, 0
    rivals = {}
    for kw, (pos, qs, doms, err) in zip(kws, results):
        table.append({"keyword": kw, "position": pos, "error": err})
        paa.extend(qs)
        for _d in (doms or []):
            rivals[_d] = rivals.get(_d, 0) + 1
        if err:
            errors += 1
            continue
        if pos is not None and pos <= top_n:
            ranked += 1
    checked = max(len(kws) - errors, 0)
    frac = (ranked / checked) if checked else 0.0
    return {"table": table, "ranked": ranked, "frac": frac,
            "checked": checked, "errors": errors,
            # Nothing was measured, so there is no evidence of zero ranking.
            "zero_ranking": bool(checked) and frac < CFG["zero_ranking_frac"],
            "paa_pool": list(dict.fromkeys(paa)),
            "rivals": [{"domain": d, "appearances": n}
                       for d, n in sorted(rivals.items(),
                                          key=lambda kv: (-kv[1], kv[0]))
                       [:int(CFG.get("serp_rival_cap", 12))]]}

# ---------------------------------------------------------------------------
# STAGE 4 — pricing
# ---------------------------------------------------------------------------
def _tier_uplift(value, tiers):
    """Given a value and a list of [threshold, uplift_pct] sorted high-to-low,
    return the uplift_pct of the first threshold the value meets (else 0)."""
    for thresh, uplift in tiers:
        if value >= thresh:
            return uplift
    return 0

def rank_location(markets, state, national=False):
    """Where to measure rankings. Under national demand the keywords carry no
    city ("energy gummies", not "energy gummies texas"), so a Texas-localised
    SERP would be reporting a local result for a national term. Brendan's own
    MPG table (2026-06-10) lists one national rank per bare keyword, which is
    the format this matches. The zero-ranking uplift keys off these positions,
    so the location has to describe the same market the keywords do.

    BUT national demand and a national SERP are two different claims, and only
    the first one follows from a storefront. National demand says "this client
    sells everywhere, so measure DEMAND nationally" — correct for a shop that
    ships. A national SERP says "measure whether they OUTRANK the whole country",
    which no regional retailer ever does: Ski Barn (4 NJ stores, skibarn.com)
    scored 0/20 in the national top 100 on bare terms like "ski shop" and drew
    the largest zero-ranking uplift, +14% on the hard base, off a test it could
    not have passed (2026-08-07). Arithmetically inevitable, not a finding.

    So: if the client named markets, measure in the primary market even under
    national demand. Volume stays national; visibility is asked where their
    customers actually search. Only a client with NO markets at all — genuine
    pure-play ecommerce, the MPG case — gets measured nationally.
    """
    if national and not markets:
        return "United States"
    return loc_string(markets, state)


def rank_location_note(markets, state, national=False):
    """Human-readable 'where this was measured', for the Step 3 panel. A 0%
    ranked result is only interpretable next to the place it was measured."""
    loc = rank_location(markets, state, national)
    # The panel renders "<b>Measured in X</b> — {note}", so the note must NOT
    # open by repeating the location: it printed "Measured in mill valley,
    # California — Measured in mill valley, California. Demand is pulled…".
    # And "a regional retailer" was Ski Barn's wording leaking onto every client;
    # NASSCO is a standards body and the sentence read as nonsense. (2026-08-10)
    if loc == "United States":
        return {"location": loc, "scope": "national",
                "note": "No markets are set, so there is nowhere local to measure "
                        "— this is the whole-country result."}
    if national:
        return {"location": loc, "scope": "local_under_national",
                "note": "Demand is pulled nationally, but visibility is measured "
                        "where this client's customers search. A business with a "
                        "real local footprint never outranks the whole country, "
                        "and scoring it that way would raise the price off a test "
                        "it cannot pass. If the client has no local footprint, "
                        "clear the markets and this measures nationally instead."}
    return {"location": loc, "scope": "local",
            "note": "The client's primary market."}


def resolve_national_demand(industry="", band="", manual=False, markets=None,
                            goal=""):
    """Should this client be priced on GEO-LESS (national) search volume?

    Three sources, any of which is sufficient:
      1. RZ industry taxonomy — a rule carrying national_demand (ecommerce and
         its product-brand siblings). Industry sets no price of its own; it
         only says "measure demand nationally," and the volume/competition/
         visibility signals then price the client on their own merits.
      2. Geo scope of nationwide — no cities, so the pull is already geo-less.
      3. Manual operator checkbox, for the cases RZ mistags.

    A fourth source — an actual storefront detected on the client's own site —
    is applied inside stage1b_refine, because it is only knowable once the
    sitemap has been read. See detect_ecommerce().

    MARKETS VETO THE INDUSTRY TAG (2026-08-07). BE on the Ski Barn quote: "They
    have specific locations, the quote didn't consider the local impact, it just
    treated them as an ecommerce/nationwide option." An RZ tag of "Retail -
    General / E-commerce" describes what the client SELLS; the market list
    describes where they TRADE, and the operator typed it deliberately. With both
    present and the scope not nationwide, the market list wins — a store with
    four premises is a local campaign that also ships, and pricing it on national
    demand quotes a different campaign: geo-less terms, national volume, and a
    rank check against the whole country.

    The manual switch and an explicit nationwide scope still force national.
    Both are direct statements of intent rather than inferences from a taxonomy.

    Returns (bool, reason_string) so the UI can show WHY it flipped.
    """
    if manual:
        return True, "manual override"
    if band == "nationwide":
        return True, "nationwide geo scope"
    # The goal is the client's own statement of what they are buying, taken off
    # the adtini order form — so it outranks the market list, which describes
    # where they trade. "Online Sales" with four stores entered is a client
    # asking to be sold ecommerce visibility; price the demand nationally and
    # say so. Rankings still get measured locally when markets exist — see
    # rank_location() — because that is a question about this client, not about
    # the size of the market.
    _g = goal_forces_national(goal)
    if _g:
        return True, f"goal: {_g}"
    ind = (industry or "").strip().lower()
    mk = [m for m in (markets or []) if str(m).strip()]
    for k, r in (CFG.get("industry_pricing") or {}).items():
        if k in ind and r.get("national_demand"):
            if mk:
                return False, (f"industry “{k}” suggests national demand, but "
                               f"{len(mk)} market{'' if len(mk) == 1 else 's'} "
                               "are entered — priced on LOCAL demand. Set Geo "
                               "scope to Nationwide, or tick Price on national "
                               "demand, if this client sells everywhere.")
            return True, f"industry: {k}"
    return False, ""


# Storefront fingerprints, checked against the sitemap URLs already collected
# by fetch_site_pages. Ordered most-specific first so the reason names the
# platform when it can. Each entry: (label, regex, weight_threshold).
_ECOM_SIGNATURES = [
    ("Shopify",      re.compile(r"/(collections|products)/", re.I)),
    ("WooCommerce",  re.compile(r"/(product-category|product-tag)/", re.I)),
    ("BigCommerce",  re.compile(r"/(categories|brands)/[^/]+/?$", re.I)),
    ("Magento",      re.compile(r"/catalog/(product|category)/", re.I)),
    ("store",        re.compile(r"/(shop|store|product|catalogue|catalog)/", re.I)),
]
# A cart or checkout route is conclusive on its own — nothing but a store has
# one — so it does not need the repetition threshold the catalog paths do.
_ECOM_CHECKOUT = re.compile(r"/(cart|checkout|basket|my-account/orders)\b", re.I)


# A meta description is written for a search snippet, not as a description of the
# business, and plenty of them open with the call to action. NASSCO's is "Click
# here to read and learn more about the education, technical resources and
# advocacy for the underground infrastructure industry" — offered verbatim as the
# business description, which then feeds the grounding filter and the AI prompt,
# so "click", "here" and "read" join the client's accepted vocabulary.
#
# Detected and REFUSED rather than rewritten. Every attempt to cut the CTA off the
# front either left a fragment ("To schedule your free estimate...") or ate real
# words, and there is a better source sitting right behind it: refusing falls back
# to the AI read of the site, which describes the business because that is what it
# was asked for. (2026-08-11)
_CTA_OPENER = re.compile(
    r"^\s*(?:please\s+)?"
    r"(?:click|tap|visit|contact|call|see|read|learn|find|discover|explore|"
    r"schedule|request|sign\s+up|subscribe|join)\b",
    re.I)


def _is_cta(text):
    """Does this meta description open by telling the reader what to do?"""
    return bool(_CTA_OPENER.match((text or "").strip()))


def detect_ecommerce(urls, min_hits=3):
    """Is the client running a storefront? Reads ONLY the sitemap URLs that
    fetch_site_pages already collected — no extra HTTP calls, no extra latency.

    Exists because national-demand pricing hangs off the RZ industry tag, and
    the tag is routinely missing or wrong: a smoke shop selling nationwide gets
    filed under a local retail category, so its volume is pulled per-city and
    the quote is built on the wrong demand figure entirely.

    A single /product/ URL is not a store — brochure sites use that path for
    one product page — so catalog patterns need `min_hits` distinct matches.
    A cart or checkout route is accepted on its own.

    Returns (bool, reason_string).
    """
    urls = [u for u in (urls or []) if u]
    if not urls:
        return False, ""
    for u in urls:
        if _ECOM_CHECKOUT.search(u):
            return True, "cart/checkout page on site"
    if any("sitemap_products" in u.lower() for u in urls):
        return True, "Shopify product sitemap"
    for label, rx in _ECOM_SIGNATURES:
        hits = {u for u in urls if rx.search(u)}
        if len(hits) >= min_hits:
            return True, f"{len(hits)} {label} product URLs on site"
    return False, ""


def _volume_dollar_add(total_volume, free_below, brackets):
    """Fixed $ added for search volume above a normalized baseline, using a
    declining marginal rate (tax-bracket style). Each bracket [lo, hi, rate]
    charges 'rate' $/search for the volume that falls within [lo, hi]; a hi of
    None means open-ended. Returns total $ added (0 if at/below the baseline)."""
    if not total_volume or total_volume <= free_below:
        return 0
    add = 0.0
    for b in brackets:
        lo, hi, rate = b[0], b[1], b[2]
        if total_volume > lo:
            top = total_volume if hi is None else min(total_volume, hi)
            band = max(0, top - lo)
            add += band * rate
    return add

def stage4_price(band, adder, zero_ranking, addon_markets=0, markup_pct=None,
                 pct_not_ranking=None, total_volume=None, base_override=None,
                 ecommerce=False, industry="", ai_search=False,
                 national_demand=False, geo_override=None, addon_override=None,
                 goal="", pageone_rank=None, site_rebuild="",
                 _formula_pass=False):
    if markup_pct is None:
        markup_pct = CFG["default_markup_pct"]
    # THE RANK CHECK MEASURED A SITE THAT IS BEING REPLACED. See the CFG note
    # on rebuild_new_domain_pct_not_ranking. Only a NEW DOMAIN moves the number
    # -- a same-domain rebuild keeps its rankings through redirects, so its
    # measurement still describes the campaign.
    rebuild = str(site_rebuild or "").strip().lower()
    rebuild_applied = False
    if rebuild == "new":
        _rp = CFG.get("rebuild_new_domain_pct_not_ranking", 100)
        if _rp is not None:
            rebuild_applied = pct_not_ranking != float(_rp)
            pct_not_ranking = float(_rp)
    # RETAIL IS CANONICAL (2026-08-05). The anchors and every hard-dollar extra
    # in CFG were back-solved from Brendan's QUOTED tier prices as
    # client / 1.35, so they are a calibration basis, not a real cost. The
    # agency split is 35% OF GROSS (see the SSG/Vici grid: agency profit =
    # retail x 0.35, Vici bills retail x 0.65), which is what the reputation
    # tool already does. Reading those anchors as true cost overstated partner
    # cost by 14% and implied a 25.9% margin instead of 35%.
    #
    # Fixed WITHOUT touching a single calibrated constant: convert the
    # calibration basis to true cost once, then divide by (1 - margin). At the
    # calibrated 35% this is algebraically identical to the old x1.35, so every
    # client price is unchanged; away from 35% it now moves the way a
    # margin-of-gross should. CAL_* are frozen history, not business inputs.
    # CONSTANTS ARE PARTNER COST NOW (2026-08-13), so there is nothing to
    # convert — this factor stays as a named 1.0 rather than disappearing,
    # because every comment below is written in terms of it and the old
    # calibration history is only readable with it in view.
    CAL_MARKUP, CAL_MARGIN = 1.35, 0.35
    cal_to_hard = 1.0                                      # was 0.8775
    mg = min(0.95, max(0.0, markup_pct / 100.0))           # margin OF GROSS
    m = cal_to_hard / (1.0 - mg)                           # basis -> retail
    to_true_hard = 1.0 - mg                                # retail -> true cost

    # PARTNER HARD COST IS THE CANONICAL, ROUNDED FIGURE (2026-08-05). The
    # ladder is now built cost-first: calibration basis -> true partner cost,
    # rounded to a clean $50 -> retail = cost / (1 - margin), rounded UP to $50.
    #
    # Cost rounds to the NEAREST $50, not up. Rounding both numbers up compounds
    # (cost ceils, then retail ceils again) and drifted retail +$0-100 above the
    # calibrated prices, mean +$37. Nearest-$50 on the cost cancels most of that:
    # retail is unchanged on 146 of 240 anchor values, mean drift +$0.40, worst
    # case one $50 step either way, and all three live anchors land exactly on
    # today's prices (2050 -> $2,800, 2100 -> $2,850, 2350 -> $3,200).
    #
    # Consequence to accept: with both figures on $50 boundaries the realised
    # margin lands 35.1-35.9% rather than exactly 35%. Two clean numbers and an
    # exact ratio cannot all three hold at once.
    def r50n(x):
        return int(round(round(x, 6) / 50.0) * 50)

    def cost_of(basis):
        return r50n(basis * cal_to_hard)

    def retail_of(cost):
        # CLIENT-FACING — rounds up. Also what the add-on market rate goes
        # through, so the two conventions cannot drift apart.
        return r50up(cost / (1.0 - mg))

    # Resolve the industry rule FIRST — national demand has to be known before
    # the anchor is picked, because it routes to the national anchor.
    rule_key, rule = None, None
    ind = (industry or "").strip().lower()
    # The industry field is multi-select (values joined with " | "), so
    # several rules can match at once. Precedence: the STRONGEST card wins
    # (largest anchor_add) — a hospital that also sells products online is
    # priced as a hospital, not as a shop.
    _matches = [(k, r) for k, r in CFG.get("industry_pricing", {}).items() if k in ind]
    if _matches:
        rule_key, rule = max(_matches, key=lambda kr: int(kr[1].get("anchor_add", 0)))
    # The legacy ecommerce checkbox no longer maps to a pricing rule (it has
    # no anchor_add as of 2026-07-25) — it is a national-demand signal only.
    # No markets= here: stage4_price never receives the market list, so passing
    # one would be a NameError. The veto is applied upstream — /api/refine,
    # /api/metrics and the rank endpoints all resolve it WITH markets, and the
    # resulting flag arrives here as `national_demand`.
    nat_demand, nat_reason = resolve_national_demand(
        industry, band, bool(ecommerce) or bool(national_demand), goal=goal)

    # A product brand priced on national demand sits on the NATIONAL anchor
    # even when the operator picked a local/statewide scope — the client's
    # cities describe where they ship, not where the demand is measured.
    anchor_band = "nationwide" if nat_demand else band
    anchor = CFG["geo_anchor"][anchor_band]                # hard cost

    # --- volume-based add: fixed $ for volume above the normalized baseline ---
    vol_add = 0
    if total_volume is not None:
        vol_add = _volume_dollar_add(total_volume, CFG.get("vol_free_below", 10000),
                                     CFG.get("volume_brackets", []))
        cap = CFG.get("volume_add_cap")
        if cap:
            vol_add = min(vol_add, cap)

    # Base before % uplift = anchor + competitive adder + volume $ add.
    base_pre = anchor + adder + vol_add
    if rule:
        base_pre += int(rule.get("anchor_add", 0))

    # WHO HOLDS PAGE ONE — see pageone_anchor_add. Added to the anchor exactly
    # like the industry rule's, and NOT scaled by the extras multiplier: this is
    # a property of the market, not one of the SERP extras a big-org card mutes.
    pageone_band = _pageone_bucket(pageone_rank)
    pageone_add = int((CFG.get("pageone_anchor_add") or {}).get(pageone_band, 0)
                      if pageone_band else 0)
    base_pre += pageone_add

    # The volume add prices UNCAPTURED demand. A client already ranking for
    # most of its terms owns that traffic; charging for it double-counts.
    ramp = CFG.get("vol_add_ramp") or None
    vol_opportunity = 1.0
    if ramp and pct_not_ranking is not None and vol_add:
        lo, hi = float(ramp[0]), float(ramp[1])
        vol_opportunity = 0.0 if hi <= lo else (float(pct_not_ranking) - lo) / (hi - lo)
        vol_opportunity = max(0.0, min(1.0, vol_opportunity))
        _prev = vol_add
        vol_add = int(round(vol_add * vol_opportunity))
        # Adjust by the DELTA — base_pre already carries the industry rule's
        # anchor_add at this point, so reassigning it from scratch silently
        # dropped the hospital/insurance premium (caught in regression).
        base_pre += (vol_add - _prev)
    vol_captured = vol_opportunity < 1.0

    # Extras suppression.
    #  - Industry rules that price on ORGANISATION size rather than SERP
    #    signals (hospital / telehealth / behavioral health) still zero both.
    #  - Nationwide scope is now governed by nationwide_service_extras, which
    #    is 1.0 as of 2026-07-25 (Brendan): volume, competition and current
    #    visibility are what separate national clients, so muting them made
    #    every national client price identically. Left as a live multiplier
    #    rather than deleted so Skidmore can be re-fit without a code change.
    nw_service = (anchor_band == "nationwide" and not (rule and rule.get("anchor_add")))
    rule_extras_off = bool(rule and rule.get("extras_off"))
    _mult = (float(CFG.get("nationwide_service_extras", 1.0)) if nw_service
             else (0.0 if rule_extras_off else 1.0))
    extras_off = _mult != 1.0
    if extras_off and vol_add:
        base_pre -= vol_add
        vol_add = int(round(vol_add * _mult))
        base_pre += vol_add

    # --- tiered zero-ranking uplift (% of head terms not ranking) ---
    zr_uplift = 0
    if pct_not_ranking is not None:
        zr_uplift = _tier_uplift(pct_not_ranking, CFG.get("zero_ranking_tiers", []))
    elif zero_ranking:
        zr_uplift = CFG.get("zero_ranking_tiers", [[0, 0]])[0][1]
    if extras_off and zr_uplift:
        zr_uplift = zr_uplift * _mult

    # MANUAL OVERRIDE: set the hard base directly; the ladder recomputes from it.
    manual_base = base_override is not None and str(base_override) != ""
    if manual_base:
        base = r50(float(base_override))
        zr_uplift = 0; vol_add = 0
    else:
        base = r50(base_pre * (1.0 + zr_uplift / 100.0))

    flat = CFG.get("tier_step_flat")
    if manual_base:
        # A manual override is the operator setting a Brendan-style base
        # directly — his premium cards ($3,950/$5,450/$6,950: Serene, Skidmore)
        # step at 38% of base, so the override ladder should too. Overriding to
        # ~$2,930 hard reproduces that card's upper tiers exactly at 35%.
        step = r50(base * CFG["step_ratio"])
    elif rule and rule.get("step_mode") == "ratio":
        # these ladders step proportionally (Brendan's ecom quote: 38% steps)
        step = r50(base * CFG["step_ratio"])
    elif anchor_band == "nationwide":
        # Brendan's national ladder steps PROPORTIONALLY — $1,500 client rungs
        # on a $3,950 base = 38% (Skidmore and MPG both). Keyed to anchor_band
        # so a product brand on national demand gets the card's shape too:
        # keeping flat $700 steps there was most of MPG's 15% shortfall (his
        # rungs are $1,500 client, the flat ladder gave $950).
        step = r50(base * CFG["step_ratio"])
    elif flat:
        # flat floor, scaling with base for premium clients: Brendan steps
        # ~$950 client on standard quotes but ~$1,300 on his biggest ladder —
        # roughly a quarter of the hard base once the base outgrows the floor.
        pct = CFG.get("tier_step_pct_of_base", 0.24)
        step = max(r50(flat), r50(base * pct))
    else:
        step = r50(base * CFG["step_ratio"])
    hard = {"base": base, "intermediate": base + step, "advanced": base + 2*step}

    # Cost first, then retail from cost.
    hard_cost = {k: cost_of(v) for k, v in hard.items()}
    client_base = retail_of(hard_cost["base"])
    floor = CFG.get("client_floor", 0)
    floored = False
    if floor and client_base < floor:
        client_base = floor
        floored = True
        cstep = (retail_of(cost_of(step)) if CFG.get("tier_step_flat")
                 else r50up(client_base * CFG["step_ratio"]))
        client = {"base": client_base,
                  "intermediate": client_base + cstep,
                  "advanced": client_base + 2*cstep}
        # Floored: retail was overridden, so cost has to be restated from it or
        # the two would describe different quotes.
        hard_cost = {k: r50n(v * to_true_hard) for k, v in client.items()}
    else:
        client = {k: retail_of(v) for k, v in hard_cost.items()}

    # ---- minimum term (applies to the whole quote, not just GEO) ----
    min_term = CFG.get("min_term_months", 6)

    # ---- Core SEO + AI Search: GEO as a % of the client's own Core SEO ----
    # Brendan: GEO averages 30-50% below SEO, rising toward parity when the
    # client has no visibility. The list price is that %; the quoted price is
    # the list less the bundle discount, applied to ALL THREE tiers.
    ai = None
    if ai_search:
        if CFG.get("geo_pricing_mode", "pct") == "card":
            card = CFG.get("geo_card", {})
            card_list = CFG.get("geo_card_list", card)
            ai = {"mode": "card",
                  "min_term_months": min_term,
                  "client_add":  {k: int(card.get(k, 0)) for k in client},
                  "client_list": {k: int(card_list.get(k, 0)) for k in client},
                  "hard_add":    {k: r50(int(card.get(k, 0)) / m) for k in client}}
        else:
            if pct_not_ranking is None:
                geo_pct = float(CFG.get("geo_pct_default", 60))
                geo_basis = "default (no ranking data)"
            else:
                geo_pct = float(CFG.get("geo_pct_default", 60))
                geo_basis = "default"
                for thresh, val in CFG.get("geo_pct_tiers", []):
                    if pct_not_ranking >= thresh:
                        geo_pct = float(val)
                        geo_basis = f"{pct_not_ranking:.0f}% of head terms not ranking"
                        break
            p_list = geo_pct / 100.0
            ai = {"mode": "pct",
                  "uplift_pct": geo_pct,
                  "geo_pct": geo_pct,
                  "geo_pct_basis": geo_basis,
                  "min_term_months": min_term,
                  "hard_add":   {k: r50(v * p_list)   for k, v in hard.items()},
                  "client_add": {k: r50up(v * p_list) for k, v in client.items()}}
        # GEO can be overridden independently of SEO. The percentage model is
        # a good default and a bad straitjacket: a client may have agreed a GEO
        # number that has nothing to do with their SEO price — a flat retainer,
        # a carried-over rate — and the alternative is overriding SEO to a
        # fiction just to move GEO. The override sets the BASE and the ladder
        # keeps its shape, scaling the upper tiers by the same ratio the SEO
        # ladder uses, so the three tiers stay proportionate to each other.
        try:
            geo_base = float(geo_override) if geo_override not in (None, "") else None
        except (TypeError, ValueError):
            geo_base = None
        if geo_base and geo_base > 0:
            ratio = {k: (hard[k] / hard["base"] if hard["base"] else 1.0) for k in hard}
            ai["hard_add"] = {k: r50(geo_base * ratio[k]) for k in hard}
            ai["client_add"] = {k: r50up(ai["hard_add"][k] * m) for k in hard}
            ai["client_list"] = dict(ai["client_add"])
            ai["manual_geo"] = True
            ai["geo_pct_basis"] = "manual override"
            ai["geo_pct"] = (round(ai["client_add"]["base"] / client["base"] * 100, 1)
                             if client.get("base") else None)
        ai["hard_total"]   = {k: hard[k] + ai["hard_add"][k] for k in hard}
        ai["client_total"] = {k: client[k] + ai["client_add"][k] for k in client}
        # THE PERCENTAGE THAT IS TRUE OF THE DOLLARS BESIDE IT.
        # geo_pct is the RULE — the tier rate the quote was built from. It is
        # not what the printed figures divide out to, because each tier rounds
        # UP independently: 59% of a $3,950 core is $2,330.50 and is quoted at
        # $2,350, which is 59.5%. The panel printed the rule next to the
        # dollars and the two did not reconcile. It is worse on a manual GEO
        # override, where the dollars have nothing to do with the rule at all.
        # So carry both: geo_pct for what the rule says, geo_pct_effective for
        # what was actually charged. (2026-08-27)
        ai["geo_pct_effective"] = {
            k: (round(ai["client_add"][k] / client[k] * 100, 1) if client.get(k) else None)
            for k in client}

    # AN ADD-ON MARKET IS A FULL CAMPAIGN IN ANOTHER CITY. There is no scope
    # fraction any more (see the RETIRED note on addon_market_ratio): the list
    # partner cost of an add-on market IS the tier's partner cost, and the
    # Add-On Market % is the only thing that comes off it.
    # ADD-ON MARKETS ARE PRICED THROUGH THE MARGIN, NOT AS A RATIO OF RETAIL
    # (2026-08-05). The client side used to be ratio x retail, computed
    # independently of the add-on's own cost. That made the per-market retail
    # figure carry its own rounding error, which then got MULTIPLIED by the
    # market count — so realised margin sagged as add-ons accumulated (34.09%
    # at five markets against a 35% target).
    #
    # Now: partner cost per market is a clean $50, and its retail price is that
    # cost / (1 - margin), rounded UP to $50. Because every component rounds up,
    # the total is automatically a $50 multiple and the margin can never land
    # below target - verified across all 1,650 anchor x ratio x count cases.
    _n_addon = max(0, int(addon_markets or 0))
    # LIST — one add-on market before the volume break.
    hard_addon_list   = {k: r50n(hard_cost[k]) for k in hard_cost}
    # THE LIST PRICE OF AN ADD-ON MARKET IS THE TIER PRICE, exactly as shown on
    # the cards — that is what "a full campaign in another city" means. Taken
    # from `client` rather than re-derived through retail_of(hard_cost): on a
    # FLOORED quote the cost was restated from the floored retail, and the trip
    # back out loses $50 to rounding, so a $2,950 campaign would have shown a
    # $2,900 list price for the same thing.
    client_addon_list = dict(client)
    # ADD-ON MARKET % — flat by bracket, applied to partner cost so the client
    # price falls out of the margin the way every other price here does.
    #
    # The DISCOUNTED cost is rounded to $10, not $50. At $50 the brackets
    # collapsed into each other: an $800 list cost is $720 at 10% off and $680
    # at 15%, and both snap to $700 — so crossing from nine markets to ten
    # displayed "15%" and changed the price by nothing. The $50 grid belongs to
    # the numbers a client reads; this one is an internal cost, and the client
    # figure derived from it still lands on $50 through retail_of.
    _ad_pct = addon_discount_pct(_n_addon)
    _ad_basis = ("%d add-on markets" % _n_addon) if _ad_pct else (
        "no add-on markets" if not _n_addon else "below the first bracket")
    if _ad_pct:
        hard_addon   = {k: int(round(round(hard_addon_list[k]
                                          * (1.0 - _ad_pct / 100.0), 6) / 10.0) * 10)
                        for k in hard_addon_list}
        client_addon = {k: retail_of(hard_addon[k]) for k in hard_addon}
    else:
        # NO DISCOUNT MEANS NO DISCOUNT. Sending the list cost back out through
        # retail_of loses $50 to rounding on a floored quote, so a 0% bracket
        # charged $2,900 against a $2,950 list — a phantom saving the operator
        # could not explain and nobody had asked for.
        hard_addon, client_addon = dict(hard_cost), dict(client_addon_list)
    # EVERY BRACKET, PRICED BY THE SERVER. The stepper moves the count without a
    # round-trip, and reimplementing r50/retail_of in the browser would put two
    # rounding rules in play — JS rounds .5 up, Python rounds it to even, so a
    # $25 remainder would show one price on screen and quote another. The panel
    # looks its bracket up in this table instead of computing anything.
    def _bracket(pct):
        h = {k: int(round(round(hard_addon_list[k] * (1.0 - pct / 100.0), 6)
                          / 10.0) * 10) for k in hard_addon_list}
        return {"min_markets": None, "pct": pct, "hard": h,
                "client": {k: retail_of(h[k]) for k in h}}

    addon_schedule = []
    for _lo, _pct in (CFG.get("addon_volume_discount_tiers") or []):
        _b = _bracket(float(_pct))
        _b["min_markets"] = int(_lo)
        addon_schedule.append(_b)
    addon_schedule.sort(key=lambda b: b["min_markets"])
    # An add-on market can be negotiated independently of the ratio. A client
    # taking eleven of them will argue the per-market rate long before the
    # primary campaign, and the alternative is distorting the whole ladder to
    # move one number. Sets the BASE; the upper tiers keep the SEO ladder's
    # shape so the three stay proportionate.
    try:
        _ao = float(addon_override) if addon_override not in (None, "") else None
    except (TypeError, ValueError):
        _ao = None
    manual_addon = bool(_ao and _ao > 0)
    if manual_addon:
        # Override sets the BASE partner cost per market; upper tiers keep the
        # ladder's shape. Retail still comes from cost through the margin.
        _ar2 = {k: (hard_cost[k] / hard_cost["base"] if hard_cost["base"] else 1.0)
                for k in hard_cost}
        hard_addon   = {k: r50n(_ao * _ar2[k]) for k in hard_cost}
        client_addon = {k: retail_of(hard_addon[k]) for k in hard_addon}
        # A NEGOTIATED RATE IS THE RATE. Taking another 20% off a number the
        # operator typed in would move a price they had already settled, and
        # they would have no way to see it happen. The bracket is reported as
        # not applied rather than silently skipped.
        hard_addon_list, client_addon_list = dict(hard_addon), dict(client_addon)
        _ad_pct, _ad_basis = 0.0, "not applied — manual per-market rate"
        addon_schedule = []
    # True partner cost is a share of RETAIL, so derive it from the client
    # tiers rather than from the calibration basis.
    hard_true = dict(hard_cost)          # already clean $50 figures
    # The COMBINED MONTHLY BUDGET — the single figure the adtini product form
    # needs. Package retail plus the per-market retail times the market count;
    # every term is already a $50 multiple, so the sum is too.
    combined = {k: client[k] + client_addon[k] * _n_addon for k in client}
    combined_hard = {k: hard_cost[k] + hard_addon[k] * _n_addon for k in hard_cost}
    # ---- WHAT THE IO PULLS FROM THIS QUOTE -------------------------------
    # One block, named the way the SEO+ ticket and the Billing PRD name things,
    # so nobody has to work out which of our internal keys maps to which of
    # their terms. Two of these exist ONLY because deriving them on the other
    # side goes wrong:
    #   * Partner Hard Cost IS here (2026-08-27), because Billing charges the
    #     partner off it and must not invent it. It is NOT the round-trip
    #     source: the IO still derives client price from Package $ and Margin %
    #     (ticket RULE 2.d.ii), because a $50-rounded cost run back through the
    #     margin comes out $50 wrong on 86 of 182 tier values. Sent as the
    #     figure to BILL, not the figure to price from. `margin_dollars` closes the
    #     gap between the two so nothing has to be reverse-engineered.
    #   * Partner Add-On Market Cost IS here, because the reverse applies:
    #     Add-On Market Price / (1 - Margin %) misses by $7.50-$17.50 a market,
    #     the partner figure having rounded to $10 and the sell figure to $50
    #     (ticket RULE 2.d.v).
    _ai_add = (ai or {}).get("client_add") or {k: 0 for k in client}
    _ai_hard = (ai or {}).get("hard_add") or {k: 0 for k in client}
    handoff = {
        "package": {k: client[k] + _ai_add.get(k, 0) for k in client},
        "core_seo_price": dict(client),
        "ai_search_price": {k: _ai_add.get(k, 0) for k in client},
        "ai_search_pct": (ai or {}).get("geo_pct") or 0,
        "ai_search_pct_effective": (ai or {}).get("geo_pct_effective") or {},
        "margin_pct": markup_pct,
        "addon_market_price": dict(client_addon),
        "partner_addon_market_cost": dict(hard_addon),
        "addon_market_discount_pct": _ad_pct,
        "addon_markets": _n_addon,
        # ---- THE PROPOSAL CHANGES THE MARKET COUNT, SO IT NEEDS EVERY RATE --
        # Add-On Market Discount % above is ONE number: the bracket THIS quote
        # landed in. That is all the IO needs, because the IO shows the quote.
        # The proposal deck is different — it lets the client move the market
        # count with an up/down, and moving it changes the bracket. So the deck
        # needs the rates for brackets the quote is not in.
        #
        # And it needs them as DOLLARS, not as a percentage to apply. Taking the
        # undiscounted client figure and knocking 10/15/20% off it lands $50 out
        # on 512 of 1260 tier values: our chain discounts the PARTNER cost,
        # rounds that to $10, and only then rounds the client figure UP to $50 —
        # two roundings in a different order from the deck's one. Every bracket
        # is priced here by the code that priced the quote, so the deck looks
        # its bracket up instead of computing one.
        #
        # THE UNDISCOUNTED LIST PRICE IS NOT SENT (2026-08-27, Kiri). It was,
        # while the deck was going to show a subtotal and then discount it. The
        # deck reads its bracket rate straight off this table and multiplies by
        # the count, so it never sees a pre-discount figure.
        "addon_market_price_by_bracket": [
            {"min_markets": b["min_markets"], "discount_pct": b["pct"],
             "client": dict(b["client"]), "partner": dict(b["hard"])}
            for b in addon_schedule],
        # MINIMUM TERM IS NOT HERE (2026-08-27, Kiri). The line item already
        # refuses a submission under six months, so sending it would be us
        # telling adtini a rule adtini enforces. It stays a pricing constant
        # (CFG.min_term_months), is shown on the quote and on the panel, and is
        # returned by /api/price at the top level — it is just not part of the
        # handoff contract.
        # ---- PARTNER COST — WHAT BILLING CHARGES THE PARTNER ---------------
        # Core + AI, per tier, the same $50 figures the Partner card shows.
        "partner_hard_cost": {k: hard_cost[k] + _ai_hard.get(k, 0) for k in hard_cost},
        "partner_core_seo_cost": dict(hard_cost),
        "partner_ai_search_cost": {k: _ai_hard.get(k, 0) for k in hard_cost},
        # PACKAGE - PARTNER HARD COST, STATED RATHER THAN DERIVED.
        # Package x (1 - Margin %) does NOT return Partner Hard Cost and is not
        # meant to: the client figure rounds UP to $50 and each partner figure
        # rounds to the NEAREST $50, so the two drift by up to ~$75 a tier. That
        # drift is margin, and it is Vici's. Sending the dollars closes the
        # books without anyone reverse-engineering the rounding:
        #     Package $  -  Partner Hard Cost  =  Margin $
        "margin_dollars": {k: (client[k] + _ai_add.get(k, 0))
                              - (hard_cost[k] + _ai_hard.get(k, 0))
                           for k in client},
    }
    # ---- THE PRICE THE FORMULA WOULD HAVE GIVEN --------------------------
    # An override REPLACES a component, so the quote it produces is partly the
    # operator's. Calibration pairs "what the formula said" against "what was
    # actually sent" — and it was reading the overridden figure as the formula
    # side, which meant it compared a human's number against itself and learnt
    # nothing from any overridden quote. Vibe check used to keep the formula
    # figure underneath and the overrides did not; with Vibe check gone
    # (2026-08-27) that discipline has to live here. Pure arithmetic, no API
    # calls, and only computed when something was actually overridden.
    _formula = None
    if not _formula_pass and any(x not in (None, "") for x in
                                 (base_override, geo_override, addon_override)):
        try:
            _fp = stage4_price(band, adder, zero_ranking, addon_markets, markup_pct,
                               pct_not_ranking=pct_not_ranking,
                               total_volume=total_volume, base_override=None,
                               ecommerce=ecommerce, industry=industry,
                               ai_search=ai_search, national_demand=national_demand,
                               geo_override=None, addon_override=None, goal=goal,
                               pageone_rank=pageone_rank, site_rebuild=site_rebuild,
                               _formula_pass=True)
            _formula = {"client_tiers": _fp["client_tiers"],
                        "ai_search": ({"client_add": _fp["ai_search"]["client_add"],
                                       "client_total": _fp["ai_search"]["client_total"]}
                                      if _fp.get("ai_search") else None),
                        "client_addon_per_market": _fp.get("client_addon_per_market"),
                        "combined_monthly": _fp.get("combined_monthly")}
        except Exception:                                   # noqa: BLE001
            app.logger.exception("formula-price pass failed")
            _formula = None
    return {"anchor": anchor, "base": base, "base_pre_uplift": base_pre, "step": step,
            "handoff": handoff,
            "formula": _formula,
            "hard_true_tiers": hard_true,
            "margin_pct_of_gross": round(mg * 100, 2),
            "agency_profit_tiers": {k: client[k] - hard_true[k] for k in client},
            "margin_realised_pct": {k: (round((client[k] - hard_true[k]) / client[k] * 100, 1)
                                        if client[k] else 0.0) for k in client},
            "combined_monthly": combined,
            "combined_hard": combined_hard,
            "combined_margin_pct": {k: (round((combined[k] - combined_hard[k]) / combined[k] * 100, 2)
                                        if combined[k] else 0.0) for k in combined},
            "national_demand": nat_demand, "national_demand_reason": nat_reason,
            "volume_captured": vol_captured,
            "volume_opportunity": round(vol_opportunity, 3),
            "min_term_months": min_term,
            # What the rebuild flag did, so the panel can say it rather than the
            # operator wondering why a client with rankings priced as if it had
            # none. rebuild_applied is False when the measurement already
            # agreed -- the flag is then a label, not a change.
            "site_rebuild": rebuild,
            "rebuild_applied": rebuild_applied,
            "extras_multiplier": _mult,
            "manual_geo": bool(ai and ai.get("manual_geo")),
            "manual_addon": manual_addon,
            "industry_rule": rule_key,
            # Every component the price is built from, so the proposal view can
            # show one chart instead of six panels. The adder was the only one
            # that lived solely as an input and never came back out.
            "competitive_adder": int(adder or 0),
            "industry_anchor_add": int(rule.get("anchor_add", 0)) if rule else 0,
            "pageone_anchor_add": pageone_add,
            "pageone_band": pageone_band,
            "pageone_rank": pageone_rank,
            # NOT MEASURED IS NOT ZERO. A quote priced before the rank check has
            # found anyone carries no page-one reading at all, and silently
            # omitting the add would make the same client price two different
            # ways depending on when the button was pressed. Said out loud so
            # step 4 can warn instead of quietly under-quoting.
            "pageone_measured": pageone_band is not None,
            "ai_search": ai,
            "floored": floored, "client_floor": floor, "manual_base": manual_base,
            "zero_ranking_uplift_pct": zr_uplift, "volume_add": vol_add,
            "pct_not_ranking": pct_not_ranking, "total_volume": total_volume,
            "hard_tiers": hard, "client_tiers": client,
            "hard_addon_per_market": hard_addon, "client_addon_per_market": client_addon,
            # The Add-On Market % and what it was applied to, so the panel can
            # print the rate rather than leave the operator to divide two
            # numbers and guess which bracket they landed in.
            "addon_discount_pct": _ad_pct,
            "addon_discount_basis": _ad_basis,
            "addon_discount_tiers": CFG.get("addon_volume_discount_tiers") or [],
            "addon_schedule": addon_schedule,
            "hard_addon_list_per_market": hard_addon_list,
            "client_addon_list_per_market": client_addon_list,
            "addon_savings_per_market": {k: client_addon_list[k] - client_addon[k]
                                         for k in client_addon},
            "markup_pct": markup_pct, "addon_markets": addon_markets,
            "tiers": client, "addon_per_market": client_addon}

# ---------------------------------------------------------------------------
# THE PROPOSAL DOCUMENT
#
# Brendan's proposals are one shape, every time: Background, an SEO section, the
# keyword table, the three keyword-set definitions, the monthly service list,
# three priced options, Case Studies. Rebuilding that by hand for every quote is
# the work this tool exists to remove, and the quote already holds every
# client-specific number in it.
#
# The standing sections live here rather than in the generator so they can be
# edited without touching code — when a case study changes it changes in one
# place. Anything that varies per client is computed, never templated.
# (2026-08-18)
# ---------------------------------------------------------------------------

PROPOSAL = {
    # -----------------------------------------------------------------------
    # BRENDAN'S WORDS, NOT OURS.
    #
    # Everything in this block that also exists in his proposals is now
    # character-for-character his, taken from the NASSCO and Media Venue
    # documents in becal/ — including the double spaces after full stops and the
    # curly quotes, because "the exact same verbiage" is the instruction and a
    # tidied-up sentence is a different sentence.
    #
    # An audit before this change found 38 paragraphs verbatim, 12 lightly
    # edited and 38 written by me. The 12 and the 38 are gone. What remains ours
    # is only what he has no equivalent for: What We Measured For You, the
    # authority-gap reading, the live results capture, and the AI Search
    # branding on his GEO section. proposal_test.py checks the boilerplate back
    # against the .docx files so it cannot drift again. (2026-08-19)
    # -----------------------------------------------------------------------
    "intro_heading": "Background Information",
    "seo_heading": "Search Engine Optimization (SEO)",
    # OURS — he has no equivalent section, so it needed a name. The first two
    # were "What We Measured For You" and, briefly, "Your Outlook": both
    # second-person, and the second promised a forecast the section does not
    # make. It is diagnosis — who holds page one, how far behind the client is,
    # what is wrong with their site — so it takes the flat noun phrase every
    # other heading in the document uses. (2026-08-19)
    "measured_heading": "Competitive Analysis",
    "campaign_heading": "Ongoing SEO Campaign:",
    "options_heading": "SEO Campaign Options:",
    "additional_heading": "Additional Keywords:",
    "case_heading": "Case Studies & Results",

    # ---- his prose, in the order it appears -------------------------------
    "intro_line": "{brand} is requesting a proposal for assistance with "
                  "improving their organic rankings, traffic, and lead flow "
                  "from search engines such as Google through strategic search "
                  "engine optimization (SEO) services.",
    "intro_close": "Please find our recommendations and proposal below.",
    "seo_intro": "The best long-term marketing strategy for driving online "
                 "leads is search engine optimization (SEO) which is the "
                 "process of improving a website’s visibility and rankings "
                 "for targeted keywords on search engines such as Google.",
    "seo_table_lead": "Based on our initial research, we came up with a "
                      "preliminary list of potential keywords and noted your "
                      "current rankings for each keyword on Google below:",
    "seo_after_table": "Based on the above data, there is significant room for "
                       "improvement in organic ranking through an organic SEO "
                       "campaign.  The goal of this campaign will be to capture "
                       "this search traffic to drive highly qualified visitors "
                       "to your website and increase leads for your business.",
    "keyword_sets_lead": "All of our SEO campaigns focus on three primary "
                         "keyword sets:",
    "keyword_sets_close": "We start every campaign with an analysis of keywords "
                          "and identify which terms map to which landing pages; "
                          "the above set of keywords are preliminary examples "
                          "of terms we would target in an organic SEO campaign.",
    "campaign_lead": "All of our ongoing SEO campaigns include the following "
                     "services each month:",
    "campaign_close": "The goal of this campaign is to rank your website as "
                      "high as possible on page 1 of search results to drive "
                      "high quality, qualified traffic and leads to your "
                      "website.",
    "options_lead": "Depending on how aggressive you wish to be with an SEO "
                    "campaign, we have included three options below following a "
                    "“good, better, best” model.  Additional campaign "
                    "options are available beyond those outlined above. Most "
                    "clients begin with one of these campaigns and later expand "
                    "their monthly investment as rankings, traffic, and lead "
                    "volume grow.",

    "keyword_sets": [
        ("Ultra-Competitive Keywords",
         "These are the most competitive terms in a particular industry and are "
         "extremely difficult to rank but yield extremely high traffic."),
        ("Competitive Keywords",
         "These are highly competitive terms which take a high amount of effort "
         "and time to rank for but can drive a tremendous amount of traffic."),
        ("Long-tail Keywords",
         "These are lower competition terms which take less time and effort to "
         "rank for but also drive less traffic however this traffic is often "
         "highly specific with high intent."),
    ],

    # (lead bullet, sub-bullets) — his shape exactly. The bold group headings
    # this used to carry were mine, and a heading changes how a sentence reads
    # even when the sentence is untouched, so they are gone too.
    "services": [
        ("Conduct full initial keyword research to identify relevant search "
         "terms in the industry.  We generally begin seeing meaningful ranking "
         "improvements within 5-10 months, however will provide reports to show "
         "progress as terms improve. We will identify three types of keywords:",
         ["Ultra-Competitive", "Competitive", "Long Tail"]),
        ("We will conduct all on-site optimization work needed.  This includes:",
         ["Full technical website audit",
          "Title tag optimization",
          "Meta information optimization",
          "Image optimization where needed",
          "Testing the site for all current SEO factors including mobile "
          "usability, speed, and other factors"]),
        ("We will engage in an off-site link building campaign to help "
         "strengthen the website domain.  Links remain one of the most "
         "important factors which determine website rankings in organic search. "
         "These links will include:",
         ["Guest posts/articles",
          "Directory type listings",
          "Citation building",
          "Content distribution/syndication",
          "Content sharing sites",
          "Micro-blogs",
          "And additional link types as relevant to the campaign"]),
        ("We will research 2-3 long tail keyword topics and write 2-3 highly "
         "SEO-focused articles/content assets, 750-1,000 words in length, and "
         "post them to the website blog each month", []),
        ("We will verify listing sites are setup properly, including Google "
         "Business Profile (formerly Google My Business) and optimize this "
         "listing as well as any supplemental listings for satellite locations",
         ["Includes on-page optimization of the Google Business Profile",
          "Also includes monthly citation building"]),
        ("We will deliver a monthly report at the start of each month outlining "
         "the work performed that month and analytical metrics for our keywords "
         "and the website’s traffic to show the progress of the campaign",
         []),
    ],

    # The per-option keyword counts Brendan quotes. Boilerplate in his documents
    # — the same three lines regardless of the grid — so they stay boilerplate.
    #
    # Three variants exist across the nine proposals and no rule separates them:
    # not list size, not vertical, not date. This is the one in six of the nine
    # (Waytek, PA Center, Nob Hill, Keller, Visit Central PA, Red Shoes), chosen
    # 2026-08-22 over Media Venue's "1-2 / 2-3 / 3-4" and the Rockingham/NASSCO
    # variant that widens competitive to 4-6 and 6-8.
    #
    # NOT derived from the grid, and it cannot be: Rockingham's table carries 40
    # ultra-competitive terms against an Advanced option promising 3-4, and Red
    # Shoes has 3 competitive terms against an Advanced option promising 8-12.
    # These lines say how many terms the campaign WORKS, not how many the table
    # lists — deriving them would print numbers Brendan never writes and would
    # make Red Shoes' Advanced smaller than its Base.
    "option_scope": {
        "base": ["1 ultra-competitive keyword", "3-5 competitive keywords",
                 "10-15 long tail keywords"],
        "intermediate": ["2 ultra-competitive keywords", "5-8 competitive keywords",
                         "12-18 long tail keywords"],
        "advanced": ["3 ultra-competitive keywords", "8-12 competitive keywords",
                     "15-20 long tail keywords"],
    },
    # Each is a list of paragraphs; the last one runs into the targeting list,
    # which is why it ends on a colon rather than a full stop.
    "option_blurb": {
        "base": ["In our base SEO campaign, we are doing SEO work to the site "
                 "to slowly improve rankings and drive more traffic to the "
                 "site.",
                 "In a base campaign, we include all SEO activities outlined "
                 "above.  We would expect to see ranking improvements for the "
                 "terms starting after 6-10 months of optimization work and "
                 "recommend targeting:"],
        "intermediate": ["An intermediate campaign includes everything "
                         "mentioned in a basic campaign, however we also "
                         "include additional types of link building which helps "
                         "build your rankings faster and more consistently as "
                         "well as our proprietary rank signaling SEO strategy "
                         "which has seen a 97% success rate in ranking keywords "
                         "on page 1 of Google.  We expect to see ranking "
                         "improvements for the terms after 5-8 months and "
                         "recommend targeting:"],
        "advanced": ["An advanced campaign includes everything mentioned in an "
                     "intermediate campaign, however we accelerate the pace at "
                     "which we optimize the site, build links, and perform "
                     "other SEO activities allowing us to accelerate ranking "
                     "growth and lead generation.  We expect to see ranking "
                     "improvements for the terms after 4-6 months and recommend "
                     "targeting:"],
    },
    # ---- what the authority gap actually costs ----------------------------
    # OURS — he has no equivalent. "A gap of 605" means nothing to a client.
    # These bands turn it into link volume and months, which is what a retainer
    # actually buys.
    #
    # DataForSEO's rank is 0-1000 and logarithmic; the published benchmarks below
    # are on the 0-100 scale, so the score is divided by ten to compare. That is
    # an approximation and is labelled as one wherever it is shown — DataForSEO's
    # own docs say their methodology differs from Moz's and Ahrefs', so this is
    # the right ORDER of effort, not a promise.
    #
    # Source: Linkscope's 2026 domain-authority ranking benchmarks —
    # 10-15 quality linking root domains a month moves 5-10 points, and each
    # 10-point step needs 2-3x the referring domains of the one before it.
    "gap_bands": [
        (20, "6-12 months", "10-20 referring domains"),
        (40, "12-18 months", "40-80 referring domains"),
        (60, "18-36 months", "100-300+ referring domains"),
        (1000, "24-48+ months", "300-1,000+ referring domains"),
    ],
    "gap_heading": "What the authority gap means",

    # HOW BIG, AND CAN WE CLOSE IT. The bands above say what the work costs;
    # these say what the number means, which is the question a client actually
    # asks. Points are on the 0-100 scale (the raw 0-1000 score divided by ten),
    # so a raw gap of 605 lands in the last band. MIRRORED in the template as
    # GAP_VERDICTS — proposal_test.py asserts the two agree.
    # AND EACH ONE NAMES THE OPTION THAT CLOSES IT. A verdict that stops at
    # "this is a wide gap" leaves the client to work out what to do about it,
    # three pages before the prices. Saying which campaign closes it faster is
    # the same sentence doing the selling it was written for. (2026-08-19)
    # EVERY VERDICT POINTS UP. Naming the base campaign as sufficient argues the
    # client down a tier inside the document that is trying to sell them one —
    # and it is not even the honest read: the difference between the tiers is
    # the PACE of link and content work, which is exactly what closes a gap.
    # These say what each gap is and which campaign closes it soonest.
    # (2026-08-19)
    "gap_verdicts": [
        (15, "a small gap",
         "That is a small gap, and the closest thing to a quick win in this "
         "proposal — an intermediate campaign would close it and move on to the "
         "harder terms inside the first year."),
        (30, "a moderate gap",
         "That is a moderate gap, and the pace of the campaign is what decides "
         "how quickly it closes — an intermediate or advanced campaign gets "
         "there materially sooner."),
        (50, "a wide gap",
         "That is a wide gap. An advanced campaign would close it considerably "
         "faster than the alternatives, because the rate of link acquisition "
         "and content production is the whole difference between them."),
        (1000, "a very large gap",
         "That is a very large gap — the sites ahead are far more established "
         "domains. An advanced campaign is what makes it a realistic target; at "
         "a slower pace the most competitive terms stay out of reach for years."),
    ],
    "gap_level_verdict":
        "{brand} already scores at or above the sites currently holding page "
        "one, which is the strongest position a campaign can start from: the "
        "authority is there and the rankings are not, so the gains come from "
        "content, on-page work and targeting the right terms. That is work that "
        "compounds with the pace it is done at, and an intermediate or advanced "
        "campaign converts the position faster.",

    # ---- AI Search (GEO) --------------------------------------------------
    # Brendan runs this as a full parallel campaign with its own three options,
    # priced IN ADDITION to SEO — see the Media Venue proposal, the only one of
    # the ten that carries it. His words throughout; the SECTION NAME is the one
    # change, because the team sells this as AI Search rather than GEO.
    "geo_heading": "AI Search (GEO)",
    "geo_intro": [
        "Generative Engine Optimization (GEO) is the process of optimizing for "
        "AI based search queries.  There are a variety of AI based search "
        "engines currently in use, with the most popular being:",
    ],
    "geo_engines": [
        "Google’s Gemini “AI Mode” & AI Overviews",
        "ChatGPT",
        "Microsoft Copilot (Bing AI Search)",
        "There are dozens of other small players in the market as well",
    ],
    "geo_context": [
        "While AI search is certainly going to be the search of the future, "
        "today AI searches only account for approximately 10% of all search "
        "volume.  Traditional search, which is optimized through search engine "
        "optimization or SEO campaigns, still dominates the marketplace with "
        "the majority of search volume.  Furthermore, studies have shown "
        "purchasing decisions are more heavily influenced still by traditional "
        "search.",
        "With that said, many brands are still wanting to get ahead of the AI "
        "trend as well as leverage AI for additional brand exposure and are "
        "beginning to engage in GEO campaigns.  Unlike SEO campaigns where a "
        "fixed keyword set is defined and tracked at a keyword level, AI search "
        "models can yield different results each time even for the same "
        "question.  This makes them much more complex for optimization efforts "
        "and much more difficult to track results/progress short of simply "
        "asking the AI model questions and analyzing the outcome/result.",
        "We have been a pioneer in the GEO and AI search industry being one of "
        "the first companies to offer optimization searches for AI models.  We "
        "launched this service in 2023 and have assisted numerous clients with "
        "ranking higher in AI search models through strategic optimization of "
        "the AI search models, strategic link and citation building, and robust "
        "analysis of what information is being utilized in AI models to render "
        "results to specific search queries.",
    ],
    "geo_campaign_heading": "On-Going AI Search Campaign:",
    "geo_campaign_lead": "All of our on-going campaigns include the following "
                         "services each month:",
    "geo_services": [
        ("Conduct full initial keyword research to identify relevant search "
         "phrases on AI models.",
         ["We generally identify 20-30 target search phrases/topics"]),
        ("We will conduct full AI model optimization to train the core AI "
         "models providing them information supporting your business being the "
         "recommended result for the given AI search queries", []),
        ("We will engage in an off-site link building, content creation and "
         "citation building campaign to help provide resources which the AI "
         "model will read and analyze when making a decision on a given search "
         "query.  This will provide outside influence to the AI model to "
         "increase the confidence in recommending your brand in given search "
         "queries.  These links will include:",
         ["Guest posts/articles",
          "Directory type listings",
          "Citation building",
          "Content distribution/creation",
          "Content sharing sites",
          "Micro-blogs",
          "And additional link types as relevant to the campaign"]),
        ("We will deliver a monthly report at the start of each month outlining "
         "the work performed that month and any available analytical metrics "
         "for our keywords and AI searches to document progress in improving "
         "your AI results.", []),
    ],
    "geo_options_heading": "AI Search Campaign Options:",
    "geo_options_lead": "Depending on how aggressive you wish to be with an AI "
                        "Search campaign, we have included three options below "
                        "following a “good, better, best” model.  "
                        "These campaign options would be in addition to one of "
                        "the proposed SEO campaign options and are discounted "
                        "accordingly in conjunction with an SEO campaign.",
    "geo_option_blurb": {
        "base": ["In our base AI Search campaign, we are engaging in GEO work "
                 "to slowly improve your brand appearing in recommendations in "
                 "AI models.  In a base campaign, we include everything "
                 "outlined above.  We would expect to see improvements starting "
                 "after 6-10 months of optimization work."],
        "intermediate": ["An intermediate campaign includes everything "
                         "mentioned in a basic campaign, however we also target "
                         "1-2 premium placements per month which are "
                         "specifically analyzed by GEO models in rendering "
                         "results and recommendations to more quickly impact "
                         "the AI search results.  We expect to see ranking "
                         "improvements after 5-8 months."],
        "advanced": ["An advanced campaign includes everything mentioned in an "
                     "intermediate campaign, however we accelerate the pace at "
                     "which we create premium placements to 3-4 per month.  We "
                     "expect to see ranking improvements after 6 months."],
    },
    "closing": [
        "Additional keywords can be added to any campaign for an additional fee "
        "upon request.",
        "We greatly appreciate the opportunity and your trust in evaluating us "
        "as your digital partner.  Everything we do is centered around "
        "providing best in class customer service and getting you results.  As "
        "such, we have included several reference projects and case studies "
        "below as a showcase of our work.",
    ],
    # ---- performance-based, Brendan verbatim (Ooten Law, 2026-08-20) -------
    "perf_heading": "Performance Based SEO",
    "perf_intro": [
        "A performance-based SEO campaign is very similar to the above SEO "
        "campaign options and is only available to select clients.  In a "
        "performance-based campaign, we will engage in the same work, however "
        "there are no fees charged until you rank for a term. Once you rank for "
        "a term, depending on the position on page 1, there is a monthly fee "
        "associated with that term ranking which is paid per month as long as "
        "the term ranks.  Should you want to cancel work for a specific term, "
        "you would continue to pay for that term ranking for a period of "
        "{tail} months post-cancellation as long as the term continues to rank.  "
        "In the event within that {tail} month period the term stops ranking, you "
        "would then stop paying when the term stops ranking.  The initial term "
        "for the campaign is {term} months which then auto-converts to a month "
        "to month thereafter with the {tail} month tail for any ranking terms.",
        "Below is a sample set of potential keywords including the term, your "
        "current Google ranking for a specific term as well as the cost "
        "associated with each keyword once it ranks on page 1 of results on "
        "Google as determined by our proprietary rank tracker data.  You do not "
        "pay for a keyword until it ranks.  We do require a minimum of {min} in "
        "potential term ranking value to initiate a performance-based SEO "
        "campaign.",
    ],
    "perf_columns": ["Keyword", "Practice Area", "Competition Tier",
                     "Google Current Organic Ranking", "Cost Page 1",
                     "Cost Top 5", "Cost Top 3", "Cost #1",
                     "Current Achieved Tier"],
    "case_projects_label": "SEO Projects:",
    "case_studies": [
        "All Year Cooling – 5+ year SEO client ranking #1 for “Miami "
        "AC Repair” and over 100 other terms on page 1 resulting in a "
        "1,000+% increase in organic traffic and leads.",
        "Engage – 7+ year SEO client ranking #1 for “Motivational "
        "Speakers”, one of the most competitive keywords on the internet, "
        "as well as dozens of other terms driving hundreds of leads per month.",
        "ERI – 10+ year SEO client ranking #1 for “Electronic "
        "Recyclers”, “Data Destruction”, and hundreds of other "
        "terms driving hundreds of leads per year.",
        "Goldstone Financial – 4+ year SEO client ranking #1 for "
        "“Financial Advisor Chicago”, “Financial Planner "
        "Chicago”, “Retirement Advisor” and dozens of other "
        "highly competitive terms.",
        "Lakeside Equipment – 10+ year SEO client ranking #1 for over 50 "
        "keywords in their industry resulting in an over 5,000% increase in "
        "organic traffic.",
    ],
    "case_closing": "Over the past two decades, we have managed hundreds of "
                    "digital marketing, web and SEO campaigns with a focus on "
                    "delivering results and forging long term client "
                    "relationships.  We have a 97+% client retention rate "
                    "because we believe in going the extra mile for each and "
                    "every client and greatly appreciate the opportunity to "
                    "show your business how we can help you grow through "
                    "digital.  When you’re ready to get started, simply "
                    "email your account representative and we will send over a "
                    "short digital contract to initiate our relationship.  "
                    "We’re looking forward to working with you and your "
                    "team!",
}


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html", build=BUILD_STR)

DEMO_MODE = os.environ.get("DEMO_MODE", "").lower() in ("1", "true", "yes")

def mock_pipeline(seeds, markets, state, domain, brand, band, addon):
    """Realistic sample data — no DataForSEO calls. Deterministic per input so
    the demo feels responsive to what the partner typed. Cannot time out."""
    market = markets[0] if markets else ""

    # Head terms: seed + market variants, descending volume
    head_terms = []
    for s in seeds:
        if market:
            head_terms.append(f"{s} {market}".strip())
        head_terms.append(s)
    seen = set(); head_terms = [h for h in head_terms if not (h in seen or seen.add(h))]
    ultra, comp = [], []
    for i, h in enumerate(head_terms):
        vol = max(40, 620 - i * 55)
        (ultra if i < 3 else comp).append({"kw": h, "vol": vol})
    comp = comp[:6]

    # Long-tail: question-shaped, longer phrases
    templates = ["how much does {s} cost in {m}", "best {s} near me",
                 "what to look for in {s} in {m}", "affordable {s} for adults in {m}",
                 "is {s} covered by insurance in {m}"]
    longtail = []
    for s in seeds:
        for t in templates:
            kw = t.format(s=s, m=market or "your area").replace("  ", " ").strip()
            longtail.append({"kw": kw, "vol": 0})
    longtail = longtail[:10]

    # Ranking table: mostly Not Found (zero-ranking demo), one ranked deep
    all_rows = ultra + comp + longtail
    table = []
    for i, r in enumerate(all_rows):
        pos = 54 if i == len(all_rows) - 1 else "Not Found"
        table.append({"kw": r["kw"], "pos": pos})
    ranked, total = 0, len(all_rows)   # 0 in top 50 -> zero-ranking fires
    zero_ranking = True
    adder, score = 300, 2              # hard-cost high-competition sample

    base = CFG["geo_anchor"][band] + adder + CFG["zero_ranking_bonus"]
    flat = CFG.get("tier_step_flat")
    if manual_base:
        # A manual override is the operator setting a Brendan-style base
        # directly — his premium cards ($3,950/$5,450/$6,950: Serene, Skidmore)
        # step at 38% of base, so the override ladder should too. Overriding to
        # ~$2,930 hard reproduces that card's upper tiers exactly at 35%.
        step = r50(base * CFG["step_ratio"])
    elif rule and rule.get("step_mode") == "ratio":
        # these ladders step proportionally (Brendan's ecom quote: 38% steps)
        step = r50(base * CFG["step_ratio"])
    elif band == "nationwide":
        # Brendan's national ladder also steps proportionally — $1,500 client
        # on a $3,950 base = the same 38% ratio (Skidmore, 2026-07-20)
        step = r50(base * CFG["step_ratio"])
    elif flat:
        # flat floor, scaling with base for premium clients: Brendan steps
        # ~$950 client on standard quotes but ~$1,300 on his biggest ladder —
        # roughly a quarter of the hard base once the base outgrows the floor.
        pct = CFG.get("tier_step_pct_of_base", 0.24)
        step = max(r50(flat), r50(base * pct))
    else:
        step = r50(base * CFG["step_ratio"])
    tiers = {"base": base, "intermediate": base + step, "advanced": base + 2*step}
    # The demo prices add-ons through the same bracket the real path does, so a
    # walkthrough does not show a number the tool would never quote.
    _dm = 1.0 - addon_discount_pct(addon) / 100.0
    addon_per = {k: r50(v * _dm) for k, v in tiers.items()}

    export_rows = (
        [{"kw": r["kw"], "rank": "Not Found", "comp": "Ultra Competitive"} for r in ultra] +
        [{"kw": r["kw"], "rank": "Not Found", "comp": "Competitive"} for r in comp] +
        [{"kw": r["kw"], "rank": "Not Found", "comp": "Long Tail"} for r in longtail])

    return {
        "demo": True,
        "stage1": {"ultra": ultra, "competitive": comp, "long_tail": longtail, "count": total},
        "stage3a": {"adder": adder, "score": score},
        "stage3b": {"ranked": ranked, "total": total, "frac": 0,
                    "zero_ranking": zero_ranking,
                    "paa": [r["kw"] for r in longtail[:6]], "table": table},
        "stage4": {"anchor": CFG["geo_anchor"][band], "adder": adder,
                   "zero_bonus": CFG["zero_ranking_bonus"], "base": base,
                   "step": step, "tiers": tiers, "addon_per_market": addon_per,
                   "addon_markets": addon, "band": band},
        "export_rows": export_rows,
    }

@app.route("/quote", methods=["POST"])
def quote():
    d = request.get_json(force=True)
    seeds   = clean_seeds(d.get("keywords", []))
    markets = usable_markets(d.get("geo_values") or [])
    state   = (d.get("state") or "").strip()
    domain  = (d.get("domain") or "").strip()
    brand   = (d.get("brand") or "").strip()
    band    = d.get("geo_scope", "single_city")
    addon   = int(d.get("addon_markets", 0) or 0)

    if not seeds:
        return jsonify({"error": "At least one keyword/vertical is required."}), 400
    if band not in CFG["geo_anchor"]:
        return jsonify({"error": f"Unknown geo scope '{band}'."}), 400

    # DEMO_MODE: serve sample data instantly, no API calls, cannot time out.
    if DEMO_MODE:
        return jsonify(mock_pipeline(seeds, markets, state, domain, brand, band, addon))

    try:
        s1 = stage1_keyword_list(seeds, markets, state, brand)
        if not s1["all"]:
            return jsonify({"error": "No keywords returned — try broader seeds or check the market/state."}), 400
        m3 = stage3_metrics(s1["head"], markets, state)
        r3 = stage3_rankcheck(s1["all"], domain, markets, state, brand)
        p  = stage4_price(band, m3["adder"], r3["zero_ranking"], addon,
                          ecommerce=bool(d.get("ecommerce")),
                          industry=(d.get("industry") or ""),
                          ai_search=bool(d.get("ai_search")),
                          national_demand=bool(d.get("national_demand")),
                          geo_override=d.get("geo_override"),
                          addon_override=d.get("addon_override"),
                          goal=(d.get("goal") or ""))
    except requests.HTTPError as e:
        return jsonify({"error": f"DataForSEO request failed: {e}. Check DFS_LOGIN / DFS_PASSWORD, or set DEMO_MODE=1 to run on sample data."}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}. Set DEMO_MODE=1 to run on sample data."}), 500

    # Fold PAA questions into the long-tail bucket (they're real long-tail queries
    # Google confirms users ask). Keep existing long-tails first, then top up with
    # PAA until we hit the target, deduping against everything already in the list.
    used = {r["keyword"].lower() for r in s1["ultra"] + s1["competitive"] + s1["long_tail"]}
    longtail = [{"kw": r["keyword"], "vol": r["volume"]} for r in s1["long_tail"]]
    for q in r3["paa_pool"]:
        if len(longtail) >= CFG["longtail_target"]:
            break
        ql = q.lower()
        if ql not in used:
            used.add(ql)
            longtail.append({"kw": q, "vol": 0})   # PAA has no volume figure

    # Build the exportable keyword table: keyword / rank / competitiveness
    rank_map = {t["keyword"]: t["position"] for t in r3["table"]}
    def comp_label(kw, tier):
        return tier
    export_rows = []
    for r in s1["ultra"]:
        pos = rank_map.get(r["keyword"]); export_rows.append(
            {"kw": r["keyword"], "rank": pos if pos is not None else "Not Found", "comp": "Ultra Competitive"})
    for r in s1["competitive"]:
        pos = rank_map.get(r["keyword"]); export_rows.append(
            {"kw": r["keyword"], "rank": pos if pos is not None else "Not Found", "comp": "Competitive"})
    for lt in longtail:
        pos = rank_map.get(lt["kw"])
        export_rows.append(
            {"kw": lt["kw"], "rank": pos if pos is not None else "Not Found", "comp": "Long Tail"})

    return jsonify({
        "stage1": {
            "ultra":       [{"kw": r["keyword"], "vol": r["volume"]} for r in s1["ultra"]],
            "competitive": [{"kw": r["keyword"], "vol": r["volume"]} for r in s1["competitive"]],
            "long_tail":   longtail,
            "count": len(s1["all"]),
        },
        "stage3a": {"adder": m3["adder"], "score": m3["median_score"]},
        "stage3b": {
            "ranked": r3["ranked"], "total": len(s1["all"]),
            "frac": round(r3["frac"]*100), "zero_ranking": r3["zero_ranking"],
            "paa": r3["paa_pool"][:15],
            "table": [{"kw": t["keyword"],
                       "pos": (t["position"] if t["position"] is not None else "Not Found")}
                      for t in r3["table"]],
        },
        "stage4": {
            "anchor": p["anchor"], "adder": m3["adder"],
            "zero_bonus": CFG["zero_ranking_bonus"] if r3["zero_ranking"] else 0,
            "base": p["base"], "step": p["step"], "tiers": p["tiers"],
            "addon_per_market": p["addon_per_market"], "addon_markets": addon,
            "band": band,
        },
        "export_rows": export_rows,
    })

@app.route("/export.csv", methods=["POST"])
def export_csv():
    """Stateless CSV: frontend posts back the rows it already has."""
    import csv, io
    d = request.get_json(force=True)
    rows = d.get("rows", [])
    client = (d.get("client") or "client").replace(" ", "_")
    buf = io.StringIO()
    w = csv.writer(buf)
    # CPC and keyword difficulty stay ON SCREEN for the reviewer but out of the
    # export — the CSV travels into proposals, and internal pricing signals
    # don't belong in a client-facing artifact.
    w.writerow(["Keyword", "Current Google Rank", "Competitiveness"])
    for r in rows:
        w.writerow([r.get("kw", ""), r.get("rank", ""), r.get("comp", "")])
    from flask import Response
    return Response(buf.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": f"attachment; filename={client}_keywords.csv"})

# ===========================================================================
# STEPPED LIVE ENDPOINTS — each is its own short request so nothing times out.
# The frontend calls them in sequence and holds state between steps.
# ===========================================================================

@app.route("/api/suggest_regions", methods=["POST"])
@_json_error_guard
def api_suggest_regions():
    """Propose vernacular region names for the entered markets, then keep only
    the ones with real search demand for the client's own service."""
    d = request.get_json(force=True) or {}
    markets = usable_markets(d.get("geo_values") or [])
    state = (d.get("state") or "").strip()
    seeds = clean_seeds(d.get("keywords") or [])
    if not markets:
        return jsonify({"regions": [], "rejected": [],
                        "note": "Add at least one geographic targeting area first."})
    if not seeds:
        return jsonify({"regions": [], "rejected": [],
                        "note": "Add a Keyword / Vertical Focus term first — a region name is "
                                "only tested against the client's own service."})
    cands = claude_region_names(markets, state, d.get("brand") or "",
                                d.get("business_desc") or "")
    if not cands:
        return jsonify({"regions": [], "rejected": [],
                        "note": "No commonly-searched regional name for these markets. That is a "
                                "normal answer — most markets don't have one."})
    svc = clean_kw(seeds[0].lower()).strip()
    kept, rejected = validate_region_names(cands, svc, markets, state)
    return jsonify({
        "regions": [{"name": n, "volume": v} for n, v in kept],
        "rejected": [{"name": n, "volume": v} for n, v in rejected],
        "tested_with": svc,
        "note": "" if kept else
                f"Suggested {', '.join(c for c in cands)} — but none reach "
                f"{CFG.get('region_min_volume', 10)} searches a month with \u201c{svc}\u201d "
                f"attached, so none are worth quoting against.",
    })


@app.route("/api/keywords", methods=["POST"])
@_json_error_guard
def api_keywords():
    """Step 1 — build + bucket the keyword list. One ideas call + parallel suggestions."""
    d = request.get_json(force=True)
    seeds   = clean_seeds(d.get("keywords", []))
    markets = usable_markets(d.get("geo_values") or [])
    state   = derive_state(markets, (d.get("state") or "").strip())
    brand   = (d.get("brand") or "").strip()
    domain  = (d.get("domain") or "").strip()
    business_desc = (d.get("business_desc") or "").strip()
    if not seeds:
        return jsonify({"error": "At least one keyword/vertical is required."}), 400
    try:
        s1 = stage1_keyword_list(seeds, markets, state, brand, domain, business_desc)
    except requests.HTTPError as e:
        return jsonify({"error": f"DataForSEO error: {e}. Check funds / credentials."}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}"}), 500
    if not s1["all"]:
        return jsonify({"error": "No keywords returned — try broader seeds or check market/state."}), 400
    conv = lambda L: [{"kw": r["keyword"], "vol": r["volume"],
                       # "broader" = the figure is a county/state/US number
                       # because this city had none of its own; "unknown" = no
                       # figure at all. Neither is this city's demand.
                       "vol_scope": r.get("vol_scope", ""),
                       "vol_area": r.get("vol_area", ""),
                       # WHICH market this row names. Step 3 measures the rank
                       # in that market rather than in the primary one for every
                       # row — see api_rankings. (2026-08-10)
                       "city": r.get("city", ""),
                       "origin": r.get("origin", "")} for r in L]
    resp = {
        "ultra": conv(s1["ultra"]), "competitive": conv(s1["competitive"]),
        "long_tail": conv(s1["long_tail"]), "head": conv(s1["head"]),
        "all": conv(s1["all"]), "refined_by_ai": s1.get("refined_by_ai", False),
        "business_desc": s1.get("business_desc", ""),
        "site_pages_found": s1.get("site_pages_found", 0),
        "site_terms": s1.get("site_terms", []),
        "market_vocab": s1.get("market_vocab", []),
        "market_pool": s1.get("market_pool", []),
    }
    # Thin-list guard: sparse/niche verticals or too few seeds produce a short
    # list. Flag it so the partner can add more seed terms for a fuller table.
    if len(s1["all"]) < 6 or len(s1["competitive"]) == 0:
        resp["thin_warning"] = ("Only a few keywords came back — this vertical may "
            "be low-volume, or try adding more seed terms (e.g. related services) "
            "for a fuller keyword table like the proposals.")
    return jsonify(resp)

@app.route("/api/refine", methods=["POST"])
@_json_error_guard
def api_refine():
    """Step 1b — AI refinement + exact-match volume, run as a SEPARATE request so
    a heavy Claude call can't time out the list build. Takes the buckets the build
    step returned (plus any user edits) and returns the refined, volume-corrected
    list. Non-fatal: on any failure, returns the input list unchanged so the flow
    continues with the rules-based buckets."""
    d = request.get_json(force=True)
    seeds   = clean_seeds(d.get("keywords", []))
    markets = usable_markets(d.get("geo_values") or [])
    state   = derive_state(markets, (d.get("state") or "").strip())
    brand   = (d.get("brand") or "").strip()
    domain  = (d.get("domain") or "").strip()
    business_desc = (d.get("business_desc") or "").strip()
    site_terms_kw = d.get("site_terms", [])
    market_vocab = d.get("market_vocab") or []
    market_pool = [x for x in (d.get("market_pool") or []) if x.get("keyword")]
    phrase_geos = [p.strip() for p in d.get("phrase_geos", []) if p and p.strip()]
    # National demand: RZ industry (ecommerce family) OR nationwide scope OR
    # the operator's manual checkbox. Flips the volume pull to geo-less; the
    # grid itself still uses the client's cities.
    nat_demand, nat_reason = resolve_national_demand(
        industry=(d.get("industry") or ""),
        band=d.get("geo_scope", d.get("band", "")),
        manual=bool(d.get("national_demand")) or bool(d.get("ecommerce")),
        markets=markets, goal=(d.get("goal") or ""))
    # rebuild bucket rows from what the frontend sends back (kw + vol)
    def rows(key):
        return [{"keyword": x["kw"], "volume": x.get("vol", 0), "src": "build"}
                for x in d.get(key, []) if x.get("kw")]
    ultra, competitive, long_tail = rows("ultra"), rows("competitive"), rows("long_tail")
    try:
        s1 = stage1b_refine(seeds, markets, state, brand, domain, business_desc,
                            ultra, competitive, long_tail, site_terms_kw, phrase_geos,
                            national_demand=nat_demand,
                            national_reason=nat_reason,
                            grid_axis=(d.get("grid_axis") or ""),
                            industry=(d.get("industry") or ""),
                            market_vocab=market_vocab, market_pool=market_pool,
                            product_demand=bool(d.get("national_demand"))
                                           or bool(d.get("ecommerce")),
                            goal=(d.get("goal") or ""),
                            suggested=d.get("suggested") or [],
                            negatives=[x.strip() for x in (d.get("negatives") or [])
                                       if str(x).strip()],
                            ranked=d.get("ranked") or [],
                            band=d.get("geo_scope", d.get("band", "")))
    except Exception as e:
        # graceful: hand back the unrefined list so the pipeline still works
        conv0 = lambda L: [{"kw": r["keyword"], "vol": r["volume"], "origin": ""} for r in L]
        app.logger.exception("stage1b_refine failed")
        return jsonify({"national_demand": nat_demand,
                        "national_demand_reason": nat_reason,
                        "ultra": conv0(ultra), "competitive": conv0(competitive),
                        "long_tail": conv0(long_tail),
                        "head": conv0(ultra + competitive),
                        "all": conv0(ultra + competitive + long_tail),
                        "refined_by_ai": False, "refine_attempted": True,
                        "business_desc": "",
                        "site_pages_found": 0, "refine_error": str(e)})
    conv = lambda L: [{"kw": r["keyword"], "vol": r["volume"],
                       # "broader" = the figure is a county/state/US number
                       # because this city had none of its own; "unknown" = no
                       # figure at all. Neither is this city's demand.
                       "vol_scope": r.get("vol_scope", ""),
                       "vol_area": r.get("vol_area", ""),
                       # WHICH market this row names. Step 3 measures the rank
                       # in that market rather than in the primary one for every
                       # row — see api_rankings. (2026-08-10)
                       "city": r.get("city", ""),
                       "origin": r.get("origin", "")} for r in L]
    return jsonify({
        "ultra": conv(s1["ultra"]), "competitive": conv(s1["competitive"]),
        "long_tail": conv(s1["long_tail"]), "head": conv(s1["head"]),
        "all": conv(s1["all"]), "refined_by_ai": s1.get("refined_by_ai", False),
        "refine_attempted": True,
        # Diagnostics from the list build. api_refine names the keys it
        # forwards, so anything stage1b_refine returns that isn't listed here
        # never reaches the browser — every one of these panels has been
        # rendering against undefined and showing nothing (2026-07-28).
        "city_selection": s1.get("city_selection") or {},
        "city_locs": s1.get("city_locs") or {},
        "city_volumes": s1.get("city_volumes") or {},
        "site_locations": s1.get("site_locations") or [],
        "tier_moves": s1.get("tier_moves") or [],
        "scope_warning": s1.get("scope_warning") or "",
        "scope_why": s1.get("scope_why") or "",
        "scope_kind": s1.get("scope_kind") or "warn",
        "scope_note": s1.get("scope_note") or "",
        "service_areas": s1.get("service_areas") or [],
        "gbp_locations": s1.get("gbp_locations"),
        "gbp_cities": s1.get("gbp_cities") or [],
        "seed_ranking": s1.get("seed_ranking") or {},
        "business_desc_inferred": s1.get("business_desc_inferred") or "",
        "pins_refused": s1.get("pins_refused") or [],
        "services_deduped": s1.get("services_deduped") or [],
        "negatives_dropped": s1.get("negatives_dropped") or [],
        "negative_conflicts": s1.get("negative_conflicts") or [],
        "seeds_demoted": s1.get("seeds_demoted") or [],
        "seeds_dropped_suggested": s1.get("seeds_dropped_suggested") or [],
        "seeds_folded": s1.get("seeds_folded") or [],
        "seed_quality": s1.get("seed_quality") or {},
        "service_slots": s1.get("service_slots") or 0,
        "min_unranked_terms": s1.get("min_unranked_terms", 0),
        "unranked_probe_max": s1.get("unranked_probe_max", 0),
        "seed_services_used": s1.get("seed_services_used", 0),
        "seed_services_total": s1.get("seed_services_total", 0),
        "seed_services_dropped": s1.get("seed_services_dropped", 0),
        "pinned_head_terms": s1.get("pinned_head_terms") or [],
        "blocked_pins": s1.get("blocked_pins") or [],
        "dropped_out_of_area": s1.get("dropped_out_of_area") or [],
        "geo_filter_off": bool(s1.get("geo_filter_off")),
        "dropped_ungrounded": s1.get("dropped_ungrounded") or [],
        "grounding_would_drop": s1.get("grounding_would_drop") or [],
        "grounding_gap_words": s1.get("grounding_gap_words") or [],
        "grounding_total": s1.get("grounding_total") or 0,
        "grounding_stood_down": bool(s1.get("grounding_stood_down")),
        # RE-EXPORTED, because the panel reads THIS payload and not stage 1's.
        # The slot-fill valve restored five terms on MPG Gummies and the note
        # naming them rendered nowhere — the key existed on the stage-1 dict and
        # was simply never copied across the boundary. Same shape of miss as the
        # `scope` tag that died here in August. (2026-08-18)
        "grounding_restored": s1.get("grounding_restored") or [],
        "grounding_slot_fill": s1.get("grounding_slot_fill"),
        "business_desc": s1.get("business_desc", ""),
        "site_pages_found": s1.get("site_pages_found", 0),
        "grid": s1.get("grid", False),
        "services": s1.get("services", []),
        "service_volume": s1.get("service_volume", {}),
        "total_volume": s1.get("total_volume", None),
        "volume_error": s1.get("volume_error"),
        "demand_frame": s1.get("demand_frame") or {},
        "grid_axis": s1.get("grid_axis") or {},
        "near_me_added": s1.get("near_me_added") or [],
        "acronym_collisions": s1.get("acronym_collisions") or {},
        "volume_location": s1.get("volume_location"),
        "volume_source": s1.get("volume_source") or "google_ads",
        "state_missing": s1.get("state_missing", False),
        "grid_cities": s1.get("grid_cities", []),
        # stage1b_refine can UPGRADE this after reading the site (a storefront
        # is national demand even when the RZ tag missed it), so its answer
        # wins over the one resolved before the site was fetched.
        "national_demand": bool(s1.get("national_demand", nat_demand)),
        "national_demand_reason": s1.get("national_demand_reason") or nat_reason,
        "ecommerce_detected": bool(s1.get("ecommerce_detected")),
        "ecommerce_reason": s1.get("ecommerce_reason") or "",
        "ecommerce_suppressed": s1.get("ecommerce_suppressed") or "",
        "market_volume_gaps": s1.get("market_volume_gaps") or [],
        "market_renames": s1.get("market_renames") or [],
        "service_swaps": s1.get("service_swaps") or [],
        "service_upgrade_ratio": s1.get("service_upgrade_ratio", 0),
        "topics": s1.get("topics") or [],
        "topic_source": s1.get("topic_source") or "",
        "topic_fixes": s1.get("topic_fixes") or [],
        "geo_forms": s1.get("geo_forms") or [],
        "service_forms": s1.get("service_forms") or [],
        "pool_dropped": s1.get("pool_dropped") or [],
        "pool_added": s1.get("pool_added") or [],
        "pool_status": s1.get("pool_status") or "",
    })

# ---------------------------------------------------------------------------
# IMPORT AN EXISTING SEO REPORT
#
# Why a FILE and not a paste box: the numbers that matter are pictures. A real
# Ski Barn July report was checked (2026-08-05) — the ranking tables are
# embedded PNGs, so text extraction returns the traffic figures and loses the
# entire keyword list. Pasting text cannot work for this format. Reading the
# images can, and does.
#
# Everything extracted is a SUGGESTION. The operator ticks what to apply; the
# quote is never changed by an import on its own. A report is a set of claims —
# sometimes months stale, sometimes framed to flatter the incumbent.
# A slide screenshot is one picture holding several tables, and table text is
# small. Shrinking the whole thing to a 1568px long edge left ~7px glyphs and the
# reader could only make out the bottom table — a Ski Barn import returned the 9
# ski keywords and missed all 10 patio/BBQ ones, saying so in its own notes
# (2026-08-05). So: TILE anything oversized instead of shrinking it, and select
# by PIXEL AREA rather than file size, because bytes say nothing about legibility.
_REPORT_MAX_IMAGES = 16                   # tiles count toward this
_REPORT_MIN_PIXELS = 60_000               # ~250x250; below this it is an icon
# The API downsamples any image to ~1568px on its long edge before reading it,
# so sending a 4000px screenshot gains nothing — it is scaled to 0.38x and a
# table's 11px text becomes 4px. The only way to keep text legible is to TILE at
# native resolution and let each tile use its own 1568px budget.
_REPORT_TILE_EDGE = 1568
_REPORT_TILE_OVERLAP = 0.18               # so a row on a seam survives in a neighbour
_REPORT_TILES_PER_SOURCE = 4              # one dense slide can't eat the whole budget
# Below this width a single scaled image beats slicing: 2030 -> 1568 is 0.77x and
# keeps every row whole, where slicing would strand the labels.
_REPORT_STITCH_ABOVE = 2600
_REPORT_LABEL_FRACTION = 0.28             # left share of a table that holds the keywords


def _emf_text(data, row_tol=None):
    """Read the text out of a Windows METAFILE (.emf / .wmf) table.

    A ranking table pasted from Excel into PowerPoint on Windows arrives as an
    EMF, not a PNG. Pillow cannot open one — "cannot find loader for this WMF
    file" — so _report_images skipped it and the reader was handed a deck with
    no ranking tables in it at all. Junk Bee Gone's slides 2, 3 and 4 were three
    EMFs holding 2,061 text runs between them, and the import returned "no
    keyword table found" (2026-08-09).

    Rendering an EMF needs LibreOffice, which is not on the deploy. But an EMF is
    a list of drawing records and the text is stored as EMR_EXTTEXTOUTW (type
    84) with its own coordinates — so the table can be read directly, and the
    x/y positions rebuild the rows and columns. Text beats a screenshot anyway:
    no tiling, no resolution limit, no OCR.

    Returns tab-separated rows, top-to-bottom then left-to-right.
    """
    import struct
    runs, off, n = [], 0, len(data or b"")
    guard = 0
    while off < n - 8 and guard < 200000:
        guard += 1
        try:
            rtype, size = struct.unpack_from("<II", data, off)
        except Exception:
            break
        if size < 8 or off + size > n:
            break
        if rtype in (83, 84):                     # EXTTEXTOUTA / EXTTEXTOUTW
            try:
                x, y = struct.unpack_from("<ii", data, off + 36)
                nchars, offstr = struct.unpack_from("<II", data, off + 44)
                if 0 < nchars < 4096 and 8 <= offstr < size:
                    raw = data[off + offstr: off + offstr + nchars * (2 if rtype == 84 else 1)]
                    txt = (raw.decode("utf-16-le", "ignore") if rtype == 84
                           else raw.decode("latin-1", "ignore"))
                    txt = txt.replace("\x00", "").strip()
                    if txt:
                        runs.append((y, x, txt))
            except Exception:
                pass
        off += size
    if not runs:
        return ""
    # Row height comes from the FILE, not a guess. A fixed tolerance of 40 was
    # double the real row gap here (median 20), so it merged pairs of rows and
    # the keyword column came out concatenated. Half the median gap between
    # distinct y values splits rows without splitting a single row's runs.
    if row_tol is None:
        ys = sorted({r[0] for r in runs})
        gaps = sorted(b - a for a, b in zip(ys, ys[1:]) if b > a)
        row_tol = max(2, (gaps[len(gaps) // 2] // 2) if gaps else 8)
    # Rebuild rows: same y (within tolerance) is one row, x orders the columns.
    runs.sort(key=lambda r: (r[0], r[1]))
    lines, cur, last_y = [], [], None
    for y, x, txt in runs:
        if last_y is not None and abs(y - last_y) > row_tol:
            if cur:
                lines.append("\t".join(t for _, t in sorted(cur)))
            cur = []
        cur.append((x, txt))
        last_y = y
    if cur:
        lines.append("\t".join(t for _, t in sorted(cur)))
    return "\n".join(lines)


def _pptx_metafile_text(zf):
    """Every EMF/WMF in the deck, read as text and labelled by slide."""
    rel_by_slide = {}
    for name in zf.namelist():
        m = re.match(r"ppt/slides/_rels/(slide\d+)\.xml\.rels$", name)
        if not m:
            continue
        try:
            xml = zf.read(name).decode("utf-8", "ignore")
        except Exception:
            continue
        for media in re.findall(r"\.\./media/([^\"']+\.(?:emf|wmf))", xml, re.I):
            rel_by_slide.setdefault(m.group(1), []).append(media)
    out = []
    for slide in sorted(rel_by_slide, key=lambda s: int(re.sub(r"\D", "", s) or 0)):
        for media in rel_by_slide[slide]:
            try:
                txt = _emf_text(zf.read("ppt/media/" + media))
            except Exception:
                txt = ""
            if txt.strip():
                out.append(f"[{slide}.xml · {media} — table read from metafile]\n{txt}")
    return "\n".join(out)


def _pptx_slide_text(zf):
    """Slide text straight out of the XML. No new dependency: <a:t> holds every
    run of visible text, which is enough for titles, labels and any figures that
    were typed rather than screenshotted."""
    out = []
    for name in sorted(n for n in zf.namelist()
                       if n.startswith("ppt/slides/slide") and n.endswith(".xml")):
        try:
            xml = zf.read(name).decode("utf-8", "ignore")
        except Exception:
            continue
        runs = re.findall(r"<a:t>(.*?)</a:t>", xml, re.S)
        if runs:
            txt = " ".join(re.sub(r"\s+", " ", r).strip() for r in runs)
            out.append(f"[{name.rsplit('/', 1)[-1]}] {txt}")
    return "\n".join(out)[:20000]


def _flatten(im):
    """Composite onto WHITE before anything else.

    Report screenshots are saved with the label column transparent — the text is
    dark pixels over alpha. .convert("RGB") maps transparent to BLACK, so the
    keyword column became black text on black and vanished. That is why an
    import returned "ranking numbers without keyword labels" and dropped 10 of
    19 keywords: the labels were always in the file, and this function was
    erasing them (2026-08-05).
    """
    from PIL import Image
    if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
        im = im.convert("RGBA")
        bg = Image.new("RGBA", im.size, (255, 255, 255, 255))
        return Image.alpha_composite(bg, im).convert("RGB")
    return im.convert("RGB") if im.mode != "RGB" else im


def _encode_jpeg(im, quality=88):
    import io
    im = _flatten(im)
    buf = io.BytesIO()
    im.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def _tile_image(raw):
    """One source image -> request-ready JPEGs.

    Two failures drove this shape (both Ski Barn, 2026-08-05):

    1. Sent whole, a 4094x3379 slide is downsampled by the API to 0.38x and an
       11px table row becomes 4px. Only one of the two keyword tables was
       readable.
    2. Split into columns, the KEYWORD COLUMN lives only in the leftmost tile.
       Every other tile is a field of numbers with nothing to attach them to,
       and the reader correctly refused them: "ranking numbers without keyword
       labels". A 2030px-wide table lost its labels on the second tile.

    So: only split horizontally when the alternative is a severe downscale, and
    when we do, STITCH the label column onto every slice so each tile is
    self-describing. Anything moderately wide is simply scaled to fit — 2030px
    to 1568px is 0.77x and perfectly legible, and keeps rows whole.
    """
    try:
        from PIL import Image
        import io
        im = _flatten(Image.open(io.BytesIO(raw)))
        w, h = im.size
        edge = _REPORT_TILE_EDGE

        def bands(src):
            """Full-width horizontal slices, scaled to fit. Rows stay intact."""
            sw, sh = src.size
            scale = min(1.0, edge / float(sw))
            band_h = int(edge / scale) if scale < 1.0 else edge
            if sh <= band_h:
                return [_encode_jpeg(_fit_width(src, edge))]
            step = max(1, int(band_h * (1.0 - _REPORT_TILE_OVERLAP)))
            ys = list(range(0, max(1, sh - band_h) + 1, step)) or [0]
            if ys[-1] + band_h < sh:
                ys.append(max(0, sh - band_h))
            return [_encode_jpeg(_fit_width(src.crop((0, y, sw, min(sh, y + band_h))), edge))
                    for y in ys][:_REPORT_TILES_PER_SOURCE]

        if w <= _REPORT_STITCH_ABOVE:
            return bands(im)

        # Very wide: slice into column groups, each carrying the label column.
        label_w = max(1, int(w * _REPORT_LABEL_FRACTION))
        slice_w = max(1, edge - label_w)
        out, x = [], label_w
        while x < w and len(out) < _REPORT_TILES_PER_SOURCE:
            right = im.crop((x, 0, min(w, x + slice_w), h))
            labels = im.crop((0, 0, label_w, h))
            stitched = Image.new("RGB", (label_w + right.size[0], h), "white")
            stitched.paste(labels, (0, 0))
            stitched.paste(right, (label_w, 0))
            out.extend(bands(stitched))
            x += max(1, int(slice_w * (1.0 - _REPORT_TILE_OVERLAP)))
        return out[:_REPORT_TILES_PER_SOURCE] or bands(im)
    except Exception:
        return [raw]


def _fit_width(im, max_w):
    from PIL import Image
    if im.size[0] <= max_w:
        return im
    ratio = max_w / float(im.size[0])
    return im.resize((max_w, max(1, int(im.size[1] * ratio))), Image.LANCZOS)


def _report_images(filename, data):
    """Candidate images from the upload, largest first. A .pptx carries its
    screenshots in ppt/media; a bare image is itself."""
    lower = (filename or "").lower()
    if lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
        return [data]
    if lower.endswith((".pptx", ".potx")):
        import io, zipfile
        try:
            from PIL import Image
        except Exception:
            Image = None
        scored = []
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for n in zf.namelist():
                if not (n.startswith("ppt/media/")
                        and n.lower().endswith((".png", ".jpg", ".jpeg"))):
                    continue
                raw = zf.read(n)
                px = 0
                if Image is not None:
                    try:
                        px = Image.open(io.BytesIO(raw)).size[0] * \
                             Image.open(io.BytesIO(raw)).size[1]
                    except Exception:
                        px = 0
                if px and px < _REPORT_MIN_PIXELS:
                    continue                      # icon, arrow, logo
                scored.append((px or len(raw), raw))
        # Biggest by pixel area first — that is where dense tables live.
        return [raw for _px, raw in sorted(scored, key=lambda x: -x[0])]
    return []


def _report_text(filename, data):
    lower = (filename or "").lower()
    if lower.endswith((".pptx", ".potx")):
        import io, zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                # Slide text first, then any table that arrived as a metafile.
                # The metafile text is the ranking table itself on Windows decks,
                # so it must not be truncated away by the slide-text cap.
                parts = [_pptx_slide_text(zf)]
                try:
                    mf = _pptx_metafile_text(zf)
                except Exception:
                    app.logger.exception("_pptx_metafile_text failed")
                    mf = ""
                if mf:
                    parts.append(mf)
                return "\n".join(x for x in parts if x)
        except Exception:
            return ""
    if lower.endswith((".txt", ".csv", ".md")):
        return data.decode("utf-8", "ignore")[:20000]
    return ""


_REPORT_SCHEMA_PROMPT = """You are reading an SEO performance report for a client.
Extract ONLY what is actually present. Never infer, never complete a pattern, never
invent a plausible keyword. An empty list is the correct answer when the report does
not show something.

Return ONLY valid JSON, no prose:
{
  "period": "the reporting period as written, e.g. July 2026",
  "monthly_spend": null,
  "monthly_revenue": null,
  "markets": ["City, ST for every location column or named market"],
  "state": "two-letter state the client operates in, or null if genuinely unclear",
  "keywords": ["exact keyword text as written, one per row of any ranking table"],
  "notes": ["anything a pricing reviewer should know, one short sentence each"],
  "rankings": [{"keyword": "...", "best_position": 3, "markets_top10": 4}]
}

EMIT THE FIELDS IN THAT ORDER. "rankings" is last on purpose: it is the longest
section and the least important, so if you run out of room the loss falls there
rather than on the keyword list.

Rules:
- Keyword tables are often SCREENSHOTS. Read them from the images.
- A block headed "table read from metafile" is a ranking table that was pasted from
  Excel and has been decoded to TEXT for you. It is tab-separated: the first column is
  the keyword, the remaining columns are that keyword's position in each market, in the
  order the header row lists them. Trust it over any image — it is exact, not OCR.
- THERE ARE USUALLY SEVERAL TABLES, sometimes stacked in one screenshot.
- LARGE SCREENSHOTS ARE SUPPLIED AS OVERLAPPING TILES of the same picture. On a
  very wide table the KEYWORD COLUMN IS COPIED ONTO THE LEFT EDGE OF EVERY TILE,
  so a tile showing positions always carries its own labels — read them together
  rather than reporting the table as unlabelled. Tiles
  repeat content on purpose: a row cut by one tile's edge appears whole in its
  neighbour. Read every tile, MERGE what they show, and do not count a keyword
  twice because it appeared in two tiles. If a tile shows a keyword column
  without its position columns, pair it with the overlapping tile that has them.
- Read EVERY table in EVERY image and merge them. A retailer often has one table per product family
  (e.g. ski gear in one, patio/BBQ in another) — returning only the clearest
  table is a failure. If a row is genuinely illegible, say so in notes rather
  than dropping the table silently.
- Keep placeholders exactly as printed: "bbq grills <cityname>" stays as-is.
- A position of 100 (or 100+) usually means NOT RANKING.
- KEEP THE OUTPUT COMPACT. Do NOT emit one row per keyword-per-market — that is
  hundreds of rows and the response gets cut off mid-JSON. For each keyword give
  ONE row: its best position across the markets, and how many markets have it in
  the top 10. Keywords and markets are the important part; positions are context.
- monthly_spend is what the CLIENT PAYS for the campaign. Revenue, traffic value
  and ad spend are NOT campaign spend — leave it null unless the report states a fee.
- markets: only real places used as columns or headings, not every city mentioned.
- ALWAYS QUALIFY A MARKET WITH ITS STATE, even when the column header does not.
  Report columns usually read bare ("Wayne", "Paramus"), and a bare city name is
  ambiguous — there is a Wayne in NJ, PA and MI — so the quote is built without a
  state suffix and the rank check looks in the wrong place. Infer the state from
  the rest of the report: the client's own name, the business description, other
  place names, geo-qualified keywords like "ski shop paramus nj". Return
  "Wayne, NJ", not "Wayne". Set "state" to the client's main state as well. If a
  market is plainly in a different state from the others (e.g. New York City
  alongside New Jersey towns), qualify it with its OWN state.
- If a table appears twice (start vs current), use the CURRENT/most recent one for
  rankings and say so in notes."""


@app.route("/api/import_report", methods=["POST"])
@_json_error_guard
def api_import_report():
    """Read an existing SEO report and return what it contains as SUGGESTIONS."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return jsonify({"error": "ANTHROPIC_API_KEY is not set on this deploy, so "
                                 "report reading is unavailable."}), 400
    f = request.files.get("file")
    if f is None:
        return jsonify({"error": "No file received."}), 400
    data = f.read()
    if not data:
        return jsonify({"error": "That file is empty."}), 400
    if len(data) > 25 * 1024 * 1024:
        return jsonify({"error": "File is larger than 25MB — export fewer slides."}), 400

    imgs = _report_images(f.filename, data)
    text = _report_text(f.filename, data)
    if not imgs and not text.strip():
        return jsonify({"error": f"Nothing readable in {f.filename!r}. Supported: "
                                 ".pptx, .png, .jpg, .txt, .csv."}), 400

    # ROUND-ROBIN across source images. Taking them in order let the first,
    # densest slide spend the entire budget, so the small single-table
    # screenshots later in the deck were never sent at all.
    per_source = [_tile_image(raw) for raw in imgs]
    content, n_sent, depth = [], 0, 0
    while n_sent < _REPORT_MAX_IMAGES and any(len(t) > depth for t in per_source):
        for tiles in per_source:
            if n_sent >= _REPORT_MAX_IMAGES:
                break
            if len(tiles) > depth:
                content.append({"type": "image",
                                "source": {"type": "base64",
                                           "media_type": "image/jpeg",
                                           "data": base64.b64encode(tiles[depth]).decode()}})
                n_sent += 1
        depth += 1
    content.append({"type": "text",
                    "text": _REPORT_SCHEMA_PROMPT
                            + ("\n\nTEXT FOUND IN THE FILE:\n" + text if text.strip()
                               else "\n\n(No usable text layer — read the images.)")})
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({"model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                             "max_tokens": 16000, "temperature": 0,
                             "messages": [{"role": "user", "content": content}]}),
            timeout=120)
        resp.raise_for_status()
        body = resp.json()
        txt = "".join(b.get("text", "") for b in body.get("content", [])
                      if b.get("type") == "text").strip()
        truncated = body.get("stop_reason") == "max_tokens"
    except Exception as e:
        return jsonify({"error": f"Report read failed: {e}"}), 502

    def _salvage(t):
        """A truncated reply is cut off in the LAST array it was writing, so the
        earlier ones are complete and worth keeping. Reading 50 keywords and
        then discarding all of them over a missing brace is the wrong trade
        (2026-08-06)."""
        got = {}
        for key in ("keywords", "markets", "notes"):
            mm = re.search(r'"%s"\s*:\s*\[(.*?)\]' % key, t, re.S)
            if mm:
                got[key] = [x.strip().strip(",").strip()
                            for x in re.findall(r'"((?:[^"\\]|\\.)*)"', mm.group(1))]
        for key in ("period",):
            mm = re.search(r'"%s"\s*:\s*"((?:[^"\\]|\\.)*)"' % key, t, re.S)
            if mm:
                got[key] = mm.group(1)
        for key in ("monthly_spend", "monthly_revenue"):
            mm = re.search(r'"%s"\s*:\s*([0-9.]+)' % key, t)
            if mm:
                try:
                    got[key] = float(mm.group(1))
                except Exception:
                    pass
        return got

    m = re.search(r"\{.*\}", txt, re.S)
    out, partial = None, False
    if m:
        try:
            out = json.loads(m.group(0))
        except Exception:
            out = None
    if out is None:
        out = _salvage(txt)
        partial = True
        if not out.get("keywords") and not out.get("markets"):
            return jsonify({"error": ("The reader's reply was cut off before "
                                      "anything usable came back."
                                      if truncated else
                                      "The reader did not return usable JSON."),
                            "raw": txt[:400]}), 502

    # Normalise, and cap so a 200-row report can't flood the seed field.
    kws = [str(x).strip() for x in (out.get("keywords") or []) if str(x).strip()][:120]
    mkts = [str(x).strip() for x in (out.get("markets") or []) if str(x).strip()][:40]
    ranks = [r for r in (out.get("rankings") or []) if isinstance(r, dict)][:600]
    if partial or truncated:
        out.setdefault("notes", []).insert(0,
            "The reader's reply was cut off, so this is a partial read — the "
            "keyword and market lists came through but position detail may be "
            "incomplete. Check the list against the report before relying on it.")
    return jsonify({
        "ok": True,
        "partial": bool(partial or truncated),
        "filename": f.filename,
        # Size travels with the extraction so a quote can say WHICH file it read
        # without storing the file (see rememberReport in the template).
        "size": len(data),
        "images_read": n_sent,
        "sources_read": len(imgs),
        "text_chars": len(text),
        "keywords": kws,
        "markets": mkts,
        "rankings": ranks,
        "monthly_spend": out.get("monthly_spend"),
        "monthly_revenue": out.get("monthly_revenue"),
        "period": (out.get("period") or "")[:120],
        "state": (str(out.get("state") or "").strip().upper()[:2] or None),
        "notes": [str(n)[:300] for n in (out.get("notes") or [])][:8],
    })


@app.route("/api/metrics", methods=["POST"])
@_json_error_guard
def api_metrics():
    """Step 2 — competitive adder from head-term bids. One search_volume call."""
    d = request.get_json(force=True)
    head    = [{"keyword": k} for k in d.get("head", [])]
    markets = usable_markets(d.get("geo_values") or [])
    # phrase geos must be strippable so bare-term metrics resolve for
    # "managed it services south jersey" -> "managed it services"
    # measure_first, not primary_first: the CPC adder is a localised lookup, so
    # it belongs in the client's home state rather than in whichever market
    # happens to carry the most demand (2026-08-07).
    markets = measure_first(markets, (d.get("state") or "").strip(),
                            d.get("primary_market"))
    markets = markets + [p.strip() for p in d.get("phrase_geos", []) if p and p.strip()]
    state   = derive_state(markets, (d.get("state") or "").strip())
    # Same national-demand basis Step 1 used, resolved the same way rather than
    # trusted from the client, so the two steps cannot disagree.
    nat, _nr = resolve_national_demand(
        industry=(d.get("industry") or ""),
        band=d.get("geo_scope", d.get("band", "")),
        manual=bool(d.get("national_demand")) or bool(d.get("ecommerce")),
        markets=markets, goal=(d.get("goal") or ""))
    try:
        m3 = stage3_metrics(head, markets, state, national=nat,
                            industry=(d.get("industry") or ""))
    except requests.HTTPError as e:
        return jsonify({"error": f"DataForSEO error: {e}."}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}"}), 500
    return jsonify({"adder": m3["adder"], "score": m3["median_score"],
                    "adder_basis": m3.get("adder_basis"), "cpc_used": m3.get("cpc_used"),
                    "cpc_low_confidence": m3.get("cpc_low_confidence"),
                    "cpc_n_bids": m3.get("cpc_n_bids"),
                    "flat_adder": m3.get("flat_adder"),
                    "bid_source": m3.get("bid_source"),
                    "bid_ideas_error": m3.get("bid_ideas_error"),
                    "bid_labs_error": m3.get("bid_labs_error"),
                    "adder_blocked": m3.get("adder_blocked"),
                    "kd_suggested_adder": m3.get("kd_suggested_adder"),
                    "kd_score": m3.get("kd_score"),
                    "restricted_vertical": m3.get("restricted_vertical"),
                    "national_demand": nat,
                    "bid_error": m3.get("bid_error"),
                    "bid_location": m3.get("bid_location"),
                    "bid_terms_queried": m3.get("bid_terms_queried"),
                    "n_markets": m3.get("n_markets"),
                    "cpc": m3.get("cpc", {}), "kd": m3.get("kd", {}),
                    "median_kd": m3.get("median_kd"), "kd_error": m3.get("kd_error"),
                    "kd_vs_cpc": m3.get("kd_vs_cpc"),
                    "bid_stats": m3.get("bid_stats"), "breaks": m3.get("breaks")})

def _serp_parse_items(items, domain_dom, brand, top_n=None):
    """Shared SERP parsing for live + task modes: first organic position for
    the client domain, People-Also-Ask questions (brand-mention filtered), and
    WHO IS ACTUALLY ON PAGE ONE.

    The domains were always in this response and were always thrown away. They
    are the one thing that distinguishes two clients the rest of the formula
    reads as identical: Amare's page one for Santa Fe rentals is Zillow,
    Apartments.com and Trulia; Nob Hill Dental's is other dentists in Salem.
    Same keyword count, same not-found share, same CPC band, same difficulty
    band — and Brendan priced them $600 apart. Free, because this SERP is
    already fetched for the rank check. (2026-08-17)
    """
    top_n = int(top_n or CFG.get("serp_competitor_depth", 10))
    pos, paa, doms = None, [], []
    for it in items or []:
        if it.get("type") == "organic":
            d = (it.get("domain") or "").lower().replace("www.", "")
            if domain_dom and domain_dom in (it.get("domain") or ""):
                if pos is None:
                    pos = it.get("rank_absolute")
            elif d and d not in doms and len(doms) < top_n:
                doms.append(d)
        if it.get("type") == "people_also_ask":
            for el in it.get("items", []):
                q = el.get("title")
                if q and (brand or "").lower() not in q.lower():
                    paa.append(q)
    return pos, paa, doms


@app.route("/api/rankings_submit", methods=["POST"])
@_json_error_guard
def api_rankings_submit():
    """Step 3, async mode — submit ALL rank lookups as DataForSEO tasks in one
    call. Task mode has no 30s wall: the platform ceiling only ever killed us
    because LIVE lookups block while Google is crawled. Tasks queue server-side
    and the frontend polls /api/rankings_collect until they land."""
    d = request.get_json(force=True)
    kws     = [k for k in d.get("keywords", []) if k]
    markets = usable_markets(d.get("geo_values") or [])
    state   = derive_state(markets, (d.get("state") or "").strip())
    markets = measure_first(markets, state, d.get("primary_market"))
    markets = rank_markets(kws, markets, state)
    top_n   = CFG["zero_ranking_top_n"]
    depth   = max(top_n, 10)
    nat, _r = resolve_national_demand(d.get("industry") or "",
                                      d.get("geo_scope") or d.get("band") or "",
                                      bool(d.get("national_demand")),
                                      markets=markets,
                                      goal=(d.get("goal") or ""))
    loc     = rank_location(markets, state, nat)
    # Same per-market rule as the live path: a row naming Morristown is measured
    # in Morristown. Task mode is what the retry button uses, so leaving it on a
    # single location would silently mix two measurement bases in one table.
    cities  = {str(k.get("kw") or ""): str(k.get("city") or "")
               for k in (d.get("rows") or []) if isinstance(k, dict)}

    def _loc_for(kw):
        if nat:
            return loc
        city = cities.get(kw, "")
        named = (next((m for m in markets
                       if parse_market(m, state)[0].strip().lower() == city.lower()),
                      city) if city else market_for_keyword(kw, markets, state))
        return rank_location([named], state, False) if named else loc

    payload = [{"keyword": kw, "location_name": _loc_for(kw), "language_code": "en",
                "depth": depth, "priority": 2, "tag": kw[:255]} for kw in kws]
    try:
        data = dfs_post("/serp/google/organic/task_post", payload, timeout=25)
    except Exception as e:
        return jsonify({"error": f"task submit failed: {e}"}), 502
    out = []
    for t in (data.get("tasks") or []):
        kw = ((t.get("data") or {}).get("keyword")) or ((t.get("data") or {}).get("tag")) or ""
        if t.get("status_code") in (20100, 20000) and t.get("id"):
            out.append({"kw": kw, "task_id": t["id"]})
        else:
            out.append({"kw": kw, "task_id": None,
                        "error": f"{t.get('status_code')}: {t.get('status_message')}"})
    # Report the SAME per-market location the live path reports. Task mode is
    # what the slow tail falls back to, and on a cold run that can be every row
    # — so a note reading "Measured in Knoxville" here made the whole per-market
    # change look like it had not shipped. (2026-08-10)
    note = rank_location_note(markets, state, nat)
    locs = sorted({_loc_for(kw) for kw in kws})
    if len(locs) > 1:
        pretty = [l.replace(",United States", "").replace(",", ", ") for l in locs]
        note = {"location": ", ".join(pretty), "scope": "per_market",
                "note": "Each keyword was measured in the market it names — "
                        + ", ".join(pretty) + "."}
    for r in out:
        r["loc"] = (_loc_for(r.get("kw") or "")
                    .replace(",United States", "").replace(",", ", "))
    return jsonify({"tasks": out, "rank_location": note})


@app.route("/api/rankings_collect", methods=["POST"])
@_json_error_guard
def api_rankings_collect():
    """Poll pending rank tasks. Returns done rows (same shape as /api/rankings)
    and the still-pending task list to poll again."""
    d = request.get_json(force=True)
    tasks  = d.get("tasks", [])
    domain = (d.get("domain") or "").strip()
    brand  = (d.get("brand") or "").strip()
    dom = domain.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    top_n = CFG["zero_ranking_top_n"]
    done, pending, paa = [], [], []
    rivals = {}

    def one(t):
        data = dfs_post(f"/serp/google/organic/task_get/regular/{t['task_id']}",
                        None, timeout=12, method="GET")
        task0 = (data.get("tasks") or [{}])[0]
        sc = task0.get("status_code")
        if sc == 20000:
            res = (task0.get("result") or [{}])[0]
            pos, qs, doms = _serp_parse_items(res.get("items") or [], dom, brand)
            return ("done", pos, qs, doms)
        if sc in (40601, 40602, 40100):      # queued / in progress
            return ("pending", None, [], [])
        return ("error", None, [], [])

    with ThreadPoolExecutor(max_workers=8) as ex:
        futs = {ex.submit(one, t): t for t in tasks if t.get("task_id")}
        results = {}
        for fut in futs:
            t = futs[fut]
            try:
                results[t["kw"]] = fut.result()
            except Exception:
                results[t["kw"]] = ("pending", None, [], [])  # transient: poll again
    for t in tasks:
        if not t.get("task_id"):
            done.append({"kw": t["kw"], "pos": "—", "ranked_top": False, "error": True})
            continue
        status, pos, qs, doms = results.get(t["kw"], ("pending", None, [], []))
        if status == "done":
            done.append({"kw": t["kw"],
                         "pos": (pos if pos is not None else "Not Found"),
                         "ranked_top": (pos is not None and pos <= top_n),
                         "error": False})
            paa.extend(qs)
            # WHO IS ALREADY THERE. Collected per keyword and counted across the
            # list below, because one SERP is an anecdote and ten is a market.
            for _d in doms:
                rivals[_d] = rivals.get(_d, 0) + 1
        elif status == "error":
            done.append({"kw": t["kw"], "pos": "—", "ranked_top": False, "error": True})
        else:
            pending.append(t)
    return jsonify({"done": done, "pending": pending, "paa": paa[:40],
                    "rivals": [{"domain": d, "appearances": n}
                               for d, n in sorted(rivals.items(),
                                                  key=lambda kv: (-kv[1], kv[0]))
                               [:int(CFG.get("serp_rival_cap", 12))]]})


@app.route("/api/market_signals", methods=["POST"])
@_json_error_guard
def api_market_signals():
    """The three things the price has never seen, measured and REPORTED.

    Four pricing thresholds in this tool sit above where its clients live:
    search volume pays nothing below 10,000/mo, the CPC adder needs a $20 click
    before rounding leaves anything, organic difficulty's first break is at 30
    and every client measured reads 0-25, and the zero-visibility bonus wants 90%
    not-ranking. So every quote computes under the floor and the floor is the
    price — which is why two clients Brendan priced $600 apart came out identical.

    This measures what is left, on the theory that what actually differs between
    a Santa Fe rental community and a Salem dental practice is WHO IS ALREADY ON
    PAGE ONE and how far the client is from them:

      1. the incumbents   — free, read off the rank-check SERP we already fetch
      2. the authority gap — one bulk_ranks call for the client and all of them
      3. technical debt    — one instant_pages call, fields we already receive

    NOTHING HERE TOUCHES THE PRICE. Difficulty was banded by analogy this morning
    and turned out to be measuring national difficulty of a bare term for a local
    client — the wrong quantity, not just the wrong band. So these are reported
    first, across a real spread of clients, and banded only once there is a range
    to band against. (2026-08-17)
    """
    d = request.get_json(force=True) or {}
    domain = (d.get("domain") or "").strip().lower()
    domain = domain.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    rivals = [str(x.get("domain") if isinstance(x, dict) else x or "").lower()
              for x in (d.get("rivals") or [])]
    rivals = [r for r in rivals if r and r != domain][:int(CFG.get("serp_rival_cap", 12))]
    counts = {str(x.get("domain") or "").lower(): int(x.get("appearances") or 0)
              for x in (d.get("rivals") or []) if isinstance(x, dict)}

    ranks, rank_err = fetch_domain_authority(([domain] if domain else []) + rivals)
    health, health_err = fetch_technical_health(domain) if domain else ({}, "")

    client_rank = ranks.get(domain)
    # A SOCIAL PROFILE IS NOT WHO YOU ARE COMPETING WITH, AND THE FILTER WAS
    # ONLY ON ONE OF THE THREE NUMBERS. pageone_strength() has excluded these
    # since the Instagram result put Pennsylvania Center in the national band —
    # but the MEDIAN and the STRONGEST were still computed over everything, so a
    # panel could read "page one strength 704 — national platforms" one line
    # above "strongest 1,000", and the proposal's authority gap was measured
    # against a Facebook page nobody is trying to out-rank. Same rule, all three
    # figures. The list still SHOWS them, because they are genuinely on the page
    # and hiding them would misdescribe it. (2026-08-19)
    _real = [r for r in rivals
             if str(r or "").lower().replace("www.", "") not in _PAGEONE_NON_RIVAL]
    rival_ranks = [ranks[r] for r in _real if r in ranks]
    # Fall back to the unfiltered set rather than reporting nothing: a page one
    # made entirely of social profiles is a real, and telling, measurement.
    if not rival_ranks:
        rival_ranks = [ranks[r] for r in rivals if r in ranks]
    rival_ranks.sort()
    median_rival = (rival_ranks[len(rival_ranks) // 2] if rival_ranks else None)
    top_rival = (rival_ranks[-1] if rival_ranks else None)
    # THE GAP DECLINES TO BE COMPUTED OFF AN UNMEASURED CLIENT. A client rank of
    # 0 means bulk_ranks had no backlink data for them, which is indistinguishable
    # from a site with genuinely none — and subtracting it produces the largest
    # gap in the book from the weakest evidence in it. Amare read client 0 against
    # a median incumbent of 565: a gap of 565 that is not comparable to Nob Hill's
    # 35 against a measured 91. The median incumbent needs no client reading at
    # all, which is why it is the number to band on. (2026-08-17)
    client_measured = bool(client_rank)
    gap = (median_rival - client_rank
           if client_measured and median_rival is not None else None)
    _rivrows = [{"domain": r, "rank": ranks.get(r), "appearances": counts.get(r, 0)}
                for r in rivals]
    # What the PRICE reads — see pageone_strength(). The median stays on the
    # panel because it is the honest summary of the page; the band keys on the
    # strongest repeated incumbent because one Zillow changes the job.
    strength = pageone_strength(_rivrows)
    band = _pageone_bucket(strength)
    band_add = int((CFG.get("pageone_anchor_add") or {}).get(band, 0) if band else 0)

    return jsonify({
        "domain": domain,
        "client_rank": client_rank,
        "client_measured": client_measured,
        "pageone_rank": strength,
        "pageone_band": band,
        "pageone_add": band_add,
        "median_rival_rank": median_rival,
        "top_rival_rank": top_rival,
        "gap": gap,
        "rivals": [{"domain": r, "rank": ranks.get(r),
                    "appearances": counts.get(r, 0)}
                   for r in rivals],
        "n_measured": len(rival_ranks),
        "health": health,
        "authority_error": rank_err,
        "health_error": health_err,
        # NO LONGER ALWAYS FALSE. As of 2026-08-18 the page-one band carries a
        # partner-dollar add, so this says which of the three signals moved the
        # quote. Authority-gap and site condition are still reported only.
        "applied_to_price": bool(band_add),
    })


# ---------------------------------------------------------------------------
# BACK-MEASURE: the signal, run against the quotes whose real price is known.
#
# Three clients measured by hand said median incumbent authority orders the way
# Brendan's prices order — 126 / 437 / 565 against $2,950 / $3,550 / $3,550 —
# and it was the first input that ever separated them. But three points contain
# exactly one boundary, and four of the eight inputs tested "fit" it, including
# organic difficulty, which had already been ruled out by experiment that
# morning. A fit on one boundary is not a finding.
#
# The evidence needed already exists: twelve saved quotes carrying a domain, a
# keyword list and the price Brendan actually sent. This runs the same
# measurement over all of them so the band is fitted on twelve points instead of
# three. Nothing here touches a price, and it never will — banding is a separate
# decision taken after looking at the table.
#
# ONE CLIENT PER REQUEST. Five SERPs a client at a 12/min ceiling is about six
# minutes across the book, and Render kills a request long before that. The
# browser drives the loop, exactly as the rank check already does. (2026-08-17)


def backmeasure_targets(payloads):
    """Saved quotes that can be back-measured, and why the rest cannot."""
    out, skipped = [], []
    for q in payloads or []:
        p = q.get("payload") or {}
        if isinstance(p, str):
            try:
                p = json.loads(p)
            except Exception:                                 # noqa: BLE001
                continue
        inp = p.get("inputs") or {}
        name = q.get("name") or q.get("client") or f"quote {q.get('id')}"
        dom = str(inp.get("domain") or (inp.get("sites") or [""])[0] or "").strip()
        actual = (p.get("actual") or {}).get("base")
        # The head terms the build already chose. Falling back to the full list
        # sorted by volume covers quotes saved before the head split existed.
        kw = p.get("kw") or {}
        head = [str(r.get("kw") or "") for r in (kw.get("head") or []) if r.get("kw")]
        if not head:
            head = [str(r.get("kw") or "") for r in
                    sorted((kw.get("all") or []),
                           key=lambda r: -int(r.get("vol") or 0)) if r.get("kw")]
        why = ("no client website saved" if not dom
               else "no keyword list saved" if not head
               else "no Actual price entered" if not actual else "")
        row = {"id": q.get("id"), "name": name, "client": q.get("client") or "",
               "domain": dom, "actual_base": actual,
               "formula_base": ((p.get("pricing") or {}).get("client_tiers")
                                or {}).get("base"),
               "markets": len(inp.get("geo_values") or []),
               "industry": _first_industry(inp.get("industry")),
               "terms": len(head)}
        if why:
            skipped.append(dict(row, why=why))
        else:
            out.append(row)
    return out, skipped


@app.route("/api/backmeasure/list")
@_json_error_guard
def api_backmeasure_list():
    """Which saved quotes the back-measure can run on, without running it."""
    if not storage.enabled():
        return jsonify({"targets": [], "skipped": [], "error": "saving is off"})
    targets, skipped = backmeasure_targets(storage.all_payloads("seo"))
    return jsonify({"targets": targets, "skipped": skipped,
                    "terms_each": int(CFG.get("backmeasure_terms", 5) or 5),
                    "applied_to_price": False})


@app.route("/api/backmeasure/one", methods=["POST"])
@_json_error_guard
def api_backmeasure_one():
    """Measure ONE saved quote's page one. Called in a loop by the browser."""
    if not storage.enabled():
        return jsonify({"error": "saving is off"}), 400
    d = request.get_json(force=True) or {}
    qid = d.get("id")
    rec = storage.load_quote(qid) if qid is not None else None
    if not rec:
        return jsonify({"error": f"quote {qid} not found"}), 404
    p = rec.get("payload") or {}
    if isinstance(p, str):
        p = json.loads(p)
    inp = p.get("inputs") or {}
    dom = str(inp.get("domain") or (inp.get("sites") or [""])[0] or "").strip()
    dom = dom.lower().replace("https://", "").replace("http://", "")
    dom = dom.replace("www.", "").strip("/")
    brand = inp.get("brand") or rec.get("client") or ""
    markets = usable_markets(inp.get("geo_values") or [])
    state = derive_state(markets, (inp.get("state") or "").strip())
    kw = p.get("kw") or {}
    head = [str(r.get("kw") or "") for r in (kw.get("head") or []) if r.get("kw")]
    if not head:
        head = [str(r.get("kw") or "") for r in
                sorted((kw.get("all") or []),
                       key=lambda r: -int(r.get("vol") or 0)) if r.get("kw")]
    head = head[:int(CFG.get("backmeasure_terms", 5) or 5)]
    if not dom or not head:
        return jsonify({"error": "nothing to measure on this quote"}), 400

    # Page one, per term, through the SAME parser the live rank check uses — the
    # one place a hand-rolled second copy has already cost a whole client's
    # incumbents once today.
    deadline = time.time() + 22
    rivals, errs, measured = {}, [], 0
    for k in head:
        if time.time() > deadline - 4:
            errs.append("ran out of request budget")
            break
        try:
            _pos, _paa, doms = _serp_one(k, dom, markets, state, brand,
                                         int(CFG.get("serp_competitor_depth", 10)),
                                         deadline=deadline)
            measured += 1
            for x in doms or []:
                rivals[x] = rivals.get(x, 0) + 1
        except Exception as e:                                # noqa: BLE001
            errs.append(f"{k}: {str(e)[:60]}")

    order = sorted(rivals.items(), key=lambda kv: (-kv[1], kv[0]))
    top = [x for x, _ in order][:int(CFG.get("serp_rival_cap", 12))]
    ranks, rank_err = fetch_domain_authority(([dom] if dom else []) + top)
    rr = sorted(ranks[r] for r in top if r in ranks)
    median_rival = rr[len(rr) // 2] if rr else None
    client_rank = ranks.get(dom)
    client_measured = bool(client_rank)
    # KEPT ON THE QUOTE, so the Calibration panel can test "who is on page one"
    # the same way it tests geo band and search volume — rather than in a
    # separate table read by eye. Five SERPs and a minute of rate-limited waiting
    # is not a thing to pay for twice, and a variable that is only ever eyeballed
    # never gets the fragility guards the other eight already have. Additive: no
    # version snapshot, no reordering, no price touched.
    _rivrows = [{"domain": x, "rank": ranks.get(x), "appearances": n}
                for x, n in order[:int(CFG.get("serp_rival_cap", 12))]]
    signals = {"pageone_rank": pageone_strength(_rivrows),
               "median_rival_rank": median_rival,
               "top_rival_rank": (rr[-1] if rr else None),
               "client_rank": client_rank, "client_measured": client_measured,
               "n_incumbents": len(rr), "terms_measured": measured,
               "rivals": [{"domain": x, "rank": ranks.get(x), "appearances": n}
                          for x, n in order[:int(CFG.get("serp_rival_cap", 12))]]}
    try:
        storage.patch_signals(rec.get("id"), signals)
    except Exception:                                         # noqa: BLE001
        app.logger.exception("could not attach signals to quote")
    return jsonify({
        "id": rec.get("id"), "name": rec.get("name") or "",
        "client": rec.get("client") or "", "domain": dom,
        "actual_base": (p.get("actual") or {}).get("base"),
        "formula_base": ((p.get("pricing") or {}).get("client_tiers") or {}).get("base"),
        "markets": len(inp.get("geo_values") or []),
        "industry": _first_industry(inp.get("industry")),
        "terms_measured": measured, "terms_asked": len(head),
        "median_rival_rank": median_rival,
        "top_rival_rank": (rr[-1] if rr else None),
        "pageone_rank": signals.get("pageone_rank"),
        "pageone_band": _pageone_bucket(signals.get("pageone_rank")),
        "pageone_add": int((CFG.get("pageone_anchor_add") or {}).get(
            _pageone_bucket(signals.get("pageone_rank")) or "", 0)),
        "client_rank": client_rank,
        "client_measured": client_measured,
        # Only when the client was actually measured — see the note in
        # api_market_signals. An unmeasured client produces the biggest gap in
        # the book off the weakest evidence in it.
        "gap": (median_rival - client_rank
                if client_measured and median_rival is not None else None),
        "n_incumbents": len(rr),
        "rivals": [{"domain": x, "rank": ranks.get(x), "appearances": n}
                   for x, n in order[:int(CFG.get("serp_rival_cap", 12))]],
        "errors": errs + ([rank_err] if rank_err else []),
        "applied_to_price": False,
    })


@app.route("/api/rank_location", methods=["POST"])
@_json_error_guard
def api_rank_location():
    """Where a rank check WOULD be measured, without spending a lookup. Lets
    the Step 3 panel state the location even for a run served entirely from
    cache or restored from a saved quote."""
    d = request.get_json(force=True) or {}
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    markets = measure_first(markets, state, d.get("primary_market"))
    # The panel must name the place the ROWS were measured in, not the place the
    # operator typed -- otherwise "Measured in Whatcom County" sits above a table
    # of Bellingham keywords and the two read as a contradiction, which they are.
    markets = rank_markets(d.get("keywords") or [], markets, state)
    nat, reason = resolve_national_demand(d.get("industry") or "",
                                          d.get("geo_scope") or d.get("band") or "",
                                          bool(d.get("national_demand")),
                                          markets=markets,
                                          goal=(d.get("goal") or ""))
    out = rank_location_note(markets, state, nat)
    out["national_demand"] = nat
    out["national_demand_reason"] = reason
    return jsonify(out)


# (kw, location, domain, top_n) -> (pos, ts). In-memory: 1 gunicorn worker,
# so every request sees it; restarts just mean a cold cache. TTL keeps a
# calibration session fast without ever serving stale-day rankings.
RANK_CACHE = {}
RANK_CACHE_TTL = 6 * 3600
RANK_CACHE_MAX = 8000
_rank_cache_lock = threading.Lock()

def _rank_cache_get(kw, loc, dom, top_n):
    """Cached (position, page-one domains), or "MISS".

    THE DOMAINS ARE CACHED TOO, AND THAT IS NOT COSMETIC. This stored only the
    position, so a re-run inside the 6-hour TTL contributed no incumbents from
    any cached row — page one was measured from whatever happened to be
    re-fetched. Harmless while the signal was reported only; not harmless now
    that the page-one band carries $390 of partner cost, because it made the
    PRICE depend on cache state. Amare re-priced with every incumbent showing
    "1 of your terms" off fifteen cached rows, which also defeated the
    two-or-more-terms guard on pageone_strength and dropped it onto the bare
    maximum — the exact fragile path that guard exists to prevent. (2026-08-18)
    """
    with _rank_cache_lock:
        ent = RANK_CACHE.get((kw, loc, dom, top_n))
    if not ent:
        return "MISS", []
    # Entries written before the domains were cached are 2-tuples. A redeploy
    # clears the dict, but a config change inside a live process does not.
    pos, doms, stamp = (ent if len(ent) == 3 else (ent[0], [], ent[1]))
    if time.time() - stamp < RANK_CACHE_TTL:
        return pos, list(doms or [])
    return "MISS", []

def _rank_cache_put(kw, loc, dom, top_n, pos, doms=None):
    with _rank_cache_lock:
        if len(RANK_CACHE) > RANK_CACHE_MAX:
            RANK_CACHE.clear()
        RANK_CACHE[(kw, loc, dom, top_n)] = (pos, list(doms or []), time.time())

def fetch_domain_authority(domains):
    """Backlink rank (0-1000) for a batch of domains, in ONE request.

    The authority gap is what the industry actually scopes on. Linkscope's 2026
    benchmarks put local-services page one at DA 28 average with the top decile
    at 45+, and translate a gap straight into work: DA 20->40 is 40-80 referring
    domains over 12-18 months, DA 40->60 is 100-300+ over 18-36. That is a
    duration and a link volume, which is what a retainer ladder actually sells —
    and nothing in this tool has ever looked at it.

    `backlinks/bulk_ranks/live` takes up to 1000 targets per call, so the client
    and every incumbent on their page one cost one request between them. Returns
    {domain: rank} and {} on any failure — no key, no network, bad JSON — so a
    dead call reports nothing rather than claiming everyone is weak.
    (2026-08-17)
    """
    doms = []
    for d in domains or []:
        d = str(d or "").strip().lower().replace("https://", "").replace("http://", "")
        d = d.replace("www.", "").strip("/")
        if d and d not in doms:
            doms.append(d)
    if not doms:
        return {}, ""
    try:
        data = dfs_post("/backlinks/bulk_ranks/live",
                        [{"targets": doms[:1000]}], timeout=25)
        task0 = ((data or {}).get("tasks") or [{}])[0] or {}
        if task0.get("status_code") not in (20000, None):
            return {}, f"{task0.get('status_code')}: {task0.get('status_message')}"
        out = {}
        for block in (task0.get("result") or []):
            for it in (block.get("items") or []):
                t = str(it.get("target") or "").lower().replace("www.", "")
                if t:
                    out[t] = int(it.get("rank") or 0)
        # A ZERO IS NOT A MEASUREMENT OF WEAKNESS. bulk_ranks answers for every
        # target it was asked about and writes 0 for one it has no backlink data
        # on, so "no authority" and "not in the index" arrive as the same number.
        # That is survivable for an incumbent — a page-one domain with no
        # backlink data is genuinely weak — but not for the CLIENT, whose rank
        # is the subtrahend in the gap. Amare read 0 against a median incumbent
        # of 565, and a gap of 565 computed against an unmeasured client is not
        # comparable to Nob Hill's 35 computed against a measured 91. Callers
        # get the zeros AND the list, so the gap can decline to be computed
        # rather than quietly inflate. (2026-08-17)
        return out, ""
    except Exception as e:                                    # noqa: BLE001
        app.logger.exception("bulk_ranks failed")
        return {}, str(e)[:120]


# The on-page checks worth counting as debt. DataForSEO returns ~60 booleans and
# most are advisory; these are the ones that cost real remediation hours and that
# every pricing guide names when it says "current site condition".
_ONPAGE_DEBT = (
    ("is_https", False, "not on HTTPS"),
    ("no_h1_tag", True, "no H1"),
    ("title_too_long", True, "title too long"),
    ("no_title", True, "missing title"),
    ("no_description", True, "missing meta description"),
    ("duplicate_title_tag", True, "duplicate title"),
    ("duplicate_meta_tags", True, "duplicate meta"),
    ("low_content_rate", True, "thin content"),
    ("large_page_size", True, "heavy page"),
    ("high_loading_time", True, "slow load"),
    ("has_render_blocking_resources", True, "render-blocking assets"),
    ("no_image_alt", True, "images without alt text"),
    ("broken_links", True, "broken links"),
    ("is_4xx_code", True, "4xx errors"),
    ("is_5xx_code", True, "5xx errors"),
    ("no_favicon", True, "no favicon"),
    ("canonical_to_broken", True, "canonical to a broken URL"),
    ("has_meta_refresh_redirect", True, "meta-refresh redirect"),
)


def _instant_pages(dom):
    """One instant_pages call that survives a parameter this endpoint rejects.

    I added `accept_language` on the strength of a docs summary and DataForSEO
    answered `40501: Invalid Field: 'accept_language'` — so the site check failed
    for EVERY client, not just the blocked one it was meant to rescue. Their docs
    render client-side and could not be read to check, and guessing a second time
    is the same move that caused this.

    So the extras are optional by construction: on an Invalid Field reply the
    named field is dropped and the call is retried, down to the bare URL, which
    is the request that worked before any of this. A parameter that helps is
    kept; one this endpoint has never heard of costs a retry instead of a
    reading. Returns the task dict, or an error string. (2026-08-18)
    """
    extras = {"custom_user_agent": CFG.get("onpage_user_agent"),
              "browser_preset": "desktop"}
    last = ""
    for _ in range(len(extras) + 1):
        payload = dict(extras)
        payload["url"] = f"https://{dom}"
        data = dfs_post("/on_page/instant_pages", [payload], timeout=25)
        task0 = ((data or {}).get("tasks") or [{}])[0] or {}
        code = task0.get("status_code")
        if code in (20000, None):
            return task0
        msg = str(task0.get("status_message") or "")
        last = f"{code}: {msg}"
        bad = re.search(r"Invalid Field:\s*'?\"?([a-z_]+)", msg, re.I)
        if not (bad and bad.group(1) in extras):
            return last
        app.logger.warning("instant_pages rejected %s — retrying without it",
                           bad.group(1))
        extras.pop(bad.group(1))
    return last or "instant_pages refused every parameter set"


def fetch_technical_health(domain):
    """The client site's own condition, from a call the tool already makes.

    `on_page/instant_pages` has been in this codebase purely as a fallback for
    scraping page TITLES when the sitemap is unreadable — the response's
    `onpage_score` and its ~60 `checks` were received and discarded on every
    build. Every pricing guide names site condition as an upfront cost driver
    ("technical debt, broken links and thin content increase upfront costs") and
    it is the one input here that needs no new judgement: the API returns a score
    out of 100 and a list of what failed.

    Returns {"score": 0-100, "failed": [labels], "checked": n} and {} on failure.
    (2026-08-17)
    """
    dom = str(domain or "").strip().lower()
    dom = dom.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    if not dom:
        return {}, ""
    try:
        # NOT AS "RSiteAuditor". instant_pages defaults its user agent to
        # `Mozilla/5.0 (compatible; RSiteAuditor)` — a string that announces
        # itself as a bot to every WAF on the internet. Amare Homes came back
        # with an empty page and the panel printed "0/100 · nothing flagged" for
        # a site nobody had read. Ask as a browser on the call we already pay
        # for. (2026-08-17)
        task0 = _instant_pages(dom)
        if isinstance(task0, str):
            return {}, task0
        item = None
        for block in (task0.get("result") or []):
            for it in (block.get("items") or []):
                item = it
                break
            if item:
                break
        if not item:
            return _lighthouse_health(dom, "no page returned")
        checks = item.get("checks") or {}
        # A SITE THAT REFUSED US IS NOT A SITE SCORING ZERO. With no checks
        # there is nothing to score, and `onpage_score or 0` turned that into
        # 0.0/100 with an empty failure list — the worst possible reading,
        # printed as though it were a measurement, for the one client in three
        # whose site blocks crawlers. The information to tell them apart was
        # already here: len(checks). (2026-08-17)
        if not checks:
            return _lighthouse_health(dom, "the page returned nothing to check")
        failed = [label for key, bad_when, label in _ONPAGE_DEBT
                  if key in checks and bool(checks[key]) is bool(bad_when)]
        score = round(float(item.get("onpage_score") or 0), 1)

        # ZERO OUT OF A HUNDRED WITH NOTHING WRONG IS A CONTRADICTION, and the
        # empty-checks guard above does not catch it: for a page it could not
        # fetch, DataForSEO returns a FULL checks object with everything at its
        # default, so `checks` is populated, every check reads as passing, and
        # the score comes back 0. Amare printed "0/100 · nothing flagged" again
        # on the build after that guard shipped.
        #
        # No new API knowledge needed to spot it — the two halves of the answer
        # disagree. A page that genuinely scores 0 has failures; a page with no
        # failures does not score 0. An HTTP error on the item says the same
        # thing more directly, when it is there to read. (2026-08-18)
        _http = item.get("status_code")
        _dead = (isinstance(_http, int) and _http >= 400) or (
            score <= 0 and not failed)
        if _dead:
            return _lighthouse_health(
                dom, (f"their site answered HTTP {_http}" if isinstance(_http, int)
                      and _http >= 400 else
                      "the page scored 0 with nothing flagged, which means it was "
                      "not read"))
        return {"score": score,
                "failed": failed, "checked": len(checks), "source": "on_page",
                "timing": (item.get("page_timing") or {}).get("dom_complete")}, ""
    except Exception as e:                                    # noqa: BLE001
        app.logger.exception("instant_pages health failed")
        return {}, str(e)[:120]


# Lighthouse audit ids that mean the same thing as the _ONPAGE_DEBT checks, so a
# site read through the fallback reports in the same vocabulary as one read
# directly. Google's audit set is not the same shape — there is no "no H1" audit
# — so this is deliberately the intersection rather than a translation.
_LH_DEBT = (
    ("is-on-https", "not on HTTPS"),
    ("document-title", "missing title"),
    ("meta-description", "missing meta description"),
    ("image-alt", "images without alt text"),
    ("crawlable-anchors", "links Google cannot follow"),
    ("hreflang", "broken hreflang"),
    ("http-status-code", "page returns an error status"),
    ("viewport", "no mobile viewport"),
)


def _lighthouse_health(dom, why):
    """Second opinion on a site instant_pages could not read.

    Lighthouse fetches as Chrome rather than as an auditor, so it clears the
    naive 403 that the default user agent trips, and DataForSEO wraps Google's
    own project — same credentials, same dfs_post, one extra call. Its audits
    cover meta descriptions, titles, hreflang, crawlable links and image alt
    text, which is most of what _ONPAGE_DEBT counts.

    OFF BY DEFAULT AND ONLY ON FAILURE: this is a second request, and doubling
    the cost of every quote to rescue the occasional blocked site is a bad
    trade. Returns ({}, why) unchanged when it is off or when it fails too, so
    the caller still learns that nothing was measured. (2026-08-17)
    """
    if not CFG.get("technical_health_fallback"):
        return {}, why
    try:
        # SAME LESSON AS instant_pages, AND I DID NOT APPLY IT HERE. That
        # endpoint answered `40501: Invalid Field` to a parameter I had taken
        # from a docs summary, so the extras there are optional by construction
        # — dropped by name and retried on rejection. This call was written the
        # same afternoon with the same unverifiable docs and NO such retry, so
        # `lighthouse 40501` is the fallback refusing the request rather than
        # the client's site refusing us: the rescue never ran, and the line
        # blamed the site for it. (2026-08-19)
        extras = {"for_mobile": False,
                  "categories": ["performance", "seo", "accessibility",
                                 "best-practices"]}
        task0, code = {}, None
        for _ in range(len(extras) + 1):
            payload = dict(extras)
            payload["url"] = f"https://{dom}"
            data = dfs_post("/on_page/lighthouse/live/json", [payload],
                            timeout=40)
            task0 = ((data or {}).get("tasks") or [{}])[0] or {}
            code = task0.get("status_code")
            if code in (20000, None):
                break
            msg = str(task0.get("status_message") or "")
            bad = re.search(r"Invalid Field:\s*'?\"?([a-z_]+)", msg, re.I)
            if not (bad and bad.group(1) in extras):
                break
            app.logger.warning("lighthouse rejected %s — retrying without it",
                               bad.group(1))
            extras.pop(bad.group(1))
        if code not in (20000, None):
            return {}, f"{why}; lighthouse {code}"
        block = ((task0.get("result") or [{}])[0] or {})
        # DataForSEO wraps Google's report, and the docs render client-side so the
        # exact envelope could not be read before shipping. Both shapes are
        # accepted rather than guessed at: a wrong guess here fails silently, and
        # a silent failure in a fallback is worse than no fallback.
        block = block.get("lighthouse_result") or block
        audits = (block.get("audits") or {})
        cats = (block.get("categories") or {})
        if not audits and not cats:
            return {}, f"{why}; lighthouse returned nothing either"

        # AND THE SCALE IS NOT ASSUMED EITHER. Raw Lighthouse scores 0-1;
        # DataForSEO's own page advertises "1-100 scores". Reading a 0-100 score
        # as 0-1 would mark every passing audit as failed — a clean site would
        # come back as the most broken in the book, which is precisely the class
        # of error this whole change exists to stop. Anything above 1 is a
        # percentage. (2026-08-17)
        def _pct(v):
            if not isinstance(v, (int, float)):
                return None
            v = float(v)
            return v if v > 1 else v * 100

        failed = []
        for aid, label in _LH_DEBT:
            v = _pct((audits.get(aid) or {}).get("score"))
            if v is not None and v < 100:
                failed.append(label)
        seo = _pct((cats.get("seo") or {}).get("score"))
        return {"score": round(seo, 1) if seo is not None else None,
                "failed": failed, "checked": len(audits), "source": "lighthouse",
                "timing": None}, ""
    except Exception as e:                                    # noqa: BLE001
        app.logger.exception("lighthouse health failed")
        return {}, f"{why}; lighthouse {str(e)[:60]}"


def serp_top_domains(kw, loc, client_dom="", top_n=5, deadline=None):
    """The first few organic domains for one query, plus where the client sits.

    The collision check can say a bare acronym's volume looks wrong; only the
    result page can say whose it is. Cheapest possible answer: one SERP, the top
    handful of domains, and whether the client appears at all. If LACP returns
    cisco.com and juniper.net, the operator has their answer without leaving the
    page. (2026-08-10)
    """
    depth = max(int(CFG.get("zero_ranking_top_n", 100)), 10)
    payload = [{"keyword": kw, "location_name": loc, "language_code": "en",
                "depth": depth}]
    # Retry once. A single 14s read timeout left LACP — the biggest number in the
    # list and the whole reason the check exists — as "lookup failed", with the
    # operator's most important question unanswered. (2026-08-10)
    last = None
    for attempt in range(2):
        remaining = (deadline - time.time()) if deadline else 20
        if remaining < 5:
            raise last or TimeoutError("acronym SERP budget exhausted")
        try:
            data = dfs_post("/serp/google/organic/live/regular", payload,
                            timeout=min(14 if attempt == 0 else remaining - 1,
                                        remaining, 20))
            break
        except Exception as e:
            last = e
    else:
        raise last
    task0 = ((data or {}).get("tasks") or [{}])[0] or {}
    if task0.get("status_code") not in (20000, None):
        raise RuntimeError(f"{task0.get('status_code')}: {task0.get('status_message')}")
    items = ((task0.get("result") or [{}])[0] or {}).get("items") or []
    doms, client_pos = [], None
    for it in items:
        if it.get("type") != "organic":
            continue
        d = (it.get("domain") or "").lower().replace("www.", "")
        if client_dom and client_dom in d and client_pos is None:
            client_pos = it.get("rank_absolute")
        if d and d not in doms and len(doms) < top_n:
            doms.append(d)
    return doms, client_pos


def claude_replacements(removed, kept_seeds, brand="", domain="",
                        business_desc="", topic="", n=6):
    """Alternative phrasings for a service whose term had to be deleted.

    Removing a colliding acronym can leave a service uncovered — "lacp" goes and
    nothing in the list speaks to lateral assessment any more. Rebuilding refills
    the slot, but that re-runs the whole step for one gap. This proposes named
    candidates instead, inside the same topic, which are then MEASURED before any
    of them is offered: a suggestion with no search volume is the exact mistake
    that produced NASSCO's first keyword list. (2026-08-10)
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or not removed:
        return []
    prompt = f"""A keyword list for {brand or domain or "a business"} had to drop the term "{removed}".
{f'The business describes itself as: {business_desc}' if business_desc else ''}
{f'That term belonged to the topic: {topic}' if topic else ''}
Its other terms are: {", ".join(str(x) for x in (kept_seeds or [])[:25])}

Propose {n} SHORT search phrases that cover the same service as "{removed}" but that
real buyers type into Google. Rules:
1. Phrases people SEARCH, not descriptions of the organisation. "manhole inspection
   certification" not "manhole assessment certification program provider".
2. 2-5 words. No city names, no brand names other than this client's own.
3. Stay on the same service as the removed term. Do not drift to its siblings.
4. Prefer the words an outsider would use over the client's internal naming.

Return ONLY JSON: {{"terms": ["...", "..."]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 600, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=25)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        terms = json.loads(text).get("terms") or []
    except Exception:
        app.logger.exception("claude_replacements failed")
        return []
    out, seen = [], {str(removed).strip().lower()}
    for t in terms:
        t = clean_kw(strip_placeholders(str(t).strip().lower()))
        if t and t not in seen and 1 < len(t.split()) <= 6:
            seen.add(t)
            out.append(t)
    return out[:n]


@app.route("/api/replacement_terms", methods=["POST"])
@_json_error_guard
def api_replacement_terms():
    """Measured candidates to fill the slot a deleted term left behind."""
    d = request.get_json(force=True) or {}
    removed = (d.get("removed") or "").strip()
    if not removed:
        return jsonify({"terms": []})
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    nat = bool(d.get("national_demand")) or not markets
    cands = claude_replacements(removed,
                               [x for x in (d.get("seeds") or []) if x],
                               d.get("brand") or "", d.get("domain") or "",
                               d.get("business_desc") or "",
                               d.get("topic") or "")
    if not cands:
        return jsonify({"terms": [], "error": "no candidates proposed"})
    # MEASURE before offering. Same rule the acronym chips follow.
    try:
        vols, _pc, verr = fetch_local_volume(cands, [] if nat else markets, state,
                                             national=nat)
    except Exception as e:
        return jsonify({"terms": [], "error": str(e)[:120]})
    floor = int(CFG.get("replacement_min_volume", 20))
    rows = [{"term": t, "volume": int((vols or {}).get(t, 0) or 0)} for t in cands]
    rows.sort(key=lambda r: -r["volume"])
    return jsonify({"terms": [r for r in rows if r["volume"] >= floor],
                    "rejected": [r for r in rows if r["volume"] < floor],
                    "basis": "US national" if nat else "targeted cities",
                    "error": verr})


@app.route("/api/acronym_serp", methods=["POST"])
@_json_error_guard
def api_acronym_serp():
    """Who actually owns each flagged acronym's result page."""
    d = request.get_json(force=True) or {}
    terms = [str(t).strip() for t in (d.get("terms") or []) if str(t).strip()][:8]
    if not terms:
        return jsonify({"results": []})
    dom = (d.get("domain") or "").replace("https://", "").replace("http://", "")
    dom = dom.replace("www.", "").strip("/")
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    nat, _r = resolve_national_demand(d.get("industry") or "",
                                      d.get("geo_scope") or "",
                                      bool(d.get("national_demand")),
                                      markets=markets,
                                      goal=(d.get("goal") or ""))
    # An acronym's ownership is a national question — "who does Google think this
    # word belongs to" — so it is asked nationally even on a local quote.
    loc = "United States"
    _ = (nat, state, markets)
    out = []
    budget = time.time() + max(18, REQUEST_BUDGET_S - 12)
    with ThreadPoolExecutor(max_workers=min(4, len(terms))) as ex:
        futs = {ex.submit(serp_top_domains, t, loc, dom, 5, budget): t
                for t in terms}
        for fut in futs:
            t = futs[fut]
            try:
                doms, pos = fut.result()
                out.append({"term": t, "domains": doms, "client_pos": pos})
            except Exception as e:
                out.append({"term": t, "domains": [], "client_pos": None,
                            "error": str(e)[:120]})
    order = {t: i for i, t in enumerate(terms)}
    out.sort(key=lambda r: order.get(r["term"], 99))
    return jsonify({"results": out, "location": loc, "client_domain": dom})


@app.route("/api/rankings", methods=["POST"])
@_json_error_guard
def api_rankings():
    """Step 3 — rank-check ONE small batch of keywords (frontend loops batches).
    Each call is short: a few parallel SERP lookups."""
    d = request.get_json(force=True)
    batch   = d.get("batch", [])
    domain  = (d.get("domain") or "").strip()
    markets = usable_markets(d.get("geo_values") or [])
    state   = derive_state(markets, (d.get("state") or "").strip())
    brand   = (d.get("brand") or "").strip()
    dom = domain.replace("https://", "").replace("http://", "").replace("www.", "").strip("/")
    markets = measure_first(markets, state, d.get("primary_market"))
    markets = rank_markets([(x.get("kw") if isinstance(x, dict) else x)
                            for x in (batch or [])], markets, state)
    top_n = CFG["zero_ranking_top_n"]
    nat, _r = resolve_national_demand(d.get("industry") or "",
                                      d.get("geo_scope") or d.get("band") or "",
                                      bool(d.get("national_demand")),
                                      markets=markets,
                                      goal=(d.get("goal") or ""))
    loc = rank_location(markets, state, nat)
    # MEASURE EACH KEYWORD IN THE MARKET IT NAMES.
    #
    # Every row used to be checked from the primary market, so Knoxville's SERP
    # answered for "junk removal morristown tn" and "junk removal sevierville tn"
    # too. That produced 32/32 ranked and, through recommend_addons, "0 add-on
    # markets recommended" — a claim that the client already has a presence in
    # all six markets, from a test that only ever looked at one of them. The
    # cheap version of the question is the wrong question. Same API cost either
    # way: one call per keyword, just pointed at the right place. (2026-08-10)
    #
    # `batch` accepts plain strings (older clients) or {kw, city} objects.
    per_kw_loc, order = {}, []
    for item in batch:
        if isinstance(item, dict):
            kw = str(item.get("kw") or "").strip()
            city = str(item.get("city") or "").strip()
        else:
            kw, city = str(item or "").strip(), ""
        if not kw:
            continue
        order.append(kw)
        # A national row, a site term or a long-tail top-up carries no market;
        # those stay on the primary one, which is what they are asking about.
        mk_named = ""
        if not nat:
            if city:
                mk_named = next((m for m in markets
                                 if parse_market(m, state)[0].strip().lower()
                                 == city.lower()), city)
            else:
                # No tag on the row — read the market off the keyword text.
                mk_named = market_for_keyword(kw, markets, state)
        per_kw_loc[kw] = (rank_location([mk_named], state, False)
                          if mk_named else loc)
    batch = order
    results, paa = [], []
    rivals = {}
    hits = {}
    to_fetch = []
    err_msgs = []
    for kw in batch:
        c, cdoms = _rank_cache_get(kw, per_kw_loc.get(kw, loc), dom, top_n)
        if c != "MISS":
            hits[kw] = c
            # A cached row still tells you who holds page one.
            for _d in cdoms:
                rivals[_d] = rivals.get(_d, 0) + 1
        else:
            to_fetch.append(kw)
    try:
        with ThreadPoolExecutor(max_workers=CFG["rank_check_workers"]) as ex:
            _budget = int(CFG.get("rank_batch_budget_s") or 0) or max(20, REQUEST_BUDGET_S - 15)
            batch_deadline = time.time() + _budget
            futs = {ex.submit(_serp_one, kw, dom, markets, state, brand, top_n,
                              batch_deadline,
                              per_kw_loc.get(kw, loc)): kw for kw in to_fetch}
            done = {}
            for fut in futs:
                kw = futs[fut]
                try:
                    pos, qs, doms = fut.result()
                    err = False
                    # Page one, per keyword, tallied across the batch — see
                    # _serp_parse_items. This endpoint is the one that runs when
                    # Google answers fast enough to skip the queue, which is how
                    # Nob Hill reported "0 incumbents measured" with its rivals
                    # sitting in the response. (2026-08-17)
                    for _d in (doms or []):
                        rivals[_d] = rivals.get(_d, 0) + 1
                except Exception as e:
                    # lookup FAILED — record it as unknown, NOT as "Not Found".
                    # Counting a failed call as not-ranking would inflate the
                    # zero-ranking percentage and therefore the price.
                    pos, qs, err = None, [], True
                    # Keep WHY. "40501 location_name not found" and "read
                    # timeout" both show as "check failed", and they call for
                    # opposite responses: fix the geo, or press retry.
                    err_msgs.append(str(e)[:160])
                done[kw] = (pos, qs, err)
                if not err:
                    _rank_cache_put(kw, per_kw_loc.get(kw, loc), dom, top_n, pos,
                                    doms)
        for kw in batch:
            if kw in hits:
                pos, qs, err = hits[kw], [], False
            else:
                pos, qs, err = done.get(kw, (None, [], True))
            results.append({"kw": kw,
                            "pos": ("—" if err else (pos if pos is not None else "Not Found")),
                            "ranked_top": (not err and pos is not None and pos <= top_n),
                            "error": err,
                            # A batch that answers instantly is a batch that
                            # never called Google. Say so, rather than leaving
                            # the operator to wonder whether the check ran.
                            "cached": kw in hits})
            paa.extend(qs)
    except requests.HTTPError as e:
        return jsonify({"error": f"DataForSEO error: {e}."}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}"}), 500
    # The most common failure reason, verbatim. A batch of "40501 Invalid Field:
    # 'location_name'" means the geo is wrong and retrying will fail identically;
    # a batch of timeouts means press retry. Same "check failed" row either way,
    # so the message has to reach the panel.
    top_err = ""
    if err_msgs:
        top_err = max(set(err_msgs), key=err_msgs.count)
    note = rank_location_note(markets, state, nat)
    # Say which markets were actually measured. "Measured in Knoxville" while
    # half the rows were checked in Morristown and Sevierville would be worse
    # than saying nothing.
    locs = sorted({v for v in per_kw_loc.values()})
    if len(locs) > 1:
        pretty = [l.replace(",United States", "").replace(",", ", ") for l in locs]
        note = {"location": ", ".join(pretty), "scope": "per_market",
                "note": "Each keyword was measured in the market it names — "
                        + ", ".join(pretty)
                        + ". Checking every row from one city answers a different "
                          "question, and the add-on recommendation reads presence "
                          "per market off this table."}
    for r in results:
        r["loc"] = (per_kw_loc.get(r["kw"], loc)
                    .replace(",United States", "").replace(",", ", "))
    return jsonify({"results": results, "paa": list(dict.fromkeys(paa)),
                    "n_cached": len(hits), "n_fetched": len(to_fetch),
                    "n_errors": len(err_msgs), "error_reason": top_err,
                    "rank_location": note,
                    "rivals": [{"domain": d, "appearances": n}
                               for d, n in sorted(rivals.items(),
                                                  key=lambda kv: (-kv[1], kv[0]))
                               [:int(CFG.get("serp_rival_cap", 12))]]})

@app.route("/api/qualify_markets", methods=["POST"])
@_json_error_guard
def api_qualify_markets():
    """Attach the right state to each bare market name.

    A report's column headers read "Wayne", "Paramus", "New York City" with no
    states. The importer used to append the report's own state to all of them,
    which produced "New York City, NJ" — a market that does not exist, pointing
    the rank check and the CPC lookup at nothing (2026-08-07).

    The ZIP index already knows. Precedence:
      1. the market already carries a state -> leave it alone
      2. the city name resolves to exactly ONE state -> use it, whatever the
         report said (this is what fixes New York City)
      3. several candidates and the report's state is one of them -> report's
      4. several candidates, report's state is NOT one -> report's, flagged
      5. no candidates at all -> report's if given, else flagged as unresolved

    No external calls; runs off bundled ZIP data.
    """
    d = request.get_json(force=True) or {}
    fallback = (d.get("state") or "").strip()
    fb_abbr = (STATE_ABBREV.get(fallback.lower(), "") or
               (fallback if len(fallback) == 2 else "")).upper()
    idx = _zip_index()
    by_city = {}
    for (city, st) in idx:
        by_city.setdefault(city, set()).add(st)

    out = []
    dropped = []
    for raw in (d.get("markets") or []):
        m = str(raw or "").strip()
        if not m:
            continue
        # A ranking table's market column carries "near me" rows, because the
        # agency tracked "junk removal near me" as its own line. Qualifying it
        # produced "near me, TN" — which then became markets[0] and therefore
        # the location_name on every SERP call. Never offer it. (2026-08-10)
        if is_non_place_geo(m):
            dropped.append(m)
            continue
        already = re.search(r",\s*([A-Za-z]{2})\s*$", m)
        if already:
            out.append({"input": m, "qualified": m,
                        "abbr": already.group(1).upper(),
                        "source": "given", "certain": True})
            continue
        city = m.lower()
        # CITY_STATE is the curated metro map and outranks the ZIP scan.
        curated = CITY_STATE.get(city, "")
        cands = sorted(by_city.get(city, set()))
        if curated:
            # STATE_ABBREV stores lowercase; a pill reads "Philadelphia, PA".
            ab = (STATE_ABBREV.get(curated.lower(), "") or "").upper()
            out.append({"input": m, "qualified": f"{m}, {ab}" if ab else m,
                        "abbr": ab, "source": "known metro", "certain": True})
        elif len(cands) == 1:
            out.append({"input": m, "qualified": f"{m}, {cands[0]}",
                        "abbr": cands[0], "source": "only one state has it",
                        "certain": True})
        elif cands and fb_abbr in cands:
            out.append({"input": m, "qualified": f"{m}, {fb_abbr}",
                        "abbr": fb_abbr, "source": "report's state",
                        "certain": True})
        elif cands and fb_abbr:
            out.append({"input": m, "qualified": f"{m}, {fb_abbr}",
                        "abbr": fb_abbr, "source": "report's state",
                        "certain": False,
                        "note": f"{m} exists in {', '.join(cands[:6])} but not "
                                f"{fb_abbr} — check this one."})
        elif fb_abbr:
            out.append({"input": m, "qualified": f"{m}, {fb_abbr}",
                        "abbr": fb_abbr, "source": "report's state",
                        "certain": False,
                        "note": f"{m} isn't in the city database — using the "
                                f"report's state, {fb_abbr}."})
        else:
            out.append({"input": m, "qualified": m, "abbr": "",
                        "source": "", "certain": False,
                        "note": f"No state for {m}, and the report didn't say."})
    return jsonify({"markets": out, "dropped": dropped})


@app.route("/api/build")
def api_build():
    """What is actually on disk, file by file. Read by the header's src chip."""
    return jsonify({"build": BUILD_STR, "source_fingerprint": SOURCE_FP,
                    "files": source_file_hashes()})


@app.route("/api/markets", methods=["POST"])
@_json_error_guard
def api_markets():
    """How many distinct MARKETS the entered geos actually cover.

    Separate question from add-on markets, and asked much earlier: this is
    "what does this footprint amount to", which decides the grid shape, the
    coverage percentages and the market count everything else is built on. It
    runs off bundled coordinates, so there is no API call and it can answer
    the moment the geos are typed rather than after a full build.
    """
    d = request.get_json(force=True) or {}
    entered = [m for m in (d.get("geo_values") or []) if m and m.strip()]
    state = (d.get("state") or "").strip()
    # National demand has FOUR sources — the manual switch, a nationwide band,
    # the campaign goal, and the RZ industry tag — and the band recommendation
    # read only the switch. An ecommerce-tagged client with no markets was priced
    # nationally by the industry rule and never got told to fix the band, which is
    # the case where the anchor is wrong and nobody ticked anything. Ask the same
    # function the pricing pipeline asks. (2026-08-10)
    _nat, _nat_why = resolve_national_demand(
        d.get("industry") or "",
        d.get("geo_scope") or d.get("band") or "",
        bool(d.get("national_demand")),
        markets=[m for m in entered if not is_non_place_geo(m)],
        goal=(d.get("goal") or ""))
    if not entered:
        # An empty geo list is exactly when the band recommendation matters most:
        # nothing is left to describe, yet the dropdown may still read
        # "Non-contiguous region" and that is what sets the pricing anchor. The
        # early return skipped the suggestion entirely. (2026-08-10)
        return jsonify({"cities": 0, "entered": 0, "markets": 0, "groups": [],
                        "unlocated": [], "non_place": [], "state_geos": [],
                        "overlaps": [], "covered": [],
                        "radius": int(CFG.get("market_radius_miles", 25)),
                        "located": 0,
                        "scope_suggestion": suggest_geo_scope(
                            [], state, _nat, _nat_why)})
    # Non-places cover nothing, so they cannot be markets. Reported back so the
    # operator can see WHY the count moved rather than watching a pill silently
    # stop mattering.
    non_place = [m for m in entered if is_non_place_geo(m)]
    # A whole state is not a market either. Reported separately from non-places
    # because the fix is different: a state usually means the operator wants
    # every city in it, which is a Statewide scope, not a pill.
    state_geos = [m for m in entered if m not in non_place and is_state_geo(m)]
    mk = [m for m in entered if m not in non_place and m not in state_geos]
    overlaps = geo_overlaps(entered, state)
    covered = sorted({c for o in overlaps if o["kind"] in ("county", "duplicate")
                      for c in o["contained"]})
    if not mk:
        return jsonify({"cities": len(entered), "markets": 0, "groups": [],
                        "unlocated": [], "non_place": non_place,
                        "state_geos": state_geos,
                        "overlaps": overlaps, "covered": covered,
                        "radius": int(CFG.get("market_radius_miles", 25)),
                        "located": 0, "scope_suggestion": {}})
    groups, located, unlocated = group_by_distance(mk, state)
    named = []
    for g in sorted(groups, key=len, reverse=True):
        # Prefer a real town over a county as the market's name: the anchor is
        # what the keyword grid and the rank check are built on, and nobody
        # searches "junk removal knox county tn".
        anchor = market_anchor(g, state)
        named.append({"anchor": anchor,
                      "members": [anchor] + [m for m in g if m != anchor],
                      "size": len(g)})
    return jsonify({
        "cities": len(mk),
        "entered": len(entered),
        "markets": len(groups) + len(unlocated),
        "groups": named,
        "unlocated": unlocated,
        "non_place": non_place,
        "state_geos": state_geos,
        "overlaps": overlaps,
        "covered": covered,
        "radius": int(CFG.get("market_radius_miles", 25)),
        "located": len(located),
        # The band the markets themselves imply. A suggestion, not an
        # assignment: the operator's dropdown still picks the pricing anchor.
        "scope_suggestion": suggest_geo_scope(mk, state, _nat, _nat_why),
        # How many service slots this many markets buys. Expand needs it to know
        # when to stop: 51 seeds for 20 slots is 31 terms nobody will ever quote.
        # (2026-08-13)
        "service_slots": services_needed(len(mk)),
        # THE CEILING, WHICH IS THE ONLY SAFE NUMBER TO CAP EXPAND ON.
        # service_slots above is computed on the markets ENTERED. The build
        # computes the same thing on the markets that SURVIVE the demand check,
        # and markets can only be dropped there, never added -- so the entered
        # count gives the FEWEST slots the grid can end up with, not the most.
        # Whatcom County: five markets entered -> 7 slots -> Expand skipped as
        # "full" against a typed list of 20; four of the five markets then
        # carried no demand, the grid crossed one city, and the same function
        # returned 20 slots. Filters removed six terms after that and nothing
        # refilled them, so a client with plenty of services was quoted on 14.
        # services_needed() only ever goes UP as cities are dropped, so capping
        # Expand here costs nothing: enforce_seed_services() trims to the real
        # slot count afterwards, and expanded terms sit at the end of the list,
        # which is the end it trims from. (2026-08-24)
        "expand_slots": int(CFG.get("grid_max_services", 20)),
    })


@app.route("/api/addon_suggestion", methods=["POST"])
@_json_error_guard
def api_addon_suggestion():
    """Suggest an add-on market count from the assembled rank table.

    Its own route because it needs the WHOLE table, which only exists on the
    frontend after step 3 has looped its batches. No external calls — pure
    arithmetic on data the quote already holds.
    """
    d = request.get_json(force=True) or {}
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    out = recommend_addons(markets, state, d.get("table") or [],
                           site_locations=d.get("site_locations") or [],
                           site_pages_found=d.get("site_pages_found"),
                           metro_groups=d.get("metro_groups") or [],
                           city_volumes=d.get("city_volumes") or {})
    out["gbp_locations"] = d.get("gbp_locations")
    # Surface HOW the markets were counted. Four rounds of this were spent
    # guessing which branch ran because nothing on screen said (2026-08-03).
    out["grouped_by"] = d.get("grouped_by") or ""
    out["city_locs"] = d.get("city_locs") or {}
    out["service_areas"] = len(d.get("service_areas") or [])
    return jsonify(out)


@app.route("/api/price", methods=["POST"])
@_json_error_guard
def api_price():
    """Step 4 — pure pricing math, instant. Returns hard cost + client (marked-up)."""
    d = request.get_json(force=True)
    band = d.get("band", "single_city")
    if band not in CFG["geo_anchor"]:
        return jsonify({"error": f"Unknown geo scope '{band}'."}), 400
    adder = int(d.get("adder", 0) or 0)
    zero  = bool(d.get("zero_ranking", False))
    addon = int(d.get("addon_markets", 0) or 0)
    markup = d.get("markup_pct", None)
    markup = float(markup) if markup not in (None, "") else None
    pct_not_ranking = d.get("pct_not_ranking", None)
    pct_not_ranking = float(pct_not_ranking) if pct_not_ranking not in (None, "") else None
    total_volume = d.get("total_volume", None)
    total_volume = int(total_volume) if total_volume not in (None, "") else None
    base_override = d.get("base_override", None)
    base_override = base_override if base_override not in (None, "") else None
    # The median incumbent's authority, measured by the rank check's own SERPs.
    # None when the check has not found anyone yet, which is a different state
    # from "nobody strong" and is carried through as such.
    pageone_rank = d.get("pageone_rank", None)
    pageone_rank = int(pageone_rank) if pageone_rank not in (None, "") else None
    p = stage4_price(band, adder, zero, addon, markup, pageone_rank=pageone_rank,
                     pct_not_ranking=pct_not_ranking, total_volume=total_volume,
                     base_override=base_override, ecommerce=bool(d.get("ecommerce")),
                     industry=(d.get("industry") or ""),
                     ai_search=bool(d.get("ai_search")),
                     national_demand=bool(d.get("national_demand")),
                     geo_override=d.get("geo_override"),
                     addon_override=d.get("addon_override"),
                     goal=(d.get("goal") or ""),
                     site_rebuild=(d.get("site_rebuild") or ""))
    return jsonify({"anchor": p["anchor"], "adder": adder,
                    "site_rebuild": p.get("site_rebuild", ""),
                    "rebuild_applied": p.get("rebuild_applied", False),
                    "national_demand": p.get("national_demand", False),
                    "national_demand_reason": p.get("national_demand_reason", ""),
                    "min_term_months": p.get("min_term_months"),
                    "extras_multiplier": p.get("extras_multiplier", 1.0),
                    "manual_geo": p.get("manual_geo", False),
                    "manual_addon": p.get("manual_addon", False),
                    "industry_rule": p.get("industry_rule"),
                    "industry_anchor_add": p.get("industry_anchor_add", 0),
                    "ai_search": p.get("ai_search"),
                    "base_pre_uplift": p["base_pre_uplift"], "manual_base": p["manual_base"],
                    "zero_ranking_uplift_pct": p["zero_ranking_uplift_pct"],
                    "volume_add": p["volume_add"],
                    "pct_not_ranking": p["pct_not_ranking"], "total_volume": p["total_volume"],
                    "base": p["base"], "step": p["step"],
                    # THE CHART IS A CLIENT-FACING BREAKDOWN AND IT WAS SHORT.
                    # competitive_adder never reached it, so Ooten's proposal
                    # printed "Competition —  bids and difficulty both below the
                    # first break" while $1,000 of adder sat inside the $5,450 it
                    # was breaking down. The rows did not add up to the total and
                    # nothing said so. (2026-08-22)
                    "competitive_adder": p.get("competitive_adder", 0),
                    "pageone_anchor_add": p.get("pageone_anchor_add", 0),
                    "pageone_band": p.get("pageone_band"),
                    "pageone_measured": p.get("pageone_measured", False),
                    "floored": p.get("floored", False),
                    "client_floor": p.get("client_floor"),
                    "hard_tiers": p["hard_tiers"], "client_tiers": p["client_tiers"],
                    "hard_addon_per_market": p["hard_addon_per_market"],
                    "client_addon_per_market": p["client_addon_per_market"],
                    "addon_discount_pct": p["addon_discount_pct"],
                    "addon_discount_basis": p["addon_discount_basis"],
                    "addon_discount_tiers": p["addon_discount_tiers"],
                    "addon_schedule": p["addon_schedule"],
                    "hard_addon_list_per_market": p["hard_addon_list_per_market"],
                    "client_addon_list_per_market": p["client_addon_list_per_market"],
                    "addon_savings_per_market": p["addon_savings_per_market"],
                    "margin_pct_of_gross": p["margin_pct_of_gross"],
                    "handoff": p.get("handoff", {}),
                    "markup_pct": p["markup_pct"], "addon_markets": addon, "band": band})

def _perf_fill_bids(d, eligible=True):
    """Give every quoted term a bid before the performance table prices it.

    Step 2 scores the HEAD terms only — that is all the competitive adder needs
    — so most rows reach this table with nothing to price on and fall to the
    floor. One Labs call covers up to a thousand keywords.

    SHARED WITH THE DOCUMENT BUILDER ON PURPOSE. This lived inside the API
    endpoint, so the panel saw filled bids and the .docx did not: the document
    priced the same table off head terms alone, came in under the $10,000
    minimum and dropped the whole section without a word. (2026-08-22)
    """
    backfill = {"asked": 0, "filled": 0, "error": ""}
    cpc = dict(d.get("cpc") or {})
    if eligible:
        _nk = lambda v: re.sub(r"\s+", " ", str(v or "").strip().lower())
        have = {_nk(k) for k, v in cpc.items() if v}
        _prows = _proposal_rows(d)
        _sfx = _nk(grid_suffix([r["kw"] for r in _prows]))
        want, bare_of = [], {}
        for r in _prows:
            k = _nk(r["kw"])
            if k in have:
                continue
            # The grid's own suffix first; the entered market only as a fallback,
            # because the two can name the same place differently.
            b = k[: -len(_sfx)].strip() if (_sfx and k.endswith(_sfx)) else ""
            if not b:
                b = _nk(_strip_markets(r["kw"],
                                       (d.get("inputs") or {}).get("geo_values")
                                       or d.get("markets") or [], d.get("state") or ""))
            bare_of[k] = b or k
            if (b or k) not in want:
                want.append(b or k)
        backfill["asked"] = len(want)
        if want:
            got, err = fetch_bids_via_labs(
                want, (d.get("inputs") or {}).get("geo_values") or d.get("markets") or [],
                d.get("state") or "", bool(d.get("national_demand")))
            backfill["error"] = err or ""
            for k, b in bare_of.items():
                hit = (got or {}).get(b) or {}
                v = hit.get("bid") or hit.get("cpc") or 0
                if v:
                    cpc[k] = v
                    backfill["filled"] += 1
        d = dict(d, cpc=cpc)
    return d, backfill


@app.route("/api/perf_quote", methods=["POST"])
@_json_error_guard
def api_perf_quote():
    """Pay-for-performance: is this client eligible, and what would it bill?"""
    d = request.get_json(force=True) or {}
    # Eligibility is judged on the GRID, not on terms added for this table —
    # the gate is about the client's existing position, and padding the list
    # with terms chosen for their bid would move it.
    elig = perf_eligibility(d.get("table") or [],
                            site_rebuild=(d.get("site_rebuild") or ""))
    d = _perf_merge_extra(d)
    # EVERY QUOTED TERM NEEDS A BID, NOT JUST THE HEAD TERMS. Step 2 scores the
    # head terms only — that is all the competitive adder needs — so on Ooten
    # eleven of twenty rows reached this table with no bid and fell to the $80
    # floor, and the potential value read $2,080 against Brendan's $25,540. The
    # floor is the right answer for a term nobody will quote a bid on; it is the
    # wrong answer for a term nobody ASKED about. One Labs call covers up to a
    # thousand keywords and returns the bid beside the volume. (2026-08-22)
    d, backfill = _perf_fill_bids(d, elig.get("eligible"))
    rows = _perf_rows(d) if elig.get("eligible") else []
    total = sum(r["page1"] for r in rows)
    floor = int(CFG.get("perf_min_monthly_value", 10000) or 0)
    return jsonify({
        "eligibility": elig,
        "rows": rows,
        "total_page1": total,
        "total_top5": sum(r["top5"] for r in rows),
        "total_top3": sum(r["top3"] for r in rows),
        "total_one": sum(r["one"] for r in rows),
        "minimum": floor,
        "meets_minimum": bool(rows) and total >= floor,
        "no_bids": len([r for r in rows if r.get("bid") in (None, 0)]),
        "backfill": backfill,
        # The column is blank on quotes built before topics carried their
        # members. Say so rather than leaving an empty column unexplained.
        "areas_known": bool([t for t in perf_topics(d) if t.get("terms")]),
        "no_area": len([r for r in rows if not r.get("area")]),
        "candidates": perf_candidates(d),
        "table_target": int(CFG.get("perf_table_terms", 50) or 50),
        "in_table": len(rows),
        "term_months": int(CFG.get("perf_initial_term_months", 12)),
        "tail_months": int(CFG.get("perf_tail_months", 6)),
    })


@app.route("/api/config", methods=["GET"])
@_json_error_guard
def api_config_get():
    """Expose the tunable pricing constants for the review panel."""
    return jsonify({
        "geo_anchor": CFG["geo_anchor"],
        "industry_pricing": CFG.get("industry_pricing", {}),
        # Operational knobs the panel reads LIVE. They also ride on the step-1
        # payload, but that copy is written when the keyword list is built and
        # restored verbatim forever after — so a quote built on Wednesday kept
        # Wednesday's call budget and every change to these was invisible until
        # someone rebuilt step 1. The panel prefers these. (2026-08-21)
        "min_unranked_terms": CFG.get("min_unranked_terms", 3),
        "unranked_probe_max": CFG.get("unranked_probe_max", 4),
        "dfs_calls_per_minute": CFG.get("dfs_calls_per_minute", 10),
        "competitive_adder": CFG["competitive_adder"],
        "bid_score_breaks": CFG["bid_score_breaks"],
        "cpc_adder_enabled": CFG.get("cpc_adder_enabled", True),
        "cpc_adder_mult": CFG.get("cpc_adder_mult", 3.0),
        "cpc_adder_cap": CFG.get("cpc_adder_cap", 1500),
        "cpc_adder_knee": CFG.get("cpc_adder_knee", 62.0),
        "cpc_adder_mult_high": CFG.get("cpc_adder_mult_high", 14.0),
        "tier_step_pct_of_base": CFG.get("tier_step_pct_of_base", 0.24),
        "ecom_anchor_add": CFG.get("ecom_anchor_add", 0),
        "geo_pricing_mode": CFG.get("geo_pricing_mode", "pct"),
        "geo_pct_tiers": CFG.get("geo_pct_tiers", []),
        "geo_pct_default": CFG.get("geo_pct_default", 60),
        "addon_volume_discount_tiers": CFG.get("addon_volume_discount_tiers", []),
        "perf_page_depth": CFG.get("perf_page_depth", 50),
        "perf_eligible_min_share": CFG.get("perf_eligible_min_share", 0.5),
        "perf_min_monthly_value": CFG.get("perf_min_monthly_value", 10000),
        "perf_page1_mult": CFG.get("perf_page1_mult", 2.1),
        "perf_page1_floor": CFG.get("perf_page1_floor", 80),
        "serp_frame_offset": CFG.get("serp_frame_offset", 0.40),
        "serp_head_px": CFG.get("serp_head_px", 190),
        "min_term_months": CFG.get("min_term_months", 6),
        "nationwide_service_extras": CFG.get("nationwide_service_extras", 1.0),
        "vol_add_ramp": CFG.get("vol_add_ramp", [40, 60]),
        # Shown so a model change is visible rather than silent: an unpinned
        # alias can move under you and quietly reshape every keyword list.
        # Pin it with the CLAUDE_MODEL env var.
        "claude_model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
        "claude_model_pinned": model_is_snapshot(os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6")),
        "source_fingerprint": SOURCE_FP,
        "dfs_timeout": DFS_TIMEOUT,
        "request_budget_s": REQUEST_BUDGET_S,
        "pin_head_terms": CFG.get("pin_head_terms", 3),
        "pin_min_volume": CFG.get("pin_min_volume", 300),
        "geo_card": CFG.get("geo_card", {}),
        "cpc_adder_free_below": CFG.get("cpc_adder_free_below", 5.0),
        "zero_ranking_bonus": CFG["zero_ranking_bonus"],
        "zero_ranking_top_n": CFG["zero_ranking_top_n"],
        "zero_ranking_frac": CFG["zero_ranking_frac"],
        "zero_ranking_tiers": CFG.get("zero_ranking_tiers", []),
        "vol_free_below": CFG.get("vol_free_below", 10000),
        "volume_brackets": CFG.get("volume_brackets", []),
        "step_ratio": CFG["step_ratio"],
        "tier_step_flat": CFG.get("tier_step_flat"),
        "volume_add_cap": CFG.get("volume_add_cap"),
        "client_floor": CFG["client_floor"],
        "addon_market_ratio": CFG["addon_market_ratio"],
        "default_markup_pct": CFG["default_markup_pct"],
        "ultra_bucket_size": CFG["ultra_bucket_size"],
        "grid_mode": CFG.get("grid_mode", True),
        "goal_options": GOAL_OPTIONS,
        "goal_scope": GOAL_SCOPE,
        "service_min_volume": CFG.get("service_min_volume", 0),
        "service_upgrade_ratio": CFG.get("service_upgrade_ratio", 0),
        "service_max_swaps": CFG.get("service_max_swaps", 3),
        "store_intent_tier_boost": CFG.get("store_intent_tier_boost", 3.0),
        "grid_target_keywords": CFG.get("grid_target_keywords", 32),
        "grid_min_services": CFG.get("grid_min_services", 4),
        "grid_max_services": CFG.get("grid_max_services", 20),
        "grid_max_cities": CFG.get("grid_max_cities", 10),
        "metro_no_suffix_zips": CFG.get("metro_no_suffix_zips", 25),
        "metro_no_suffix_share": CFG.get("metro_no_suffix_share", 0.6),
        "grid_state_suffix": CFG.get("grid_state_suffix", True),
        "competitive_bucket_size": CFG["competitive_bucket_size"],
        "longtail_target": CFG["longtail_target"],
    })

@app.route("/api/config", methods=["POST"])
@_json_error_guard
def api_config_set():
    """Apply edited constants to the running session (not persisted to disk —
    a restart reverts to the file defaults). Lets Brendan tune and re-quote live."""
    d = request.get_json(force=True)
    # A CONSTANT THAT QUIETLY REVERTS IS WORSE THAN ONE THAT NEVER MOVED, because
    # by then nobody is watching it. The source fingerprint tells you which CODE
    # is running; this is the same idea for the CONFIG, so a session tuned away
    # from the file says so in the header. (2026-08-12)
    CFG_EDITS.append({"keys": sorted(k for k in d if k != "_note"),
                      "note": str(d.get("_note") or "")[:120]})
    try:
        if "geo_anchor" in d:
            for k, v in d["geo_anchor"].items():
                if k in CFG["geo_anchor"]:
                    CFG["geo_anchor"][k] = int(v)
        if "competitive_adder" in d:
            for k, v in d["competitive_adder"].items():
                CFG["competitive_adder"][int(k)] = int(v)
        if "bid_score_breaks" in d:
            CFG["bid_score_breaks"] = [float(x) for x in d["bid_score_breaks"]]
        # zero_ranking_tiers: [[pct_not_ranking, uplift_pct], ...] sorted high-to-low
        if "zero_ranking_tiers" in d and isinstance(d["zero_ranking_tiers"], list):
            tiers = []
            for pair in d["zero_ranking_tiers"]:
                if isinstance(pair, (list, tuple)) and len(pair) == 2:
                    tiers.append([float(pair[0]), float(pair[1])])
            tiers.sort(key=lambda t: t[0], reverse=True)
            CFG["zero_ranking_tiers"] = tiers
        # geo_pct_tiers: [[min_pct_not_ranking, geo_pct_of_seo], ...] high-to-low
        if "geo_pct_tiers" in d and isinstance(d["geo_pct_tiers"], list):
            gt = []
            for pair in d["geo_pct_tiers"]:
                if isinstance(pair, (list, tuple)) and len(pair) == 2:
                    gt.append([float(pair[0]), float(pair[1])])
            gt.sort(key=lambda t: t[0], reverse=True)
            CFG["geo_pct_tiers"] = gt
        # addon_volume_discount_tiers: [[min_market_count, pct_off], ...]
        # high-to-low, first match wins. Counts are whole markets; a fractional
        # threshold would make "10 markets" mean different things on two runs.
        if ("addon_volume_discount_tiers" in d
                and isinstance(d["addon_volume_discount_tiers"], list)):
            at = []
            for pair in d["addon_volume_discount_tiers"]:
                if isinstance(pair, (list, tuple)) and len(pair) == 2:
                    try:
                        lo, pct = int(float(pair[0])), float(pair[1])
                    except (TypeError, ValueError):
                        continue
                    # A 100%-off bracket prices add-on markets at nothing and a
                    # negative one charges more for buying more. Neither is a
                    # discount, and both would price silently.
                    if lo >= 1 and 0.0 <= pct < 100.0:
                        at.append([lo, pct])
            at.sort(key=lambda t: t[0], reverse=True)
            if at:
                CFG["addon_volume_discount_tiers"] = at
        if "vol_add_ramp" in d and isinstance(d["vol_add_ramp"], list) and len(d["vol_add_ramp"]) == 2:
            CFG["vol_add_ramp"] = [float(d["vol_add_ramp"][0]), float(d["vol_add_ramp"][1])]
        if "geo_pricing_mode" in d and d["geo_pricing_mode"] in ("pct", "card"):
            CFG["geo_pricing_mode"] = d["geo_pricing_mode"]
        # volume_brackets: [[lo, hi, dollars_per_search], ...]; hi may be null/"".
        if "volume_brackets" in d and isinstance(d["volume_brackets"], list):
            brs = []
            for b in d["volume_brackets"]:
                if isinstance(b, (list, tuple)) and len(b) >= 3:
                    lo = float(b[0])
                    hi = None if b[1] in (None, "", "null") else float(b[1])
                    rate = float(b[2])
                    brs.append([lo, hi, rate])
            brs.sort(key=lambda x: x[0])
            CFG["volume_brackets"] = brs
        if "vol_free_below" in d and d["vol_free_below"] not in (None, ""):
            CFG["vol_free_below"] = float(d["vol_free_below"])
        if "cpc_adder_enabled" in d:
            CFG["cpc_adder_enabled"] = bool(d["cpc_adder_enabled"])
        if "grid_mode" in d:
            CFG["grid_mode"] = bool(d["grid_mode"])
        if "grid_state_suffix" in d:
            CFG["grid_state_suffix"] = bool(d["grid_state_suffix"])
        for key, caster in [("grid_target_keywords", int), ("grid_min_services", int),
                            ("grid_max_services", int), ("grid_max_cities", int),
                            ("metro_no_suffix_zips", int)]:
            if key in d and d[key] not in (None, ""):
                CFG[key] = caster(d[key])
        for key, caster in [("service_min_volume", int), ("service_max_swaps", int),
                            ("service_upgrade_ratio", float),
                            ("metro_no_suffix_share", float),
                            ("store_intent_tier_boost", float),
                            ("zero_ranking_bonus", int), ("zero_ranking_top_n", int),
                            ("zero_ranking_frac", float), ("step_ratio", float),
                            ("client_floor", int), ("addon_market_ratio", float),
                            ("default_markup_pct", float), ("ultra_bucket_size", int),
                            ("competitive_bucket_size", int), ("longtail_target", int),
                            ("cpc_adder_mult", float), ("cpc_adder_cap", int),
                            ("cpc_adder_free_below", float), ("cpc_adder_knee", float),
                            ("cpc_adder_mult_high", float), ("tier_step_pct_of_base", float),
                            ("ecom_anchor_add", int),
                            ("pin_head_terms", int),
                            ("pin_min_volume", int),
                            ("geo_pct_default", float),
                            ("min_term_months", int),
                            ("nationwide_service_extras", float)]:
            if key in d and d[key] not in (None, ""):
                CFG[key] = caster(d[key])
        # Nullable knobs: empty/0 disables (flat step falls back to step_ratio;
        # no cap means volume brackets run uncapped).
        for key in ("tier_step_flat", "volume_add_cap"):
            if key in d:
                v = d[key]
                CFG[key] = None if v in (None, "", "null", 0, "0") else int(float(v))
    except (ValueError, TypeError) as e:
        return jsonify({"error": f"Invalid value: {e}"}), 400
    return jsonify({"ok": True})

@app.route("/api/serp_recommend", methods=["POST"])
@_json_error_guard
def api_serp_recommend():
    """Pick the most persuasive head term to screenshot for a proposal:
    prefer a 'Not Found' term, then most competitive, then geo-modified."""
    d = request.get_json(force=True)
    head = d.get("head", [])          # [{"kw":..., "comp":"Ultra"/"Competitive"}]
    ranks = d.get("ranks", {})        # {kw: "Not Found" | position}
    markets = usable_markets(d.get("geo_values") or [])
    def is_geo(kw):
        return any(m.lower() in kw.lower() for m in markets)
    # ABSENT ON PURPOSE, NOT ABSENT BY ACCIDENT. The point of the exhibit is a
    # page one WITHOUT the client on it — that is the whole argument for the
    # retainer. This used to treat a term with no measurement at all the same as
    # a term measured and genuinely missing (`ranks.get(kw, "Not Found")`), so a
    # term the rank check had simply never reached could win the slot and the
    # screenshot could come back with the client sitting at the top of it. Ski
    # Barn's proposal shipped a capture of "ski jackets wayne nj" with Ski Barn
    # first in the local pack. Measured-and-absent now outranks unmeasured, and
    # both outrank a term the client already holds. (2026-08-19)
    def tier(kw):
        r = ranks.get(kw, None)
        if r == "Not Found" or r == 0:
            return 3                          # measured, and they are not there
        if r is None or r == "":
            return 1                          # never measured — a guess
        try:
            pos = int(r)
        except (TypeError, ValueError):
            return 1
        return 2 if pos > 20 else 0           # deep enough to argue, or theirs
    def score(item):
        kw = item.get("kw", "")
        comp_rank = 2 if item.get("comp", "").lower().startswith("ultra") else 1
        return (tier(kw),                     # absent first
                comp_rank,                    # most competitive
                1 if is_geo(kw) else 0)       # geo-modified
    if not head:
        return jsonify({"recommended": None, "options": []})
    ordered = sorted(head, key=score, reverse=True)
    best = ordered[0]
    return jsonify({"recommended": best["kw"],
                    # Say WHY, so a weak pick is visible rather than silent: if
                    # the strongest term available is one they already rank for,
                    # the operator should know before it goes in a document.
                    "basis": {3: "measured and not ranking",
                              2: "ranking below the first two pages",
                              1: "no ranking measured for it",
                              0: "they already rank for it"}[tier(best["kw"])],
                    "weak": tier(best["kw"]) <= 1,
                    "options": [h["kw"] for h in head]})

@app.route("/api/serp_queue", methods=["POST"])
@_json_error_guard
def api_serp_queue():
    """Step A — queue the SERP task and return immediately with the task_id.
    Short request (no waiting). The frontend then polls /api/serp_fetch."""
    d = request.get_json(force=True)
    keyword = (d.get("keyword") or "").strip()
    markets = usable_markets(d.get("geo_values") or [])
    state   = derive_state(markets, (d.get("state") or "").strip())
    device  = d.get("device", "desktop")
    if not keyword:
        return jsonify({"error": "No keyword provided."}), 400
    try:
        tp = dfs_post("/serp/google/organic/task_post", [{
            "keyword": keyword, "location_name": loc_string(markets, state),
            "language_code": "en", "device": device, "priority": 2}])
        task = (tp.get("tasks") or [{}])[0]
        task_id = task.get("id")
        if not task_id:
            return jsonify({"error": f"Task not created: {task.get('status_message')}"}), 502
        # pass display params through so the fetch step can size the screenshot
        return jsonify({"task_id": task_id, "keyword": keyword, "device": device,
                        "width": d.get("width"), "height": d.get("height"),
                        "scale": d.get("scale")})
    except requests.HTTPError as e:
        return jsonify({"error": f"DataForSEO error: {e}"}), 502
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}"}), 500

def _trim_serp_image(png_bytes, max_h=None, blank_thresh=245, collapse_over=110, keep=36, aspect=None):
    """Collapse tall near-blank horizontal bands in a SERP screenshot (the AI
    Mode 'Thinking' placeholder leaves hundreds of empty pixels), optionally
    cap the final height, and re-encode as JPEG. Blank detection samples a
    40px-wide downscale per row, so it's fast even on 8000px pages."""
    import io
    from PIL import Image, ImageOps
    im = Image.open(io.BytesIO(png_bytes)).convert("RGB")
    w, h = im.size
    # Dark-mode guard: DFS occasionally renders Google in dark theme. Detect a
    # dark page background (sample corners + center-top) and convert to a
    # light-mode look: invert luminance, then rotate hue 180\u00b0 so brand
    # colors (blue links, logo) come back approximately correct.
    _pts = [(4, 4), (w - 5, 4), (w // 2, 4), (4, min(h - 5, 200))]
    if sum(sum(im.getpixel(p)) for p in _pts) / (len(_pts) * 3) < 100:
        inv = ImageOps.invert(im)
        hsv = inv.convert("HSV")
        ch = list(hsv.split())
        ch[0] = ch[0].point(lambda x: (x + 128) % 256)
        im = Image.merge("HSV", ch).convert("RGB")
        # White-point fix: the inverted background is light grey; scale so it
        # reads as true white without blowing out text or brand colors.
        bg = max(1, max(sum(im.getpixel(p)) // 3 for p in _pts))
        if bg < 250:
            scale = 255.0 / bg
            im = im.point(lambda x: min(255, int(x * scale)))
    strip = im.resize((40, h))
    px = strip.load()
    blank = []
    for y in range(h):
        row_ok = True
        for x in range(40):
            r, g, b = px[x, y]
            if r < blank_thresh or g < blank_thresh or b < blank_thresh:
                row_ok = False
                break
        blank.append(row_ok)
    segments, y = [], 0
    while y < h:
        if blank[y]:
            start = y
            while y < h and blank[y]:
                y += 1
            run = y - start
            if run > collapse_over:
                segments.append((start, start + keep))       # keep a sliver
            else:
                segments.append((start, y))
        else:
            start = y
            while y < h and not blank[y]:
                y += 1
            segments.append((start, y))
    new_h = sum(b - a2 for a2, b in segments)
    out = Image.new("RGB", (w, new_h), (255, 255, 255))
    cy = 0
    for a2, b in segments:
        band = im.crop((0, a2, w, b))
        out.paste(band, (0, cy))
        cy += b - a2
    # Right-edge crop (post-collapse, so blank rows can't dilute the sample):
    # sample 128 row-bands per column; a column counts as content if ANY band
    # is non-blank. Mirror the left margin so the crop looks intentional.
    w, h2 = out.size
    strip2 = out.resize((w, 128))
    spx = strip2.load()
    def _col_has_content(x):
        for yy in range(128):
            r, g, b = spx[x, yy]
            if r < blank_thresh or g < blank_thresh or b < blank_thresh:
                return True
        return False
    right = w - 1
    while right > 0 and not _col_has_content(right):
        right -= 1
    left_margin = 0
    while left_margin < w - 1 and not _col_has_content(left_margin):
        left_margin += 1
    new_w = min(w, right + 1 + max(24, left_margin))
    if 0 < new_w < w * 0.97:
        out = out.crop((0, 0, new_w, h2))
    if aspect:
        # Enforce a fixed aspect ratio (e.g. "3:2" landscape) for consistent
        # proposal layout: crop the bottom if too tall, pad white if too short.
        try:
            aw, ah = (float(x) for x in str(aspect).split(":"))
        except (ValueError, TypeError):
            aw = ah = 0
        if aw > 0 and ah > 0:
            w2, h2 = out.size
            target = int(round(w2 * ah / aw))
            if h2 > target:
                out = out.crop((0, 0, w2, target))
            elif h2 < target:
                canvas = Image.new("RGB", (w2, target), (255, 255, 255))
                canvas.paste(out, (0, 0))
                out = canvas
    if max_h and out.size[1] > max_h:
        out = out.crop((0, 0, out.size[0], max_h))
    buf = io.BytesIO()
    out.save(buf, "JPEG", quality=85)
    return buf.getvalue()


@app.route("/api/serp_fetch", methods=["POST"])
@_json_error_guard
def api_serp_fetch():
    """Step B — try to fetch the screenshot for a queued task_id. Returns the
    image if ready, or {ready:false} if still processing. Frontend polls this.
    Each call is short, so no request-timeout risk."""
    d = request.get_json(force=True)
    task_id = (d.get("task_id") or "").strip()
    device  = d.get("device", "desktop")
    keyword = d.get("keyword", "")
    if not task_id:
        return jsonify({"error": "No task_id."}), 400
    # build screenshot params, including optional sizing
    shot = {"task_id": task_id, "browser_preset": device}
    if d.get("width"):  shot["browser_screen_width"]  = int(d["width"])
    if d.get("height"): shot["browser_screen_height"] = int(d["height"])
    if d.get("scale"):  shot["browser_screen_scale_factor"] = float(d["scale"])
    try:
        sc = dfs_post("/serp/screenshot", [shot])
        try:
            image_url = sc["tasks"][0]["result"][0]["items"][0]["image"]
        except (KeyError, IndexError, TypeError):
            image_url = None
        if not image_url:
            msg = (sc.get("tasks") or [{}])[0].get("status_message", "")
            return jsonify({"ready": False, "status": msg})
        login = os.environ.get("DFS_LOGIN", ""); pw = os.environ.get("DFS_PASSWORD", "")
        tok = base64.b64encode(f"{login}:{pw}".encode()).decode()
        img = requests.get(image_url, headers={"Authorization": f"Basic {tok}"}, timeout=60)
        img.raise_for_status()
        content, mime = img.content, "image/png"
        if d.get("trim"):
            # Rep-tool proposal shots: collapse blank bands (AI-overview
            # placeholder renders as a huge white gap), cap height for a
            # landscape-ish exhibit, JPEG to keep saved-quote payloads sane.
            try:
                content = _trim_serp_image(content, max_h=int(d.get("max_h") or 0) or None,
                                           aspect=d.get("aspect"))
                mime = "image/jpeg"
            except Exception as _te:
                print(f"[serp trim] skipped: {_te}")
        b64 = base64.b64encode(content).decode()
        return jsonify({"ready": True, "keyword": keyword,
                        "data_url": f"data:{mime};base64,{b64}"})
    except requests.HTTPError as e:
        # screenshot endpoint returns an error while the task is still running;
        # treat as not-ready rather than a hard failure so the poll continues
        return jsonify({"ready": False, "status": f"processing ({e})"})
    except Exception as e:
        return jsonify({"error": f"Unexpected error: {e}"}), 500

# ---------------------------------------------------------------------------
# SAVED QUOTES — persistence with version history (like the Meta forecast tool).
# Degrades gracefully: if no DATABASE_URL, /api/quotes/status reports disabled
# and the UI shows "attach a database to enable" instead of the Save controls.
# ---------------------------------------------------------------------------
_LOCATIONS_CACHE = {"names": None, "codes": None}


def us_location_code(name):
    """Numeric Google Ads location_code for a DataForSEO location_name.

    NOT usable for Labs endpoints — Labs keys off its own location set and
    rejects Google Ads city codes ("40501 Invalid Field: 'location_code'").
    Kept for Google-Ads-side lookups. Returns None when unrecognised.
    """
    if not name:
        return None
    if _LOCATIONS_CACHE.get("codes") is None:
        try:
            data = dfs_get("/keywords_data/google_ads/locations/us")
            items = (data.get("tasks") or [{}])[0].get("result") or []
            _LOCATIONS_CACHE["codes"] = {
                (it.get("location_name") or "").strip().lower(): it.get("location_code")
                for it in items
                if it.get("location_name") and it.get("location_code")}
        except Exception:
            _LOCATIONS_CACHE["codes"] = {}
    return _LOCATIONS_CACHE["codes"].get(str(name).strip().lower())

def dfs_get(path, timeout=60):
    login = os.environ.get("DFS_LOGIN", "")
    pw    = os.environ.get("DFS_PASSWORD", "")
    token = base64.b64encode(f"{login}:{pw}".encode()).decode()
    resp = requests.get(BASE + path,
                        headers={"Authorization": f"Basic {token}"}, timeout=timeout)
    resp.raise_for_status()
    return resp.json()


def us_location_names():
    """All US location_names DataForSEO recognises, cached for the process.
    Used to validate the cities a partner typed BEFORE spending API calls on a
    misspelling (e.g. 'Kakuana' should be 'Kaukauna')."""
    if _LOCATIONS_CACHE["names"] is not None:
        return _LOCATIONS_CACHE["names"]
    try:
        data = dfs_get("/keywords_data/google_ads/locations/us")
        items = (data.get("tasks") or [{}])[0].get("result") or []
        names = [it.get("location_name", "") for it in items if it.get("location_name")]
        _LOCATIONS_CACHE["names"] = names
        return names
    except Exception:
        _LOCATIONS_CACHE["names"] = []
        return []


def validate_cities(cities, state):
    """Check each entered city resolves to a real DataForSEO location in the
    chosen state. Returns [{city, ok, resolved, suggestions[]}]. Suggestions use
    close-match scoring so a typo surfaces the intended city."""
    import difflib
    names = us_location_names()
    out = []
    if not names:
        return [{"city": c, "ok": None, "resolved": "", "suggestions": []} for c in cities]
    for c in cities:
        c_name, c_state = parse_market(c, state)
        state_l = (c_state or "").strip().lower()
        in_state = [n for n in names if state_l and f",{state_l}," in n.lower()] if state_l else names
        city_only = {}
        for n in in_state:
            first = n.split(",")[0].strip().lower()
            city_only.setdefault(first, n)
        cl = c_name.strip().lower()
        if cl in STATE_ABBREV:            # a state used AS a geo ("delaware")
            out.append({"city": c, "ok": True, "kind": "state",
                        "resolved": f"{cl.title()},United States", "suggestions": []})
        elif cl in city_only:
            out.append({"city": c, "ok": True, "kind": "city",
                        "resolved": city_only[cl], "suggestions": []})
        else:
            close = difflib.get_close_matches(cl, list(city_only.keys()), n=3, cutoff=0.72)
            if close:                     # probably a typo of a real city
                out.append({"city": c, "ok": False, "kind": "typo", "resolved": "",
                            "suggestions": [city_only[m] for m in close]})
            else:                         # regional phrase ("south jersey") —
                                          # legit in keyword TEXT, not a location
                out.append({"city": c, "ok": True, "kind": "phrase",
                            "resolved": "", "suggestions": []})
    return out


@app.route("/api/validate_geo", methods=["POST"])
@_json_error_guard
def api_validate_geo():
    d = request.get_json(force=True)
    cities = usable_markets(d.get("geo_values") or [])
    state  = (d.get("state") or "").strip()
    if not cities:
        return jsonify({"error": "No cities to check."}), 400
    try:
        return jsonify({"state": state, "results": validate_cities(cities, state)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def claude_menu_to_terms(labels, brand, domain, seeds, business_desc):
    """Convert raw nav-menu labels into search-phrase service terms. A menu
    says 'Healthcare' or 'Warehouse'; a searcher types 'healthcare construction
    company'. Returns {label: term_or_None} — None means drop it (careers,
    press, process pages). Empty dict when the AI isn't available, so the
    caller can fall back to raw labels."""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or not labels:
        return {}
    prompt = f"""These are navigation menu labels scraped from a business's website. Convert each into the search phrase a potential CUSTOMER would type into Google when looking for that service from this kind of business.

BUSINESS: {brand or "(unknown)"} — {domain}
WHAT THEY DO: {business_desc or "(infer from the labels and any seeds)"}
EXISTING SEED TERMS: {", ".join(seeds) if seeds else "(none)"}

MENU LABELS: {json.dumps(labels, ensure_ascii=False)}

Rules:
- Sector/industry labels get the core service appended: for a commercial builder, "Healthcare" -> "healthcare construction company", "Self-Storage" -> "self storage construction".
- Labels that already read like a service ("Commercial Construction") may pass through nearly as-is, normalized to how people search.
- USE THE CUSTOMER'S VOCABULARY, not the site's page template. If most labels share one template word (a menu of "X Treatment & Therapy" condition pages, "Y Repair Services" pages), do NOT echo that word into every term — a person with anxiety types "anxiety therapist" or "anxiety therapy", not "anxiety treatment therapy". Vary the phrasing to match real searches.
- When the labels are all variations of ONE parent service (conditions, specialties, sub-services), ALSO make sure the parent's everyday head terms are represented — the bread-and-butter words customers actually type ("therapist", "therapy", "counseling", "mental health clinic" for a behavioral-health practice) — by mapping the most general labels to those instead of to another templated variant.
- Map to null anything that is NOT a purchasable service: careers, press, blog, media, "our process", team pages, generic CTAs.
- NEVER add "near me", "nearby", "closest" or any other proximity phrase. Every term is crossed with
  a city later, and "mattress store near me acworth ga" is not a phrase any human types — "near me"
  IS the location. Write the bare service and let the grid add the place.
- Lowercase, no geo, 2-5 words each.

Return ONLY a JSON object mapping every input label to its search phrase or null. No preamble, no markdown fences."""
    try:
        resp = requests.post("https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({"model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 1500, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=25)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.M).strip()
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            out = {}
            for k, v in parsed.items():
                if not (isinstance(v, str) and v.strip()):
                    out[k] = None
                    continue
                # Strip proximity here too. The prompt asks; this guarantees.
                t = clean_kw(strip_placeholders(strip_proximity(v.strip().lower()))).strip()
                out[k] = t or None
            return out
    except Exception:
        pass
    return {}


class _NavLinkParser(HTMLParser):
    """Collect anchor text from the page, tracking whether each link sits inside
    menu context. Menu structure is the signal: businesses list the services
    they actually sell in their navigation. Menu context means a semantic
    <nav>/<header> OR any element whose class/id contains nav|menu — WordPress
    themes and page builders routinely skip the semantic tags and ship
    <div class="menu">/<ul id="main-menu"> instead."""
    _NAV = {"nav", "header"}
    _MENUISH = re.compile(r"(?:^|[\s_-])(?:nav|menu)(?:$|[\s_-])|nav(?:bar|igation)|menu[-_]", re.I)
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._stack = []          # per open tag: True if it opened menu context
        self.nav_depth = 0
        self._in_a = False
        self._href = ""
        self._buf = []
        self.nav_links, self.other_links = [], []
    _VOID = {"br","img","input","meta","link","hr","area","base","col","embed","source","track","wbr"}
    def _is_menuish(self, tag, attrs):
        if tag in self._NAV: return True
        d = dict(attrs)
        blob = (d.get("class") or "") + " " + (d.get("id") or "") + " " + (d.get("role") or "")
        return bool(self._MENUISH.search(blob)) or (d.get("role") or "").lower() == "navigation"
    def handle_starttag(self, tag, attrs):
        if tag in self._VOID:
            return
        menuish = self._is_menuish(tag, attrs)
        self._stack.append((tag, menuish))
        if menuish: self.nav_depth += 1
        if tag == "a":
            self._in_a = True; self._buf = []
            self._href = (dict(attrs).get("href") or "")
    def handle_endtag(self, tag):
        # pop to the matching open tag (tolerates unclosed tags in the wild)
        for i in range(len(self._stack) - 1, -1, -1):
            if self._stack[i][0] == tag:
                for _t, m in self._stack[i:]:
                    if m and self.nav_depth: self.nav_depth -= 1
                del self._stack[i:]
                break
        if tag == "a" and self._in_a:
            self._in_a = False
            text = " ".join("".join(self._buf).split())
            rec = (text, self._href)
            (self.nav_links if self.nav_depth else self.other_links).append(rec)
    def handle_data(self, data):
        if self._in_a: self._buf.append(data)

_MENU_GENERIC = {
    "home","about","about us","contact","contact us","blog","news","careers",
    "gallery","portfolio","testimonials","reviews","faq","faqs","privacy policy",
    "privacy","terms","terms of use","team","our team","meet the team","locations",
    "location","sitemap","login","log in","search","services","our services",
    "projects","our projects","our work","work","resources","get a quote",
    "request a quote","free quote","free estimate","get started","learn more",
    "read more","view all","see all","menu","español","facebook","instagram",
    "linkedin","twitter","youtube","x",
    "start my career","start my project","our process","approach","our approach","media","blogs",
    "press releases","press","join our team","apply now","employment",
    "history","our history","our story","leadership","safety","awards",
}
_SERVICE_PATH_HINT = re.compile(
    r"/[a-z0-9-]*(?:services?|markets?|sectors?|industr(?:y|ies)|what-we-do|"
    r"capabilit(?:y|ies)|specialt(?:y|ies)|divisions?|expertise|solutions?)"
    r"[a-z0-9-]*(?:/|$)", re.I)

# Generic fallbacks, only used when the client's own vocabulary yields nothing.
_QUALIFIER_FALLBACK = ("certification", "training", "course")
_QUALIFIER_STOP = {
    "the", "and", "for", "with", "your", "our", "near", "best", "top", "new",
    "services", "service", "company", "companies", "local", "cheap", "free",
    "professional", "commercial", "residential", "industry", "industrial",
}


def qualifier_words(seeds, site_terms=None, limit=3):
    """The words THIS client attaches to things, commonest first.

    The qualified form of an acronym was hardcoded to "certification", which is
    right for a standards body and wrong for most clients — "adr" wants
    "mediation", "pmp" wants "exam", "hipaa" wants "compliance". Read it off the
    client's own vocabulary instead: the words already recurring in their seeds
    ARE how they describe what they sell. (2026-08-10)
    """
    counts = {}
    for t in list(seeds or []) + list(site_terms or []):
        words = re.sub(r"[^a-z0-9 ]", " ", str(t).lower()).split()
        # Skip 1-word seeds: an acronym on its own contributes no qualifier.
        if len(words) < 2:
            continue
        for w in words:
            if len(w) < 4 or w in _QUALIFIER_STOP:
                continue
            counts[w] = counts.get(w, 0) + 1
    ranked = sorted(counts, key=lambda w: (-counts[w], w))
    out = [w for w in ranked if counts[w] >= 2][:limit]
    for f in _QUALIFIER_FALLBACK:
        if len(out) >= limit:
            break
        if f not in out:
            out.append(f)
    return out[:limit]


_ACRONYM_STOP = {
    "USA", "US", "FAQ", "FAQS", "PDF", "HTML", "HTTP", "HTTPS", "WWW", "URL",
    "CEO", "CFO", "COO", "CTO", "HR", "IT", "PR", "AI", "SEO", "PPC", "ROI",
    "AM", "PM", "EST", "CST", "PST", "EDT", "CDT", "PDT", "MST", "UTC",
    "JAN", "FEB", "MAR", "APR", "JUN", "JUL", "AUG", "SEP", "SEPT", "OCT",
    "NOV", "DEC", "MON", "TUE", "WED", "THU", "FRI", "SAT", "SUN",
    "AND", "THE", "FOR", "ALL", "NEW", "TOP", "OUR", "YOU", "NOW", "MORE",
    "LOGIN", "SIGN", "JOIN", "HOME", "ABOUT", "BLOG", "NEWS", "SHOP", "CART",
    "MENU", "NEXT", "PREV", "BACK", "SEND", "OPEN", "CLOSE", "VIEW", "READ",
    "COVID", "ADA", "GDPR", "CCPA", "TM", "LLC", "INC", "LTD", "CO",
    # Currency and unit codes: a CLOSED list, unlike English words, so a blocklist
    # is the right shape here. Ski Barn's site says "prices in USD" in caps and
    # nowhere in lowercase, so the lowercase-word test cannot see it. (2026-08-12)
    "USD", "EUR", "GBP", "CAD", "AUD", "NZD", "JPY", "CHF", "MXN", "INR",
    "MPG", "MPH", "KPH", "PSI", "BTU", "KWH", "SQFT", "LBS", "OZS", "GAL",
    "MIN", "MAX", "AVG", "QTY", "SKU", "UPC", "ISBN", "VIN", "ZIP",
}


def mine_acronyms(html, brand="", limit=14):
    """Find the SHORT branded names a client's buyers actually search.

    NASSCO's seeds read "bsdi pipeline inspection certification" and
    "underground infrastructure industry association" — descriptions of the
    organisation, produced by expanding nav labels into prose. Total measured
    demand across sixteen of them: 10/mo. What a contractor types is "PACP",
    "MACP", "ITCP", "CIPP". The acronyms were on the client's own site the whole
    time and nothing looked for them, because the menu converter only ever made
    labels LONGER. (2026-08-10)

    Three sources, most reliable first:
      1. "Pipeline Assessment Certification Program (PACP)" — the expansion is
         right there in the parentheses, so the meaning is certain.
      2. PACP(tm) / MACP(R) — a trademark marker means the client owns the term.
      3. A bare ALL-CAPS token that recurs, which is how a program gets
         referred to once the page has introduced it.

    Returns [{"acronym", "expansion", "hits", "source"}], commonest first.
    Measurement happens in the caller — an acronym with no search volume is a
    filing code, not a keyword.
    """
    txt = re.sub(r"(?is)<(script|style|noscript)[^>]*>.*?</\1>", " ", html or "")
    txt = re.sub(r"(?s)<[^>]+>", " ", txt)
    txt = (txt.replace("&amp;", "&").replace("&nbsp;", " ")
              .replace("&#8482;", "\u2122").replace("&trade;", "\u2122")
              .replace("&reg;", "\u00ae"))
    txt = re.sub(r"\s+", " ", txt)

    brand_up = re.sub(r"[^A-Z]", "", (brand or "").upper())
    found = {}

    # A HAND-MAINTAINED BLOCKLIST ALWAYS MISSES THE NEXT ONE. Ski Barn offered
    # "free" (550,000/mo), "here" (246,000/mo) and "usd" (135,000/mo) as its
    # industry acronyms, because FREE / HERE / USD appear in caps somewhere on the
    # page and none of them were on the list. Seeding one of those would have put
    # half a million searches a month into the volume total.
    #
    # The page answers this itself: an industry program is referred to in CAPS and
    # essentially never written out in lowercase running text on the same site.
    # "free" and "here" are all over a retail page in lowercase; "PACP" never is.
    # Measured against the document instead of guessed from a list. (2026-08-12)
    lower_words = set(re.findall(r"\b[a-z]{3,6}\b", txt))

    def add(ac, expansion, source):
        ac = ac.strip().upper()
        if (len(ac) < 3 or len(ac) > 6 or ac in _ACRONYM_STOP
                or not ac.isalpha() or (brand_up and ac == brand_up)
                or ac.lower() in lower_words):
            return
        e = found.setdefault(ac, {"acronym": ac, "expansion": "", "hits": 0,
                                  "source": source})
        if expansion and not e["expansion"]:
            # The regex greedily walks back over capitalised words, so it picks
            # up the sentence lead-in ("NASSCO The Pipeline Assessment...").
            # Trim the brand and any leading article — this is a tooltip.
            ex = expansion.strip()
            if brand_up:
                ex = re.sub(r"^\s*" + re.escape(brand or "") + r"\b", "", ex,
                            flags=re.I).strip()
            ex = re.sub(r"^(?:the|our|a|an|and|its|their)\s+", "", ex, flags=re.I)
            e["expansion"] = ex.strip()[:70]
            e["source"] = source

    # 1. Expansion followed by the abbreviation in brackets.
    for m in re.finditer(r"((?:[A-Z][A-Za-z-]+\s+){1,6})\(\s*([A-Z]{3,6})\s*[\u2122\u00ae]?\s*\)", txt):
        add(m.group(2), m.group(1), "expansion in brackets")
    # 2. Trademark marker.
    for m in re.finditer(r"\b([A-Z]{3,6})\s*[\u2122\u00ae]", txt):
        add(m.group(1), "", "trademarked by the client")
    # 3. Recurring bare capitals.
    for m in re.finditer(r"(?<![A-Za-z])([A-Z]{3,6})(?![A-Za-z])", txt):
        ac = m.group(1).upper()
        if len(ac) < 3 or ac in _ACRONYM_STOP:
            continue
        if ac in found:
            found[ac]["hits"] += 1
        else:
            add(ac, "", "repeated on the page")
            if ac in found:
                found[ac]["hits"] = 1

    out = [v for v in found.values()
           if v["hits"] >= 2 or v["source"] != "repeated on the page"]
    out.sort(key=lambda v: (v["source"] == "repeated on the page", -v["hits"],
                            v["acronym"]))
    return out[:limit]


def claude_industry_services(brand="", domain="", industry="", business_desc="",
                             site_pages=None, seeds=None, geo="", n=None):
    """What ELSE does a business of this kind sell that people search for.

    The list can only ever choose among the seeds it is handed, so when those
    seeds come from an incumbent's ranking report they inherit that incumbent's
    blind spots. Junk Bee Gone is the case (2026-08-10): 69 seeds from the old
    agency's report, and the tool spent eight slots on synonyms — haul away,
    haul away junk, haul away service, junk haulers, hauling services, junk
    remover, remove junk, junk. Brendan's list for the same client covered
    hoarding cleanup, demolition, shed demolition, paper shredding, estate
    cleanout, house cleanout, construction debris removal and appliance removal:
    eight DIFFERENT services, none of them in the seed list, several of them
    named on the client's own website.

    So ask for the GAPS rather than a generic industry list — the model is told
    what is already covered and asked what is missing. Output is offered as seed
    chips, never fed straight into the grid: seeds are never blocked by the
    grounding check, so a term the operator accepts is trusted, and one they
    ignore costs nothing.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return []
    n = int(n if n else CFG.get("industry_gap_n", 22) or 22)
    have = [str(x).strip().lower() for x in (seeds or []) if str(x).strip()]
    pages = [str(p) for p in (site_pages or [])][:40]
    prompt = f"""A search campaign is being scoped for {brand or domain or "a business"}.
Industry: {industry or "(not given)"}
{f'They describe themselves as: {business_desc}' if business_desc else ''}
{f'Pages on their website: {", ".join(pages)}' if pages else ''}
{f'Area served: {geo}' if geo else ''}

The keyword list ALREADY covers these services:
{", ".join(have[:60]) if have else "(nothing yet)"}

Name up to {n} ADDITIONAL service lines a business of this type sells that people
search for, and that are NOT already covered above. Rules:
1. A DIFFERENT service, never a synonym of one already listed. "haul away
   service" when "junk removal" is present is a synonym — do not return it.
   This applies WITHIN your own answer too: {n} slots means {n} different things
   this business sells, not one thing worded {n} ways. A dental list runs
   cleanings, crowns, whitening, root canals, extractions, dentures, implants,
   emergency care — it does NOT run overlay onlay, crown overlay, overlay crown,
   overlays on teeth, tooth overlay. Return fewer terms rather than pad with
   rewordings; anything past three variants of one procedure is discarded
   unread.
2. The phrase a customer types, 2-4 words, no city, no brand.
3. Only services this business plausibly sells. If the website pages or the
   description name something, prefer it.
4. Order by how commonly the service is bought, most common first.
5. USE THE CLIENT'S OWN NOUN FOR WHAT THEY SELL, taken from the description
   above — not the category's noun. A description saying "single-family rental
   homes" means the searched noun is "homes for rent" / "houses for rent" /
   "rental homes", NOT "apartments": those are a neighbouring industry and the
   client ranks for none of them. Getting this wrong replaces the client's entire
   vocabulary, which is worse than returning nothing.
6. If they sell INVENTORY rather than services — rentals, property, vehicles,
   stock — the equivalent of a service line is what a buyer types: their noun
   plus one distinguishing attribute. Bedroom count, size, feature, condition.
   "3 bedroom homes for rent", "homes for rent with garage", "pet friendly
   rentals".

Return ONLY JSON: {{"services": [{{"term": "hoarding cleanup", "why": "named on their site"}}]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                # ASKING FOR 22 AND BUDGETING FOR 14 IS A SILENT ZERO. Each item
                # carries a `why` string, the reply is one JSON object, and a
                # truncated object does not parse — so the whole pass would
                # return [] and the panel would say "nothing was proposed" for a
                # model that answered fine. Scaled with n. (2026-08-17)
                "max_tokens": max(1200, 110 * n), "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=45)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        raw = json.loads(text).get("services") or []
        if body.get("stop_reason") == "max_tokens":
            app.logger.warning("claude_industry_services truncated at max_tokens")
    except Exception:
        app.logger.exception("claude_industry_services failed")
        return []
    out, seen = [], set(have)
    for it in raw:
        t = clean_kw(strip_placeholders(str((it or {}).get("term") or "").lower())).strip()
        if not t or t in seen or not (1 < len(t.split()) <= 5):
            continue
        seen.add(t)
        out.append({"term": t, "why": str((it or {}).get("why") or "")[:80]})
    return out[:n]



def claude_seed_kinds(seeds, brand="", domain="", industry="",
                      business_desc="", site_pages=None, sells_products=False):
    """Is each seed a SERVICE this business performs, or just a THING?

    "old tvs" (80/mo) ranked fourth on Junk Bee Gone's focus list and pushed
    hoarding cleanup, appliance removal and demolition below the cut. It is not a
    service: it is an object, and most of that volume is people SHOPPING for old
    televisions, not paying to have one taken away — the same bare-versus-
    qualified collision the acronym check exists for, in a different costume.
    "furniture movers" is the other kind of miss: a real service, but a different
    trade from the one being quoted.

    A word list cannot do this job. "old tvs" has no action verb and neither does
    "invisalign", and one is an item while the other is the whole practice. So it
    is asked, per client, with the business in front of it.

    Only ever DEMOTES, never reorders and never deletes: the caller shows every
    verdict with its reason and the operator can keep the lot. Returns
    {term: {"kind": service|item|other_business, "why": str}} and {} on any
    failure — no key, no network, bad JSON — so the ranking degrades to
    volume-only rather than breaking. (2026-08-11)

    "reference" is the fourth bucket and the newest: a term the client's own site
    ranks for because it PUBLISHES about it. NPAIHB, a tribal health board, came
    back with twelve of twenty focus terms reading "haddon matrix", "fancy shawl
    dance" and "double balls" — every one a real page on their site, every one a
    ranking they already hold and earn nothing from, and none of them a service
    anybody could buy. The other three buckets could not name it, so all twelve
    defaulted to "service" and six reached Ultra Competitive. (2026-08-16)
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    terms = []
    for s in seeds or []:
        t = str(s or "").strip().lower()
        if t and t not in terms:
            terms.append(t)
    if not api_key or not terms:
        return {}
    pages = [str(p) for p in (site_pages or [])][:30]
    prompt = f"""{brand or domain or "A business"} is being quoted for SEO.
Industry: {industry or "(not given)"}
{f'They describe themselves as: {business_desc}' if business_desc else ''}
{f'Pages on their website: {", ".join(pages)}' if pages else ''}
{'This client SELLS PRODUCTS (a storefront was detected, or they are priced on national product demand), so the objects they stock are their offer, not things they are hired to work on.' if sells_products else ''}

Below is the list of terms the partner wants this client to rank for. Classify
EVERY term as exactly one of:

  "service"        - the person searching this wants what this client provides.
                     THE NORMAL CASE, and it covers products as well as work: a
                     bare object IS the offer when the client sells that object.
                     "invisalign" for a dentist, "trex decking" for a deck
                     builder, "old tvs" for a used-electronics dealer,
                     "mattress" for a mattress store.
  "item"           - the phrase names an object, and this client does not SELL
                     that object - they are paid to do something TO it. The
                     person searching it is at least as likely to be shopping
                     for the thing as hiring this client, so the volume is not
                     theirs. "old tvs" and "mattress" for a JUNK REMOVAL company
                     (they would be paid to take one away; the searcher wants to
                     buy one). The same two phrases are "service" for a used-
                     electronics dealer and a mattress store.
  "other_business" - a real service, but a DIFFERENT trade from this client's.
                     "furniture movers" for a junk removal company; "roofing"
                     for a plumber. Also use this for a competitor's company name.
  "reference"      - THE CLIENT PUBLISHES ABOUT THIS; nobody is hiring anyone.
                     The searcher wants to read a definition, an article or a
                     fact sheet, and this client happens to host one. A tribal
                     health board's site ranks for "haddon matrix" because its
                     EpiCenter publishes injury-prevention worksheets using it,
                     for "fancy shawl dance" and "double ball" because its youth
                     programme publishes articles on powwow dancing and
                     traditional games. Real pages, real rankings, and no work to
                     sell against any of them: the client already holds the
                     position and earns nothing from it.
                     This fires most on organisations whose website is mostly
                     resources — a health board, a museum, a trade association,
                     a university department — where the ranked-keywords pass
                     returns their library rather than their offer.
                     It is NOT "reference" merely because the phrase is a noun,
                     or because a searcher might be doing research. Ask whether
                     the client could be HIRED, FUNDED or PARTNERED WITH off the
                     back of that search. "tribal epidemiology center" is a
                     service line; "double ball" is an article.

Then, SEPARATELY, mark each term's vocabulary as "client" or "adjacent":

  "client"   - it uses the words this client uses for what they sell. THE NORMAL
               CASE.
  "adjacent" - a real, searched term for the NEIGHBOURING product type, which
               this client does not actually offer. They may still be worth
               quoting, so this is not a rejection - it decides ORDER when there
               are more terms than slots.

Amare Homes is the case: a build-for-rent community of detached single-family
homes. "3 bedroom homes for rent" is client vocabulary; "studio apartments" and
"apartment leasing" are adjacent - more searched, but they do not rent
apartments and will never rank for them. Ranking on volume alone put fourteen
apartment terms in the quote and dropped every homes-for-rent term the partner
had typed. (2026-08-13)

Rules:
- Default to "service". These buckets are not balanced: most terms are
  services; "item" and "reference" are the rare cases.
- THE SAME PHRASE GOES BOTH WAYS depending on the client, so decide from the
  business above, never from the words alone. Ask: would this searcher be happy
  to land on this client's website? If the client sells the object, yes - that
  is "service".
- If the client is a retailer, dealer, manufacturer or product brand, the object
  is their whole offer and "item" should essentially never fire. Shopping intent
  is exactly the audience they want.
- Never say "item" when the phrase already contains the work ("mattress removal",
  "tv disposal", "appliance pickup") - that is a service whatever the client is.
- "why" must be under 12 words and say what the term actually is.
- "vocab" is about WORDING, not quality. Most terms are "client". Only mark
  "adjacent" when the term names a different product type from the one the
  business above describes.

TERMS: {json.dumps(terms, ensure_ascii=False)}

Return ONLY JSON:
{{"terms": [{{"term": "old tvs", "kind": "item", "vocab": "client", "why": "an object, not work the company performs"}}]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 4000, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=60)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        raw = json.loads(text).get("terms") or []
    except Exception:
        app.logger.exception("claude_seed_kinds failed")
        return {}
    # A PRODUCT SELLER'S OBJECTS ARE ITS OFFER. For a used-electronics dealer
    # "old tvs" is the whole business and shopping intent is exactly the audience
    # wanted, so the item verdict is refused outright rather than trusted — a
    # wrong call there deletes the client's best terms. Judgement is still used
    # for other_business, which is wrong for a retailer too but not fatal.
    # (2026-08-11)
    out = {}
    for it in raw:
        t = str((it or {}).get("term") or "").strip().lower()
        k = str((it or {}).get("kind") or "").strip().lower()
        if not t:
            continue
        # Vocabulary is recorded for EVERY term, including the ones whose kind is
        # unremarkable — it decides seed order, not membership.
        if str((it or {}).get("vocab") or "").strip().lower() == "adjacent":
            out[t] = {"kind": "adjacent",
                      "why": str((it or {}).get("why") or "")[:90]}
        if k not in ("item", "other_business", "reference"):
            continue
        if k == "item" and sells_products:
            continue
        out[t] = {"kind": k, "why": str((it or {}).get("why") or "")[:90]}
    return out


# --- CALIBRATION: formula versus what was actually sent ----------------------
# Every quote already stores the Actual column beside the formula's own numbers,
# and until now nothing ever read them back. One client being 15% low is an
# anecdote; the same 15% across every quote of that SHAPE is a rule, and the
# difference between those two is the only thing standing between calibration and
# overfitting. So the deltas are grouped by the things that plausibly drive them —
# geo band and demand basis — and a recommendation is only offered once a group
# has enough quotes to mean something. (2026-08-12)

# Every session edit to CFG, in order. Empty means the running config is exactly
# what the file says.
CFG_EDITS = []

CALIB_MIN_N = 3            # below this a group is an anecdote, not a pattern
CALIB_MIN_GAP_PCT = 4.0    # below this the formula is already within noise
_TIERS = ("base", "intermediate", "advanced")


def _median(xs):
    xs = sorted(x for x in xs if x is not None)
    if not xs:
        return None
    n = len(xs)
    return xs[n // 2] if n % 2 else (xs[n // 2 - 1] + xs[n // 2]) / 2.0


def calibration_rows(payloads):
    """One row per saved quote that has an Actual price filled in."""
    rows = []
    for q in payloads or []:
        p = q.get("payload") or {}
        if isinstance(p, str):
            try:
                p = json.loads(p)
            except Exception:
                continue
        actual = p.get("actual") or {}
        pricing = p.get("pricing") or {}
        # The tier cards show the COMBINED price on a Core SEO + AI Search quote,
        # which is what the operator typed the actual against.
        ai = pricing.get("ai_search") or {}
        # AN OVERRIDDEN QUOTE IS NOT A FORMULA PRICE. stage4_price carries the
        # un-overridden figure alongside; prefer it, or calibration compares an
        # operator's own number against itself. (2026-08-27)
        _f = pricing.get("formula") or {}
        _fai = _f.get("ai_search") or {}
        formula = (_fai.get("client_total") or _f.get("client_tiers")
                   or ai.get("client_total") or pricing.get("client_tiers") or {})
        pairs = {t: (formula.get(t), actual.get(t)) for t in _TIERS}
        if not any(a and f for f, a in pairs.values()):
            continue
        inp = p.get("inputs") or {}
        kw = p.get("kw") or {}
        band = (inp.get("geo_scope") or pricing.get("band") or "").strip() or "unset"
        nat = bool(inp.get("national_demand") or pricing.get("national_demand"))
        rows.append({
            "id": q.get("id"), "name": q.get("name") or "", "client": q.get("client") or "",
            "updated_at": q.get("updated_at"),
            "band": band, "national_demand": nat,
            "strategy": inp.get("strategy") or "Core SEO",
            # The RZ category list is multi-select; the FIRST tag is the one the
            # pricing rules key on, so that is the one to group by.
            "industry": _first_industry(inp.get("industry")),
            "industry_rule": pricing.get("industry_rule") or "",
            "industry_anchor_add": pricing.get("industry_anchor_add") or 0,
            "markets": len(inp.get("geo_values") or []),
            "pct_not_ranking": pricing.get("pct_not_ranking"),
            # Attached by the back-measure, absent on any quote it has not run
            # on yet — which is the normal state and reads as "not measured"
            # rather than as a zero.
            # What the formula said BEFORE this quote was last rebuilt. A row
            # showing "the formula runs 26% low" is measuring TODAY's keyword
            # list against the price sent in July, and if the list changed in
            # between that is a fact about the tool, not about the price.
            "formula_was": p.get("formulaWas") or None,
            "median_rival_rank": ((p.get("signals") or {}).get("pageone_rank")
                                  if (p.get("signals") or {}).get("pageone_rank")
                                  is not None
                                  else (p.get("signals") or {}).get("median_rival_rank")),
            "top_rival_rank": (p.get("signals") or {}).get("top_rival_rank"),
            "total_volume": kw.get("total_volume") or pricing.get("total_volume"),
            "formula": {t: pairs[t][0] for t in _TIERS},
            "actual": {t: pairs[t][1] for t in _TIERS},
            "gap": {t: (pairs[t][1] - pairs[t][0])
                    if (pairs[t][0] and pairs[t][1]) else None for t in _TIERS},
            "gap_pct": {t: round(100.0 * (pairs[t][1] - pairs[t][0]) / pairs[t][0], 1)
                        if (pairs[t][0] and pairs[t][1]) else None for t in _TIERS},
            # Step = what each tier adds. Base and step are different constants,
            # and a single quote cannot tell you which one is wrong.
            "step_formula": (pairs["intermediate"][0] - pairs["base"][0])
            if (pairs["base"][0] and pairs["intermediate"][0]) else None,
            "step_actual": (pairs["intermediate"][1] - pairs["base"][1])
            if (pairs["base"][1] and pairs["intermediate"][1]) else None,
        })
    return rows


def calibration_groups(rows):
    """Group the deltas by band + demand basis and summarise each."""
    buckets = {}
    for r in rows:
        key = (r["band"], bool(r["national_demand"]))
        buckets.setdefault(key, []).append(r)
    out = []
    for (band, nat), rs in buckets.items():
        base_gap = _median([r["gap_pct"]["base"] for r in rs])
        out.append({
            "band": band, "national_demand": nat, "n": len(rs),
            "median_gap_pct": {t: _median([r["gap_pct"][t] for r in rs]) for t in _TIERS},
            "median_gap": {t: _median([r["gap"][t] for r in rs]) for t in _TIERS},
            "median_base_formula": _median([r["formula"]["base"] for r in rs]),
            "median_base_actual": _median([r["actual"]["base"] for r in rs]),
            "median_step_formula": _median([r["step_formula"] for r in rs]),
            "median_step_actual": _median([r["step_actual"] for r in rs]),
            "quotes": [r["name"] for r in rs][:8],
            "enough": len(rs) >= CALIB_MIN_N,
            "sort": abs(base_gap or 0),
        })
    out.sort(key=lambda g: (-g["n"], -g["sort"]))
    return out


def calibration_advice(groups, rows, drivers=None):
    """Which constant to touch, and by how much — or explicitly nothing.

    Deliberately conservative. A recommendation needs a group of at least
    CALIB_MIN_N quotes and a median gap past CALIB_MIN_GAP_PCT, and it names the
    constant rather than applying it: these move the price on every future quote
    of that shape.
    """
    tips = []
    strong = [g for g in groups if g["enough"]]

    def edit_line(key, value):
        """The literal line to change in app.py, because the session copy reverts.

        api_config_set() only tunes the RUNNING process — a redeploy restores the
        file. So a recommendation that can only be applied to the session is half
        a recommendation; it ships with the permanent version attached.
        """
        return f'    "{key}": {value},'
    if not strong:
        n = len(rows)
        tips.append({
            "kind": "wait", "constant": "", "from": None, "to": None,
            "text": (f"{n} quote{'' if n == 1 else 's'} with an actual price filled in — "
                     f"not enough to change anything. A group needs {CALIB_MIN_N} quotes "
                     "of the same shape before a gap is a pattern rather than one "
                     "client's judgement call."),
        })
        return tips

    # Is the gap universal, or specific to a shape? If EVERY strong group leans
    # the same way, the floor/anchor is the suspect. If one leans and the others
    # do not, it is that shape's own rate.
    leans = [g for g in strong if abs(g["median_gap_pct"]["base"] or 0) >= CALIB_MIN_GAP_PCT]
    flat = [g for g in strong if g not in leans]

    # THE FLOOR IS GLOBAL. It lifts every quote in every shape, so it can only be
    # the answer when no OTHER shape is already correct AT the floor — otherwise
    # raising it to fix nationwide would push the local quotes that currently match
    # BE exactly straight past him. This is precisely the overfit the view exists
    # to catch, so the check belongs here and not in whoever reads the output.
    floor = CFG.get("client_floor", 2950)
    floor_locked = [g for g in flat
                    if (g["median_base_formula"] or 0) <= floor + 1]
    floor_wants = []

    for g in leans:
        gap = g["median_gap_pct"]["base"]
        d = g["median_gap"]["base"]
        shape = _calib_shape(g)
        on_floor = (g["median_base_formula"] or 0) <= floor + 1 and d and d > 0
        if on_floor and floor_locked:
            tips.append({
                "kind": "blocked", "constant": "", "from": None, "to": None,
                "text": (f"{shape}: median {gap:+.1f}% on base across {g['n']} quotes, and "
                         f"the formula is sitting on the floor (${floor:,}) — but so is "
                         + ", ".join(_calib_shape(x) for x in floor_locked)
                         + ", which already matches. Raising client_floor would fix this "
                           "shape and push those past the mark, so the floor is NOT the "
                           "answer. This needs a rate that only applies to this shape "
                           "(the geo_pct_tiers rung, or a national-demand anchor) — say "
                           "the word and I'll add one."),
            })
        # The band rate is what sets the base on a geo-priced quote; the floor is
        # what sets it when the formula lands under BE's observed minimum.
        elif on_floor:
            # Collected, not emitted: client_floor is ONE number, and two shapes
            # asking it for different values is not two recommendations — it is
            # evidence the floor is the wrong lever for at least one of them.
            floor_wants.append((shape, g,
                                int(round((g["median_base_actual"] or 0) / 50.0) * 50)))
        else:
            tips.append({
                "kind": "band", "constant": "geo_pct_tiers",
                "from": None, "to": None,
                "text": (f"{shape}: median {gap:+.1f}% on base across {g['n']} quotes "
                         f"(${d:+,.0f}). This shape is priced off the band rate, so the "
                         "rung of geo_pct_tiers these quotes land on is the constant to "
                         "move — not the floor, which is not what set these prices."),
            })
        # Base and step are separate constants and can be wrong independently.
        sf, sa = g["median_step_formula"], g["median_step_actual"]
        if sf and sa and abs(sa - sf) >= 100:
            tips.append({
                "kind": "step", "constant": "tier_step_flat",
                "from": CFG.get("tier_step_flat"),
                "to": None,
                "text": (f"{shape}: each tier adds ${sf:,.0f} in the formula and "
                         f"${sa:,.0f} in the actual quotes. That is the tier STEP, not the "
                         "base — a separate constant (tier_step_flat / "
                         "tier_step_pct_of_base). Fixing the base alone would leave "
                         "intermediate and advanced still off."),
            })
    # ---- resolve the floor once ----
    if len(floor_wants) == 1:
        shape, g, want = floor_wants[0]
        tips.append({
            "kind": "floor", "constant": "client_floor", "from": floor, "to": want,
            "apply": {"client_floor": want}, "edit": edit_line("client_floor", want),
            "text": (f"{shape}: the formula is landing ON the floor (${floor:,}) and the "
                     f"actual is ${g['median_base_actual']:,.0f} — median "
                     f"{g['median_gap_pct']['base']:+.1f}% across {g['n']} quotes. The floor "
                     "only ever lifts a price, so raising it cannot disturb any quote "
                     "already above it."),
        })
    elif len(floor_wants) > 1:
        lo = min(w for _s, _g, w in floor_wants)
        detail = "; ".join(f"{sh} wants ${w:,}" for sh, _g, w in floor_wants)
        tips.append({
            "kind": "floor", "constant": "client_floor", "from": floor, "to": lo,
            "apply": {"client_floor": lo}, "edit": edit_line("client_floor", lo),
            "text": (f"Every leaning shape is sitting on the floor, but they disagree on "
                     f"what it should be — {detail}. client_floor is one number, so it can "
                     f"only go to the LOWEST of them (${lo:,}), which is safe for all of "
                     "them and fixes none of them completely. The rest of each gap has to "
                     "come from a rate that applies to that shape alone — raising the floor "
                     "to the highest would overprice the others."),
        })

    # ---- the ranking rung IS the geo rate table ----
    top = next((d for d in (drivers or []) if d.get("spread") is not None), None)
    if top and top["variable"] == "ranking rung" and (top["spread"] or 0) >= CALIB_MIN_GAP_PCT:
        bits = " · ".join(f"{b['value']} {b['median_gap_pct']:+.1f}% (n={b['n']})"
                          for b in top["buckets"] if b["n"] >= 2)
        tips.append({
            "kind": "band", "constant": "geo_pct_tiers", "from": None, "to": None,
            "text": ("The variable that best separates the gaps is how much of the client "
                     f"already RANKS — {bits}. That is not a coincidence: geo_pct_tiers is "
                     "literally a table keyed on that percentage, so these are its rungs "
                     "reading wrong rather than the base being wrong for everyone. Move the "
                     "rungs the leaning quotes land on."
                     + (" Fragile — at least one bucket has fewer than four quotes."
                        if top.get("fragile") else "")),
        })

    if top and top["variable"] in ("industry", "industry rule") \
            and (top["spread"] or 0) >= CALIB_MIN_GAP_PCT:
        bits = " · ".join(f"{b['value']} {b['median_gap_pct']:+.1f}% (n={b['n']})"
                          for b in top["buckets"] if b["n"] >= 2)
        tips.append({
            "kind": "band", "constant": "industry_anchor_add", "from": None, "to": None,
            "text": ("The gaps separate by VERTICAL more than by anything else — "
                     f"{bits}. There is already a per-industry lever for exactly this "
                     "(industry_anchor_add, keyed on the RZ category), so this is that "
                     "rule reading wrong for these categories rather than the base being "
                     "wrong for everyone. Worth checking against BE first: a vertical "
                     "pattern can also just be which clients happened to come in that "
                     "month."
                     + (" Fragile — at least one bucket has fewer than four quotes."
                        if top.get("fragile") else "")),
        })

    for g in flat:
        tips.append({
            "kind": "ok", "constant": "", "from": None, "to": None,
            "text": (f"{_calib_shape(g)}: median {g['median_gap_pct']['base']:+.1f}% across "
                     f"{g['n']} quotes — inside noise. Leave it alone."),
        })
    if leans and flat:
        tips.insert(0, {
            "kind": "note", "constant": "", "from": None, "to": None,
            "text": ("The gap is not universal: "
                     + ", ".join(_calib_shape(g) for g in leans)
                     + " lean, while " + ", ".join(_calib_shape(g) for g in flat)
                     + " are already right. That is the signal to change a "
                       "shape-specific rate rather than the base for everyone."),
        })
    return tips


def _calib_shape(g):
    band = (g["band"] or "unset").replace("_", " ")
    return f"{band}{' + national demand' if g['national_demand'] else ''}"


# --- WHICH VARIABLE EXPLAINS THE SPREAD -------------------------------------
# The groups above slice on band and demand basis, because that is where the
# price anchor comes from. But those are two of at least eight things that could
# be driving a gap, and the rest were being collected and ignored.
#
# The wrong fix is grouping by all of them: with a couple of dozen quotes, a
# four-way slice puts n=1 in every cell and nothing clears the threshold. So each
# candidate is tested ON ITS OWN, against the whole set, which keeps n as large as
# it can be. The question per variable is narrow: split the quotes by it, and do
# the halves disagree? A variable whose buckets all show the same gap explains
# nothing, however plausible it sounded.
#
# It also reports what it CANNOT separate. If every nationwide quote is also
# priced on national demand — which is the normal case — then those two variables
# partition the set identically and no amount of data will tell them apart. Saying
# so is the difference between a finding and a coincidence. (2026-08-12)

def _first_industry(v):
    """The primary RZ category, however it was stored (list, string, blank)."""
    if isinstance(v, (list, tuple)):
        v = next((x for x in v if str(x or "").strip()), "")
    t = str(v or "").strip()
    if not t:
        return "no industry set"
    # "Waste Management/Utilities - Trash / Dumpster Rental" -> the top level, so
    # two dumpster clients group together instead of splitting on the sub-category.
    return t.split(" - ")[0].strip()[:48]


def _rank_rung(pct):
    """Which geo_pct_tiers rung a quote landed on — the table that sets the rate."""
    if pct is None:
        return "no ranking data"
    for lo, rate in CFG.get("geo_pct_tiers") or []:
        if pct >= lo:
            return f"{lo}%+ not ranking ({rate}%)"
    return "unknown"


def _vol_bucket(v):
    """Above or below the first volume bracket — below it, volume adds $0."""
    brackets = CFG.get("volume_brackets") or []
    first = brackets[0][0] if brackets else 10000
    if not v:
        return "no volume data"
    return f"volume over {first:,}/mo" if v >= first else f"volume under {first:,}/mo"


def _mkt_bucket(n):
    return "1 market" if n <= 1 else ("2-5 markets" if n <= 5 else "6+ markets")


# A BUSINESS'S OWN SOCIAL PROFILE IS NOT AN INCUMBENT. These are 1000-authority
# domains that turn up on ordinary local SERPs constantly — an Instagram page, a
# Facebook page, a YouTube video — and none of them is who you outrank by being a
# better dentist. Pennsylvania Center for Dental Excellence has five Philadelphia
# dental sites on its page one, all 143-240, and instagram.com on two terms.
# instagram.com set the strength to 1,000, which put a plainly local practice in
# the national-platform band and quoted it $350 over what Brendan sent.
#
# NOT A GENERAL DIRECTORY BAN. Zillow, Trulia, Apartments.com and Yelp stay in:
# they take the click and they are what a local business is actually competing
# with for it. The line is a profile page you could own yourself. (2026-08-18)
_PAGEONE_NON_RIVAL = frozenset("""
instagram.com facebook.com m.facebook.com youtube.com m.youtube.com
linkedin.com tiktok.com twitter.com x.com pinterest.com reddit.com
threads.net tumblr.com flickr.com vimeo.com
""".split())


def pageone_strength(rivals, min_appearances=2):
    """How strong the strongest REAL incumbent on page one is.

    A MAX, NOT A MEDIAN, because that is the mechanism: one Zillow changes the
    job. NPAIHB's page one is ihs.gov, Wikipedia and Facebook and its median is
    437; Amare's is Zillow, Trulia and Redfin and its median is 597. Brendan
    priced both at $3,550, so a median with a cut between them splits two clients
    he priced identically, while the strongest reads 1,000 for both.

    THE 2+ TERMS IS THE GUARD ON THE MAX. A bare maximum is exactly what a stray
    facebook.com or wikipedia.org result — a SERP artifact rather than a
    competitor — would swing, and that would push a local client into the
    national band on one keyword's noise. A platform holding page one across
    several of the client's terms is not an artifact. Falls back to the plain
    maximum when nothing repeats, so a thin measurement reports something rather
    than nothing. (2026-08-18)
    """
    ranked = [(int(r.get("rank") or 0), int(r.get("appearances") or 0))
              for r in (rivals or [])
              if r.get("rank") is not None
              and str(r.get("domain") or "").lower().replace("www.", "")
              not in _PAGEONE_NON_RIVAL]
    if not ranked:
        return None
    repeated = [v for v, n in ranked if n >= int(min_appearances)]
    return max(repeated) if repeated else max(v for v, _ in ranked)


def _pageone_bucket(rank):
    """Who already holds page one, by the median incumbent's backlink authority.

    THE MEDIAN, NOT THE GAP. The gap needs the client's own authority, and
    bulk_ranks writes 0 both for a domain with no backlinks and for one it has no
    data on — Amare read 0 and would have produced the largest gap in the book
    off the weakest evidence in it. The median incumbent needs no reading of the
    client at all.

    THE BANDS ARE NOT FITTED TO THE THREE CLIENTS MEASURED BY HAND. 126, 437 and
    565 would each land in their own bucket at almost any cut, which is how a
    variable looks like a perfect driver on n=3 — four of the eight inputs
    already tested "fit" that same single boundary, including organic difficulty,
    which had been ruled out by experiment the same morning. These are round
    thirds of DataForSEO's 0-1000 scale, and the panel's existing guards
    (min_bucket, the fragile flag) decide whether they have earned anything.
    (2026-08-17)
    """
    # None, not a label. A quote the back-measure has not run on is not a bucket
    # of its own — it would be the biggest bucket in the panel for weeks, and a
    # "spread" between measured and unmeasured quotes measures nothing but which
    # ones got measured first.
    if rank is None:
        return None
    if rank < 200:
        return "page one: local businesses (under 200)"
    if rank < 400:
        return "page one: regional or institutional (200-399)"
    return "page one: national platforms (400+)"


CALIB_DRIVERS = [
    ("geo band", lambda r: (r["band"] or "unset").replace("_", " ")),
    ("demand basis", lambda r: "national" if r["national_demand"] else "local"),
    ("ranking rung", lambda r: _rank_rung(r.get("pct_not_ranking"))),
    ("strategy", lambda r: r.get("strategy") or "Core SEO"),
    ("search volume", lambda r: _vol_bucket(r.get("total_volume"))),
    ("market count", lambda r: _mkt_bucket(int(r.get("markets") or 0))),
    # WHO IS ON PAGE ONE — the ninth variable, tested exactly like the other
    # eight. Three clients measured by hand put it in the right order against
    # Brendan's prices (126/437/565 against $2,950/$3,550/$3,550) and it was the
    # first input that ever separated them. That is a reason to MEASURE it across
    # the book, not to believe it: this panel already knows how to say "n is too
    # small" and how to spot one outlier wearing several hats, and a variable
    # read off a table by eye gets neither. (2026-08-17)
    ("page one", lambda r: _pageone_bucket(r.get("median_rival_rank"))),
    # INDUSTRY, tested the same way as everything else rather than assumed. There
    # is already an industry pricing rule (industry_anchor_add), so if BE's prices
    # move by vertical this is the variable that will show it — and if they do not,
    # this says that too, which is the more useful answer for a rule nobody wants
    # to maintain per category. (2026-08-12)
    ("industry", lambda r: r.get("industry") or "no industry set"),
    ("industry rule", lambda r: r.get("industry_rule") or "no rule matched"),
]


def calibration_drivers(rows, min_bucket=2):
    """Per variable: does splitting the quotes on it separate the gaps?

    `spread` is the distance between the best and worst bucket median, counting
    only buckets with at least `min_bucket` quotes. Big spread = this variable is
    where the gap lives. Near zero = it explains nothing.
    """
    out = []
    for label, fn in CALIB_DRIVERS:
        buckets = {}
        for r in rows:
            if r["gap_pct"]["base"] is None:
                continue
            # A driver returns None for a quote it cannot classify, and that
            # quote sits the test out rather than forming an "unknown" bucket —
            # the difference between measured and not-yet-measured quotes is a
            # fact about the measuring, not about the price.
            key = fn(r)
            if key is None:
                continue
            buckets.setdefault(key, []).append(r["gap_pct"]["base"])
        shown = [{"value": k, "n": len(v), "median_gap_pct": _median(v)}
                 for k, v in buckets.items()]
        # The ranking rung is an ORDERED scale, so listing it by gap printed the
        # rungs 0, 40, 90, 70 and the trend in it was unreadable. Ordered variables
        # sort by their own value; the rest by size of gap.
        if label == "ranking rung":
            shown.sort(key=lambda b: -_leading_num(b["value"]))
        else:
            shown.sort(key=lambda b: -(b["median_gap_pct"] or 0))
        solid = [b for b in shown if b["n"] >= min_bucket]
        spread = (max(b["median_gap_pct"] for b in solid)
                  - min(b["median_gap_pct"] for b in solid)) if len(solid) > 1 else None
        # A median over three quotes moves if one of them moves, so a variable can
        # look like a perfect driver on composition alone. Flagged rather than
        # hidden: the split may still be real, it just is not yet evidence.
        fragile = bool(solid) and min(b["n"] for b in solid) < 4
        out.append({"variable": label, "buckets": shown,
                    "testable": len(solid) > 1, "spread": spread,
                    "fragile": fragile,
                    # One bucket holding everything means the variable does not
                    # vary in this data at all — nothing to learn either way.
                    "constant_here": len(shown) <= 1})
    # SOLID BEFORE FRAGILE. Sorting on spread alone put a three-quote split at the
    # top of the panel above a ten-quote one, and the loudest line was one outlier
    # wearing several hats — Ski Barn's single statewide quote at -41.9% was
    # driving four of the six splits. The fragile flag said so in small text under
    # a heading that had already made the point. Order says it instead.
    out.sort(key=lambda d: (d["spread"] is None, bool(d.get("fragile")),
                            -(d["spread"] or 0)))
    return out


def _leading_num(text):
    m = re.match(r"\s*(-?\d+(?:\.\d+)?)", str(text or ""))
    return float(m.group(1)) if m else -1e9


def calibration_outliers(rows, factor=2.5):
    """Quotes whose gap sits far from every other quote's.

    A median is robust to an outlier only when the bucket is big enough to have
    one. With three quotes in a split, one extreme quote IS the median, so a
    single unusual deal can appear as evidence in four different variables at
    once. Named here with the splits it lands in, so it can be checked or
    excluded rather than silently believed. (2026-08-12)
    """
    gaps = [(r, abs(r["gap_pct"]["base"])) for r in rows
            if r["gap_pct"]["base"] is not None]
    if len(gaps) < 4:
        return []
    typical = _median([g for _r, g in gaps]) or 0.0
    # When most quotes match exactly the median gap is ~0, and dividing by it
    # produces "20x the typical gap" against a typical of nothing. Floor it for the
    # comparison and drop the multiple from the output when it is meaningless.
    ratio_meaningful = typical >= 2.0
    if typical < 1.0:
        typical = 1.0
    out = []
    for r, g in gaps:
        if g >= typical * factor and g >= CALIB_MIN_GAP_PCT:
            out.append({
                "name": r["name"] or r["client"],
                "gap_pct": r["gap_pct"]["base"],
                "shape": f"{(r['band'] or 'unset').replace('_', ' ')}"
                         + (" + national demand" if r["national_demand"] else ""),
                "times_typical": round(g / typical, 1) if ratio_meaningful else None,
                # WHICH BUCKET, per variable. Four stacked warning cards repeating
                # the same forty words is not four findings — it is one fact the
                # reader has to carry to the numbers themselves. Sent as a map so
                # the panel can mark the exact bucket each outlier sits in.
                # (2026-08-12)
                "buckets": {label: fn(r) for label, fn in CALIB_DRIVERS},
            })
    out.sort(key=lambda o: -abs(o["gap_pct"]))
    return out


def calibration_confounds(rows):
    """Pairs of variables that partition these quotes identically.

    Two variables that always move together cannot be told apart no matter how
    many quotes arrive — the data has no case where one changes and the other
    does not. Reported so a finding attributed to one is not silently a finding
    about the other.
    """
    if len(rows) < 2:
        return []
    sigs = {}
    for label, fn in CALIB_DRIVERS:
        vals = [fn(r) for r in rows]
        # Partition signature: which rows share a value, independent of the names.
        groups = {}
        for i, v in enumerate(vals):
            groups.setdefault(v, []).append(i)
        if len(groups) <= 1:
            continue                     # constant here, nothing to confound with
        sigs[label] = frozenset(frozenset(g) for g in groups.values())
    # Reported as CLASSES, not pairs. Five variables that all split the set the
    # same way produce ten pairs and one fact, and the ten pairs bury it.
    classes = {}
    for label, sig in sigs.items():
        classes.setdefault(sig, []).append(label)
    return [{"variables": sorted(v), "n_groups": len(set(sig))}
            for sig, v in classes.items() if len(v) > 1]

@app.route("/api/config_state")
@_json_error_guard
def api_config_state():
    """Has the running config been tuned away from the file this session?"""
    keys = sorted({k for e in CFG_EDITS for k in e["keys"]})
    return jsonify({"edited": bool(CFG_EDITS), "n": len(CFG_EDITS), "keys": keys,
                    "note": (CFG_EDITS[-1]["note"] if CFG_EDITS else ""),
                    "values": {k: CFG.get(k) for k in keys}})


@app.route("/api/calibration")
@_json_error_guard
def api_calibration():
    """Formula vs what was actually sent, across every saved quote."""
    if not storage.enabled():
        return jsonify({"rows": [], "groups": [], "advice": [], "error": "saving is off"})
    rows = calibration_rows(storage.all_payloads("seo"))
    groups = calibration_groups(rows)
    return jsonify({"rows": rows, "groups": groups,
                    "drivers": calibration_drivers(rows),
                    "confounds": calibration_confounds(rows),
                    "outliers": calibration_outliers(rows),
                    "advice": calibration_advice(groups, rows,
                                                  calibration_drivers(rows)),
                    "min_n": CALIB_MIN_N,
                    "constants": {"client_floor": CFG.get("client_floor"),
                                  "tier_step_flat": CFG.get("tier_step_flat"),
                                  "tier_step_pct_of_base": CFG.get("tier_step_pct_of_base"),
                                  "geo_pct_tiers": CFG.get("geo_pct_tiers")}})


@app.route("/api/rank_seeds", methods=["POST"])
@_json_error_guard
def api_rank_seeds():
    """Fold synonym seeds and rank the rest by measured demand. Preview only —
    the operator applies it."""
    d = request.get_json(force=True) or {}
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    nat = bool(d.get("national_demand")) or not markets
    seeds = [x for x in (d.get("seeds") or []) if x]
    kinds = {}
    if not d.get("skip_kinds"):
        dom = (d.get("domain") or "").strip()
        pages, urls = [], []
        if dom:
            try:
                pages = fetch_site_pages(dom, collect_urls=urls) or []
            except Exception:
                pages, urls = [], []
        # Two independent reads on "do they sell things": the operator's own
        # national-product-demand switch, and a storefront in the URL structure.
        ecom, _why = detect_ecommerce(urls)
        kinds = claude_seed_kinds(seeds, d.get("brand") or "", dom,
                                  d.get("industry") or "",
                                  d.get("business_desc") or "", pages,
                                  # NOT `nat` — that is true whenever no geos
                                  # are entered, which would switch the check off
                                  # for a local hauler mid-setup. Only the
                                  # operator's explicit product-demand tick counts.
                                  sells_products=bool(d.get("national_demand")
                                                      or ecom))
    return jsonify(rank_seeds(seeds, markets, state, national=nat,
                              limit=d.get("limit"), kinds=kinds))


@app.route("/api/expand_services", methods=["POST"])
@_json_error_guard
def api_expand_services():
    """Service lines the seed list is missing — proposed, then measured."""
    d = request.get_json(force=True) or {}
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    nat = bool(d.get("national_demand")) or not markets
    dom = (d.get("domain") or "").strip()
    pages = []
    if dom:
        try:
            pages = fetch_site_pages(dom) or []
        except Exception:
            pages = []
    cands = claude_industry_services(
        d.get("brand") or "", dom, d.get("industry") or "",
        d.get("business_desc") or "", pages,
        [x for x in (d.get("seeds") or []) if x],
        ", ".join(markets[:4]))
    if not cands:
        # NOT the same as "nothing cleared the volume floor" — the floor was never
        # reached. Amare Homes read that line while the gap-finder was the ONLY
        # source left (their site 403s), so the one thing that could still help
        # failed silently. Name the input that is missing. (2026-08-13)
        missing = [lbl for lbl, v in (("Industry", d.get("industry")),
                                      ("Business description", d.get("business_desc")))
                   if not (v or "").strip()]
        return jsonify({"services": [], "no_candidates": True,
                        "missing_inputs": missing,
                        "error": ("nothing was proposed"
                                  + (" — no " + " or ".join(missing) + " to go on"
                                     if missing else ""))})
    terms = [c["term"] for c in cands]
    try:
        vols, _pc, verr = fetch_local_volume(terms, [] if nat else markets, state,
                                             national=nat)
    except Exception as e:
        return jsonify({"services": [], "error": str(e)[:120]})
    floor = int(CFG.get("expand_min_volume", 20))
    rows = [{"term": c["term"], "why": c["why"],
             "volume": int((vols or {}).get(c["term"], 0) or 0)} for c in cands]
    rows.sort(key=lambda r: -r["volume"])
    good, folded_rows = fold_proposals([r for r in rows if r["volume"] >= floor],
                                       seeds=(d.get("seeds") or []),
                                       markets=([] if nat else markets), state=state)
    # The prompt is now asked for 22 rather than 14, and a model asked for more
    # than the vertical has will pad with rewordings — which is the one thing this
    # pass exists NOT to do. fold_proposals catches near-duplicates by
    # containment; this catches the fourth distinct wording of one procedure,
    # which containment does not. Checked in code because a prompt rule cannot
    # be. (2026-08-17)
    _gkept, _gfam = cap_service_family([r["term"] for r in good],
                                       seeds=(d.get("seeds") or []),
                                       markets=([] if nat else markets), state=state)
    if _gfam:
        _gset = set(_gkept)
        good = [r for r in good if r["term"] in _gset]
    # IS THIS A THIN MARKET, OR A THIN PROPOSAL SET? The panel could already relax
    # the floor when the CLIENT'S OWN built list came in under it — but on Amare
    # Homes no build existed yet, which is exactly when the operator is assembling
    # seeds and the floor is silently withholding the whole vertical. It offered
    # one term (apartments for rent, 90/mo) and held back the ten Brendan actually
    # quoted on.
    #
    # The candidate pool answers it with no build needed. Two conditions, and the
    # second is what stops a FAT client's junk tail triggering this: the typical
    # candidate must be under the floor, AND the best term the vertical has in
    # this market must itself be small. Santa Fe apartments top out at 90/mo, so
    # 10/mo is the tail of a small market. Junk removal in Knoxville tops out at
    # 2,400/mo, so 10/mo is noise and the floor stands. (2026-08-13)
    _v = sorted(r["volume"] for r in rows)
    _typ = _best = 0
    thin_market = False
    if _v:
        _typ = _v[len(_v) // 2]
        _best = _v[-1]
        _mult = float(CFG.get("expand_thin_market_mult", 10) or 10)
        thin_market = bool(_typ < floor and _best < floor * _mult)
    return jsonify({"services": good,
                    "folded": [r["term"] for r in folded_rows],
                    "family_capped": [{"term": t, "on": w} for t, w in _gfam],
                    "asked_for": int(CFG.get("industry_gap_n", 22) or 22),
                    "proposed": len(rows),
                    "rejected": [r for r in rows if r["volume"] < floor],
                    # The panel has to be able to ask "is this floor stricter than
                    # the quote it is protecting?" — NPAIHB's own terms run at
                    # 8/mo and this refused eight at under 20. (2026-08-12)
                    "floor": floor,
                    "thin_market": thin_market,
                    "market_typical": _typ, "market_best": _best,
                    "basis": "US national" if nat else "targeted cities",
                    "error": verr})



def _split_proposal_kinds(items, d, dom, pages=None, ecom=False, alias=None):
    """Run claude_seed_kinds over PROPOSED chips and split them in two.

    A chip is a suggestion, so a wrong verdict costs nothing here — the operator
    still sees the term and its reason and can add it anyway. What it buys is that
    "+ Add all 9" cannot sweep nine member tribes into the focus list, which is
    exactly how NPAIHB's quote was built. Same classifier the build now runs, so
    the two cannot disagree. (2026-08-12)

    `alias` maps a chip's term to a LONGER name to be judged under. "CTWS" means
    nothing to a classifier; "Confederated Tribes of Warm Springs" is the whole
    answer. The acronym miner reads expansions off the page already and was the
    one proposal source nothing vetted — it offered NPAIHB `ctws` at 260/mo, the
    operator added it, and the build then set it aside as a member tribe. Both
    calls were right and the operator did the work twice. (2026-08-12)

    Returns (services, not_services). Degrades to (items, []) on any failure.
    """
    alias = alias or {}
    terms, back = [], {}
    for x in items or []:
        t = (x.get("term") or x.get("label") or "")
        judged = alias.get(t) or alias.get(str(t).strip().lower()) or t
        if judged and str(judged).strip().lower() not in back:
            terms.append(judged)
        back.setdefault(str(t).strip().lower(), str(judged).strip().lower())
    if not terms:
        return list(items or []), []
    try:
        kinds = claude_seed_kinds(terms, d.get("brand") or "", dom,
                                  d.get("industry") or "",
                                  d.get("business_desc") or "", pages or [],
                                  sells_products=bool(d.get("national_demand")
                                                      or ecom))
    except Exception:                                     # noqa: BLE001
        app.logger.exception("claude_seed_kinds failed on proposals")
        return list(items or []), []
    if not kinds:
        return list(items or []), []
    good, bad = [], []
    for x in items or []:
        t = str(x.get("term") or x.get("label") or "").strip().lower()
        k = kinds.get(t) or kinds.get(back.get(t, t))
        if k:
            bad.append(dict(x, kind=k.get("kind", "item"),
                            kind_why=k.get("why", "")))
        else:
            good.append(x)
    # Every chip rejected means the classifier is what is wrong — a menu read that
    # returns nothing is worse than a menu read with a caveat on it.
    if not good:
        return list(items or []), []
    return good, bad


def serp_snippets(query, loc, limit=8):
    """Top organic titles + snippets for a query.

    What everyone who has already described this business wrote about it — their
    Google Business listing, the directories, the local press. One SERP call.
    (2026-08-13)
    """
    payload = [{"keyword": query, "location_name": loc, "language_code": "en",
                "depth": 10}]
    data = dfs_post("/serp/google/organic/live/regular", payload, timeout=18)
    task0 = ((data or {}).get("tasks") or [{}])[0] or {}
    if task0.get("status_code") not in (20000, None):
        raise RuntimeError(f"{task0.get('status_code')}: {task0.get('status_message')}")
    items = ((task0.get("result") or [{}])[0] or {}).get("items") or []
    out = []
    for it in items:
        if it.get("type") != "organic":
            continue
        out.append({"title": (it.get("title") or "")[:160],
                    "snippet": (it.get("description") or "")[:320],
                    "domain": (it.get("domain") or "").lower(),
                    "url": (it.get("url") or "")[:300]})
        if len(out) >= int(limit):
            break
    return out


def claude_business_desc(brand, markets, industry, snippets, domain=""):
    """What this client IS, read off the SERP for their own name.

    Amare Homes' site 403s every request, and the business description is the one
    input the industry pass cannot work without — so a client whose server blocks
    us had no route to a keyword list at all. Their name and their city are
    already on the order, and the answer is on the first page of Google.

    GROUNDED, NOT RECALLED. The model is given real current search results and
    told to describe only what they say. Asking it to remember a specific small
    business invents a plausible one, which is worse than nothing here: the
    description feeds service expansion, so a confident wrong sentence produces a
    confident wrong quote. It must return low confidence and say so rather than
    guess, and the caller only ever OFFERS the result. (2026-08-13)

    Returns {"text", "confidence": high|low, "why", "sources": [url]} or {}.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or not snippets:
        return {}
    where = ", ".join(markets[:3]) if markets else ""
    listing = "\n".join(
        f"- {s.get('domain','')}: {s.get('title','')} — {s.get('snippet','')}"
        for s in snippets[:8])
    prompt = f"""Below are the top Google results for a business's own name.
Write ONE sentence saying what the business is and what it sells, for use as the
"business description" field on an SEO quote.

Business name: {brand}
{f'Website: {domain}' if domain else ''}
{f'Market: {where}' if where else ''}
{f'Industry (as tagged by the partner): {industry}' if industry else ''}

SEARCH RESULTS
{listing}

Rules:
- Use ONLY what these results say. Do not add anything you happen to know about
  this company or infer from its name.
- If the results are not clearly about THIS business in THIS market - they name a
  different company, a directory page with no detail, or a similarly-named
  business elsewhere - set confidence "low" and say what is missing in "why".
  A wrong description produces a wrong keyword list, so an honest "not sure" is
  worth more than a plausible sentence.
- Name the specific things sold, not the category. "1-3 bedroom single-family
  rental homes with garages" beats "residential property services".
- One sentence, under 40 words, plain and factual. No marketing adjectives.

Return ONLY JSON:
{{"text": "...", "confidence": "high", "why": "", "sources": ["domain.com"]}}"""
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={"x-api-key": api_key, "anthropic-version": "2023-06-01",
                     "content-type": "application/json"},
            data=json.dumps({
                "model": os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-6"),
                "max_tokens": 900, "temperature": 0,
                "messages": [{"role": "user", "content": prompt}]}), timeout=60)
        resp.raise_for_status()
        body = resp.json()
        text = "".join(b.get("text", "") for b in body.get("content", [])
                       if b.get("type") == "text").strip()
        text = re.sub(r"^```(?:json)?|```$", "", text, flags=re.M).strip()
        got = json.loads(text) or {}
    except Exception:
        app.logger.exception("claude_business_desc failed")
        return {}
    desc = re.sub(r"\s+", " ", str(got.get("text") or "")).strip()[:400]
    if not desc:
        return {}
    # A CTA is not a description — same refusal the site meta-description path uses.
    if _is_cta(desc):
        return {}
    return {"text": desc,
            "confidence": ("high" if str(got.get("confidence") or "").lower()
                           == "high" else "low"),
            "why": str(got.get("why") or "")[:200],
            "sources": [str(x)[:120] for x in (got.get("sources") or [])][:5]}


def fetch_ranked_keywords(domain, markets=None, state="", limit=None):
    """What the client ALREADY RANKS FOR — no keywords needed as input.

    The whole tool has run the other way round: guess terms, then check whether
    they rank. Brendan's Amare list has "gated community homes for rent santa fe
    nm" at #21 and "homes for rent with garage santa fe nm" at #19, and nobody
    guesses those — they are read off a ranked-keywords report. Four of his twenty
    already ranked, which is also why his nouns were right: the client's real
    positions say "homes for rent", not "apartments".

    Labs, so location_code and not location_name — 2840 is the US, and the city
    stays INSIDE the returned phrase, which is the shape a local list wants.
    Returns the phrase both as-found and geo-stripped, because the grid appends a
    city of its own and would otherwise double it. (2026-08-13)

    Returns [{"term", "bare", "position", "volume", "url"}], best position first.
    """
    dom = re.sub(r"^https?://", "", (domain or "").strip()).strip("/")
    dom = re.sub(r"^www\.", "", dom)
    if not dom:
        return []
    lim = int(limit or CFG.get("ranked_keywords_limit", 80))
    payload = [{"target": dom, "location_code": 2840, "language_code": "en",
                "limit": lim, "load_rank_absolute": True,
                "order_by": ["ranked_serp_element.serp_item.rank_group,asc"],
                "filters": [["ranked_serp_element.serp_item.rank_group", "<=",
                             int(CFG.get("zero_ranking_top_n", 100))]]}]
    data = dfs_post("/dataforseo_labs/google/ranked_keywords/live", payload,
                    timeout=25)
    task0 = ((data or {}).get("tasks") or [{}])[0] or {}
    if task0.get("status_code") not in (20000, None):
        raise RuntimeError(f"{task0.get('status_code')}: {task0.get('status_message')}")
    out, seen = [], set()
    for block in (task0.get("result") or []):
        for it in (block.get("items") or []):
            kd = (it or {}).get("keyword_data") or {}
            kw = str(kd.get("keyword") or "").strip().lower()
            if not kw or kw in seen:
                continue
            seen.add(kw)
            se = ((it.get("ranked_serp_element") or {}).get("serp_item") or {})
            pos = se.get("rank_group") or se.get("rank_absolute")
            bare = seed_norm(kw, markets or [], state) or kw
            out.append({"term": kw, "bare": bare,
                        "position": int(pos) if pos else None,
                        "volume": int(((kd.get("keyword_info") or {})
                                       .get("search_volume")) or 0),
                        "url": str(se.get("url") or "")[:300]})
    out.sort(key=lambda r: (r["position"] or 999, -r["volume"]))
    return out


@app.route("/api/ranked_keywords", methods=["POST"])
@_json_error_guard
def api_ranked_keywords():
    """Seeds read off the client's OWN positions. Needs a domain and nothing else."""
    d = request.get_json(force=True) or {}
    dom = (d.get("domain") or "").strip()
    if not dom:
        return jsonify({"error": "Add the client website first."}), 400
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    try:
        rows = fetch_ranked_keywords(dom, markets, state, d.get("limit"))
    except Exception as e:                                # noqa: BLE001
        return jsonify({"error": str(e)[:160]}), 502
    have = {seed_norm(x, markets, state) for x in (d.get("seeds") or []) if x}
    # THEIR OWN NAME IS NOT A KEYWORD TO WIN. Amare ranks #4 for "amare santa fe"
    # — there is no work to sell there, and seeding it would add its volume to a
    # total the price is computed from. (2026-08-13)
    _bw = {w for w in re.split(r"[^a-z0-9]+", (d.get("brand") or "").lower())
           if len(w) > 2 and w not in _GROUNDING_STOP}
    def _is_brand(term):
        toks = [t for t in re.split(r"[^a-z0-9]+", term) if t]
        return bool(toks) and bool(_bw) and all(t in _bw for t in toks)
    fresh = [r for r in rows if r["bare"] not in have and not _is_brand(r["bare"])]
    brandy = [r["bare"] for r in rows if _is_brand(r["bare"])]
    # NINE OVERLAY VARIANTS, ALL OF THEM REAL POSITIONS. This is the source the
    # family cap was missing, and it is the worst place to miss it: a position is
    # treated as proof of ownership, so these chips sort AHEAD of everything the
    # other two panels found, and Nob Hill Dental spent six of thirteen grid
    # slots on one procedure. The client genuinely ranks for all nine — that is
    # not the question. Selling the fourth wording of a service you are already
    # selling the first three wordings of is not a campaign.
    #
    # The operator's own typed terms are exempt and are what mark a family as the
    # business rather than a runaway, so a rental community whose whole list is
    # homes-for-rent variants loses nothing. (2026-08-17)
    _kept, _famdrop = cap_service_family(
        [r["bare"] for r in fresh], seeds=(d.get("seeds") or []),
        markets=markets, state=state)
    if _famdrop:
        _keepset = set(_kept)
        fresh = [r for r in fresh if r["bare"] in _keepset]
    if d.get("own"):
        # THE OWNERSHIP QUESTION WANTS THE RAW SET. No family cap, no seed
        # filter, no brand filter — every term the domain ranks for, bare, so
        # the caller can ask "is this one of them" and get a true answer.
        return jsonify({"owned": sorted({r["bare"] for r in rows if r.get("bare")}),
                        "total": len(rows)})
    return jsonify({"keywords": fresh, "total": len(rows),
                    "brand_terms": brandy,
                    "family_capped": [{"term": t, "on": w} for t, w in _famdrop],
                    "already_seeded": len([r for r in rows
                                           if r["bare"] in have]),
                    "top20": sum(1 for r in rows
                                 if (r["position"] or 999) <= 20)})


@app.route("/api/describe_client", methods=["POST"])
@_json_error_guard
def api_describe_client():
    """Business description from the client's NAME and MARKET — no site access.

    The fallback for a site that refuses to be read. Offered to the operator with
    its sources and its own confidence attached; never written into the field by
    this endpoint.
    """
    d = request.get_json(force=True) or {}
    brand = (d.get("brand") or "").strip()
    if not brand:
        return jsonify({"error": "Add the client name first."}), 400
    markets = usable_markets(d.get("geo_values") or [])
    state = derive_state(markets, (d.get("state") or "").strip())
    dom = re.sub(r"^https?://", "", (d.get("domain") or "").strip()).strip("/")
    loc = loc_string(markets, state) if markets else "United States"
    # The name plus the place, which is what a person would type. The market goes
    # in the QUERY as well as the location, because a national location with a
    # bare brand name finds the biggest company sharing it.
    query = f"{brand} {markets[0].split(',')[0]}" if markets else brand
    try:
        snippets = serp_snippets(query, loc)
    except Exception as e:                                # noqa: BLE001
        return jsonify({"error": f"the search lookup failed: {str(e)[:140]}"}), 502
    if not snippets:
        return jsonify({"query": query, "location": loc, "results": 0,
                        "error": "nothing came back for that name and market."})
    got = claude_business_desc(brand, markets, d.get("industry") or "",
                              snippets, dom)
    if not got:
        return jsonify({"query": query, "location": loc,
                        "results": len(snippets),
                        "error": "the results didn't describe a business clearly."})
    got.update({"query": query, "location": loc, "results": len(snippets),
                "domains": [s["domain"] for s in snippets[:5]]})
    return jsonify(got)


@app.route("/api/site_services", methods=["POST"])
@_json_error_guard
def api_site_services():
    """Parse the client site's navigation into candidate service terms. Menu
    items are how the business describes what it sells — often a better seed
    list than anything a partner types in freehand."""
    d = request.get_json(force=True) or {}
    dom = re.sub(r"^https?://", "", (d.get("domain") or "").strip()).strip("/")
    pasted = (d.get("pasted") or "").strip()
    if not dom and not pasted:
        return jsonify({"error": "Add the client website first."}), 400

    if pasted:
        # Manual escape hatch for sites that block automated access entirely:
        # the partner pastes the menu / service list (one per line or
        # comma-separated) and it runs through the same cleanup + AI conversion
        # as a parsed nav would.
        raw = [p.strip(" \t•·-–—>") for chunk in pasted.splitlines()
               for p in chunk.split(",")]
        out, seen = [], set()
        for t in raw:
            t = re.sub(r"[»›→▸▾▼+]+$", "", t).strip()
            tl = t.lower()
            if not t or tl in _MENU_GENERIC or len(t) > 48 or len(t.split()) > 6:
                continue
            if tl in seen:
                continue
            seen.add(tl)
            out.append({"label": t, "source": "pasted", "service_path": False})
        out = out[:40]
        seeds = [x for x in (d.get("seeds") or []) if isinstance(x, str)]
        mapping = claude_menu_to_terms([x["label"] for x in out],
                                       d.get("brand") or "", dom or "(pasted list)",
                                       seeds, d.get("business_desc") or "")
        ai_used = bool(mapping)
        if ai_used:
            conv, seen_t = [], set()
            for x in out:
                term = mapping.get(x["label"], x["label"].lower())
                if term is None or term in seen_t:
                    continue
                seen_t.add(term); x["term"] = term; conv.append(x)
            out = conv
        else:
            for x in out:
                x["term"] = x["label"].lower()
        # 39 chips off a retail menu is not 39 services — see fold_proposals().
        out, folded_out = fold_proposals(out, seeds=(d.get("seeds") or []))
        # NO MORE THAN A FEW VARIANTS OF ONE PROCEDURE — the prompt has
        # asked for this since July and a prompt cannot be checked.
        # See cap_service_family. (2026-08-17)
        out, family_out = cap_service_family(out, seeds=(d.get("seeds") or []))
        out, not_svc = _split_proposal_kinds(out, d, dom or "(pasted list)")
        return jsonify({"domain": dom, "services": out,
                        "not_services": not_svc,
                        "folded": [(x.get("term") or x.get("label") or "")
                                   for x in folded_out],
                        "family_capped": [{"term": t, "on": w}
                                          for t, w in (family_out or [])],
                        "ai_refined": ai_used,
                        "from_sitemap": False, "pasted": True, "n_nav_links": 0})
    # Two identities: some servers stub out bots, others' WAFs block a Chrome UA
    # that lacks full browser fingerprints while allowing honest bots through.
    # Try both per URL and keep whichever returns a page with real links.
    # A BARE UA STRING IS NOT A BROWSER FINGERPRINT. The comment above has said
    # for weeks that WAFs block a Chrome UA "that lacks full browser
    # fingerprints", and both identities sent exactly that: one header. Cloudflare
    # 403'd both on santafe.amare-homes.com while the page served fine to other
    # clients. The third identity sends what a real Chrome navigation sends —
    # Sec-Fetch-*, Upgrade-Insecure-Requests, a full Accept and Accept-Encoding —
    # which is the cheap half of getting past a default WAF ruleset. (2026-08-13)
    _CHROME = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
               "(KHTML, like Gecko) Chrome/126.0 Safari/537.36")
    _UAS = [
        ("browser", {"User-Agent": _CHROME}),
        ("bot", {"User-Agent": "Mozilla/5.0 (compatible; adtini-seo-quote/1.0)"}),
        ("chrome-full", {
            "User-Agent": _CHROME,
            "Accept": ("text/html,application/xhtml+xml,application/xml;q=0.9,"
                       "image/avif,image/webp,*/*;q=0.8"),
            "Accept-Encoding": "gzip, deflate, br",
            "Accept-Language": "en-US,en;q=0.9",
            "Upgrade-Insecure-Requests": "1",
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "none",
            "Sec-Fetch-User": "?1",
            "Sec-Ch-Ua": '"Chromium";v="126", "Not:A-Brand";v="24"',
            "Sec-Ch-Ua-Mobile": "?0",
            "Sec-Ch-Ua-Platform": '"Windows"',
            "Cache-Control": "max-age=0",
            "Connection": "keep-alive",
        }),
    ]
    html = ""
    fetch_err = None
    diag = []
    statuses = []                 # every HTTP code seen, so the panel can explain
    # try both host variants regardless of how the pill was entered — and never
    # double the www. prefix (www.www.example.org is how that bug looks)
    bare = re.sub(r"^www\.", "", dom)
    for url in dict.fromkeys([f"https://{dom}", f"https://{bare}", f"https://www.{bare}"]):
        for ua_name, hdrs in _UAS:
            try:
                r, insecure = get_client_site(url, timeout=10, allow_redirects=True,
                                              headers=dict(hdrs))
                candidate = r.text[:800_000]
                nlinks = candidate.lower().count("<a")
                statuses.append(r.status_code)
                diag.append(f"{url} [{ua_name}] -> HTTP {r.status_code}, {nlinks} links"
                            + (" (TLS chain incomplete — read without verifying)"
                               if insecure else ""))
                # A 403 BODY IS THE WAF'S PAGE, NOT THE CLIENT'S. It used to be
                # accepted as `html`, so the heading miner scraped Cloudflare's
                # block page and reported "no menu items found" — describing a
                # document the client never wrote. Only a 2xx is the site.
                # (2026-08-13)
                if not (200 <= r.status_code < 300):
                    continue
                if nlinks >= 5:
                    html = candidate
                    break
                if not html:
                    html = candidate
            except Exception as e:
                fetch_err = e
                diag.append(f"{url} [{ua_name}] -> {type(e).__name__}")
        if html and html.lower().count("<a") >= 5:
            break
    if not html:
        # THE ADVICE HAS TO MATCH THE FAILURE. A 403 is an explicit refusal and no
        # header trick the operator can perform will change it; the panel was
        # telling them "links showing 0 with HTTP 200 means the site cloaks bots"
        # under four lines that said 403 and SSLError. (2026-08-13)
        # ONE PATH, AND IT ALWAYS KEEPS THE ESCAPE HATCH. A 403 used to 502 into a
        # red box with no way forward; the panel that offers "paste their menu
        # instead" only renders on a 200 with an empty list. So a server that
        # ANSWERED — refused, missing, erroring — comes back 200 with no services
        # and the codes attached, and the operator still gets the paste box. Only a
        # domain that could not be reached at all is a hard failure. (2026-08-13)
        if statuses:
            return jsonify({"domain": dom, "services": [], "folded": [],
                            "acronyms": [], "n_nav_links": 0,
                            "statuses": statuses, "diag": diag})
        _fe = str(fetch_err)
        if "CERTIFICATE_VERIFY_FAILED" in _fe or "SSLError" in _fe:
            _fe = ("the site's HTTPS certificate chain is incomplete, and reading it "
                   "without verification also failed. Browsers hide this; a server does "
                   "not. Type the services into Keyword / Vertical Focus by hand — "
                   "nothing else in the quote depends on reading the site.")
        elif "Max retries" in _fe or "timed out" in _fe.lower():
            _fe = ("the site didn't respond in time — it may be blocking automated "
                   "requests. Enter the services by hand and continue.")
        return jsonify({"error": f"Couldn't fetch the site: {_fe}",
                        "diag": diag}), 502
    p = _NavLinkParser()
    try:
        p.feed(html)
    except Exception:
        pass

    # The homepage's own self-description — meta description (or og:description)
    # — is the business's one-line answer to "what are you?", which is exactly
    # what the business-description field wants. Offered as a prefill, never
    # silently applied: it's marketing copy, so a human should glance at it.
    def _meta(name_attr, name_val):
        m = re.search(
            r'<meta[^>]+' + name_attr + r'\s*=\s*["\']' + name_val +
            r'["\'][^>]*content\s*=\s*["\']([^"\']+)["\']', html, re.I)
        if not m:
            m = re.search(
                r'<meta[^>]+content\s*=\s*["\']([^"\']+)["\'][^>]*' + name_attr +
                r'\s*=\s*["\']' + name_val + r'["\']', html, re.I)
        return (m.group(1).strip() if m else "")
    site_desc = _meta("name", "description") or _meta("property", "og:description")
    site_desc = re.sub(r"\s+", " ", site_desc)[:400]
    # A CTA is not a description — fall back to the AI read of the site.
    if _is_cta(site_desc):
        site_desc = ""

    def _clean(t):
        t = re.sub(r"[»›→▸▾▼+]+$", "", t).strip()
        return t
    def _keep(t, href, require_hint):
        tl = t.lower().strip()
        if not tl or tl in _MENU_GENERIC: return False
        if len(tl) > 48 or len(tl.split()) > 6: return False
        if not re.search(r"[a-z]", tl): return False
        if re.search(r"\d{3}", tl): return False          # phone numbers
        if require_hint and not _SERVICE_PATH_HINT.search(href or ""): return False
        return True

    out, seen = [], set()
    # Pass 1 — links inside <nav>/<header>. Pass 2 — links anywhere on the page
    # whose URL path looks service-ish (/services/, /markets/, /industries/...),
    # which catches sites that render menus without semantic nav tags.
    for links, need_hint, src in ((p.nav_links, False, "menu"),
                                  (p.other_links, True, "page")):
        for text, href in links:
            t = _clean(text)
            if not _keep(t, href, need_hint): continue
            key = t.lower()
            if key in seen: continue
            seen.add(key)
            hinted = bool(_SERVICE_PATH_HINT.search(href or ""))
            out.append({"label": t, "source": src, "service_path": hinted})

    # Pass 2.5 — HEADINGS. Portfolio-style sites (design studios, agencies)
    # run deliberately minimal navs — Work / About / Contact, all generic — and
    # put the actual service taxonomy in on-page section headings ("01.Branding",
    # "02.Packaging Design"). Links can't see those, so when the nav yields
    # nothing, harvest heading + <strong>/<b> text instead. The AI conversion
    # step already drops non-service items, so this can afford to over-collect.
    used_headings = False
    if len(out) < 3:
        raw_heads = re.findall(r"<(h[1-6]|strong|b)\b[^>]*>(.*?)</\1>",
                               html, re.I | re.S)
        import html as _htmlmod
        n_before = len(out)
        for _tag, inner in raw_heads:
            t = re.sub(r"<[^>]+>", " ", inner)          # strip nested tags
            t = _htmlmod.unescape(t)
            t = " ".join(t.split())
            t = re.sub(r"^\s*\d{1,2}\s*[.):\-–—]?\s*", "", t)  # "01.Branding" -> "Branding"
            t = t.strip(" :·•|")
            tl = t.lower()
            if not t or tl in _MENU_GENERIC or tl in seen: continue
            if len(t) > 48 or len(t.split()) > 6: continue
            if not re.search(r"[a-z]", tl): continue
            if re.search(r"\d{3}", tl): continue          # phone numbers
            if re.search(r"[.!?]$", t): continue          # sentences, not labels
            seen.add(tl)
            out.append({"label": t, "source": "heading", "service_path": False})
            if len(out) - n_before >= 15: break
        used_headings = len(out) > n_before

    # Pass 3 — JS-built navs render no anchors in raw HTML. The sitemap is
    # static XML that JavaScript can't hide, and page slugs map to the same
    # service taxonomy a menu would. Same crawler used for business-desc
    # inference; capped at ~5s internally.
    used_sitemap = False
    if len(out) < 3:
        try:
            for topic in fetch_site_pages(dom, limit=30):
                key = topic.lower()
                if key in seen or key in _MENU_GENERIC: continue
                if len(topic) > 48 or len(topic.split()) > 6: continue
                seen.add(key)
                out.append({"label": topic.title(), "source": "sitemap",
                            "service_path": False})
            used_sitemap = len(out) > 0
        except Exception:
            pass
    # service-path links first (strongest signal), then menu order
    out.sort(key=lambda x: (not x["service_path"], x["source"] != "menu"))
    out = out[:40]

    # Convert raw labels into search-phrase terms. "Healthcare" is a menu item,
    # not a search — a customer types "healthcare construction company". Claude
    # sees the whole label set plus business context, so it also drops
    # non-service items the static filter missed. On any failure, raw labels
    # pass through so the feature degrades instead of breaking.
    seeds = [s for s in (d.get("seeds") or []) if isinstance(s, str)]
    mapping = claude_menu_to_terms([s["label"] for s in out],
                                   d.get("brand") or "", dom, seeds,
                                   d.get("business_desc") or site_desc or "")
    ai_used = bool(mapping)
    if ai_used:
        converted, seen_terms = [], set()
        for s in out:
            term = mapping.get(s["label"], s["label"].lower())
            if term is None:
                continue                      # AI says: not a service
            if term in seen_terms:
                continue                      # two labels -> same phrase
            seen_terms.add(term)
            s["term"] = term
            converted.append(s)
        out = converted
    else:
        for s in out:
            s["term"] = s["label"].lower()

    # ---- the SHORT names, measured -----------------------------------------
    # The menu converter only makes labels longer, which is how a set of seeds
    # with 10/mo between them got built. Mine the client's own acronyms and
    # price each one, so what comes back is not "here are some capitals we saw"
    # but "these are searched, these are not". One volume call. (2026-08-10)
    acronyms = []
    try:
        mined = mine_acronyms(html, d.get("brand") or "")
        if mined:
            quals = qualifier_words(d.get("seeds") or [],
                                    [x.get("term") or x.get("label") for x in out])
            probes, seen_p = [], set()
            for a in mined:
                lo = a["acronym"].lower()
                for form in [lo] + [f"{lo} {q}" for q in quals]:
                    if form not in seen_p:
                        seen_p.add(form)
                        probes.append(form)
            vols, _pc, verr = fetch_local_volume(probes, [], "", national=True)
            floor = int(CFG.get("acronym_min_volume", 20))
            ceil = int(CFG.get("acronym_max_volume", 50000))
            for a in mined:
                lo = a["acronym"].lower()
                forms = {f: int((vols or {}).get(f, 0) or 0)
                         for f in [lo] + [f"{lo} {q}" for q in quals]}
                best = max(forms, key=lambda f: forms[f])
                a["volume"] = forms[best]
                a["term"] = best
                a["forms"] = forms
                # Too big to be this client's program. Reported with a reason
                # rather than dropped silently, so a genuinely huge one can be
                # argued with instead of vanishing.
                a["too_big"] = forms[lo] > ceil
                a["worth_it"] = ((not verr) and forms[best] >= floor
                                 and not a["too_big"])
            # A token doing six figures on its own is a common word the miner
            # mistook for an acronym; it never reaches the chips.
            acronyms = [a for a in mined if not a.get("too_big")]
            for a in acronyms:
                a["qualifiers"] = quals
    except Exception as _ae:
        acronyms = [{"error": str(_ae)[:120]}]
    # 39 chips off a retail menu is not 39 services — see fold_proposals().
    out, folded_out = fold_proposals(out, seeds=(d.get("seeds") or []))
    # NO MORE THAN A FEW VARIANTS OF ONE PROCEDURE — the prompt has
    # asked for this since July and a prompt cannot be checked.
    # See cap_service_family. (2026-08-17)
    out, family_out = cap_service_family(out, seeds=(d.get("seeds") or []))
    # Headings are the risky source by construction: it fires when the nav gave
    # nothing, so it over-collects on purpose, and on npaihb.org it collected
    # member tribes. Classify whatever came back, whatever the source.
    # The client's own menu labels ARE the "what does this site sell" context, and
    # a storefront in the link structure is the second read on whether their
    # objects are their offer — same two signals the build uses.
    try:
        _ecom_chips, _ = detect_ecommerce(
            [h for _t, h in (list(p.nav_links) + list(p.other_links)) if h])
    except Exception:                                     # noqa: BLE001
        _ecom_chips = False
    # ONE call for BOTH proposal groups. The acronyms were exempt, which is how a
    # member tribe's initials became a seed; and a second call would let the two
    # groups disagree about the same client.
    _acr_items = [dict(a, term=(a.get("term") or a.get("acronym") or "").lower())
                  for a in acronyms if a.get("acronym")]
    _alias = {}
    for a in _acr_items:
        ex = str(a.get("expansion") or "").strip()
        if ex and len(ex.split()) > 1:
            _alias[a["term"]] = ex.lower()
    _all, _not = _split_proposal_kinds(
        out + _acr_items, d, dom,
        pages=[x.get("label") for x in out][:30], ecom=_ecom_chips, alias=_alias)
    _menu_terms = {str(x.get("term") or x.get("label") or "").lower() for x in out}
    _no_terms = {str(x.get("term") or "").lower() for x in _not}
    out = [x for x in _all if str(x.get("term") or x.get("label") or "").lower() in _menu_terms]
    acronyms = [a for a in acronyms
                if (a.get("term") or a.get("acronym") or "").lower() not in _no_terms]
    not_svc = _not
    return jsonify({"domain": dom, "services": out,
                    "not_services": not_svc,
                    "folded": [(x.get("term") or x.get("label") or "")
                               for x in folded_out],
                    "family_capped": [{"term": t, "on": w}
                                      for t, w in (family_out or [])],
                    "ai_refined": ai_used, "from_sitemap": used_sitemap,
                    "from_headings": used_headings,
                    "site_description": site_desc,
                    "acronyms": acronyms,
                    "n_nav_links": len(p.nav_links), "diag": diag})


@app.route("/api/quotes/status")
@_json_error_guard
def api_quotes_status():
    # Diagnostic detail so "saving is off" isn't a black box: report whether the
    # URL is present and whether the Postgres driver imported.
    return jsonify({
        "enabled": storage.enabled(),
        "has_database_url": bool(os.environ.get("DATABASE_URL", "")),
        "driver_installed": getattr(storage, "_HAVE_DRIVER", False),
        "detail": storage.status_detail(),
    })

@app.route("/api/quotes", methods=["GET"])
@_json_error_guard
def api_quotes_list():
    if not storage.enabled():
        return jsonify({"enabled": False, "quotes": []})
    search = (request.args.get("q") or "").strip()
    tool = (request.args.get("tool") or "seo").strip()
    return jsonify({"enabled": True, "quotes": storage.list_quotes(search, tool)})

@app.route("/api/quotes", methods=["POST"])
@_json_error_guard
def api_quotes_save():
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled — attach a Postgres database in Render."}), 400
    d = request.get_json(force=True)
    name = (d.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Give the quote a name."}), 400
    client = (d.get("client") or "").strip()
    payload = d.get("payload") or {}
    tool = (d.get("tool") or "seo").strip()
    qid = storage.save_quote(name, client, payload, tool)
    return jsonify({"ok": True, "id": qid})

@app.route("/api/quotes/<int:qid>", methods=["GET"])
@_json_error_guard
def api_quotes_load(qid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    q = storage.load_quote(qid)
    if not q:
        return jsonify({"error": "Not found."}), 404
    return jsonify(q)

@app.route("/api/quotes/<int:qid>", methods=["PUT"])
@_json_error_guard
def api_quotes_update(qid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    d = request.get_json(force=True)
    payload = d.get("payload") or {}
    name = d.get("name"); client = d.get("client")
    ok, version_saved = storage.update_quote(
        qid, payload,
        name=name.strip() if isinstance(name, str) else None,
        client=client.strip() if isinstance(client, str) else None)
    if not ok:
        return jsonify({"error": "Not found."}), 404
    return jsonify({"ok": True, "id": qid,
                    "version_saved": version_saved,
                    "unchanged": not version_saved})

@app.route("/api/quotes/<int:qid>/share", methods=["POST"])
@_json_error_guard
def api_quotes_share(qid):
    """Mint (or return the existing) read-only review link for a saved quote."""
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled — attach Postgres first."}), 400
    token = storage.get_or_create_share_token(qid)
    if not token:
        return jsonify({"error": "Quote not found."}), 404
    return jsonify({"token": token,
                    "url": request.host_url.rstrip("/") + "/review/" + token})

@app.route("/api/review/<token>")
@_json_error_guard
def api_review(token):
    """Read-only quote fetch for the review page. Token is the credential;
    no edit endpoints accept it."""
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    q = storage.load_by_token(token)
    if not q:
        return jsonify({"error": "This review link is invalid or the quote was deleted."}), 404
    return jsonify(q)

@app.route("/review/<token>")
def review_page(token):
    """Same template as the owning tool; the frontend sees /review/ in the
    path and switches to read-only review mode."""
    tool = "seo"
    if storage.enabled():
        tool = storage.get_tool_by_token(token) or "seo"
    return render_template("reputation.html" if tool == "rep" else "index.html", build=BUILD_STR)

@app.route("/favicon.svg")
def favicon():
    svg = """<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 64 64'>
<rect width='64' height='64' rx='14' fill='#002D58'/>
<circle cx='28' cy='27' r='13' fill='none' stroke='#F1B434' stroke-width='5'/>
<line x1='37.5' y1='36.5' x2='50' y2='49' stroke='#F1B434' stroke-width='6' stroke-linecap='round'/>
<text x='28' y='32.5' font-family='Arial,Helvetica,sans-serif' font-size='15' font-weight='bold'
      fill='#FDFBF7' text-anchor='middle'>$</text>
</svg>"""
    from flask import Response
    return Response(svg, mimetype="image/svg+xml",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/q/<int:qid>")
def edit_link_page(qid):
    """Edit deep-link: opens the owning tool with this saved quote loaded,
    ready to revise. Version history makes collaborative edits safe — every
    save snapshots the prior state."""
    tool = "seo"
    if storage.enabled():
        q = storage.load_quote(qid)
        if q:
            tool = (q.get("tool") if isinstance(q, dict) else None) or "seo"
    return render_template("reputation.html" if tool == "rep" else "index.html", build=BUILD_STR)

@app.route("/api/quotes/version/<int:vid>", methods=["DELETE"])
@_json_error_guard
def api_quotes_version_delete(vid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    storage.delete_version(vid)
    return jsonify({"ok": True})

@app.route("/api/quotes/<int:qid>", methods=["DELETE"])
@_json_error_guard
def api_quotes_delete(qid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    storage.delete_quote(qid)
    return jsonify({"ok": True})

@app.route("/api/quotes/<int:qid>/versions", methods=["GET"])
@_json_error_guard
def api_quotes_versions(qid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    return jsonify({"versions": storage.list_versions(qid)})

@app.route("/api/quotes/version/<int:vid>", methods=["GET"])
@_json_error_guard
def api_quotes_version_load(vid):
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    v = storage.load_version(vid)
    if not v:
        return jsonify({"error": "Not found."}), 404
    return jsonify(v)

# ---------------------------------------------------------------------------
# Reputation Management tab — separate template + pricing module (rep_pricing).
# Shares this Render service and the DFS credentials; nothing else overlaps
# with the SEO pipeline.
# ---------------------------------------------------------------------------
import rep_pricing
import rep_scan
rep_scan.init(dfs_post)

@app.route("/api/rep_scan_terms", methods=["POST"])
@_json_error_guard
def api_rep_scan_terms():
    """Brand term universe + negative-modifier volumes (one KFK live call)."""
    d = request.get_json(force=True)
    brand = (d.get("brand") or "").strip()
    if not brand:
        return jsonify({"error": "Brand name required."}), 400
    try:
        return jsonify(rep_scan.scan_terms(brand))
    except Exception as e:
        return jsonify({"error": f"Term scan failed: {e}"}), 502

@app.route("/api/rep_scan_serp", methods=["POST"])
@_json_error_guard
def api_rep_scan_serp():
    """'{brand} reviews' top-10 threat table + related searches + autosuggest."""
    d = request.get_json(force=True)
    brand = (d.get("brand") or "").strip()
    if not brand:
        return jsonify({"error": "Brand name required."}), 400
    try:
        return jsonify(rep_scan.scan_serp(brand, (d.get("domain") or "").strip()))
    except Exception as e:
        return jsonify({"error": f"SERP scan failed: {e}"}), 502

@app.route("/api/rep_scan_autocomplete", methods=["POST"])
@_json_error_guard
def api_rep_scan_autocomplete():
    """Auto-suggest flags — separate endpoint so its latency never stacks
    onto the SERP call (Render's proxy cuts requests around 100s)."""
    d = request.get_json(force=True)
    brand = (d.get("brand") or "").strip()
    if not brand:
        return jsonify({"error": "Brand name required."}), 400
    try:
        return jsonify(rep_scan.scan_autocomplete(brand))
    except Exception as e:
        return jsonify({"error": f"Autocomplete scan failed: {e}"}), 502

@app.route("/api/rep_scan_locations", methods=["POST"])
@_json_error_guard
def api_rep_scan_locations():
    """Google Business location discovery (instant, database-backed)."""
    d = request.get_json(force=True)
    brand = (d.get("brand") or "").strip()
    if not brand:
        return jsonify({"error": "Brand name required."}), 400
    try:
        return jsonify(rep_scan.scan_locations(brand, domain=(d.get("domain") or "")))
    except Exception as e:
        return jsonify({"error": f"Location scan failed: {e}"}), 502

@app.route("/api/rep_reviews_submit", methods=["POST"])
@_json_error_guard
def api_rep_reviews_submit():
    """Queue worst-first review pulls for selected locations (priority ~1min)."""
    d = request.get_json(force=True)
    pids = [p for p in (d.get("place_ids") or []) if p]
    if not pids:
        return jsonify({"error": "No locations selected."}), 400
    try:
        return jsonify(rep_scan.reviews_submit(
            pids, int(d.get("depth") or rep_pricing.SCAN_SETTINGS["review_pull_depth"])))
    except Exception as e:
        return jsonify({"error": f"Review submit failed: {e}"}), 502

@app.route("/api/rep_reviews_collect", methods=["POST"])
@_json_error_guard
def api_rep_reviews_collect():
    d = request.get_json(force=True)
    tids = [t for t in (d.get("task_ids") or []) if t]
    if not tids:
        return jsonify({"error": "No task ids."}), 400
    try:
        return jsonify(rep_scan.reviews_collect(tids))
    except Exception as e:
        return jsonify({"error": f"Review collect failed: {e}"}), 502


@app.route("/reputation")
def reputation():
    return render_template("reputation.html", build=BUILD_STR)

@app.route("/api/rep_config", methods=["GET"])
@_json_error_guard
def api_rep_config_get():
    """Expose the rep-tool tunables for the pricing panel — mirror of the
    SEO tool's /api/config."""
    rc = rep_pricing.REP_CFG
    return jsonify({
        "review_margin_pct": rc["review_removal"]["default_margin_pct"],
        "review_brackets": rc["review_removal"]["brackets"],
        "site_brackets": rc["article_removal"]["brackets"],
        "site_premium_per": rc["article_removal"]["premium_per"],
        "bundle": {k: rep_pricing.SEARCH_BUNDLE[k]
                   for k in ("supp_base", "as_base", "comp_per_1k", "floor", "cap")},
        "shield_monthly": rc["shield"]["monthly_hard"],
        "shield_per_extra_location": rc["shield"]["per_extra_location_hard"],
        "geo": {p: rep_pricing.GEO[p]["monthly"] for p in ("setup", "scale")},
        "bundle_discount_pct": rc["bundle"]["recurring_discount_pct"],
        "internal_cost_pct": rep_pricing.INTERNAL_COST_PCT["pct"],
        "review_pull_depth": rep_pricing.SCAN_SETTINGS["review_pull_depth"],
    })

@app.route("/api/rep_config", methods=["POST"])
@_json_error_guard
def api_rep_config_set():
    """Apply edited constants to the running session (not persisted — a
    restart reverts to file defaults). Same live-tuning model as the SEO tool."""
    d = request.get_json(force=True)
    rc = rep_pricing.REP_CFG
    try:
        if "review_margin_pct" in d:
            rc["review_removal"]["default_margin_pct"] = min(0.90, max(0.0, float(d["review_margin_pct"])))
        if "review_brackets" in d and isinstance(d["review_brackets"], list):
            for i, b in enumerate(d["review_brackets"]):
                if i < len(rc["review_removal"]["brackets"]) and "hard" in b:
                    rc["review_removal"]["brackets"][i]["hard"] = float(b["hard"])
        if "site_brackets" in d and isinstance(d["site_brackets"], list):
            for i, b in enumerate(d["site_brackets"]):
                if i < len(rc["article_removal"]["brackets"]) and "per" in b:
                    rc["article_removal"]["brackets"][i]["per"] = int(float(b["per"]))
        if "site_premium_per" in d:
            rc["article_removal"]["premium_per"] = int(float(d["site_premium_per"]))
        if "bundle" in d and isinstance(d["bundle"], dict):
            for k in ("supp_base", "as_base", "comp_per_1k", "floor", "cap"):
                if k in d["bundle"]:
                    rep_pricing.SEARCH_BUNDLE[k] = int(float(d["bundle"][k])) if k != "comp_per_1k" else float(d["bundle"][k])
        if "shield_monthly" in d:
            rc["shield"]["monthly_hard"] = int(float(d["shield_monthly"]))
        if "shield_per_extra_location" in d:
            rc["shield"]["per_extra_location_hard"] = int(float(d["shield_per_extra_location"]))
        if "geo" in d and isinstance(d["geo"], dict):
            for p in ("setup", "scale"):
                if p in d["geo"]:
                    rep_pricing.GEO[p]["monthly"] = int(float(d["geo"][p]))
        if "bundle_discount_pct" in d:
            rc["bundle"]["recurring_discount_pct"] = min(0.9, max(0.0, float(d["bundle_discount_pct"])))
        if "internal_cost_pct" in d:
            rep_pricing.INTERNAL_COST_PCT["pct"] = min(1.0, max(0.0, float(d["internal_cost_pct"])))
        if "review_pull_depth" in d:
            rep_pricing.SCAN_SETTINGS["review_pull_depth"] = min(4490, max(10, int(float(d["review_pull_depth"]))))
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": f"Config apply failed: {e}"}), 400

@app.route("/api/rep_quote", methods=["POST"])
@_json_error_guard
def api_rep_quote():
    d = request.get_json(force=True)
    try:
        return jsonify(rep_pricing.build_rep_quote(d))
    except Exception as e:
        return jsonify({"error": f"Quote build failed: {e}"}), 500

@app.route("/api/rep_volume", methods=["POST"])
@_json_error_guard
def api_rep_volume():
    """US-national exact-match volume for the brand terms — drives the
    Search Protection base+multiplier formula. Reuses fetch_exact_volume
    (Labs keyword_overview, per-term exact volume)."""
    d = request.get_json(force=True)
    terms = [t.strip() for t in (d.get("terms") or []) if t and t.strip()]
    if not terms:
        return jsonify({"error": "No brand terms provided."}), 400
    vols = fetch_exact_volume(terms, [], "")
    if not vols:
        return jsonify({"error": "DataForSEO returned no volume — check terms "
                                 "or DFS credentials."}), 502
    per_term = {t: vols.get(t.lower(), 0) for t in terms}
    return jsonify({"per_term": per_term, "total": sum(per_term.values())})


# initialize the DB tables on startup (no-op when saving isn't enabled)
try:
    storage.init_db()
except Exception as _e:
    print(f"[storage] init skipped: {_e}")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)


# ---------------------------------------------------------------------------
# THE .DOCX EXPORT
# ---------------------------------------------------------------------------

def gap_effort(client_rank, rival_rank):
    """The authority gap, in link volume and months rather than in points.

    Reads the band the client has to climb INTO — the target, not where they
    start — because that is what sets the cost of the next ten points. A local
    dentist thirty points behind other local dentists is a different campaign
    from a supplement brand six hundred behind Amazon, and the number alone does
    not say which you are looking at. (2026-08-18)
    """
    if not rival_rank:
        return None
    tgt = float(rival_rank) / 10.0
    cur = (float(client_rank) / 10.0) if client_rank else 0.0
    months, links = PROPOSAL["gap_bands"][-1][1], PROPOSAL["gap_bands"][-1][2]
    for ceiling, m, l in PROPOSAL["gap_bands"]:
        if tgt <= ceiling:
            months, links = m, l
            break
    pts = round(tgt - cur, 1)
    # A BAND IS AN ALTITUDE, NOT A DISTANCE. Pennsylvania Center for Dental
    # Excellence sits at 23.5 against a page one of 24.0 — half a point — and
    # reading the band it is climbing INTO returned "12-18 months", which is the
    # cost of the whole band rather than of the half point. When the client is
    # already level, the honest answer is that links are not the constraint.
    if pts <= 5:
        return {"client": round(cur, 1), "target": round(tgt, 1), "points": pts,
                "months": "", "links": "", "level": True}
    return {"client": round(cur, 1), "target": round(tgt, 1),
            "points": pts, "months": months, "links": links, "level": False}


def gap_verdict(g, brand=""):
    """Is the gap big or small, and can we close it — in one sentence.

    "Their typical authority score is 313, against 380 for Grav" is a
    measurement, and a client cannot tell from it whether they are nearly there
    or nowhere near. This says which, and says it in terms of what our help
    does about it. MIRRORED as GAP_VERDICTS in the template. (2026-08-19)
    """
    if not g:
        return None
    b = brand or "This site"
    if g.get("level"):
        return {"label": "no gap",
                "line": PROPOSAL["gap_level_verdict"].format(brand=b)}
    pts = float(g.get("points") or 0)
    for ceiling, label, line in PROPOSAL["gap_verdicts"]:
        if pts <= ceiling:
            return {"label": label, "line": line}
    ceiling, label, line = PROPOSAL["gap_verdicts"][-1]
    return {"label": label, "line": line}


def _p_money(v):
    try:
        return "$" + format(int(round(float(v))), ",")
    except Exception:                                         # noqa: BLE001
        return "—"


def _proposal_rows(d):
    """The keyword table, in Brendan's column order plus the volume we measure.

    His table is Keyword / Google Current Rank / Keyword Type. The volume column
    is ours — it is the one number in the document nobody else in the pitch can
    produce, because it comes from a live search-volume lookup rather than an
    estimate, and it is why the tiers are where they are.
    """
    tier_label = {"ultra": "Ultra Competitive", "competitive": "Competitive",
                  "long_tail": "Long Tail"}
    # THE SAME KEY THE PANEL USES. Trim and collapse whitespace as well as
    # lower-casing: the two sides matched differently and the document printed
    # positions for five terms the tool was calling "Not checked". (2026-08-19)
    def _nk(v):
        return re.sub(r"\s+", " ", str(v or "").strip().lower())
    by_kw = {}
    for r in (d.get("table") or []):
        k = _nk(r.get("kw"))
        if k and k not in by_kw:
            by_kw[k] = r
    rows = []
    for tier in ("ultra", "competitive", "long_tail"):
        for r in (d.get("kw") or {}).get(tier, []) or []:
            kw = str(r.get("kw") or r.get("keyword") or "").strip()
            if not kw:
                continue
            live = by_kw.get(_nk(kw)) or {}
            # AN UNMEASURED TERM DOES NOT GO IN A CLIENT DOCUMENT. A failed or
            # still-queued lookup carries pos "—", and a term with no row at all
            # carries nothing — both used to print as a rank, the first as a
            # bare dash and the second as "Not Found". "Not Found" is a positive
            # claim that the client does not rank for a term, made about a term
            # nobody checked. Three of Ski Barn's rows were 40101 errors and
            # would have gone to the client as three terms they are missing
            # from. Left out entirely and counted separately. (2026-08-19)
            if live.get("error") or live.get("queued") or live.get("expired"):
                continue
            pos = live.get("pos")
            if not live:
                continue
            if not isinstance(pos, int) and pos not in ("Not Found", None):
                continue
            rows.append({
                "kw": kw,
                "rank": (str(pos) if isinstance(pos, int) else "Not Found"),
                "tier": tier_label[tier],
                "vol": int(r.get("vol") or 0),
            })
    # Ranked first and best-ranked at the top, exactly as his tables read: the
    # document opens on the terms the client can verify in one search.
    def _key(x):
        try:
            return (0, int(x["rank"]))
        except (TypeError, ValueError):
            return (1, -x["vol"])
    rows.sort(key=_key)
    return rows


def handoff_meta(d):
    """THE NON-PRICING HALF OF WHAT THE IO PULLS — AND IT IS SHORT ON PURPOSE.

    Two fields. Everything else adtini already has or does not render.

    The first cut of this block sent brand, website, industry, market names,
    page-one incumbents, authority scores and site condition. Kiri struck all of
    it (2026-08-27) and was right on both counts:

      * Brand, website, industry and the markets come FROM adtini — order
        details, the product card, client details. Sending them back is a round
        trip that can only introduce disagreement.
      * Incumbents, authority and site condition are ours and are genuinely
        interesting, but no slide in the proposal deck draws them. A field
        nobody renders is a field that goes stale without anyone noticing.

    THE ONE THING TO WATCH. The market NAME can drift: a client entered as
    "Whatcom County, WA" is quoted on "bellingham wa", so a slide printing the
    entered market beside our keywords will disagree with them. If that shows up,
    the fix is one field here, not a lookup on their side.

    Reads the same accessor the .docx reads (_proposal_rows), so the deck and the
    document can never describe different keywords. Nothing is recomputed: a
    saved quote hands back exactly what was sent to the client.
    """
    d = d or {}
    serp = d.get("serp") or {}

    return {
        # ---- the terms ----------------------------------------------------
        # THREE COLUMNS, NAMED THE WAY THE SLIDE NAMES THEM: term / google rank
        # / competitiveness. Per-row volume is deliberately NOT sent — the slide
        # does not show it. Counts are not sent either: the deck has the rows and
        # can count them, and a stored count is a second thing to keep in step.
        # "Not Found" is a MEASURED miss. An unmeasured term is not in the list —
        # a failed or still-queued lookup is not a ranking claim.
        "keywords": [{"term": r["kw"], "google_rank": r["rank"],
                      "competitiveness": r["tier"]} for r in _proposal_rows(d)],
        # ---- the SERP -----------------------------------------------------
        "serp_screenshot_keyword": serp.get("kw") or "",
        "serp_screenshot_captured": bool(serp.get("img")),
    }


@app.route("/api/handoff", methods=["POST"])
@_json_error_guard
def api_handoff():
    """EVERYTHING THE IO PULLS, IN ONE CALL.

    Pricing comes from the quote as it was saved rather than being re-priced —
    re-running stage4_price against today's config would quietly reprice a quote
    the client has already seen.
    """
    d = request.get_json(force=True) or {}
    return jsonify({"pricing": (d.get("pricing") or {}).get("handoff") or {},
                    "proposal": handoff_meta(d)})


@app.route("/api/quotes/<int:qid>/handoff", methods=["GET"])
@_json_error_guard
def api_quote_handoff(qid):
    """The same block for a SAVED quote, so adtini can pull by id."""
    if not storage.enabled():
        return jsonify({"error": "Saving isn't enabled."}), 400
    q = storage.load_quote(qid)
    if not q:
        return jsonify({"error": "Not found."}), 404
    d = q.get("payload") or q.get("data") or q
    return jsonify({"quote_id": qid,
                    "pricing": (d.get("pricing") or {}).get("handoff") or {},
                    "proposal": handoff_meta(d)})


def _perf_merge_extra(d):
    """Fold operator-measured extra terms into the shapes _proposal_rows reads.

    They join the performance table ONLY — not the priced grid, not the SEO
    keyword table, and not total_volume. The campaign still commits to the
    twenty terms the options describe; these are terms the client could be
    billed for if they rank, which is a different promise.
    """
    extra, _seen_x = [], set()
    for x in (d.get("perf_extra") or []):
        if not isinstance(x, dict) or not x.get("kw"):
            continue
        _k = re.sub(r"\s+", " ", str(x["kw"]).strip().lower())
        if _k in _seen_x:
            continue                      # measured twice — auto run, then by hand
        _seen_x.add(_k)
        extra.append(x)
    if not extra:
        return d
    kw = {k: list(v or []) for k, v in (d.get("kw") or {}).items()
          if k in ("ultra", "competitive", "long_tail")}
    kw.setdefault("long_tail", [])
    table = list(d.get("table") or [])
    seen = {re.sub(r"\s+", " ", str(r.get("kw") or "").strip().lower()) for r in table}
    for x in extra:
        # AN UNMEASURED TERM DOES NOT GO IN A CLIENT DOCUMENT. A row with no
        # position reads as "Not Found" downstream — a positive claim that the
        # client does not rank for a term nobody checked. Same rule the keyword
        # table follows, and these terms arrive one rank check at a time, so a
        # half-finished batch is the normal case rather than the odd one.
        pos = x.get("pos")
        if not isinstance(pos, int) and str(pos).strip().lower() not in ("not found",):
            continue
        k = re.sub(r"\s+", " ", str(x["kw"]).strip().lower())
        kw.setdefault(x.get("tier") or "long_tail", []).append(
            {"kw": x["kw"], "vol": int(x.get("vol") or 0)})
        if k not in seen:
            table.append({"kw": x["kw"], "pos": pos})
            seen.add(k)
    return dict(d, kw=kw, table=table)


def _perf_rows(d):
    """The pay-for-performance table: every quoted term with its four prices.

    Same rows as the keyword table — an unmeasured term still does not go in a
    client document — plus the term's own measured top-of-page bid, which is
    what the Page-1 cost is derived from. Sorted by Cost Page 1 descending,
    which is how Brendan's reads: the terms worth the most money first.

    An operator override for a term wins over the derived figure. His four most
    expensive rows look hand-set and the tool cannot reproduce a hand.
    """
    cpc = {re.sub(r"\s+", " ", str(k or "").strip().lower()): v
           for k, v in (d.get("cpc") or {}).items()}
    over = {re.sub(r"\s+", " ", str(k or "").strip().lower()): v
            for k, v in (d.get("perf_override") or {}).items()}
    area = d.get("practice_area") or {}
    topics = perf_topics(d)
    rows, seen_kw = [], set()
    for r in _proposal_rows(_perf_merge_extra(d)):
        key = re.sub(r"\s+", " ", r["kw"].strip().lower())
        bid = cpc.get(key)
        price = perf_term_price(bid)
        try:
            man = float(over.get(key) or 0)
        except (TypeError, ValueError):
            man = 0.0
        if man > 0:
            price = perf_term_price(man / float(CFG.get("perf_page1_mult", 2.1) or 2.1))
            price["page1"] = int(round(man))
        pos = None
        try:
            pos = int(r["rank"])
        except (TypeError, ValueError):
            pos = None
        # ONE ROW PER TERM. A term the grid cut and the extra pool measured can
        # arrive from both sides in different tiers, and the document printed
        # "car accident lawyer knoxville tn" twice. (2026-08-22)
        if key in seen_kw:
            continue
        seen_kw.add(key)
        rows.append({
            "kw": r["kw"],
            # An operator override wins; otherwise the build's own topics.
            "area": (str(area.get(key) or "").strip()
                     or perf_area(r["kw"], topics)),
            "tier": r["tier"],
            "rank": r["rank"],
            "bid": (round(float(bid), 2) if bid not in (None, "") else None),
            "manual": man > 0,
            "achieved": perf_tier_label(pos if pos is not None else 0),
            **price,
        })
    rows.sort(key=lambda x: (-x["page1"], x["kw"]))
    return rows


# Openers that mark a sentence as an advert rather than a description.
_DESC_CTA = re.compile(
    r"\b(shop|buy|order|call|visit|browse|save|get)\s+(today|now|online|us|here)\b"
    r"|\bfree (shipping|quote|estimate|consultation)\b|\bcontact us\b"
    r"|\bclick here\b|\bshop now\b|\blearn more\b", re.I)
_DESC_IMPERATIVE = re.compile(
    r"^(get|shop|buy|find|discover|explore|save|order|call|visit|browse|"
    r"experience|choose|trust)\b", re.I)


def _clean_desc(text):
    """The client's own words, or nothing — never their meta description.

    business_desc is auto-filled from the site, which usually means the meta
    description, which is written to sell to a searcher. Ski Barn's proposal
    quoted "Get the best ski &amp; snowboard gear from Ski Barn in New Jersey -
    ... - Shop Today!" back at Ski Barn, in their own proposal, with the HTML
    entity still in it — the escape ran twice, once on the way into the field
    and again on the way into the document.

    So: unescape once, drop a trailing call to action, and if what is left still
    reads as advertising, drop the paragraph entirely. Brendan's own proposals
    carry no such line, so losing it costs nothing; keeping a bad one costs the
    first impression. (2026-08-19)
    """
    t = html.unescape(html.unescape(str(text or ""))).strip()
    if not t:
        return ""
    # Trailing CTA clause, after the last dash or full stop.
    parts = re.split(r"\s+[-\u2013\u2014]\s+", t)
    while len(parts) > 1 and _DESC_CTA.search(parts[-1]):
        parts.pop()
    t = " - ".join(parts).strip(" -\u2013\u2014")
    if not t:
        return ""
    if _DESC_CTA.search(t) or _DESC_IMPERATIVE.match(t):
        return ""
    # A description, not a slogan: needs a verb phrase and a reasonable length.
    if len(t) < 25 or t.count("!") > 0:
        return ""
    return t


def build_proposal_docx(d, _notes=None):
    """One SSG-shaped proposal, built from the quote the tool already holds."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    ATLAS = RGBColor(0x00, 0x2D, 0x58)
    P = PROPOSAL
    brand = (d.get("brand") or "").strip() or "Your Business"
    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.9)
        sec.top_margin = sec.bottom_margin = Inches(0.8)
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    def head(text, size=15):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = ATLAS
        p.space_before = Pt(14)
        return p

    def body(text, bold=False, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    def bullet(text):
        return doc.add_paragraph(text, style="List Bullet")

    def subbullet(text):
        """Brendan nests his detail under the sentence that introduces it."""
        b = doc.add_paragraph(text, style="List Bullet 2")
        b.paragraph_format.left_indent = Inches(0.62)
        b.paragraph_format.space_after = Pt(1)
        return b

    def card_grid(cards, per_row=3):
        """The page-one incumbents as bordered cards, three across.

        A flat bullet list of six domains with an authority number after each
        reads as an appendix; the same six as cards reads as a competitive
        landscape, which is what it is and how the tool shows it on screen. A
        borderless outer table does the layout and each cell carries its own
        border, because that is the one construct Word and LibreOffice both
        render identically. (2026-08-19)
        """
        if not cards:
            return None
        rows = (len(cards) + per_row - 1) // per_row
        tb = doc.add_table(rows=rows, cols=per_row)
        tb.autofit = False
        # KEEP A ROW OF CARDS WHOLE — AND THEN KEEP THE ROWS TOGETHER. cantSplit
        # stopped Word splitting a row down the middle, which had put three
        # domain names at the foot of one page and their three authority lines
        # at the head of the next. It did nothing about the rows themselves:
        # three cards then a page break then three more, with half a page of
        # white between them, still reads as two unrelated exhibits.
        #
        # There is no "keep this table together" in the format, so this is
        # keep-with-next on every paragraph in every row but the last, which is
        # what Word actually honours: each row is glued to the one after it, so
        # the block moves to the next page as a unit or not at all. (2026-08-19)
        for _i, _r in enumerate(tb.rows):
            _r._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
            if _i < len(tb.rows) - 1:
                for _c in _r.cells:
                    for _p in _c.paragraphs:
                        _p.paragraph_format.keep_with_next = True
        for i, (title, sub) in enumerate(cards):
            cell = tb.rows[i // per_row].cells[i % per_row]
            cell.width = Inches(6.6 / per_row)
            _box(cell, sz="6")
            h = cell.paragraphs[0]
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(1)
            rh = h.add_run(title)
            rh.bold = True
            rh.font.size = Pt(9.5)
            sp_ = cell.add_paragraph()
            sp_.paragraph_format.space_before = Pt(0)
            sp_.paragraph_format.space_after = Pt(0)
            rs = sp_.add_run(sub)
            rs.font.size = Pt(8.5)
            rs.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7C)
        # Trailing cells in a short last row stay empty AND unbordered, so five
        # incumbents do not print as five cards and one empty box.
        for j in range(len(cards), rows * per_row):
            tb.rows[j // per_row].cells[j % per_row].width = Inches(6.6 / per_row)
        doc.add_paragraph()
        return tb

    def _box(cell, sz="10"):
        """A navy border and interior padding on one table cell."""
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), sz)
            e.set(qn("w:color"), "002D58")
            borders.append(e)
        tcPr.append(borders)
        # Breathing room inside the border — a box with text against its edges
        # reads as a rendering accident rather than a design.
        mar = OxmlElement("w:tcMar")
        for edge, val in (("top", "110"), ("start", "150"),
                          ("bottom", "110"), ("end", "150")):
            m = OxmlElement("w:" + edge)
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            mar.append(m)
        tcPr.append(mar)
        return cell

    def option_box(title, blurbs, scope_lines, price_line):
        """One priced option, in a bordered box.

        A single-cell TABLE, because that is the only construct Word and
        LibreOffice both render a reliable border around. Paragraph borders look
        right in Word and merge into one another in LibreOffice, and these
        documents get opened in whatever the client has. (2026-08-18)
        """
        tb = doc.add_table(rows=1, cols=1)
        tb.autofit = False
        # KEEP THE BOX WHOLE. Left to itself Word broke Option 3 across the page
        # boundary and drew a border on each half, so one option read as two
        # empty-bottomed panels. A priced option is a unit; it moves to the next
        # page or it does not move. (2026-08-18)
        _trPr = tb.rows[0]._tr.get_or_add_trPr()
        _cant = OxmlElement("w:cantSplit")
        _trPr.append(_cant)
        cell = tb.rows[0].cells[0]
        cell.width = Inches(6.6)
        _box(cell)

        def _tight(par, before=0, after=2):
            par.paragraph_format.space_before = Pt(before)
            par.paragraph_format.space_after = Pt(after)
            return par

        h = _tight(cell.paragraphs[0])
        rh = h.add_run(title)
        rh.bold = True
        rh.font.size = Pt(12)
        rh.font.color.rgb = ATLAS
        # HIS PARAGRAPHS, HIS RUN-IN. The last one ends on "recommend
        # targeting:" and the list follows it, so the bold "This option
        # targets:" label that used to sit here is both redundant and words he
        # never wrote. (2026-08-19)
        for para in ([blurbs] if isinstance(blurbs, str) else list(blurbs)):
            _tight(cell.add_paragraph(para))
        if scope_lines:
            for line in scope_lines:
                b = _tight(cell.add_paragraph(line, style="List Bullet"), 0, 0)
                b.paragraph_format.left_indent = Inches(0.28)
        pr = _tight(cell.add_paragraph(), 4, 0)
        rp = pr.add_run(price_line)
        rp.bold = True
        # A thin gap between boxes, not a full empty paragraph — three of those
        # is most of the reason Option 3 fell off the page.
        _tight(doc.add_paragraph(), 0, 4)
        return tb

    # ---- cover ------------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("SEO Proposal")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = ATLAS
    n = doc.add_paragraph()
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = n.add_run(brand)
    rn.font.size = Pt(16)
    rn.bold = True

    # ---- background -------------------------------------------------------
    head(P["intro_heading"])
    desc = _clean_desc(d.get("business_desc"))
    body(P["intro_line"].format(brand=brand))
    if desc:
        body(desc)
    body(P["intro_close"])

    # ---- the SEO section --------------------------------------------------
    head(P["seo_heading"])
    body(P["seo_intro"])
    body(P["seo_table_lead"])

    # NO VOLUME COLUMN. A list where most terms read 10/mo argues against the
    # campaign, and Brendan's table does not carry one either. The demand total
    # stays in the prose, where it describes the opportunity rather than
    # itemising its weakest rows. (2026-08-18)
    rows = _proposal_rows(d)
    if rows:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(("Keyword", "Google Current Rank",
                                   "Keyword Type")):
            hdr[i].text = ""
            rr = hdr[i].paragraphs[0].add_run(label)
            rr.bold = True
        for row in rows:
            c = tbl.add_row().cells
            c[0].text = row["kw"]
            c[1].text = row["rank"]
            c[2].text = row["tier"]
        doc.add_paragraph()

    # OURS, AND IT BELONGS AGAINST THE TABLE. This is the one sentence that
    # reads the table back — how many of those rows are already ranked, and how
    # much demand the list carries — so it sits directly under it rather than
    # after a screenshot. Brendan's own conclusion follows it. (2026-08-19)
    ranked = len([r for r in rows if r["rank"].isdigit()])
    demand = sum(r["vol"] for r in rows)
    if rows:
        body(f"Of the {len(rows)} terms above, {brand} currently ranks in the "
             f"top 100 for {ranked}"
             + (f", against roughly {format(demand, ',')} searches a month "
                f"across the list" if demand else "")
             + ".")
    body(P["seo_after_table"])

    body(P["keyword_sets_lead"])
    for name, blurb in P["keyword_sets"]:
        p = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(name + ": ")
        rb.bold = True
        p.add_run(blurb)
    body(P["keyword_sets_close"])

    # ---- the live results page, if one was captured ------------------------
    # Sits between the keyword work and what we measured, which is where it
    # argues for both: the table says where they rank, this shows what that
    # actually looks like, and the section below explains who is in the way.
    sp = d.get("serp") or {}
    if sp.get("img"):
        # THE CAPTION IS ROLLED BACK IF THE PICTURE DOES NOT LAND. Written the
        # obvious way — caption, then add_picture in a try — a corrupt capture
        # left "Google results today for ..." sitting above nothing, which is
        # worse than no image at all because it reads as a rendering failure in
        # a document going to a client. (2026-08-18)
        cap = body(f"Google results today for “{sp.get('kw', '')}”:", True)
        try:
            raw = sp["img"].split(",", 1)[1]
            doc.add_picture(io.BytesIO(base64.b64decode(raw)), width=Inches(6.4))
            doc.add_paragraph()
        except Exception:                                     # noqa: BLE001
            app.logger.exception("serp image could not be embedded")
            cap._element.getparent().remove(cap._element)

    # ---- what we measured — the part a generic proposal cannot show --------
    sig = d.get("signals") or {}
    if sig.get("rivals") or sig.get("pageone_rank"):
        head(P["measured_heading"])
        body("These figures were pulled live for this campaign. Search volumes "
             "come from live keyword data, the rankings above from live result "
             "pages, and the competitive picture below from the sites currently "
             "holding those positions.")
        riv = [r for r in (sig.get("rivals") or []) if r.get("domain")][:6]
        if riv:
            _lead = body("Who currently holds page one for your terms:", True)
            _lead.paragraph_format.keep_with_next = True
            cards = []
            for r in riv:
                bits = []
                if r.get("rank") is not None:
                    bits.append(f"authority {format(int(r['rank']), ',')}")
                if r.get("appearances"):
                    bits.append(f"{r['appearances']} of your terms")
                cards.append((r["domain"], " · ".join(bits)))
            card_grid(cards)
        if sig.get("median_rival_rank") is not None:
            gapline = (f"Their typical authority score is "
                       f"{format(int(sig['median_rival_rank']), ',')}")
            if sig.get("client_measured") and sig.get("client_rank") is not None:
                gapline += (f", against {format(int(sig['client_rank']), ',')} for "
                            f"{brand}")
            body(gapline + ".")
            # HOW BIG, AND CAN WE CLOSE IT — then what closing it costs. The
            # measurement on its own does not tell a client whether they are
            # nearly there or nowhere near, which is the only thing they want
            # to know from it. (2026-08-19)
            g = gap_effort(sig.get("client_rank") if sig.get("client_measured")
                           else None, sig.get("median_rival_rank"))
            v = gap_verdict(g, brand)
            if v:
                body(v["line"], True)
            if g and not g.get("level"):
                body(f"In practical terms that is roughly "
                     f"{int(round(g['points']))} points "
                     f"of authority. Published industry benchmarks put a move of "
                     f"that size at {g['links']} acquired over {g['months']} of "
                     f"sustained work — which is what the link building and "
                     f"content elements of the campaign below are for.")
        health = sig.get("health") or {}
        if health.get("failed"):
            body("On-site issues found on your site, which we address in the "
                 "first months of the campaign:", True)
            for f in health["failed"][:8]:
                bullet(f)

    # ---- the monthly campaign ---------------------------------------------
    head(P["campaign_heading"])
    body(P["campaign_lead"])
    for lead, items in P["services"]:
        bullet(lead)
        for it in items:
            subbullet(it)
    body(P["campaign_close"])

    # ---- the options ------------------------------------------------------
    # THE THREE OPTIONS BELONG ON ONE PAGE. They are read against each other —
    # that is the entire point of a good/better/best ladder — and Option 3
    # landing alone on the next page turns a comparison into two separate asks.
    # A page break here, and tight spacing inside the boxes, fits all three.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    head(P["options_heading"])
    body(P["options_lead"])
    tiers = dict(((d.get("pricing") or {}).get("client_tiers") or {}))
    term = int((d.get("pricing") or {}).get("min_term_months") or 6)
    for i, key in enumerate(("base", "intermediate", "advanced"), start=1):
        label = {"base": "Base", "intermediate": "Intermediate",
                 "advanced": "Advanced"}[key]
        option_box(f"Option {i}: {label} SEO Campaign",
                   P["option_blurb"][key],
                   P["option_scope"][key],
                   f"This option would be a monthly investment of "
                   f"{_p_money(tiers.get(key))} with a {term} month term "
                   f"which then becomes a month-to-month commitment.")
    head(P["additional_heading"], size=12)
    body(P["closing"][0])

    # ---- AI Search (GEO), when the quote carries it ------------------------
    # A parallel campaign with its own three options, priced IN ADDITION to the
    # SEO options above — which is how Brendan writes it and how the tool
    # computes it. Absent entirely on a Core SEO quote rather than showing zeros.
    ai = (d.get("pricing") or {}).get("ai_search") or {}
    ai_add = dict(ai.get("client_add") or {})
    if any(ai_add.get(k) for k in ("base", "intermediate", "advanced")):
        head(P["geo_heading"])
        for para in P["geo_intro"]:
            body(para)
        for e in P["geo_engines"]:
            bullet(e)
        for para in P["geo_context"]:
            body(para)
        head(P["geo_campaign_heading"], size=12)
        body(P["geo_campaign_lead"])
        for lead, items in P["geo_services"]:
            bullet(lead)
            for it in items:
                subbullet(it)
        head(P["geo_options_heading"], size=12)
        body(P["geo_options_lead"])
        for i, key in enumerate(("base", "intermediate", "advanced"), start=1):
            label = {"base": "Base", "intermediate": "Intermediate",
                     "advanced": "Advanced"}[key]
            option_box(f"Option {i}: {label} AI Search Campaign",
                       P["geo_option_blurb"][key], [],
                       f"This option would be a monthly cost of "
                       f"{_p_money(ai_add.get(key))} with a month-to-month "
                       f"commitment and 12 month minimum term.")
        tot = dict(ai.get("client_total") or {})
        if tot.get("base"):
            body(f"Combined with the SEO campaign above, a base engagement is "
                 f"{_p_money(tot.get('base'))} per month, intermediate "
                 f"{_p_money(tot.get('intermediate'))}, advanced "
                 f"{_p_money(tot.get('advanced'))}.", True)

    # ---- performance-based SEO, when the client qualifies -------------------
    # A GATE, NOT A CHOICE. Brendan only offers this where the terms are already
    # inside the first five pages; from scratch the first rankings are 6-12+
    # months out and there is nothing to bill against. The section is absent
    # rather than shown empty, and the operator sees the reason in step 4.
    # OURS TO ADD, NOT THE PARTNER'S TO REQUEST. Pay for performance never
    # arrives on an order, so it is not read off the strategy field at all: the
    # gate decides and the operator can veto. perf_on False means off whatever
    # the numbers say; True or absent both mean "follow the gate", so a quote
    # saved before any of this existed behaves like a new one.
    _notes = _notes if isinstance(_notes, dict) else {}
    perf_omitted = "not requested"
    if d.get("perf_on") is not False:
        elig = perf_eligibility(d.get("table") or [],
                            site_rebuild=(d.get("site_rebuild") or ""))
        # WHAT THE PANEL SHOWED IS WHAT THE DOCUMENT CARRIES. The table used to
        # be recomputed here, which meant a SECOND live bid lookup deciding a
        # client document: if that one answered differently — or not at all —
        # the rows floored, the total fell under the minimum and the section
        # removed itself with nothing on screen to say so. The panel's rows were
        # server-computed moments earlier; use them, and only recompute when
        # they are absent. (2026-08-22)
        prows = [dict(r) for r in (d.get("perf_rows") or [])
                 if isinstance(r, dict) and r.get("kw") and r.get("page1")]
        # THE ROWS ARE TRUSTED FOR PRICES, NOT FOR EVERYTHING. Sending the
        # panel's rows is what stops a second live lookup deciding whether the
        # section exists — but they are also SAVED with the quote and restored,
        # so a row computed by an older build carried its empty Practice Area
        # into every document built since. Six of them. The column is cheap and
        # deterministic, so re-derive it here on every build and let the money
        # be the only thing the panel gets to pin. (2026-08-23)
        if prows:
            _tp = perf_topics(d)
            for _r in prows:
                if not str(_r.get("area") or "").strip():
                    _r["area"] = perf_area(_r.get("kw") or "", _tp)
        if not prows and elig.get("eligible"):
            _pd, _ = _perf_fill_bids(d, True)
            prows = _perf_rows(_pd)
        total = sum(int(r.get("page1") or 0) for r in prows)
        floor = int(CFG.get("perf_min_monthly_value", 10000) or 0)
        perf_omitted = ("" if (prows and total >= floor and elig.get("eligible"))
                        else (elig.get("reason") or "")
                        if not elig.get("eligible")
                        else "no priced terms" if not prows
                        else f"potential value {_p_money(total)} is under the "
                             f"{_p_money(floor)} minimum")
        if prows and total >= floor and elig.get("eligible"):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            head(P["perf_heading"])
            _fmt = {"tail": int(CFG.get("perf_tail_months", 6)),
                    "term": int(CFG.get("perf_initial_term_months", 12)),
                    "min": _p_money(floor)}
            for para in P["perf_intro"]:
                body(para.format(**_fmt))
            cols = P["perf_columns"]
            tb = doc.add_table(rows=1, cols=len(cols))
            tb.style = "Light Grid Accent 1"
            for i, label in enumerate(cols):
                tb.rows[0].cells[i].text = ""
                rr = tb.rows[0].cells[i].paragraphs[0].add_run(label)
                rr.bold = True
                rr.font.size = Pt(8)
            for r in prows:
                c = tb.add_row().cells
                vals = [r.get("kw", ""), r.get("area", ""), r.get("tier", ""),
                        r.get("rank", ""),
                        _p_money(r.get("page1")), _p_money(r.get("top5")),
                        _p_money(r.get("top3")), _p_money(r.get("one")),
                        r.get("achieved", "")]
                for i, v in enumerate(vals):
                    c[i].text = ""
                    rr = c[i].paragraphs[0].add_run(str(v))
                    rr.font.size = Pt(8)
            doc.add_paragraph()

    # ---- case studies -----------------------------------------------------
    head(P["case_heading"])
    body(P["closing"][1])
    body(P["case_projects_label"], True)
    for c in P["case_studies"]:
        bullet(c)
    body(P["case_closing"])

    _notes["perf_omitted"] = perf_omitted
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/api/proposal.docx", methods=["POST"])
def api_proposal_docx():
    """The quote as an SSG-shaped Word document."""
    d = request.get_json(force=True) or {}
    notes = {}
    try:
        buf = build_proposal_docx(d, notes)
    except ImportError:
        return jsonify({"error": "python-docx is not installed on this server."}), 500
    except Exception as e:                                    # noqa: BLE001
        app.logger.exception("proposal docx failed")
        return jsonify({"error": str(e)[:200]}), 500
    name = re.sub(r"[^A-Za-z0-9]+", "_",
                  (d.get("brand") or "proposal")).strip("_") or "proposal"
    resp = send_file(buf, as_attachment=True,
                     download_name=f"{name}_SEO_Proposal.docx",
                     mimetype="application/vnd.openxmlformats-officedocument."
                              "wordprocessingml.document")
    # A SECTION THAT REMOVES ITSELF HAS TO SAY SO. The performance table
    # dropped out of two built documents while the panel showed it on screen,
    # and nothing anywhere reported the decision. The header rides back with
    # the file and the panel prints it. (2026-08-22)
    if notes.get("perf_omitted"):
        resp.headers["X-Perf-Omitted"] = str(notes["perf_omitted"])[:200]
    return resp
